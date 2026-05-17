from __future__ import annotations

import argparse
import importlib
import importlib.metadata
import json
import logging
import math
import os
import platform
import random
import sys
import warnings
from dataclasses import dataclass
from io import BytesIO
from typing import Any, Callable, Dict, Iterable, List, Optional, Sequence, Tuple

import matplotlib
matplotlib.use("Agg")
import matplotlib.image as mpimg
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from scipy import stats
from scipy.stats import norm
from sklearn.base import BaseEstimator, TransformerMixin
from sklearn.exceptions import ConvergenceWarning
from sklearn.experimental import enable_iterative_imputer
from sklearn.impute import IterativeImputer
from sklearn.linear_model import BayesianRidge
from sklearn.model_selection import GridSearchCV, StratifiedKFold, StratifiedShuffleSplit
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import StandardScaler
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

warnings.filterwarnings("ignore", category=ConvergenceWarning)
LOGGER = logging.getLogger("survival_reanalysis_v3")

STUDY_ID_COL = "studyid"
EVENT_COL = "event_status"
TIME_COL = "time_to_event"
SEX_COL = "sex"

PREDICTOR_COLUMNS: List[str] = [
    "age",
    "sex",
    "smoke",
    "fhxcvd",
    "dmmed",
    "htnmed",
    "dldmed",
    "sbp",
    "dbp",
    "bmi",
    "chol",
    "hdl",
    "trig",
    "hba1c",
    "egfr",
]

CONTINUOUS_COLUMNS: List[str] = [
    "age",
    "sbp",
    "dbp",
    "bmi",
    "chol",
    "hdl",
    "trig",
    "hba1c",
    "egfr",
]

BINARY_COLUMNS: List[str] = [
    "sex",
    "smoke",
    "fhxcvd",
    "dmmed",
    "htnmed",
    "dldmed",
]

DISPLAY_NAMES: Dict[str, str] = {
    "age": "Age",
    "sex": "Sex",
    "smoke": "History of smoking",
    "fhxcvd": "Family history of CVD",
    "dmmed": "Glucose-lowering medication",
    "htnmed": "Blood pressure-lowering medication",
    "dldmed": "Lipid-lowering medication",
    "sbp": "SBP",
    "dbp": "DBP",
    "bmi": "BMI",
    "chol": "TC",
    "hdl": "HDL-C",
    "trig": "TG",
    "hba1c": "HbA1c",
    "egfr": "eGFR",
    EVENT_COL: "ASCVD event",
    TIME_COL: "Follow-up time (months)",
}

SHAP_DISPLAY_NAMES: Dict[str, str] = {
    "age": "Age",
    "sex": "Sex",
    "smoke": "Smoking history",
    "fhxcvd": "Family history of CVD",
    "dmmed": "Glucose-lowering medication",
    "htnmed": "BP-lowering medication",
    "dldmed": "Lipid-lowering medication",
    "sbp": "SBP",
    "dbp": "DBP",
    "bmi": "BMI",
    "chol": "TC",
    "hdl": "HDL-C",
    "trig": "TG",
    "hba1c": "HbA1c",
    "egfr": "eGFR",
}

TABLE1_LABELS: Dict[str, str] = {
    "age": "Age (years)",
    "sex": "Sex (Women)",
    "smoke": "History of smoking",
    "fhxcvd": "Family history of CVD",
    "dmmed": "Glucose-lowering medications",
    "htnmed": "Blood pressure-lowering medications",
    "dldmed": "Lipid-lowering medications",
    "sbp": "SBP (mmHg)",
    "dbp": "DBP (mmHg)",
    "bmi": "BMI (kg/m2)",
    "chol": "TC (mmol/L)",
    "hdl": "HDL-C (mmol/L)",
    "trig": "TG (mmol/L)",
    "hba1c": "HbA1c (%)",
    "egfr": "eGFR (mL/min/1.73m2)",
}

POSITIVE_LEVELS: Dict[str, int] = {
    "sex": 0,
    "smoke": 1,
    "fhxcvd": 1,
    "dmmed": 1,
    "htnmed": 1,
    "dldmed": 1,
    EVENT_COL: 1,
}

DEFAULT_CSV_PATH = "data.csv"
DEFAULT_OUTDIR = "outputs"

MODEL_ORDER: List[str] = [
    "CoxPH",
    "ElasticNetCox",
    "RSF",
    "GBSA",
    "DeepSurv",
    "CoxTime",
    "XGBoost",
    "SVM",
]

CALIBRATION_MODELS = ["CoxPH", "ElasticNetCox", "RSF", "GBSA", "DeepSurv", "CoxTime", "XGBoost"]
SHAP_MODELS = MODEL_ORDER.copy()

MODEL_DESCRIPTIONS: Dict[str, str] = {
    "CoxPH": "Standard Cox model.",
    "ElasticNetCox": "Regularized Cox model.",
    "RSF": "Random survival forest.",
    "GBSA": "Boosted survival trees.",
    "DeepSurv": "Neural Cox model.",
    "CoxTime": "Neural Cox-Time model.",
    "XGBoost": "XGBoost survival model.",
    "SVM": "Survival SVM.",
}

@dataclass
class FittedModelBundle:
    model_name: str
    fitted_model: Any
    best_params: Dict[str, Any]
    best_cv_score: Optional[float]
    metrics: Dict[str, Any]
    calibration_table: Optional[pd.DataFrame] = None
    pred_event_prob_horizon: Optional[np.ndarray] = None

class TabularSurvivalPreprocessor(BaseEstimator, TransformerMixin):
    def __init__(
        self,
        continuous_columns: Sequence[str],
        binary_columns: Sequence[str],
        scale_continuous: bool = False,
    ) -> None:
        self.continuous_columns = continuous_columns
        self.binary_columns = binary_columns
        self.scale_continuous = scale_continuous

    def fit(self, X: pd.DataFrame, y: Optional[Any] = None) -> "TabularSurvivalPreprocessor":
        self.columns_ = list(self.continuous_columns) + list(self.binary_columns)
        if self.scale_continuous:
            self.scaler_ = StandardScaler()
            self.scaler_.fit(X[list(self.continuous_columns)].astype(float))
        return self

    def transform(self, X: pd.DataFrame) -> pd.DataFrame:
        out = X[self.columns_].copy()
        out[list(self.continuous_columns)] = out[list(self.continuous_columns)].astype(float)
        out[list(self.binary_columns)] = out[list(self.binary_columns)].astype(float)
        if self.scale_continuous:
            out.loc[:, list(self.continuous_columns)] = self.scaler_.transform(
                out[list(self.continuous_columns)].astype(float)
            )
        return out

    def get_feature_names_out(self, input_features: Optional[Sequence[str]] = None) -> np.ndarray:
        return np.asarray(self.columns_, dtype=object)

class SimpleStepFunction:
    def __init__(self, x: np.ndarray, y: np.ndarray):
        self.x = np.asarray(x, dtype=float)
        self.y = np.asarray(y, dtype=float)

    def __call__(self, t: float) -> float:
        idx = np.searchsorted(self.x, t, side="right") - 1
        if idx < 0:
            return 1.0
        idx = min(idx, len(self.y) - 1)
        return float(self.y[idx])

class BreslowBaselineEstimator:
    def __init__(self) -> None:
        self.event_times_: Optional[np.ndarray] = None
        self.cumhaz_: Optional[np.ndarray] = None

    def fit(self, durations: np.ndarray, events: np.ndarray, log_risk: np.ndarray) -> "BreslowBaselineEstimator":
        durations = np.asarray(durations, dtype=float)
        events = np.asarray(events, dtype=int)
        log_risk = np.asarray(log_risk, dtype=float)
        risk = np.exp(log_risk)
        unique_event_times = np.sort(np.unique(durations[events == 1]))
        cumhaz = []
        running = 0.0
        for t in unique_event_times:
            at_risk = risk[durations >= t].sum()
            d_t = int(((durations == t) & (events == 1)).sum())
            if at_risk <= 0:
                increment = 0.0
            else:
                increment = d_t / at_risk
            running += increment
            cumhaz.append(running)
        self.event_times_ = unique_event_times
        self.cumhaz_ = np.asarray(cumhaz, dtype=float)
        return self

    def cumulative_hazard_at(self, times: np.ndarray) -> np.ndarray:
        if self.event_times_ is None or self.cumhaz_ is None:
            raise RuntimeError("Baseline estimator must be fit before prediction.")
        times = np.asarray(times, dtype=float)
        idx = np.searchsorted(self.event_times_, times, side="right") - 1
        out = np.zeros_like(times, dtype=float)
        valid = idx >= 0
        out[valid] = self.cumhaz_[idx[valid]]
        return out

    def survival_matrix(self, log_risk: np.ndarray, eval_times: np.ndarray) -> np.ndarray:
        base_haz = self.cumulative_hazard_at(eval_times)
        risk = np.exp(np.asarray(log_risk, dtype=float)).reshape(-1, 1)
        return np.exp(-risk * base_haz.reshape(1, -1))

    def survival_functions(self, log_risk: np.ndarray) -> List[SimpleStepFunction]:
        if self.event_times_ is None:
            raise RuntimeError("Baseline estimator must be fit before prediction.")
        surv = self.survival_matrix(log_risk, self.event_times_)
        return [SimpleStepFunction(self.event_times_, surv[i, :]) for i in range(surv.shape[0])]

class LifelinesElasticNetCox(BaseEstimator):
    def __init__(self, penalizer: float = 0.01, l1_ratio: float = 0.5) -> None:
        self.penalizer = penalizer
        self.l1_ratio = l1_ratio

    def fit(self, X: pd.DataFrame, y: Any) -> "LifelinesElasticNetCox":
        from lifelines import CoxPHFitter
        event_name, time_name = y.dtype.names
        df = X.copy()
        df[TIME_COL] = y[time_name].astype(float)
        df[EVENT_COL] = y[event_name].astype(int)
        self.model_ = CoxPHFitter(penalizer=self.penalizer, l1_ratio=self.l1_ratio)
        self.model_.fit(df, duration_col=TIME_COL, event_col=EVENT_COL)
        self.feature_names_in_ = list(X.columns)
        return self

    def predict(self, X: pd.DataFrame) -> np.ndarray:
        partial_hazard = self.model_.predict_partial_hazard(X[self.feature_names_in_])
        return np.log(np.asarray(partial_hazard, dtype=float).reshape(-1) + 1e-12)

    def predict_survival_function(self, X: pd.DataFrame) -> List[SimpleStepFunction]:
        surv_df = self.model_.predict_survival_function(X[self.feature_names_in_])
        times = surv_df.index.to_numpy(dtype=float)
        return [SimpleStepFunction(times, surv_df.iloc[:, i].to_numpy(dtype=float)) for i in range(surv_df.shape[1])]

class XGBoostCoxWrapper(BaseEstimator):
    def __init__(
        self,
        n_estimators: int = 300,
        learning_rate: float = 0.05,
        max_depth: int = 3,
        subsample: float = 0.8,
        colsample_bytree: float = 0.8,
        min_child_weight: float = 1.0,
        reg_lambda: float = 1.0,
        random_state: int = 0,
        n_jobs: int = 1,
    ) -> None:
        self.n_estimators = n_estimators
        self.learning_rate = learning_rate
        self.max_depth = max_depth
        self.subsample = subsample
        self.colsample_bytree = colsample_bytree
        self.min_child_weight = min_child_weight
        self.reg_lambda = reg_lambda
        self.random_state = random_state
        self.n_jobs = n_jobs

    def fit(self, X: pd.DataFrame, y: Any) -> "XGBoostCoxWrapper":
        import xgboost as xgb
        event_name, time_name = y.dtype.names
        durations = y[time_name].astype(float)
        events = y[event_name].astype(int)
        labels = np.where(events == 1, durations, -durations)
        self.feature_names_in_ = list(X.columns)
        self.model_ = xgb.XGBRegressor(
            objective="survival:cox",
            eval_metric="cox-nloglik",
            n_estimators=self.n_estimators,
            learning_rate=self.learning_rate,
            max_depth=self.max_depth,
            subsample=self.subsample,
            colsample_bytree=self.colsample_bytree,
            min_child_weight=self.min_child_weight,
            reg_lambda=self.reg_lambda,
            random_state=self.random_state,
            n_jobs=self.n_jobs,
            tree_method="hist",
        )
        self.model_.fit(X[self.feature_names_in_].to_numpy(dtype=float), labels)
        pred_hr = np.asarray(self.model_.predict(X[self.feature_names_in_].to_numpy(dtype=float)), dtype=float).reshape(-1)
        log_risk = np.log(np.clip(pred_hr, 1e-12, None))
        self.baseline_ = BreslowBaselineEstimator().fit(durations=durations, events=events, log_risk=log_risk)
        return self

    def predict(self, X: pd.DataFrame) -> np.ndarray:
        pred_hr = np.asarray(self.model_.predict(X[self.feature_names_in_].to_numpy(dtype=float)), dtype=float).reshape(-1)
        return np.log(np.clip(pred_hr, 1e-12, None))

    def predict_survival_function(self, X: pd.DataFrame) -> List[SimpleStepFunction]:
        log_risk = self.predict(X)
        return self.baseline_.survival_functions(log_risk)

def _clean_pycox_surv_df(surv_df: pd.DataFrame) -> pd.DataFrame:
    surv_df = surv_df.copy()
    surv_df.index = pd.Index(np.asarray(surv_df.index, dtype=float), name="time")
    surv_df = surv_df[~surv_df.index.duplicated(keep="last")].sort_index()
    return surv_df

def _risk_from_surv_df(surv_df: pd.DataFrame, time_point: float) -> np.ndarray:
    surv_df = _clean_pycox_surv_df(surv_df)
    times = surv_df.index.to_numpy(dtype=float)
    surv_mat = surv_df.to_numpy(dtype=float)
    pos = int(np.searchsorted(times, float(time_point), side="right") - 1)
    if pos < 0:
        surv = np.ones(surv_mat.shape[1], dtype=float)
    else:
        surv = surv_mat[pos, :]
    return 1.0 - np.asarray(surv, dtype=float).reshape(-1)

class PyCoxAdapter:
    def __init__(self, model: Any, preprocessor: TabularSurvivalPreprocessor, feature_names: Sequence[str],
                 predict_from_surv: bool = False, risk_horizon: float = 120.0) -> None:
        self.model = model
        self.preprocessor = preprocessor
        self.feature_names = list(feature_names)
        self.predict_from_surv = predict_from_surv
        self.risk_horizon = float(risk_horizon)

    def predict(self, X: Any) -> np.ndarray:
        X_df = as_dataframe(X, self.feature_names)
        x_mat = self.preprocessor.transform(X_df).to_numpy(dtype="float32")
        if self.predict_from_surv:
            surv_df = self.model.predict_surv_df(x_mat)
            return _risk_from_surv_df(surv_df, self.risk_horizon)
        preds = self.model.predict(x_mat)
        return np.asarray(preds, dtype=float).reshape(-1)

    def predict_survival_matrix(self, X: Any, eval_times: np.ndarray) -> np.ndarray:
        X_df = as_dataframe(X, self.feature_names)
        x_mat = self.preprocessor.transform(X_df).to_numpy(dtype="float32")
        surv_df = _clean_pycox_surv_df(self.model.predict_surv_df(x_mat))
        augmented_index = surv_df.index.union(pd.Index(eval_times)).sort_values()
        augmented = surv_df.reindex(augmented_index).ffill().fillna(1.0)
        return augmented.loc[eval_times].to_numpy().T

    def predict_survival_function(self, X: Any) -> List[SimpleStepFunction]:
        X_df = as_dataframe(X, self.feature_names)
        x_mat = self.preprocessor.transform(X_df).to_numpy(dtype="float32")
        surv_df = _clean_pycox_surv_df(self.model.predict_surv_df(x_mat))
        times = surv_df.index.to_numpy(dtype=float)
        return [SimpleStepFunction(times, surv_df.iloc[:, i].to_numpy(dtype=float)) for i in range(surv_df.shape[1])]

def setup_logging(log_path: str) -> None:
    os.makedirs(os.path.dirname(log_path), exist_ok=True)
    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s | %(levelname)s | %(message)s",
        handlers=[logging.FileHandler(log_path, encoding="utf-8"), logging.StreamHandler(sys.stdout)],
    )

def set_global_seed(seed: int) -> None:
    random.seed(seed)
    np.random.seed(seed)
    try:
        import torch
        torch.manual_seed(seed)
        if torch.cuda.is_available():
            torch.cuda.manual_seed_all(seed)
        try:
            torch.use_deterministic_algorithms(True)
        except Exception:
            pass
    except Exception:
        pass

def ensure_dir(path: str) -> str:
    os.makedirs(path, exist_ok=True)
    return path

def save_json(obj: Any, path: str) -> None:
    with open(path, "w", encoding="utf-8") as f:
        json.dump(obj, f, indent=2, default=str)

def as_dataframe(x: Any, columns: Sequence[str]) -> pd.DataFrame:
    if isinstance(x, pd.DataFrame):
        return x.loc[:, list(columns)].copy()
    if isinstance(x, np.ndarray):
        return pd.DataFrame(x, columns=list(columns))
    return pd.DataFrame(np.asarray(x), columns=list(columns))

def format_mean_sd(series: pd.Series, decimals: int = 1) -> str:
    return f"{series.mean():.{decimals}f} ({series.std(ddof=1):.{decimals}f})"

def format_median_iqr(series: pd.Series, decimals: int = 1) -> str:
    q1, q3 = series.quantile([0.25, 0.75])
    return f"{series.median():.{decimals}f} [{q1:.{decimals}f}, {q3:.{decimals}f}]"

def format_n_pct(series: pd.Series, positive_value: int = 1) -> str:
    n = int((series == positive_value).sum())
    pct = 100.0 * n / len(series)
    return f"{n} ({pct:.1f}%)"

def smd_continuous(x1: pd.Series, x2: pd.Series) -> float:
    x1 = x1.astype(float)
    x2 = x2.astype(float)
    v1 = x1.var(ddof=1)
    v2 = x2.var(ddof=1)
    pooled_sd = np.sqrt((v1 + v2) / 2.0)
    if pooled_sd == 0 or np.isnan(pooled_sd):
        return 0.0
    return float((x1.mean() - x2.mean()) / pooled_sd)

def smd_binary(x1: pd.Series, x2: pd.Series, positive_value: int = 1) -> float:
    p1 = (x1 == positive_value).mean()
    p2 = (x2 == positive_value).mean()
    denom = np.sqrt((p1 * (1 - p1) + p2 * (1 - p2)) / 2.0)
    if denom == 0 or np.isnan(denom):
        return 0.0
    return float((p1 - p2) / denom)

def get_package_versions() -> Dict[str, Optional[str]]:
    packages = [
        "numpy",
        "pandas",
        "scikit-learn",
        "scikit-survival",
        "lifelines",
        "xgboost",
        "shap",
        "matplotlib",
        "scipy",
        "torch",
        "torchtuples",
        "pycox",
        "python-docx",
        "Pillow",
    ]
    versions: Dict[str, Optional[str]] = {
        "python": platform.python_version(),
        "platform": platform.platform(),
    }
    for pkg in packages:
        try:
            versions[pkg] = importlib.metadata.version(pkg)
        except importlib.metadata.PackageNotFoundError:
            versions[pkg] = None
    return versions

def load_dataset(csv_path: str) -> pd.DataFrame:
    df = pd.read_csv(csv_path)
    df = df.copy()
    if TIME_COL not in df.columns or EVENT_COL not in df.columns:
        raise ValueError(f"Dataset must contain {TIME_COL!r} and {EVENT_COL!r}.")
    missing_predictors = [c for c in PREDICTOR_COLUMNS if c not in df.columns]
    if missing_predictors:
        raise ValueError(f"Missing predictor columns: {missing_predictors}")
    df[EVENT_COL] = df[EVENT_COL].astype(int)
    before = len(df)
    df = df[df[TIME_COL] > 0].reset_index(drop=True)
    dropped = before - len(df)
    if dropped > 0:
        LOGGER.info("Dropped %d rows with non-positive follow-up time.", dropped)
    return df

def create_missingness_table(df: pd.DataFrame) -> pd.DataFrame:
    out = pd.DataFrame({
        "Variable": df.columns,
        "n_missing": df.isna().sum().values,
        "pct_missing": 100.0 * df.isna().mean().values,
        "dtype": [str(df[c].dtype) for c in df.columns],
    })
    return out[out["n_missing"] > 0].sort_values(["pct_missing", "Variable"], ascending=[False, True]).reset_index(drop=True)

def make_strata(df: pd.DataFrame) -> pd.Series:
    candidate = df[EVENT_COL].astype(str) + "_" + df[SEX_COL].astype(str)
    if candidate.value_counts().min() >= 2:
        return candidate
    return df[EVENT_COL].astype(str)

def primary_split(df: pd.DataFrame, seed: int, test_size: float) -> Tuple[np.ndarray, np.ndarray]:
    splitter = StratifiedShuffleSplit(n_splits=1, test_size=test_size, random_state=seed)
    strata = make_strata(df)
    train_idx, test_idx = next(splitter.split(df, strata))
    return train_idx, test_idx

def get_surv_array(df: pd.DataFrame) -> Any:
    from sksurv.util import Surv
    try:
        return Surv.from_arrays(
            event=df[EVENT_COL].astype(bool).to_numpy(),
            time=df[TIME_COL].astype(float).to_numpy(),
            name_event="event",
            name_time="time",
        )
    except TypeError:
        return Surv.from_arrays(
            event=df[EVENT_COL].astype(bool).to_numpy(),
            time=df[TIME_COL].astype(float).to_numpy(),
        )

def get_structured_field_names(y: Any) -> Tuple[str, str]:
    event_name, time_name = y.dtype.names
    return event_name, time_name

def harrell_c_index(y_true: Any, risk_scores: np.ndarray) -> float:
    from sksurv.metrics import concordance_index_censored
    event_name, time_name = get_structured_field_names(y_true)
    return float(concordance_index_censored(y_true[event_name], y_true[time_name], np.asarray(risk_scores, dtype=float))[0])

def uno_c_index(y_train: Any, y_test: Any, risk_scores: np.ndarray, tau: float) -> float:
    from sksurv.metrics import concordance_index_ipcw
    risk_scores = np.asarray(risk_scores, dtype=float)
    try:
        return float(concordance_index_ipcw(y_train, y_test, risk_scores, tau=tau)[0])
    except Exception:
        return np.nan

def get_eval_times(train_df: pd.DataFrame, test_df: pd.DataFrame, start_month: float, horizon_months: float) -> np.ndarray:
    upper = min(horizon_months, float(train_df[TIME_COL].max()), float(test_df[TIME_COL].max()))
    upper = float(np.floor(upper))
    if upper < start_month:
        raise ValueError(f"Requested start_month={start_month} exceeds usable upper bound {upper}.")
    return np.arange(start_month, upper + 1.0, 1.0, dtype=float)

def predict_survival_matrix(fitted_model: Any, X: pd.DataFrame, eval_times: np.ndarray) -> np.ndarray:
    if hasattr(fitted_model, "predict_survival_matrix"):
        return np.asarray(fitted_model.predict_survival_matrix(X, eval_times), dtype=float)
    surv_fns = fitted_model.predict_survival_function(X)
    return np.asarray([[fn(t) for t in eval_times] for fn in surv_fns], dtype=float)

def compute_brier_metrics(y_train: Any, y_test: Any, surv_matrix: np.ndarray, eval_times: np.ndarray, horizon_months: float) -> Dict[str, float]:
    from sksurv.metrics import brier_score, integrated_brier_score
    ibs = float(integrated_brier_score(y_train, y_test, surv_matrix, eval_times))
    idx = int(np.argmin(np.abs(eval_times - horizon_months)))
    horizon_used = float(eval_times[idx])
    _, bs = brier_score(y_train, y_test, surv_matrix[:, [idx]], np.asarray([horizon_used], dtype=float))
    return {
        "brier_horizon_months": horizon_used,
        "brier_score_at_horizon": float(bs[0]),
        "integrated_brier_score": ibs,
    }

def km_risk_at_time(durations: Sequence[float], events: Sequence[int], time_point: float) -> float:
    from lifelines import KaplanMeierFitter
    kmf = KaplanMeierFitter()
    kmf.fit(np.asarray(durations, dtype=float), event_observed=np.asarray(events, dtype=int))
    surv = kmf.predict(time_point)
    if hasattr(surv, "iloc"):
        surv = float(surv.iloc[0])
    return float(1.0 - float(surv))

def grouped_calibration_table(test_df: pd.DataFrame, predicted_event_prob: np.ndarray, time_point: float, n_groups: int = 5) -> pd.DataFrame:
    cal_df = pd.DataFrame({
        "predicted_event_prob": np.asarray(predicted_event_prob, dtype=float),
        TIME_COL: test_df[TIME_COL].astype(float).values,
        EVENT_COL: test_df[EVENT_COL].astype(int).values,
    })
    n_groups = max(2, min(n_groups, cal_df.shape[0]))
    cal_df["group"] = pd.qcut(
        cal_df["predicted_event_prob"].rank(method="first"),
        q=n_groups,
        labels=False,
        duplicates="drop",
    )
    rows: List[Dict[str, Any]] = []
    for grp, grp_df in cal_df.groupby("group"):
        observed = km_risk_at_time(grp_df[TIME_COL].values, grp_df[EVENT_COL].values, time_point=time_point)
        rows.append({
            "group": int(grp) + 1,
            "n": int(len(grp_df)),
            "events": int(grp_df[EVENT_COL].sum()),
            "mean_predicted_risk": float(grp_df["predicted_event_prob"].mean()),
            "observed_km_risk": float(observed),
        })
    return pd.DataFrame(rows).sort_values("group").reset_index(drop=True)

def plot_calibration_panel(ax: Any, calibration_df: pd.DataFrame, model_name: str, time_point: float) -> None:
    ax.plot([0, 1], [0, 1], linestyle="--", linewidth=1)
    ax.plot(calibration_df["mean_predicted_risk"], calibration_df["observed_km_risk"], marker="o", linewidth=1)
    for _, row in calibration_df.iterrows():
        ax.annotate(str(int(row["group"])), (row["mean_predicted_risk"], row["observed_km_risk"]),
                    textcoords="offset points", xytext=(4, 4), fontsize=8)
    ax.set_title(model_name, fontsize=11)
    ax.set_xlabel(f"Predicted {int(time_point)}-month risk", fontsize=9)
    ax.set_ylabel(f"Observed KM risk", fontsize=9)
    ax.set_xlim(0, max(0.05, calibration_df[["mean_predicted_risk", "observed_km_risk"]].to_numpy().max() * 1.15))
    ax.set_ylim(0, max(0.05, calibration_df[["mean_predicted_risk", "observed_km_risk"]].to_numpy().max() * 1.15))

def calibration_slope_from_probs(test_df: pd.DataFrame, predicted_event_prob: np.ndarray) -> Dict[str, float]:
    from lifelines import CoxPHFitter
    probs = np.asarray(predicted_event_prob, dtype=float)
    probs = np.clip(probs, 1e-6, 1 - 1e-6)
    lp = np.log(-np.log(1.0 - probs))
    cal_df = pd.DataFrame({TIME_COL: test_df[TIME_COL].astype(float).values, EVENT_COL: test_df[EVENT_COL].astype(int).values, "lp": lp})
    cph = CoxPHFitter()
    cph.fit(cal_df, duration_col=TIME_COL, event_col=EVENT_COL)
    slope = float(cph.params_["lp"])
    ci_low, ci_high = map(float, cph.confidence_intervals_.loc["lp"].values.tolist())
    p_value = float(cph.summary.loc["lp", "p"])
    return {
        "calibration_slope": slope,
        "calibration_slope_ci_low": ci_low,
        "calibration_slope_ci_high": ci_high,
        "calibration_slope_p": p_value,
    }

def format_mean_ci(series: pd.Series, decimals: int = 3) -> str:
    s = series.dropna()
    if s.empty:
        return "NA"
    low, high = np.percentile(s, [2.5, 97.5])
    return f"{s.mean():.{decimals}f} ({low:.{decimals}f} to {high:.{decimals}f})"

def summarise_performance(perf_df: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame]:
    summary_rows = []
    formatted_rows = []
    metric_map = {
        "Harrell_C": "harrell_c",
        "Uno_C_tau": "uno_c_tau",
        "Brier_120m": "brier_score_at_horizon",
        "IBS_12_120m": "integrated_brier_score",
        "Calibration_slope_120m": "calibration_slope",
    }
    for model_name, sub in perf_df.groupby("model"):
        row = {"model": model_name, "n_repeats": int(sub["repeat"].nunique())}
        fmt_row = {"Model": model_name}
        for label, col in metric_map.items():
            values = sub[col].dropna()
            if values.empty:
                row[f"{col}_mean"] = np.nan
                row[f"{col}_ci_low"] = np.nan
                row[f"{col}_ci_high"] = np.nan
                fmt_row[label] = "NA"
            else:
                low, high = np.percentile(values, [2.5, 97.5])
                row[f"{col}_mean"] = float(values.mean())
                row[f"{col}_ci_low"] = float(low)
                row[f"{col}_ci_high"] = float(high)
                fmt_row[label] = format_mean_ci(values)
        summary_rows.append(row)
        formatted_rows.append(fmt_row)
    summary_df = pd.DataFrame(summary_rows).set_index("model").loc[MODEL_ORDER].reset_index()
    formatted_df = pd.DataFrame(formatted_rows).set_index("Model").loc[MODEL_ORDER].reset_index()
    return summary_df, formatted_df

def compare_models_vs_reference(perf_df: pd.DataFrame, reference_model: str = "CoxPH") -> pd.DataFrame:
    rows: List[Dict[str, Any]] = []
    metrics = ["harrell_c", "uno_c_tau", "brier_score_at_horizon", "integrated_brier_score", "calibration_slope"]
    models = [m for m in MODEL_ORDER if m != reference_model and m in perf_df["model"].unique()]
    for metric in metrics:
        wide = perf_df.pivot(index="repeat", columns="model", values=metric)
        if reference_model not in wide.columns:
            continue
        for model_name in models:
            if model_name not in wide.columns:
                continue
            diff = (wide[model_name] - wide[reference_model]).dropna()
            if diff.empty:
                continue
            low, high = np.percentile(diff, [2.5, 97.5])
            rows.append({
                "Model": model_name,
                "Reference_model": reference_model,
                "Metric": metric,
                "Mean_difference": float(diff.mean()),
                "CI_low": float(low),
                "CI_high": float(high),
                "n_repeats": int(diff.shape[0]),
            })
    out = pd.DataFrame(rows)
    if not out.empty:
        metric_order = {m: i for i, m in enumerate(metrics)}
        out["_metric_order"] = out["Metric"].map(metric_order)
        out["_model_order"] = out["Model"].map({m: i for i, m in enumerate(MODEL_ORDER)})
        out = out.sort_values(["_metric_order", "_model_order"]).drop(columns=["_metric_order", "_model_order"]).reset_index(drop=True)
    return out

def summarise_tuning_records(tuning_df: pd.DataFrame) -> pd.DataFrame:
    if tuning_df.empty:
        return tuning_df
    summary = (
        tuning_df.groupby(["model", "best_params_json"], dropna=False)
        .size()
        .reset_index(name="n_times_selected")
        .sort_values(["model", "n_times_selected"], ascending=[True, False])
        .reset_index(drop=True)
    )
    return summary

def make_harrell_scorer() -> Callable[[Any, pd.DataFrame, Any], float]:
    def _scorer(estimator: Any, X: pd.DataFrame, y: Any) -> float:
        risk = estimator.predict(X)
        return harrell_c_index(y, risk)
    return _scorer

def nelson_aalen_feature(df: pd.DataFrame) -> np.ndarray:
    data = df[[TIME_COL, EVENT_COL]].copy().astype({TIME_COL: float, EVENT_COL: int})
    event_times = np.sort(data.loc[data[EVENT_COL] == 1, TIME_COL].unique())
    if len(event_times) == 0:
        return np.zeros(len(df), dtype=float)
    cumulative = []
    running = 0.0
    for t in event_times:
        d_t = int(((data[TIME_COL] == t) & (data[EVENT_COL] == 1)).sum())
        n_t = int((data[TIME_COL] >= t).sum())
        if n_t > 0:
            running += d_t / n_t
        cumulative.append(running)
    event_times = np.asarray(event_times, dtype=float)
    cumulative = np.asarray(cumulative, dtype=float)
    row_times = data[TIME_COL].to_numpy(dtype=float)
    idx = np.searchsorted(event_times, row_times, side="right") - 1
    out = np.zeros(len(df), dtype=float)
    valid = idx >= 0
    out[valid] = cumulative[idx[valid]]
    return out

def build_imputation_matrix(df: pd.DataFrame) -> pd.DataFrame:
    out = df[PREDICTOR_COLUMNS].copy()
    out["_aux_event"] = df[EVENT_COL].astype(float)
    out["_aux_time"] = np.log1p(df[TIME_COL].astype(float))
    out["_aux_nelson_aalen"] = nelson_aalen_feature(df)
    return out

def postprocess_imputed_predictors(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    for col in BINARY_COLUMNS:
        vals = np.asarray(out[col], dtype=float)
        out[col] = (vals >= 0.5).astype(int)
    lower_bounds = {
        "age": 18.0,
        "sbp": 40.0,
        "dbp": 20.0,
        "bmi": 10.0,
        "chol": 0.1,
        "hdl": 0.1,
        "trig": 0.05,
        "hba1c": 2.0,
        "egfr": 1.0,
    }
    for col, lb in lower_bounds.items():
        out[col] = np.clip(np.asarray(out[col], dtype=float), lb, None)
    return out

def generate_imputed_split_datasets(train_df: pd.DataFrame, test_df: pd.DataFrame, m: int, seed: int, max_iter: int) -> List[Dict[str, pd.DataFrame]]:
    train_matrix = build_imputation_matrix(train_df)
    test_matrix = build_imputation_matrix(test_df)
    cols = list(train_matrix.columns)
    out: List[Dict[str, pd.DataFrame]] = []
    for i in range(m):
        imputer = IterativeImputer(
            estimator=BayesianRidge(),
            max_iter=max_iter,
            sample_posterior=True,
            random_state=seed + i,
            initial_strategy="median",
            imputation_order="ascending",
        )
        train_completed = pd.DataFrame(imputer.fit_transform(train_matrix[cols]), columns=cols, index=train_df.index)
        test_completed = pd.DataFrame(imputer.transform(test_matrix[cols]), columns=cols, index=test_df.index)
        train_pred = postprocess_imputed_predictors(train_completed[PREDICTOR_COLUMNS])
        test_pred = postprocess_imputed_predictors(test_completed[PREDICTOR_COLUMNS])
        imputed_train = pd.concat([train_df[[STUDY_ID_COL, TIME_COL, EVENT_COL]].reset_index(drop=True), train_pred.reset_index(drop=True)], axis=1)
        imputed_test = pd.concat([test_df[[STUDY_ID_COL, TIME_COL, EVENT_COL]].reset_index(drop=True), test_pred.reset_index(drop=True)], axis=1)
        imputed_combined = pd.concat([imputed_train, imputed_test], axis=0, ignore_index=True)
        out.append({"train": imputed_train, "test": imputed_test, "combined": imputed_combined})
    return out

def aggregate_table1_over_imputations(imputed_splits: List[Dict[str, pd.DataFrame]], raw_train_df: pd.DataFrame, raw_test_df: pd.DataFrame, raw_df: pd.DataFrame) -> pd.DataFrame:
    per_imp_rows: List[pd.DataFrame] = []
    for split in imputed_splits:
        df = split["combined"]
        train_df = split["train"]
        test_df = split["test"]
        rows = []
        for col in PREDICTOR_COLUMNS:
            label = TABLE1_LABELS[col]
            if col in CONTINUOUS_COLUMNS:
                rows.append({
                    "Variable": label,
                    "Overall_mean": float(df[col].mean()),
                    "Overall_sd": float(df[col].std(ddof=1)),
                    "Training_mean": float(train_df[col].mean()),
                    "Training_sd": float(train_df[col].std(ddof=1)),
                    "Testing_mean": float(test_df[col].mean()),
                    "Testing_sd": float(test_df[col].std(ddof=1)),
                    "SMD_train_vs_test": float(abs(smd_continuous(train_df[col], test_df[col]))),
                    "is_binary": 0,
                })
            else:
                pos = POSITIVE_LEVELS.get(col, 1)
                rows.append({
                    "Variable": label,
                    "Overall_n": float((df[col] == pos).sum()),
                    "Overall_pct": float(100.0 * (df[col] == pos).mean()),
                    "Training_n": float((train_df[col] == pos).sum()),
                    "Training_pct": float(100.0 * (train_df[col] == pos).mean()),
                    "Testing_n": float((test_df[col] == pos).sum()),
                    "Testing_pct": float(100.0 * (test_df[col] == pos).mean()),
                    "SMD_train_vs_test": float(abs(smd_binary(train_df[col], test_df[col], positive_value=pos))),
                    "is_binary": 1,
                })
        per_imp_rows.append(pd.DataFrame(rows))

    combined = pd.concat(per_imp_rows, axis=0, ignore_index=True)
    out_rows = []
    for variable, sub in combined.groupby("Variable"):
        is_binary = int(sub["is_binary"].iloc[0])
        if is_binary == 0:
            out_rows.append({
                "Variable": variable,
                "Overall": f"{sub['Overall_mean'].mean():.1f} ({sub['Overall_sd'].mean():.1f})",
                "Training": f"{sub['Training_mean'].mean():.1f} ({sub['Training_sd'].mean():.1f})",
                "Testing": f"{sub['Testing_mean'].mean():.1f} ({sub['Testing_sd'].mean():.1f})",
                "SMD_train_vs_test": round(float(sub["SMD_train_vs_test"].mean()), 3),
            })
        else:
            out_rows.append({
                "Variable": variable,
                "Overall": f"{int(round(sub['Overall_n'].mean()))} ({sub['Overall_pct'].mean():.1f}%)",
                "Training": f"{int(round(sub['Training_n'].mean()))} ({sub['Training_pct'].mean():.1f}%)",
                "Testing": f"{int(round(sub['Testing_n'].mean()))} ({sub['Testing_pct'].mean():.1f}%)",
                "SMD_train_vs_test": round(float(sub["SMD_train_vs_test"].mean()), 3),
            })
    out_rows.extend([
        {
            "Variable": "ASCVD events, n (%)",
            "Overall": format_n_pct(raw_df[EVENT_COL], positive_value=1),
            "Training": format_n_pct(raw_train_df[EVENT_COL], positive_value=1),
            "Testing": format_n_pct(raw_test_df[EVENT_COL], positive_value=1),
            "SMD_train_vs_test": round(abs(smd_binary(raw_train_df[EVENT_COL], raw_test_df[EVENT_COL], positive_value=1)), 3),
        },
        {
            "Variable": "Duration (months), median [IQR]",
            "Overall": format_median_iqr(raw_df[TIME_COL]),
            "Training": format_median_iqr(raw_train_df[TIME_COL]),
            "Testing": format_median_iqr(raw_test_df[TIME_COL]),
            "SMD_train_vs_test": "",
        },
    ])
    out_df = pd.DataFrame(out_rows)
    order = [TABLE1_LABELS[c] for c in PREDICTOR_COLUMNS] + ["ASCVD events, n (%)", "Duration (months), median [IQR]"]
    out_df["_order"] = out_df["Variable"].map({v: i for i, v in enumerate(order)})
    out_df = out_df.sort_values("_order").drop(columns="_order").reset_index(drop=True)
    return out_df

def fit_lifelines_cox_and_ph_test(train_df: pd.DataFrame) -> Tuple[Any, pd.DataFrame]:
    from lifelines import CoxPHFitter
    from lifelines.statistics import proportional_hazard_test
    cox_df = train_df[PREDICTOR_COLUMNS + [TIME_COL, EVENT_COL]].copy()
    cph = CoxPHFitter()
    cph.fit(cox_df, duration_col=TIME_COL, event_col=EVENT_COL)
    summary = cph.summary.reset_index().rename(columns={"covariate": "variable"})
    summary["Variable"] = summary["variable"].map(DISPLAY_NAMES).fillna(summary["variable"])
    ph_test = proportional_hazard_test(cph, cox_df, time_transform="rank")
    ph_df = ph_test.summary.reset_index().rename(columns={"index": "variable"})
    ph_df = ph_df.rename(columns={"p": "ph_test_p", "test_statistic": "ph_test_statistic"})
    merged = summary.merge(ph_df[["variable", "ph_test_statistic", "ph_test_p"]], on="variable", how="left")
    out = merged[["Variable", "variable", "coef", "se(coef)", "z", "p", "ph_test_statistic", "ph_test_p"]].copy()
    out = out.rename(columns={"coef": "beta", "se(coef)": "SE_beta", "p": "p_value"})
    return cph, out

def pool_cox_results(primary_imputed_trains: List[pd.DataFrame]) -> pd.DataFrame:
    per_imp = []
    ph_ref: Optional[pd.DataFrame] = None
    for i, train_df in enumerate(primary_imputed_trains):
        _, tbl = fit_lifelines_cox_and_ph_test(train_df)
        tbl = tbl.copy()
        tbl["imputation"] = i + 1
        per_imp.append(tbl)
        if i == 0:
            ph_ref = tbl[["variable", "ph_test_statistic", "ph_test_p"]].copy()
    all_tbl = pd.concat(per_imp, axis=0, ignore_index=True)
    rows = []
    m = len(primary_imputed_trains)
    for variable, sub in all_tbl.groupby("variable"):
        q = sub["beta"].to_numpy(dtype=float)
        u = (sub["SE_beta"].to_numpy(dtype=float)) ** 2
        qbar = float(np.mean(q))
        ubar = float(np.mean(u))
        b = float(np.var(q, ddof=1)) if len(q) > 1 else 0.0
        total_var = ubar + (1.0 + 1.0 / max(m, 1)) * b
        se = float(np.sqrt(total_var))
        z = qbar / se if se > 0 else np.nan
        p = float(2 * (1 - norm.cdf(abs(z)))) if np.isfinite(z) else np.nan
        ci_low = qbar - 1.96 * se
        ci_high = qbar + 1.96 * se
        rows.append({
            "Variable": DISPLAY_NAMES.get(variable, variable),
            "variable": variable,
            "beta": qbar,
            "HR": float(np.exp(qbar)),
            "HR_95CI_lower": float(np.exp(ci_low)),
            "HR_95CI_upper": float(np.exp(ci_high)),
            "SE_beta": se,
            "z": z,
            "p_value": p,
        })
    pooled = pd.DataFrame(rows)
    if ph_ref is not None:
        pooled = pooled.merge(ph_ref, on="variable", how="left")
    pooled["_order"] = pooled["variable"].map({c: i for i, c in enumerate(PREDICTOR_COLUMNS)})
    pooled = pooled.sort_values("_order").drop(columns="_order").reset_index(drop=True)
    return pooled

def make_shap_outputs(model_name: str, fitted_model: Any, train_df: pd.DataFrame, test_df: pd.DataFrame, out_dir: str,
                      make_force_plot: bool = False, max_background: int = 100, max_explain: int = 150) -> Optional[pd.DataFrame]:
    try:
        import shap
    except Exception as exc:
        LOGGER.warning("Skipping SHAP for %s because shap is unavailable: %s", model_name, exc)
        return None

    X_train = train_df[PREDICTOR_COLUMNS].copy()
    X_test = test_df[PREDICTOR_COLUMNS].copy()
    feature_names = list(X_test.columns)
    shap_feature_names = [SHAP_DISPLAY_NAMES.get(c, c) for c in feature_names]
    reverse_name_map = {SHAP_DISPLAY_NAMES.get(c, c): c for c in feature_names}

    background_plot = X_train.sample(n=min(max_background, len(X_train)), random_state=0).copy()
    background_plot.columns = shap_feature_names
    explain_df_plot = X_test.iloc[: min(max_explain, len(X_test))].copy()
    explain_df_plot.columns = shap_feature_names

    def predict_fn(x: Any) -> np.ndarray:
        x_df = as_dataframe(x, shap_feature_names)
        x_df = x_df.rename(columns=reverse_name_map)
        x_df = x_df.loc[:, feature_names]
        preds = fitted_model.predict(x_df)
        return np.asarray(preds, dtype=float).reshape(-1)

    try:
        masker = shap.maskers.Independent(background_plot)
        explainer = shap.Explainer(predict_fn, masker=masker, feature_names=shap_feature_names, algorithm="permutation")
        shap_values = explainer(explain_df_plot, max_evals=max(2 * len(shap_feature_names) + 1, 31))
    except Exception as exc:
        LOGGER.warning("SHAP failed for %s: %s", model_name, exc)
        return None

    mean_abs = np.abs(shap_values.values).mean(axis=0)
    shap_table = pd.DataFrame({
        "variable": feature_names,
        "Variable": shap_feature_names,
        "mean_abs_shap": mean_abs,
    }).sort_values("mean_abs_shap", ascending=False).reset_index(drop=True)

    summary_png = os.path.join(out_dir, f"shap_summary_{model_name}.png")
    try:
        plt.figure()
        shap.summary_plot(shap_values, explain_df_plot, max_display=len(shap_feature_names), show=False)
        ax = plt.gca()
        ax.tick_params(axis="y", labelsize=9)
        ax.tick_params(axis="x", labelsize=9)
        ax.set_xlabel("SHAP value (impact on model output)", fontsize=10)
        plt.tight_layout()
        plt.savefig(summary_png, dpi=400, bbox_inches="tight")
        plt.close()
    except Exception as exc:
        LOGGER.warning("SHAP summary plot failed for %s: %s", model_name, exc)

    if make_force_plot and len(explain_df_plot) > 0:
        try:
            force_png = os.path.join(out_dir, f"shap_force_{model_name}.png")
            shap.force_plot(
                shap_values.base_values[0],
                shap_values.values[0],
                explain_df_plot.iloc[0],
                matplotlib=True,
                show=False,
                contribution_threshold=0.05,
                text_rotation=45,
            )
            plt.tight_layout()
            plt.savefig(force_png, dpi=400, bbox_inches="tight")
            plt.close()
        except Exception as exc:
            LOGGER.warning("SHAP force plot failed for %s: %s", model_name, exc)

    shap_table.to_csv(os.path.join(out_dir, f"shap_importance_{model_name}.csv"), index=False)
    return shap_table

def get_deepsurv_configs() -> List[Dict[str, Any]]:
    return [
        {"num_nodes": [32, 16], "dropout": 0.1, "lr": 1e-2, "batch_size": 64, "epochs": 256},
        {"num_nodes": [32, 16], "dropout": 0.2, "lr": 1e-2, "batch_size": 64, "epochs": 256},
        {"num_nodes": [32, 32], "dropout": 0.1, "lr": 1e-2, "batch_size": 64, "epochs": 256},
        {"num_nodes": [64, 32], "dropout": 0.2, "lr": 1e-3, "batch_size": 64, "epochs": 256},
        {"num_nodes": [64, 32], "dropout": 0.1, "lr": 1e-3, "batch_size": 128, "epochs": 256},
        {"num_nodes": [16, 16], "dropout": 0.2, "lr": 1e-2, "batch_size": 128, "epochs": 256},
    ]

def get_coxtime_configs() -> List[Dict[str, Any]]:
    return [
        {"num_nodes": [32, 16], "dropout": 0.1, "lr": 1e-2, "batch_size": 64, "epochs": 256},
        {"num_nodes": [32, 16], "dropout": 0.2, "lr": 1e-2, "batch_size": 64, "epochs": 256},
        {"num_nodes": [32, 32], "dropout": 0.1, "lr": 1e-2, "batch_size": 64, "epochs": 256},
        {"num_nodes": [64, 32], "dropout": 0.2, "lr": 1e-3, "batch_size": 64, "epochs": 256},
        {"num_nodes": [64, 32], "dropout": 0.1, "lr": 1e-3, "batch_size": 128, "epochs": 256},
        {"num_nodes": [16, 16], "dropout": 0.2, "lr": 1e-2, "batch_size": 128, "epochs": 256},
    ]

def fit_deepsurv(train_df: pd.DataFrame, random_state: int) -> Tuple[PyCoxAdapter, Dict[str, Any], float]:
    import torch
    import torchtuples as tt
    from pycox.models import CoxPH as PyCoxCoxPH

    set_global_seed(random_state)
    feature_names = list(PREDICTOR_COLUMNS)
    prep = TabularSurvivalPreprocessor(CONTINUOUS_COLUMNS, BINARY_COLUMNS, scale_continuous=True)
    prep.fit(train_df[feature_names])
    splitter = StratifiedShuffleSplit(n_splits=1, test_size=0.2, random_state=random_state)
    strata = make_strata(train_df)
    tr_idx, val_idx = next(splitter.split(train_df, strata))
    tr_df = train_df.iloc[tr_idx].copy()
    val_df = train_df.iloc[val_idx].copy()

    x_tr = prep.transform(tr_df[feature_names]).to_numpy(dtype="float32")
    x_val = prep.transform(val_df[feature_names]).to_numpy(dtype="float32")
    y_tr = (tr_df[TIME_COL].astype("float32").to_numpy(), tr_df[EVENT_COL].astype("int64").to_numpy())
    y_val = (val_df[TIME_COL].astype("float32").to_numpy(), val_df[EVENT_COL].astype("int64").to_numpy())
    y_val_surv = get_surv_array(val_df)

    best_model = None
    best_score = -np.inf
    best_cfg: Dict[str, Any] = {}
    best_adapter = None
    for i, cfg in enumerate(get_deepsurv_configs(), start=1):
        set_global_seed(random_state + i)
        net = tt.practical.MLPVanilla(
            in_features=x_tr.shape[1],
            num_nodes=cfg["num_nodes"],
            out_features=1,
            batch_norm=True,
            dropout=cfg["dropout"],
            output_bias=False,
        )
        model = PyCoxCoxPH(net, tt.optim.Adam)
        model.optimizer.set_lr(cfg["lr"])
        callbacks = [tt.callbacks.EarlyStopping()]
        model.fit(
            x_tr,
            y_tr,
            batch_size=cfg["batch_size"],
            epochs=cfg["epochs"],
            callbacks=callbacks,
            verbose=False,
            val_data=(x_val, y_val),
            val_batch_size=cfg["batch_size"],
        )
        model.compute_baseline_hazards()
        risk_val = _risk_from_surv_df(model.predict_surv_df(x_val), 120.0)
        score = harrell_c_index(y_val_surv, risk_val)
        if score > best_score:
            best_score = score
            best_model = model
            best_cfg = dict(cfg)
            best_adapter = PyCoxAdapter(model=best_model, preprocessor=prep, feature_names=feature_names)
    if best_adapter is None:
        raise RuntimeError("DeepSurv tuning did not produce a valid model.")
    return best_adapter, best_cfg, float(best_score)

def fit_coxtime(train_df: pd.DataFrame, random_state: int) -> Tuple[PyCoxAdapter, Dict[str, Any], float]:
    import torchtuples as tt
    from pycox.models import CoxTime
    from pycox.models.cox_time import MLPVanillaCoxTime

    set_global_seed(random_state)
    feature_names = list(PREDICTOR_COLUMNS)
    prep = TabularSurvivalPreprocessor(CONTINUOUS_COLUMNS, BINARY_COLUMNS, scale_continuous=True)
    prep.fit(train_df[feature_names])
    splitter = StratifiedShuffleSplit(n_splits=1, test_size=0.2, random_state=random_state)
    strata = make_strata(train_df)
    tr_idx, val_idx = next(splitter.split(train_df, strata))
    tr_df = train_df.iloc[tr_idx].copy()
    val_df = train_df.iloc[val_idx].copy()

    x_tr = prep.transform(tr_df[feature_names]).to_numpy(dtype="float32")
    x_val = prep.transform(val_df[feature_names]).to_numpy(dtype="float32")
    durations_tr = tr_df[TIME_COL].astype("float32").to_numpy()
    events_tr = tr_df[EVENT_COL].astype("float32").to_numpy()
    durations_val = val_df[TIME_COL].astype("float32").to_numpy()
    events_val = val_df[EVENT_COL].astype("float32").to_numpy()
    y_val_surv = get_surv_array(val_df)

    best_model = None
    best_score = -np.inf
    best_cfg: Dict[str, Any] = {}
    best_adapter = None
    for i, cfg in enumerate(get_coxtime_configs(), start=1):
        set_global_seed(random_state + 100 + i)
        labtrans = CoxTime.label_transform()
        y_tr = labtrans.fit_transform(durations_tr, events_tr)
        y_val = labtrans.transform(durations_val, events_val)
        net = MLPVanillaCoxTime(
            in_features=x_tr.shape[1],
            num_nodes=cfg["num_nodes"],
            batch_norm=True,
            dropout=cfg["dropout"],
        )
        model = CoxTime(net, tt.optim.Adam, labtrans=labtrans)
        model.optimizer.set_lr(cfg["lr"])
        callbacks = [tt.callbacks.EarlyStopping()]
        val_data = tt.tuplefy(x_val, y_val)
        model.fit(
            x_tr,
            y_tr,
            batch_size=cfg["batch_size"],
            epochs=cfg["epochs"],
            callbacks=callbacks,
            verbose=False,
            val_data=val_data,
            val_batch_size=cfg["batch_size"],
        )
        model.compute_baseline_hazards()
        risk_val = _risk_from_surv_df(model.predict_surv_df(x_val), 120.0)
        score = harrell_c_index(y_val_surv, risk_val)
        if score > best_score:
            best_score = score
            best_model = model
            best_cfg = dict(cfg)
            best_adapter = PyCoxAdapter(model=best_model, preprocessor=prep, feature_names=feature_names, predict_from_surv=True, risk_horizon=120.0)
    if best_adapter is None:
        raise RuntimeError("CoxTime tuning did not produce a valid model.")
    return best_adapter, best_cfg, float(best_score)

def get_classical_model_specs(random_state: int, n_jobs: int) -> Dict[str, Dict[str, Any]]:
    from sksurv.ensemble import GradientBoostingSurvivalAnalysis, RandomSurvivalForest
    from sksurv.linear_model import CoxPHSurvivalAnalysis
    from sksurv.svm import FastSurvivalSVM
    return {
        "CoxPH": {
            "pipeline": Pipeline([
                ("prep", TabularSurvivalPreprocessor(CONTINUOUS_COLUMNS, BINARY_COLUMNS, scale_continuous=False)),
                ("model", CoxPHSurvivalAnalysis()),
            ]),
            "param_grid": {
                "model__ties": ["breslow", "efron"],
                "model__tol": [1e-9, 1e-7],
                "model__n_iter": [100, 300],
            },
            "supports_survival_probs": True,
        },
        "ElasticNetCox": {
            "pipeline": Pipeline([
                ("prep", TabularSurvivalPreprocessor(CONTINUOUS_COLUMNS, BINARY_COLUMNS, scale_continuous=True)),
                ("model", LifelinesElasticNetCox()),
            ]),
            "param_grid": {
                "model__penalizer": [0.001, 0.01, 0.1, 1.0],
                "model__l1_ratio": [0.0, 0.5, 1.0],
            },
            "supports_survival_probs": True,
        },
        "RSF": {
            "pipeline": Pipeline([
                ("prep", TabularSurvivalPreprocessor(CONTINUOUS_COLUMNS, BINARY_COLUMNS, scale_continuous=False)),
                ("model", RandomSurvivalForest(random_state=random_state, n_jobs=n_jobs)),
            ]),
            "param_grid": {
                "model__n_estimators": [200, 500],
                "model__min_samples_split": [10, 20],
                "model__min_samples_leaf": [5, 10],
                "model__max_features": ["sqrt", 0.5],
                "model__max_depth": [None, 3],
            },
            "supports_survival_probs": True,
        },
        "GBSA": {
            "pipeline": Pipeline([
                ("prep", TabularSurvivalPreprocessor(CONTINUOUS_COLUMNS, BINARY_COLUMNS, scale_continuous=False)),
                ("model", GradientBoostingSurvivalAnalysis(random_state=random_state)),
            ]),
            "param_grid": {
                "model__learning_rate": [0.01, 0.05, 0.1],
                "model__n_estimators": [100, 300],
                "model__min_samples_leaf": [3, 10],
                "model__max_depth": [1, 2],
                "model__subsample": [0.7, 1.0],
            },
            "supports_survival_probs": True,
        },
        "XGBoost": {
            "pipeline": Pipeline([
                ("prep", TabularSurvivalPreprocessor(CONTINUOUS_COLUMNS, BINARY_COLUMNS, scale_continuous=False)),
                ("model", XGBoostCoxWrapper(random_state=random_state, n_jobs=n_jobs)),
            ]),
            "param_grid": {
                "model__n_estimators": [200, 500],
                "model__learning_rate": [0.01, 0.05, 0.1],
                "model__max_depth": [2, 3],
                "model__subsample": [0.7, 1.0],
                "model__colsample_bytree": [0.7, 1.0],
                "model__min_child_weight": [1.0, 5.0],
                "model__reg_lambda": [1.0, 5.0],
            },
            "supports_survival_probs": True,
        },
        "SVM": {
            "pipeline": Pipeline([
                ("prep", TabularSurvivalPreprocessor(CONTINUOUS_COLUMNS, BINARY_COLUMNS, scale_continuous=True)),
                ("model", FastSurvivalSVM(random_state=random_state)),
            ]),
            "param_grid": {
                "model__alpha": [0.01, 0.1, 1.0, 10.0],
                "model__optimizer": ["avltree", "rbtree"],
                "model__max_iter": [100, 500],
                "model__rank_ratio": [1.0],
            },
            "supports_survival_probs": False,
        },
    }

def save_search_spaces(out_path: str, specs: Dict[str, Dict[str, Any]]) -> None:
    payload = {}
    for model_name, spec in specs.items():
        payload[model_name] = {
            "description": MODEL_DESCRIPTIONS.get(model_name),
            "param_grid": spec["param_grid"],
            "supports_survival_probs": spec["supports_survival_probs"],
        }
    payload["DeepSurv"] = {"description": MODEL_DESCRIPTIONS["DeepSurv"], "configs": get_deepsurv_configs(), "supports_survival_probs": True}
    payload["CoxTime"] = {"description": MODEL_DESCRIPTIONS["CoxTime"], "configs": get_coxtime_configs(), "supports_survival_probs": True}
    save_json(payload, out_path)

def fit_gridsearched_model(model_name: str, train_df: pd.DataFrame, inner_folds: int, random_state: int, n_jobs: int,
                           specs: Dict[str, Dict[str, Any]]) -> Tuple[Any, Dict[str, Any], float]:
    X_train = train_df[PREDICTOR_COLUMNS].copy()
    y_train = get_surv_array(train_df)
    strata = make_strata(train_df)
    inner_cv = StratifiedKFold(n_splits=inner_folds, shuffle=True, random_state=random_state)
    cv_splits = list(inner_cv.split(X_train, strata))
    scorer = make_harrell_scorer()
    grid = GridSearchCV(
        estimator=specs[model_name]["pipeline"],
        param_grid=specs[model_name]["param_grid"],
        scoring=scorer,
        cv=cv_splits,
        refit=True,
        n_jobs=n_jobs,
        return_train_score=False,
        error_score="raise",
    )
    grid.fit(X_train, y_train)
    return grid.best_estimator_, grid.best_params_, float(grid.best_score_)

def evaluate_model(model_name: str, fitted_model: Any, train_df: pd.DataFrame, test_df: pd.DataFrame,
                   horizon_months: float, ibs_start_month: float, calibration_groups: int) -> Tuple[Dict[str, Any], Optional[pd.DataFrame], Optional[np.ndarray]]:
    y_train = get_surv_array(train_df)
    y_test = get_surv_array(test_df)
    X_test = test_df[PREDICTOR_COLUMNS].copy()
    risk_test = np.asarray(fitted_model.predict(X_test), dtype=float)
    tau = min(float(horizon_months), float(train_df[TIME_COL].max()), float(test_df[TIME_COL].max()))
    metrics: Dict[str, Any] = {
        "harrell_c": harrell_c_index(y_test, risk_test),
        "uno_c_tau": uno_c_index(y_train, y_test, risk_test, tau=tau),
        "tau_months": tau,
        "brier_score_at_horizon": np.nan,
        "brier_horizon_months": np.nan,
        "integrated_brier_score": np.nan,
        "calibration_slope": np.nan,
        "calibration_slope_ci_low": np.nan,
        "calibration_slope_ci_high": np.nan,
        "calibration_slope_p": np.nan,
    }
    calibration_df: Optional[pd.DataFrame] = None
    pred_event_prob: Optional[np.ndarray] = None
    if model_name != "SVM":
        eval_times = get_eval_times(train_df, test_df, start_month=ibs_start_month, horizon_months=horizon_months)
        surv_matrix = predict_survival_matrix(fitted_model, X_test, eval_times)
        metrics.update(compute_brier_metrics(y_train, y_test, surv_matrix, eval_times, horizon_months))
        idx = int(np.argmin(np.abs(eval_times - horizon_months)))
        horizon_used = float(eval_times[idx])
        pred_event_prob = 1.0 - surv_matrix[:, idx]
        calibration_df = grouped_calibration_table(test_df, pred_event_prob, time_point=horizon_used, n_groups=calibration_groups)
        metrics.update(calibration_slope_from_probs(test_df, pred_event_prob))
    return metrics, calibration_df, pred_event_prob

def combine_shap_summary_images(image_paths: List[str], titles: List[str], out_path: str, ncols: int = 2) -> None:
    valid = [(p, t) for p, t in zip(image_paths, titles) if os.path.exists(p)]
    if not valid:
        return
    n = len(valid)
    ncols = min(ncols, n)
    nrows = int(math.ceil(n / ncols))
    fig, axes = plt.subplots(nrows, ncols, figsize=(11.5, 4.2 * nrows))
    axes = np.asarray(axes).reshape(-1)
    for ax in axes:
        ax.axis("off")
    for ax, (img_path, title) in zip(axes, valid):
        ax.imshow(mpimg.imread(img_path))
        ax.set_title(title, fontsize=11)
        ax.axis("off")
    plt.tight_layout()
    plt.savefig(out_path, dpi=300, bbox_inches="tight")
    plt.close()

def combine_force_plots(image_paths: List[str], titles: List[str], out_path: str) -> None:
    valid = [(p, t) for p, t in zip(image_paths, titles) if os.path.exists(p)]
    if not valid:
        return
    fig, axes = plt.subplots(1, len(valid), figsize=(7 * len(valid), 5))
    axes = np.atleast_1d(axes)
    for ax, (img_path, title) in zip(axes, valid):
        ax.imshow(mpimg.imread(img_path))
        ax.set_title(title, fontsize=11)
        ax.axis("off")
    plt.tight_layout()
    plt.savefig(out_path, dpi=300, bbox_inches="tight")
    plt.close()

def make_combined_calibration_figure(calibration_tables: Dict[str, pd.DataFrame], out_path: str, horizon_months: float) -> None:
    fig, axes = plt.subplots(2, 4, figsize=(13, 7.5))
    axes = axes.ravel()
    models = ["CoxPH", "ElasticNetCox", "RSF", "GBSA", "DeepSurv", "CoxTime", "XGBoost", "SVM"]
    for ax, model_name in zip(axes, models):
        if model_name in calibration_tables:
            plot_calibration_panel(ax, calibration_tables[model_name], model_name, time_point=horizon_months)
        else:
            ax.axis("off")
            ax.text(0.5, 0.5, f"{model_name}\nNot available\n(relative risk only)", ha="center", va="center", fontsize=11)
    plt.tight_layout()
    plt.savefig(out_path, dpi=300, bbox_inches="tight")
    plt.close()

def run_primary_split_analysis(df: pd.DataFrame, out_dir: str, seed: int, test_size: float, inner_folds: int,
                               horizon_months: float, ibs_start_month: float, calibration_groups: int,
                               n_jobs: int, m_imputations: int, impute_max_iter: int, make_shap: bool,
                               shap_max_background: int, shap_max_explain: int) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    LOGGER.info("Running primary split analysis with nested multiple imputation.")
    tables_dir = ensure_dir(os.path.join(out_dir, "tables"))
    figs_dir = ensure_dir(os.path.join(out_dir, "figures"))
    meta_dir = ensure_dir(os.path.join(out_dir, "metadata"))

    train_idx, test_idx = primary_split(df, seed=seed, test_size=test_size)
    raw_train_df = df.iloc[train_idx].copy().reset_index(drop=True)
    raw_test_df = df.iloc[test_idx].copy().reset_index(drop=True)
    imputed_splits = generate_imputed_split_datasets(raw_train_df, raw_test_df, m=m_imputations, seed=seed, max_iter=impute_max_iter)

    table1 = aggregate_table1_over_imputations(imputed_splits, raw_train_df, raw_test_df, df)
    table1.to_csv(os.path.join(tables_dir, "table1_primary_split.csv"), index=False)

    pooled_cox_table = pool_cox_results([x["train"] for x in imputed_splits])
    pooled_cox_table.to_csv(os.path.join(tables_dir, "table2_cox_hazard_ratios_primary_split.csv"), index=False)

    model_specs = get_classical_model_specs(random_state=seed, n_jobs=n_jobs)
    save_search_spaces(os.path.join(meta_dir, "hyperparameter_search_spaces.json"), model_specs)

    primary_metrics_records: List[Dict[str, Any]] = []
    tuning_rows: List[Dict[str, Any]] = []
    calibration_pred_store: Dict[str, List[np.ndarray]] = {m: [] for m in CALIBRATION_MODELS}
    shap_reference_models: Dict[str, Any] = {}
    shap_reference_train: Optional[pd.DataFrame] = None
    shap_reference_test: Optional[pd.DataFrame] = None

    for imp_idx, split in enumerate(imputed_splits, start=1):
        LOGGER.info("Primary split imputation %d / %d", imp_idx, m_imputations)
        train_df = split["train"]
        test_df = split["test"]
        split_seed = seed + imp_idx
        specs = get_classical_model_specs(random_state=split_seed, n_jobs=n_jobs)

        for model_name in ["CoxPH", "ElasticNetCox", "RSF", "GBSA", "XGBoost", "SVM"]:
            best_estimator, best_params, best_cv_score = fit_gridsearched_model(model_name, train_df, inner_folds, split_seed, n_jobs, specs)
            metrics, _, pred_event_prob = evaluate_model(model_name, best_estimator, train_df, test_df, horizon_months, ibs_start_month, calibration_groups)
            primary_metrics_records.append({"imputation": imp_idx, "model": model_name, **metrics})
            tuning_rows.append({
                "split": "primary",
                "repeat": 0,
                "imputation": imp_idx,
                "model": model_name,
                "best_cv_score": best_cv_score,
                "best_params_json": json.dumps(best_params, sort_keys=True),
            })
            if pred_event_prob is not None:
                calibration_pred_store[model_name].append(pred_event_prob)
            if imp_idx == 1 and model_name in SHAP_MODELS:
                shap_reference_models[model_name] = best_estimator

        deepsurv_model, ds_params, ds_score = fit_deepsurv(train_df, random_state=split_seed)
        ds_metrics, _, ds_pred = evaluate_model("DeepSurv", deepsurv_model, train_df, test_df, horizon_months, ibs_start_month, calibration_groups)
        primary_metrics_records.append({"imputation": imp_idx, "model": "DeepSurv", **ds_metrics})
        tuning_rows.append({"split": "primary", "repeat": 0, "imputation": imp_idx, "model": "DeepSurv", "best_cv_score": ds_score, "best_params_json": json.dumps(ds_params, sort_keys=True)})
        calibration_pred_store["DeepSurv"].append(ds_pred)
        if imp_idx == 1:
            shap_reference_models["DeepSurv"] = deepsurv_model

        coxtime_model, ct_params, ct_score = fit_coxtime(train_df, random_state=split_seed)
        ct_metrics, _, ct_pred = evaluate_model("CoxTime", coxtime_model, train_df, test_df, horizon_months, ibs_start_month, calibration_groups)
        primary_metrics_records.append({"imputation": imp_idx, "model": "CoxTime", **ct_metrics})
        tuning_rows.append({"split": "primary", "repeat": 0, "imputation": imp_idx, "model": "CoxTime", "best_cv_score": ct_score, "best_params_json": json.dumps(ct_params, sort_keys=True)})
        calibration_pred_store["CoxTime"].append(ct_pred)
        if imp_idx == 1:
            shap_reference_models["CoxTime"] = coxtime_model
            shap_reference_train = train_df.copy()
            shap_reference_test = test_df.copy()

    primary_metrics_long = pd.DataFrame(primary_metrics_records)
    primary_metrics_long.to_csv(os.path.join(tables_dir, "table3_primary_split_by_imputation.csv"), index=False)

    avg_metrics_rows = []
    for model_name, sub in primary_metrics_long.groupby("model"):
        row = {"Model": model_name}
        for col in ["harrell_c", "uno_c_tau", "brier_score_at_horizon", "integrated_brier_score", "calibration_slope"]:
            vals = sub[col].dropna()
            row[col] = float(vals.mean()) if not vals.empty else np.nan
        avg_metrics_rows.append(row)
    primary_metrics_df = pd.DataFrame(avg_metrics_rows).set_index("Model").loc[MODEL_ORDER].reset_index()
    primary_metrics_df = primary_metrics_df.rename(columns={
        "harrell_c": "Harrell_C",
        "uno_c_tau": "Uno_C_tau",
        "brier_score_at_horizon": "Brier_120m",
        "integrated_brier_score": "IBS_12_120m",
        "calibration_slope": "Calibration_slope_120m",
    })
    primary_metrics_df.to_csv(os.path.join(tables_dir, "table3_primary_split_metrics.csv"), index=False)

    calibration_tables: Dict[str, pd.DataFrame] = {}
    for model_name, pred_list in calibration_pred_store.items():
        if len(pred_list) == 0:
            continue
        mean_pred = np.mean(np.vstack(pred_list), axis=0)
        calibration_df = grouped_calibration_table(raw_test_df, mean_pred, time_point=horizon_months, n_groups=calibration_groups)
        calibration_tables[model_name] = calibration_df
        calibration_df.to_csv(os.path.join(tables_dir, f"calibration_points_primary_{model_name}.csv"), index=False)
    make_combined_calibration_figure(calibration_tables, os.path.join(figs_dir, "figure1_calibration_all_models.png"), horizon_months=horizon_months)

    shap_tables: List[pd.DataFrame] = []
    if make_shap and shap_reference_train is not None and shap_reference_test is not None:
        for model_name in SHAP_MODELS:
            fitted_model = shap_reference_models.get(model_name)
            if fitted_model is None:
                continue
            shap_table = make_shap_outputs(
                model_name=model_name,
                fitted_model=fitted_model,
                train_df=shap_reference_train,
                test_df=shap_reference_test,
                out_dir=figs_dir,
                make_force_plot=(model_name in ["CoxPH", "DeepSurv"]),
                max_background=shap_max_background,
                max_explain=shap_max_explain,
            )
            if shap_table is not None:
                shap_table.insert(0, "Model", model_name)
                shap_tables.append(shap_table)
        shap_pngs = [os.path.join(figs_dir, f"shap_summary_{m}.png") for m in SHAP_MODELS]
        combine_shap_summary_images(shap_pngs, SHAP_MODELS, os.path.join(figs_dir, "figure2_shap_summary_all_models.png"), ncols=2)
        combine_force_plots(
            [os.path.join(figs_dir, "shap_force_CoxPH.png"), os.path.join(figs_dir, "shap_force_DeepSurv.png")],
            ["CoxPH", "DeepSurv"],
            os.path.join(figs_dir, "supplementary_figure_s1_local_shap_force.png"),
        )
    if shap_tables:
        pd.concat(shap_tables, axis=0, ignore_index=True).to_csv(os.path.join(tables_dir, "shap_importance_primary_split_all_models.csv"), index=False)

    tuning_df = pd.DataFrame(tuning_rows)
    tuning_df.to_csv(os.path.join(meta_dir, "tuning_records_primary_split.csv"), index=False)
    return table1, pooled_cox_table, primary_metrics_df, tuning_df

def run_repeated_resampling(df: pd.DataFrame, out_dir: str, seed: int, n_repeats: int, test_size: float, inner_folds: int,
                            horizon_months: float, ibs_start_month: float, calibration_groups: int, n_jobs: int,
                            m_imputations: int, impute_max_iter: int) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    LOGGER.info("Running repeated stratified train/test resampling with nested multiple imputation.")
    tables_dir = ensure_dir(os.path.join(out_dir, "tables"))
    meta_dir = ensure_dir(os.path.join(out_dir, "metadata"))

    splitter = StratifiedShuffleSplit(n_splits=n_repeats, test_size=test_size, random_state=seed)
    strata = make_strata(df)
    split_level_records: List[Dict[str, Any]] = []
    tuning_rows: List[Dict[str, Any]] = []

    for repeat, (train_idx, test_idx) in enumerate(splitter.split(df, strata), start=1):
        LOGGER.info("Outer repeat %d / %d", repeat, n_repeats)
        raw_train_df = df.iloc[train_idx].copy().reset_index(drop=True)
        raw_test_df = df.iloc[test_idx].copy().reset_index(drop=True)
        imputed_splits = generate_imputed_split_datasets(raw_train_df, raw_test_df, m=m_imputations, seed=seed + 1000 * repeat, max_iter=impute_max_iter)
        imp_metrics: List[Dict[str, Any]] = []
        for imp_idx, split in enumerate(imputed_splits, start=1):
            LOGGER.info("  repeat %d imputation %d / %d", repeat, imp_idx, m_imputations)
            train_df = split["train"]
            test_df = split["test"]
            split_seed = seed + 1000 * repeat + imp_idx
            specs = get_classical_model_specs(random_state=split_seed, n_jobs=n_jobs)
            for model_name in ["CoxPH", "ElasticNetCox", "RSF", "GBSA", "XGBoost", "SVM"]:
                try:
                    best_estimator, best_params, best_cv_score = fit_gridsearched_model(model_name, train_df, inner_folds, split_seed, n_jobs, specs)
                    metrics, _, _ = evaluate_model(model_name, best_estimator, train_df, test_df, horizon_months, ibs_start_month, calibration_groups)
                    imp_metrics.append({"repeat": repeat, "imputation": imp_idx, "model": model_name, **metrics})
                    tuning_rows.append({
                        "split": "resampling",
                        "repeat": repeat,
                        "imputation": imp_idx,
                        "model": model_name,
                        "best_cv_score": best_cv_score,
                        "best_params_json": json.dumps(best_params, sort_keys=True),
                    })
                except Exception as exc:
                    LOGGER.exception("Repeat %d imputation %d failed for %s: %s", repeat, imp_idx, model_name, exc)
                    imp_metrics.append({"repeat": repeat, "imputation": imp_idx, "model": model_name, "harrell_c": np.nan, "uno_c_tau": np.nan,
                                        "tau_months": np.nan, "brier_score_at_horizon": np.nan, "brier_horizon_months": np.nan,
                                        "integrated_brier_score": np.nan, "calibration_slope": np.nan,
                                        "calibration_slope_ci_low": np.nan, "calibration_slope_ci_high": np.nan, "calibration_slope_p": np.nan})
            try:
                deepsurv_model, ds_params, ds_score = fit_deepsurv(train_df, random_state=split_seed)
                ds_metrics, _, _ = evaluate_model("DeepSurv", deepsurv_model, train_df, test_df, horizon_months, ibs_start_month, calibration_groups)
                imp_metrics.append({"repeat": repeat, "imputation": imp_idx, "model": "DeepSurv", **ds_metrics})
                tuning_rows.append({"split": "resampling", "repeat": repeat, "imputation": imp_idx, "model": "DeepSurv", "best_cv_score": ds_score, "best_params_json": json.dumps(ds_params, sort_keys=True)})
            except Exception as exc:
                LOGGER.exception("Repeat %d imputation %d failed for DeepSurv: %s", repeat, imp_idx, exc)
                imp_metrics.append({"repeat": repeat, "imputation": imp_idx, "model": "DeepSurv", "harrell_c": np.nan, "uno_c_tau": np.nan,
                                    "tau_months": np.nan, "brier_score_at_horizon": np.nan, "brier_horizon_months": np.nan,
                                    "integrated_brier_score": np.nan, "calibration_slope": np.nan,
                                    "calibration_slope_ci_low": np.nan, "calibration_slope_ci_high": np.nan, "calibration_slope_p": np.nan})
            try:
                coxtime_model, ct_params, ct_score = fit_coxtime(train_df, random_state=split_seed)
                ct_metrics, _, _ = evaluate_model("CoxTime", coxtime_model, train_df, test_df, horizon_months, ibs_start_month, calibration_groups)
                imp_metrics.append({"repeat": repeat, "imputation": imp_idx, "model": "CoxTime", **ct_metrics})
                tuning_rows.append({"split": "resampling", "repeat": repeat, "imputation": imp_idx, "model": "CoxTime", "best_cv_score": ct_score, "best_params_json": json.dumps(ct_params, sort_keys=True)})
            except Exception as exc:
                LOGGER.exception("Repeat %d imputation %d failed for CoxTime: %s", repeat, imp_idx, exc)
                imp_metrics.append({"repeat": repeat, "imputation": imp_idx, "model": "CoxTime", "harrell_c": np.nan, "uno_c_tau": np.nan,
                                    "tau_months": np.nan, "brier_score_at_horizon": np.nan, "brier_horizon_months": np.nan,
                                    "integrated_brier_score": np.nan, "calibration_slope": np.nan,
                                    "calibration_slope_ci_low": np.nan, "calibration_slope_ci_high": np.nan, "calibration_slope_p": np.nan})
        imp_metrics_df = pd.DataFrame(imp_metrics)
        for model_name, sub in imp_metrics_df.groupby("model"):
            row = {"repeat": repeat, "model": model_name}
            for col in ["harrell_c", "uno_c_tau", "brier_score_at_horizon", "integrated_brier_score", "calibration_slope"]:
                vals = sub[col].dropna()
                row[col] = float(vals.mean()) if not vals.empty else np.nan
            split_level_records.append(row)

    perf_df = pd.DataFrame(split_level_records)
    perf_df.to_csv(os.path.join(tables_dir, "performance_by_repeat.csv"), index=False)

    tuning_df = pd.DataFrame(tuning_rows)
    tuning_df.to_csv(os.path.join(meta_dir, "tuning_records_resampling.csv"), index=False)

    summary_df, table4_df = summarise_performance(perf_df)
    summary_df.to_csv(os.path.join(tables_dir, "performance_summary_repeated_resampling.csv"), index=False)
    table4_df.to_csv(os.path.join(tables_dir, "table4_repeated_resampling_formatted.csv"), index=False)

    comparisons_df = compare_models_vs_reference(perf_df, reference_model="CoxPH")
    comparisons_df.to_csv(os.path.join(tables_dir, "supplementary_table_s3_paired_differences_vs_coxph.csv"), index=False)

    tuning_summary_df = summarise_tuning_records(tuning_df)
    tuning_summary_df.to_csv(os.path.join(meta_dir, "tuning_summary_resampling.csv"), index=False)
    return perf_df, table4_df, comparisons_df

def build_supplementary_tuning_table(primary_tuning_df: pd.DataFrame) -> pd.DataFrame:
    rows = [
        {
            "Model": "CoxPH",
            "Implementation and role": "CoxPHSurvivalAnalysis (semiparametric benchmark)",
            "Prespecified tuning strategy": "ties = {breslow, efron}; tol = {1e-9, 1e-7}; n_iter = {100, 300}",
        },
        {
            "Model": "ElasticNetCox",
            "Implementation and role": "Penalized CoxPHFitter (regularized Cox benchmark)",
            "Prespecified tuning strategy": "penalizer = {0.001, 0.01, 0.1, 1.0}; l1_ratio = {0.0, 0.5, 1.0}; continuous predictors standardized within each training split",
        },
        {
            "Model": "RSF",
            "Implementation and role": "RandomSurvivalForest (tree-based survival ensemble)",
            "Prespecified tuning strategy": "n_estimators = {200, 500}; min_samples_split = {10, 20}; min_samples_leaf = {5, 10}; max_features = {sqrt, 0.5}; max_depth = {None, 3}",
        },
        {
            "Model": "GBSA",
            "Implementation and role": "GradientBoostingSurvivalAnalysis (boosted survival-tree model)",
            "Prespecified tuning strategy": "learning_rate = {0.01, 0.05, 0.1}; n_estimators = {100, 300}; min_samples_leaf = {3, 10}; max_depth = {1, 2}; subsample = {0.7, 1.0}",
        },
        {
            "Model": "DeepSurv",
            "Implementation and role": "DeepSurv (neural-network extension of Cox regression; continuous predictors standardized within each training split)",
            "Prespecified tuning strategy": "Six prespecified multilayer perceptron configurations were evaluated: (1) num_nodes = [32,16], dropout = 0.1, lr = 1e-2, batch_size = 64, max_epochs = 256; (2) num_nodes = [32,16], dropout = 0.2, lr = 1e-2, batch_size = 64, max_epochs = 256; (3) num_nodes = [32,32], dropout = 0.1, lr = 1e-2, batch_size = 64, max_epochs = 256; (4) num_nodes = [64,32], dropout = 0.2, lr = 1e-3, batch_size = 64, max_epochs = 256; (5) num_nodes = [64,32], dropout = 0.1, lr = 1e-3, batch_size = 128, max_epochs = 256; and (6) num_nodes = [16,16], dropout = 0.2, lr = 1e-2, batch_size = 128, max_epochs = 256; internal stratified 80/20 validation split with early stopping",
        },
        {
            "Model": "CoxTime",
            "Implementation and role": "CoxTime (neural non-proportional-hazards Cox extension; continuous predictors standardized within each training split)",
            "Prespecified tuning strategy": "Six prespecified multilayer perceptron configurations were evaluated: (1) num_nodes = [32,16], dropout = 0.1, lr = 1e-2, batch_size = 64, max_epochs = 256; (2) num_nodes = [32,16], dropout = 0.2, lr = 1e-2, batch_size = 64, max_epochs = 256; (3) num_nodes = [32,32], dropout = 0.1, lr = 1e-2, batch_size = 64, max_epochs = 256; (4) num_nodes = [64,32], dropout = 0.2, lr = 1e-3, batch_size = 64, max_epochs = 256; (5) num_nodes = [64,32], dropout = 0.1, lr = 1e-3, batch_size = 128, max_epochs = 256; and (6) num_nodes = [16,16], dropout = 0.2, lr = 1e-2, batch_size = 128, max_epochs = 256; internal stratified 80/20 validation split with early stopping",
        },
        {
            "Model": "XGBoost",
            "Implementation and role": "XGBRegressor with survival:cox objective (boosted tree survival model)",
            "Prespecified tuning strategy": "n_estimators = {200, 500}; learning_rate = {0.01, 0.05, 0.1}; max_depth = {2, 3}; subsample = {0.7, 1.0}; colsample_bytree = {0.7, 1.0}; min_child_weight = {1.0, 5.0}; reg_lambda = {1.0, 5.0}",
        },
        {
            "Model": "Survival SVM",
            "Implementation and role": "FastSurvivalSVM (linear ranking model; rank_ratio = 1.0)",
            "Prespecified tuning strategy": "alpha = {0.01, 0.1, 1.0, 10.0}; optimizer = {avltree, rbtree}; max_iter = {100, 500}; continuous predictors standardized within each training split",
        },
    ]
    out = pd.DataFrame(rows)
    if not primary_tuning_df.empty:
        deep_primary = primary_tuning_df[primary_tuning_df["model"] == "DeepSurv"]
        if not deep_primary.empty:
            mode_params = deep_primary["best_params_json"].mode().iloc[0]
            out.attrs["primary_deepsurv"] = mode_params
        ct_primary = primary_tuning_df[primary_tuning_df["model"] == "CoxTime"]
        if not ct_primary.empty:
            out.attrs["primary_coxtime"] = ct_primary["best_params_json"].mode().iloc[0]
    return out

def set_repeat_table_header(row: Any) -> None:
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)

def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def style_docx_table(table: Any, header_fill: str = "D9EAF7") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                if row_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if len(cell.text) <= 18:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if row_idx == 0:
            for cell in row.cells:
                set_cell_shading(cell, header_fill)
            set_repeat_table_header(row)

def add_dataframe_as_table(doc: Document, df: pd.DataFrame, title: str, note: Optional[str] = None,
                           landscape: bool = False, font_size: float = 9.0) -> None:
    if landscape:
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
    else:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.bold = True
    table = doc.add_table(rows=1, cols=len(df.columns))
    for j, col in enumerate(df.columns):
        table.cell(0, j).text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            val = row[col]
            if isinstance(val, float) and np.isnan(val):
                txt = "NA"
            else:
                txt = str(val)
            cells[j].text = txt
    style_docx_table(table)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size)
    if note:
        note_p = doc.add_paragraph()
        note_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        note_run = note_p.add_run(note)
        note_run.italic = True
        note_run.font.size = Pt(8.5)

def save_single_table_docx(df: pd.DataFrame, title: str, path: str, note: Optional[str] = None, landscape: bool = False) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Inches(0.6)
    sec.right_margin = Inches(0.6)
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    title_p = doc.add_paragraph()
    title_p.style = doc.styles["Title"]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.add_run(title)
    add_dataframe_as_table(doc, df, title="", note=note, landscape=landscape)
    doc.save(path)

def write_tables_docx(out_dir: str, table1: pd.DataFrame, table2: pd.DataFrame, table3: pd.DataFrame, table4: pd.DataFrame,
                      supp_s1: pd.DataFrame, supp_s2: pd.DataFrame, supp_s3: pd.DataFrame) -> None:
    docs_dir = ensure_dir(os.path.join(out_dir, "docx_tables"))
    combined = Document()
    sec = combined.sections[0]
    sec.left_margin = Inches(0.6)
    sec.right_margin = Inches(0.6)
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    title = combined.add_paragraph()
    title.style = combined.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("ASCVD Survival Models - Recreated Tables")
    main_tables = [
        (table1, "Table 1. Baseline characteristics of the overall cohort and the primary training and test sets.",
         "Imputed values are averaged across datasets. Event counts and follow-up come from the raw data.", True),
        (table2, "Table 2. Multivariable Cox proportional hazards model for incident ASCVD in the primary training set.",
         "Cox estimates were pooled across imputations. PH test values are from the first completed training set.", True),
        (table3, "Table 3. Primary-split performance of CoxPH and machine-learning survival models.",
         "Values are averaged across imputations.", False),
        (table4, "Table 4. Performance of CoxPH and machine-learning survival models across 30 repeated stratified train-test splits.",
         "Values are mean (2.5th to 97.5th percentile) across repeats after averaging within split.", True),
    ]
    supp_tables = [
        (supp_s1, "Supplementary Table S1. Variables with missing baseline data before imputation.", None, False),
        (supp_s2, "Supplementary Table S2. Prespecified tuning strategies used in the repeated internal validation framework.",
         "Classical models used stratified five-fold CV. DeepSurv and CoxTime used six preset networks with an internal 80/20 validation split and early stopping.", True),
        (supp_s3, "Supplementary Table S3. Paired repeated-split performance differences relative to CoxPH.",
         "Differences are model minus CoxPH.", True),
    ]

    for df, title_txt, note_txt, landscape in main_tables + supp_tables:
        add_dataframe_as_table(combined, df, title_txt, note=note_txt, landscape=landscape)
    combined.save(os.path.join(docs_dir, "ASCVD_recreated_tables_combined.docx"))

    all_tables = main_tables + supp_tables
    for idx, (df, title_txt, note_txt, landscape) in enumerate(all_tables, start=1):
        file_name = f"table_{idx:02d}.docx"
        safe_name = title_txt.split(". ", 1)[0].replace(" ", "_").replace("/", "_") + ".docx"
        save_single_table_docx(df, title_txt, os.path.join(docs_dir, safe_name), note=note_txt, landscape=landscape)


def build_safe_analysis_config(args: argparse.Namespace) -> Dict[str, Any]:
    safe_keys = [
        "seed",
        "test_size",
        "n_repeats",
        "inner_folds",
        "n_jobs",
        "m_imputations",
        "impute_max_iter",
        "horizon_months",
        "ibs_start_month",
        "calibration_groups",
        "skip_shap",
        "shap_max_background",
        "shap_max_explain",
    ]
    return {key: getattr(args, key) for key in safe_keys}

def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Nested-imputation survival model comparison pipeline.")
    parser.add_argument("--csv", default=DEFAULT_CSV_PATH, help="Path to the raw CSV with missing baseline predictors.")
    parser.add_argument("--outdir", default=DEFAULT_OUTDIR, help="Output directory.")
    parser.add_argument("--seed", type=int, default=20260320, help="Master random seed.")
    parser.add_argument("--test-size", type=float, default=0.2, help="Outer test fraction.")
    parser.add_argument("--n-repeats", type=int, default=30, help="Number of repeated outer splits.")
    parser.add_argument("--inner-folds", type=int, default=5, help="Number of inner CV folds for classical models.")
    parser.add_argument("--n-jobs", type=int, default=-1, help="Parallel workers for grid search.")
    parser.add_argument("--m-imputations", type=int, default=10, help="Number of imputations within each split.")
    parser.add_argument("--impute-max-iter", type=int, default=20, help="Maximum iterations for the iterative imputer.")
    parser.add_argument("--horizon-months", type=float, default=120.0, help="Prediction horizon in months.")
    parser.add_argument("--ibs-start-month", type=float, default=12.0, help="Start month for integrated Brier score.")
    parser.add_argument("--calibration-groups", type=int, default=5, help="Number of risk groups for grouped calibration.")
    parser.add_argument("--skip-shap", action="store_true", help="Skip SHAP outputs for the primary split.")
    parser.add_argument("--shap-max-background", type=int, default=100, help="Maximum SHAP background size.")
    parser.add_argument("--shap-max-explain", type=int, default=150, help="Maximum SHAP explanation rows.")
    return parser.parse_args()

def main() -> None:
    args = parse_args()
    out_dir = ensure_dir(args.outdir)
    ensure_dir(os.path.join(out_dir, "tables"))
    ensure_dir(os.path.join(out_dir, "figures"))
    ensure_dir(os.path.join(out_dir, "metadata"))
    setup_logging(os.path.join(out_dir, "metadata", "run.log"))
    set_global_seed(args.seed)

    if not os.path.exists(args.csv):
        raise FileNotFoundError(
            "Could not find the analysis CSV. "
            "Update DEFAULT_CSV_PATH near the top of the script or pass --csv explicitly."
        )

    LOGGER.info("Loading dataset.")
    df = load_dataset(args.csv)
    LOGGER.info("Loaded %d rows after filtering non-positive follow-up times.", len(df))

    missingness_df = create_missingness_table(df)
    missingness_df.to_csv(os.path.join(out_dir, "tables", "supplementary_table_s1_missingness.csv"), index=False)

    package_versions = get_package_versions()
    save_json(package_versions, os.path.join(out_dir, "metadata", "package_versions.json"))
    config_dump = build_safe_analysis_config(args)
    config_dump["paths_saved"] = False
    config_dump["predictors"] = PREDICTOR_COLUMNS
    config_dump["continuous_columns"] = CONTINUOUS_COLUMNS
    config_dump["binary_columns"] = BINARY_COLUMNS
    config_dump["model_descriptions"] = MODEL_DESCRIPTIONS
    save_json(config_dump, os.path.join(out_dir, "metadata", "analysis_config.json"))

    table1, table2, table3, primary_tuning_df = run_primary_split_analysis(
        df=df,
        out_dir=out_dir,
        seed=args.seed,
        test_size=args.test_size,
        inner_folds=args.inner_folds,
        horizon_months=args.horizon_months,
        ibs_start_month=args.ibs_start_month,
        calibration_groups=args.calibration_groups,
        n_jobs=args.n_jobs,
        m_imputations=args.m_imputations,
        impute_max_iter=args.impute_max_iter,
        make_shap=not args.skip_shap,
        shap_max_background=args.shap_max_background,
        shap_max_explain=args.shap_max_explain,
    )

    perf_df, table4, supp_s3 = run_repeated_resampling(
        df=df,
        out_dir=out_dir,
        seed=args.seed,
        n_repeats=args.n_repeats,
        test_size=args.test_size,
        inner_folds=args.inner_folds,
        horizon_months=args.horizon_months,
        ibs_start_month=args.ibs_start_month,
        calibration_groups=args.calibration_groups,
        n_jobs=args.n_jobs,
        m_imputations=args.m_imputations,
        impute_max_iter=args.impute_max_iter,
    )

    supp_s2 = build_supplementary_tuning_table(primary_tuning_df)
    supp_s2.to_csv(os.path.join(out_dir, "tables", "supplementary_table_s2_tuning_strategies.csv"), index=False)

    write_tables_docx(
        out_dir=out_dir,
        table1=table1,
        table2=table2,
        table3=table3,
        table4=table4,
        supp_s1=missingness_df,
        supp_s2=supp_s2,
        supp_s3=supp_s3,
    )
    LOGGER.info("Primary split metrics:\n%s", table3.to_string(index=False))
    LOGGER.info("Repeated resampling summary saved.")
    LOGGER.info("Done.")

if __name__ == "__main__":
    main()
