from __future__ import annotations

import argparse
import importlib.metadata
import os
import subprocess
import sys
import tempfile
from pathlib import Path


_SOURCES = {
    '_analysis_main.py': r'''from __future__ import annotations
import argparse
import hashlib
import importlib
import importlib.metadata
import json
import logging
import math
import os
import platform
import random
import sys
import tempfile
import time
import warnings
from dataclasses import dataclass
from io import BytesIO
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple
import matplotlib
matplotlib.use('Agg')
import matplotlib.image as mpimg
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from joblib import Parallel, delayed
from PIL import Image, ImageDraw, ImageFont
from scipy import stats
from scipy.stats import norm
from sklearn.base import BaseEstimator, TransformerMixin, clone
from sklearn.exceptions import ConvergenceWarning
from sklearn.experimental import enable_iterative_imputer
from sklearn.impute import IterativeImputer
from sklearn.linear_model import BayesianRidge, LogisticRegression
from sklearn.model_selection import ParameterGrid, StratifiedKFold, StratifiedShuffleSplit
from sklearn.pipeline import Pipeline
from sklearn.preprocessing import SplineTransformer, StandardScaler
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
warnings.filterwarnings('ignore', category=ConvergenceWarning)
LOGGER = logging.getLogger('survival_repeated_200')
STUDY_ID_COL = 'studyid'
EVENT_COL = 'event_status'
TIME_COL = 'time_to_event'
SEX_COL = 'sex'
PREDICTOR_COLUMNS: List[str] = ['age', 'sex', 'smoke', 'fhxcvd', 'dmmed', 'htnmed', 'dldmed', 'sbp', 'dbp', 'bmi', 'chol', 'hdl', 'trig', 'hba1c', 'egfr']
CONTINUOUS_COLUMNS: List[str] = ['age', 'sbp', 'dbp', 'bmi', 'chol', 'hdl', 'trig', 'hba1c', 'egfr']
BINARY_COLUMNS: List[str] = ['sex', 'smoke', 'fhxcvd', 'dmmed', 'htnmed', 'dldmed']
DISPLAY_NAMES: Dict[str, str] = {'age': 'Age', 'sex': 'Sex', 'smoke': 'History of smoking', 'fhxcvd': 'Family history of CVD', 'dmmed': 'Glucose-lowering medication', 'htnmed': 'Blood pressure-lowering medication', 'dldmed': 'Lipid-lowering medication', 'sbp': 'SBP', 'dbp': 'DBP', 'bmi': 'BMI', 'chol': 'TC', 'hdl': 'HDL-C', 'trig': 'TG', 'hba1c': 'HbA1c', 'egfr': 'eGFR', EVENT_COL: 'ASCVD event', TIME_COL: 'Follow-up time (months)'}
SHAP_DISPLAY_NAMES: Dict[str, str] = {'age': 'Age', 'sex': 'Sex', 'smoke': 'Smoking history', 'fhxcvd': 'Family history of CVD', 'dmmed': 'Glucose-lowering medication', 'htnmed': 'BP-lowering medication', 'dldmed': 'Lipid-lowering medication', 'sbp': 'SBP', 'dbp': 'DBP', 'bmi': 'BMI', 'chol': 'TC', 'hdl': 'HDL-C', 'trig': 'TG', 'hba1c': 'HbA1c', 'egfr': 'eGFR'}
TABLE1_LABELS: Dict[str, str] = {'age': 'Age (years)', 'sex': 'Sex (Women)', 'smoke': 'History of smoking', 'fhxcvd': 'Family history of CVD', 'dmmed': 'Glucose-lowering medications', 'htnmed': 'Blood pressure-lowering medications', 'dldmed': 'Lipid-lowering medications', 'sbp': 'SBP (mmHg)', 'dbp': 'DBP (mmHg)', 'bmi': 'BMI (kg/m2)', 'chol': 'TC (mmol/L)', 'hdl': 'HDL-C (mmol/L)', 'trig': 'TG (mmol/L)', 'hba1c': 'HbA1c (%)', 'egfr': 'eGFR (mL/min/1.73m2)'}
POSITIVE_LEVELS: Dict[str, int] = {'sex': 0, 'smoke': 1, 'fhxcvd': 1, 'dmmed': 1, 'htnmed': 1, 'dldmed': 1, EVENT_COL: 1}
DEFAULT_CSV_PATH = 'analysis_data.csv'
DEFAULT_OUTDIR = 'main_outputs'
PRESPECIFIED_SPLITTER_SEED = 20260320
PRESPECIFIED_N_SPLITS = 200
PRESPECIFIED_TEST_SIZE = 0.2
EXPECTED_SOURCE_ROWS = 1247
EXPECTED_FILTERED_ROWS = 1241
EXPECTED_ASCVD_EVENTS = 116
EXPECTED_UNIQUE_STUDY_IDS = 1241
PRESPECIFIED_HORIZON_MONTHS = 120.0
PRESPECIFIED_IBS_START_MONTH = 12.0
EXPECTED_PYTHON_VERSION = '3.11.7'
EXPECTED_PACKAGE_VERSIONS: Dict[str, str] = {'numpy': '1.26.4', 'pandas': '2.2.3', 'scipy': '1.14.1', 'matplotlib': '3.9.2', 'Pillow': '11.0.0', 'python-docx': '1.1.2', 'scikit-learn': '1.5.2', 'joblib': '1.4.2', 'scikit-survival': '0.23.0', 'lifelines': '0.29.0', 'xgboost': '2.1.2', 'torch': '2.5.1', 'pycox': '0.3.0', 'torchtuples': '0.2.2', 'shap': '0.46.0', 'openpyxl': '3.1.5'}
MODEL_ORDER: List[str] = ['CoxPH', 'ElasticNetCox', 'RSF', 'GBSA', 'DeepSurv', 'CoxTime', 'XGBoost', 'SVM']
CALIBRATION_MODELS = ['CoxPH', 'ElasticNetCox', 'RSF', 'GBSA', 'DeepSurv', 'CoxTime', 'XGBoost']
SHAP_MODELS = MODEL_ORDER.copy()
IBS_SELECTED_MODELS = {'ElasticNetCox', 'RSF', 'GBSA', 'DeepSurv', 'CoxTime', 'XGBoost'}
FIXED_SPECIFICATION_MODELS = {'CoxPH'}
HARRELL_SELECTED_MODELS = {'SVM'}
IBS_SELECTION_METRIC = 'integrated_brier_score_12_120_months'
HARRELL_SELECTION_METRIC = 'harrell_c'
PRESPECIFIED_CANDIDATE_COUNTS = {'CoxPH': 1, 'ElasticNetCox': 8, 'RSF': 6, 'GBSA': 6, 'DeepSurv': 6, 'CoxTime': 6, 'XGBoost': 8, 'SVM': 4}

def model_selection_metadata(model_name: str) -> Dict[str, Any]:
    if model_name in IBS_SELECTED_MODELS:
        validation_design = 'internal_stratified_80_20_validation' if model_name in {'DeepSurv', 'CoxTime'} else 'stratified_3_fold_inner_cross_validation'
        return {'selection_metric': IBS_SELECTION_METRIC, 'selection_direction': 'minimize', 'model_selection_performed': True, 'validation_design': validation_design, 'selection_window_start_months': PRESPECIFIED_IBS_START_MONTH, 'selection_window_end_months': PRESPECIFIED_HORIZON_MONTHS, 'selection_grid_step_months': 1.0, 'selection_grid_points': 109, 'n_prespecified_configurations': PRESPECIFIED_CANDIDATE_COUNTS[model_name]}
    if model_name in HARRELL_SELECTED_MODELS:
        return {'selection_metric': HARRELL_SELECTION_METRIC, 'selection_direction': 'maximize', 'model_selection_performed': True, 'validation_design': 'stratified_3_fold_inner_cross_validation', 'selection_window_start_months': None, 'selection_window_end_months': None, 'selection_grid_step_months': None, 'selection_grid_points': None, 'n_prespecified_configurations': PRESPECIFIED_CANDIDATE_COUNTS[model_name]}
    if model_name in FIXED_SPECIFICATION_MODELS:
        return {'selection_metric': 'not_applicable_fixed_specification', 'selection_direction': 'not_applicable', 'model_selection_performed': False, 'validation_design': 'not_applicable_fixed_specification', 'selection_window_start_months': None, 'selection_window_end_months': None, 'selection_grid_step_months': None, 'selection_grid_points': None, 'n_prespecified_configurations': PRESPECIFIED_CANDIDATE_COUNTS[model_name]}
    raise ValueError(f'Unknown model selection rule for {model_name!r}.')

def main_tuning_method(model_name: str) -> str:
    if model_name in FIXED_SPECIFICATION_MODELS:
        return 'fixed_specification_no_hyperparameter_selection'
    if model_name in {'DeepSurv', 'CoxTime'}:
        return 'nested_imputation_internal_stratified_80_20_validation_ibs_12_120'
    if model_name in IBS_SELECTED_MODELS:
        return 'nested_imputation_stratified_3_fold_mean_ibs_12_120'
    if model_name in HARRELL_SELECTED_MODELS:
        return 'nested_imputation_stratified_3_fold_mean_harrell_c'
    raise ValueError(f'Unknown tuning method for {model_name!r}.')

def tuning_output_metadata(model_name: str, selection_score: Optional[float]) -> Dict[str, Any]:
    metadata = model_selection_metadata(model_name)
    if metadata['model_selection_performed']:
        if selection_score is None or not np.isfinite(float(selection_score)):
            raise ValueError(f'{model_name} requires a finite selection score.')
        score_value = float(selection_score)
        if not 0.0 <= score_value <= 1.0:
            raise ValueError(f'{model_name} selection score lies outside [0, 1].')
    else:
        if selection_score is not None and np.isfinite(float(selection_score)):
            raise ValueError(f'{model_name} must not report a selection score.')
        score_value = np.nan
    return {'selection_score': score_value, **metadata}
MODEL_DESCRIPTIONS: Dict[str, str] = {'CoxPH': 'Semiparametric proportional hazards benchmark with interpretable hazard ratios.', 'ElasticNetCox': 'Regularized Cox model with elastic-net penalty.', 'RSF': 'Tree-based survival ensemble that can capture non-linearity and interactions.', 'GBSA': 'Boosted survival-tree ensemble that optimizes a survival loss function.', 'DeepSurv': 'Neural-network extension of Cox regression for nonlinear risk modeling.', 'CoxTime': 'Neural Cox-Time model that relaxes proportional hazards.', 'XGBoost': 'Gradient-boosted survival trees fitted with the Cox objective.', 'SVM': 'Linear ranking-based survival SVM used primarily for discrimination.'}

@dataclass
class FittedModelBundle:
    model_name: str
    fitted_model: Any
    best_params: Dict[str, Any]
    selection_score: Optional[float]
    metrics: Dict[str, Any]
    calibration_table: Optional[pd.DataFrame] = None
    pred_event_prob_horizon: Optional[np.ndarray] = None

class SafeStopRequested(RuntimeError):

    def __init__(self, boundary: str):
        super().__init__(f'Session limit reached after {boundary}.')
        self.boundary = boundary

def check_session_deadline(deadline_monotonic: Optional[float], boundary: str) -> None:
    if deadline_monotonic is not None and time.monotonic() >= deadline_monotonic:
        raise SafeStopRequested(boundary)

class TabularSurvivalPreprocessor(BaseEstimator, TransformerMixin):

    def __init__(self, continuous_columns: Sequence[str], binary_columns: Sequence[str], scale_continuous: bool=False) -> None:
        self.continuous_columns = continuous_columns
        self.binary_columns = binary_columns
        self.scale_continuous = scale_continuous

    def fit(self, X: pd.DataFrame, y: Optional[Any]=None) -> 'TabularSurvivalPreprocessor':
        self.columns_ = list(self.continuous_columns) + list(self.binary_columns)
        if self.scale_continuous:
            self.scaler_ = StandardScaler()
            self.scaler_.fit(X[list(self.continuous_columns)].astype(float))
        return self

    def transform(self, X: pd.DataFrame) -> pd.DataFrame:
        out = X[self.columns_].copy()
        out[list(self.continuous_columns)] = out[list(self.continuous_columns)].astype(float)
        out[list(self.binary_columns)] = out[list(self.binary_columns)].astype(float)
        if self.scale_continuous:
            out.loc[:, list(self.continuous_columns)] = self.scaler_.transform(out[list(self.continuous_columns)].astype(float))
        return out

    def get_feature_names_out(self, input_features: Optional[Sequence[str]]=None) -> np.ndarray:
        return np.asarray(self.columns_, dtype=object)

class SimpleStepFunction:

    def __init__(self, x: np.ndarray, y: np.ndarray):
        self.x = np.asarray(x, dtype=float)
        self.y = np.asarray(y, dtype=float)

    def __call__(self, t: float) -> float:
        idx = np.searchsorted(self.x, t, side='right') - 1
        if idx < 0:
            return 1.0
        idx = min(idx, len(self.y) - 1)
        return float(self.y[idx])

class BreslowBaselineEstimator:

    def __init__(self) -> None:
        self.event_times_: Optional[np.ndarray] = None
        self.cumhaz_: Optional[np.ndarray] = None

    def fit(self, durations: np.ndarray, events: np.ndarray, log_risk: np.ndarray) -> 'BreslowBaselineEstimator':
        durations = np.asarray(durations, dtype=float)
        events = np.asarray(events, dtype=int)
        log_risk = np.asarray(log_risk, dtype=float)
        risk = np.exp(log_risk)
        unique_event_times = np.sort(np.unique(durations[events == 1]))
        cumhaz = []
        running = 0.0
        for t in unique_event_times:
            at_risk = risk[durations >= t].sum()
            d_t = int(((durations == t) & (events == 1)).sum())
            if at_risk <= 0:
                increment = 0.0
            else:
                increment = d_t / at_risk
            running += increment
            cumhaz.append(running)
        self.event_times_ = unique_event_times
        self.cumhaz_ = np.asarray(cumhaz, dtype=float)
        return self

    def cumulative_hazard_at(self, times: np.ndarray) -> np.ndarray:
        if self.event_times_ is None or self.cumhaz_ is None:
            raise RuntimeError('Baseline estimator must be fit before prediction.')
        times = np.asarray(times, dtype=float)
        idx = np.searchsorted(self.event_times_, times, side='right') - 1
        out = np.zeros_like(times, dtype=float)
        valid = idx >= 0
        out[valid] = self.cumhaz_[idx[valid]]
        return out

    def survival_matrix(self, log_risk: np.ndarray, eval_times: np.ndarray) -> np.ndarray:
        base_haz = self.cumulative_hazard_at(eval_times)
        risk = np.exp(np.asarray(log_risk, dtype=float)).reshape(-1, 1)
        return np.exp(-risk * base_haz.reshape(1, -1))

    def survival_functions(self, log_risk: np.ndarray) -> List[SimpleStepFunction]:
        if self.event_times_ is None:
            raise RuntimeError('Baseline estimator must be fit before prediction.')
        surv = self.survival_matrix(log_risk, self.event_times_)
        return [SimpleStepFunction(self.event_times_, surv[i, :]) for i in range(surv.shape[0])]

class LifelinesElasticNetCox(BaseEstimator):

    def __init__(self, penalizer: float=0.01, l1_ratio: float=0.5) -> None:
        self.penalizer = penalizer
        self.l1_ratio = l1_ratio

    def fit(self, X: pd.DataFrame, y: Any) -> 'LifelinesElasticNetCox':
        from lifelines import CoxPHFitter
        event_name, time_name = y.dtype.names
        df = X.copy()
        df[TIME_COL] = y[time_name].astype(float)
        df[EVENT_COL] = y[event_name].astype(int)
        self.model_ = CoxPHFitter(penalizer=self.penalizer, l1_ratio=self.l1_ratio)
        self.model_.fit(df, duration_col=TIME_COL, event_col=EVENT_COL)
        self.feature_names_in_ = list(X.columns)
        return self

    def predict(self, X: pd.DataFrame) -> np.ndarray:
        partial_hazard = self.model_.predict_partial_hazard(X[self.feature_names_in_])
        return np.log(np.asarray(partial_hazard, dtype=float).reshape(-1) + 1e-12)

    def predict_survival_function(self, X: pd.DataFrame) -> List[SimpleStepFunction]:
        surv_df = self.model_.predict_survival_function(X[self.feature_names_in_])
        times = surv_df.index.to_numpy(dtype=float)
        return [SimpleStepFunction(times, surv_df.iloc[:, i].to_numpy(dtype=float)) for i in range(surv_df.shape[1])]

    def predict_survival_matrix(self, X: pd.DataFrame, eval_times: np.ndarray) -> np.ndarray:
        surv_df = self.model_.predict_survival_function(X[self.feature_names_in_])
        return stepwise_survival_matrix(surv_df.index.to_numpy(dtype=float), surv_df.to_numpy(dtype=float).T, eval_times)

class XGBoostCoxWrapper(BaseEstimator):

    def __init__(self, n_estimators: int=300, learning_rate: float=0.05, max_depth: int=3, subsample: float=0.8, colsample_bytree: float=0.8, min_child_weight: float=1.0, reg_lambda: float=1.0, random_state: int=0, n_jobs: int=1) -> None:
        self.n_estimators = n_estimators
        self.learning_rate = learning_rate
        self.max_depth = max_depth
        self.subsample = subsample
        self.colsample_bytree = colsample_bytree
        self.min_child_weight = min_child_weight
        self.reg_lambda = reg_lambda
        self.random_state = random_state
        self.n_jobs = n_jobs

    def fit(self, X: pd.DataFrame, y: Any) -> 'XGBoostCoxWrapper':
        import xgboost as xgb
        event_name, time_name = y.dtype.names
        durations = y[time_name].astype(float)
        events = y[event_name].astype(int)
        labels = np.where(events == 1, durations, -durations)
        self.feature_names_in_ = list(X.columns)
        self.model_ = xgb.XGBRegressor(objective='survival:cox', eval_metric='cox-nloglik', n_estimators=self.n_estimators, learning_rate=self.learning_rate, max_depth=self.max_depth, subsample=self.subsample, colsample_bytree=self.colsample_bytree, min_child_weight=self.min_child_weight, reg_lambda=self.reg_lambda, random_state=self.random_state, n_jobs=self.n_jobs, tree_method='hist')
        self.model_.fit(X[self.feature_names_in_].to_numpy(dtype=float), labels)
        pred_hr = np.asarray(self.model_.predict(X[self.feature_names_in_].to_numpy(dtype=float)), dtype=float).reshape(-1)
        log_risk = np.log(np.clip(pred_hr, 1e-12, None))
        self.baseline_ = BreslowBaselineEstimator().fit(durations=durations, events=events, log_risk=log_risk)
        return self

    def predict(self, X: pd.DataFrame) -> np.ndarray:
        pred_hr = np.asarray(self.model_.predict(X[self.feature_names_in_].to_numpy(dtype=float)), dtype=float).reshape(-1)
        return np.log(np.clip(pred_hr, 1e-12, None))

    def predict_survival_function(self, X: pd.DataFrame) -> List[SimpleStepFunction]:
        log_risk = self.predict(X)
        return self.baseline_.survival_functions(log_risk)

    def predict_survival_matrix(self, X: pd.DataFrame, eval_times: np.ndarray) -> np.ndarray:
        return self.baseline_.survival_matrix(self.predict(X), eval_times)

def _clean_pycox_surv_df(surv_df: pd.DataFrame) -> pd.DataFrame:
    surv_df = surv_df.copy()
    surv_df.index = pd.Index(np.asarray(surv_df.index, dtype=float), name='time')
    surv_df = surv_df[~surv_df.index.duplicated(keep='last')].sort_index()
    times = surv_df.index.to_numpy(dtype=float)
    values = surv_df.to_numpy(dtype=float)
    if surv_df.empty or times.size == 0 or (not np.isfinite(times).all()) or np.any(times < 0.0) or np.any(np.diff(times) <= 0.0) or (not np.isfinite(values).all()) or np.any(values < -1e-08) or np.any(values > 1.0 + 1e-08) or np.any(np.diff(values, axis=0) > 1e-06):
        raise RuntimeError('PyCox produced an invalid survival-probability table.')
    return surv_df

def _risk_from_surv_df(surv_df: pd.DataFrame, time_point: float) -> np.ndarray:
    surv_df = _clean_pycox_surv_df(surv_df)
    times = surv_df.index.to_numpy(dtype=float)
    surv_mat = surv_df.to_numpy(dtype=float)
    pos = int(np.searchsorted(times, float(time_point), side='right') - 1)
    if pos < 0:
        surv = np.ones(surv_mat.shape[1], dtype=float)
    else:
        surv = surv_mat[pos, :]
    return 1.0 - np.asarray(surv, dtype=float).reshape(-1)

class PyCoxAdapter:

    def __init__(self, model: Any, preprocessor: TabularSurvivalPreprocessor, feature_names: Sequence[str], predict_from_surv: bool=False, risk_horizon: float=120.0) -> None:
        self.model = model
        self.preprocessor = preprocessor
        self.feature_names = list(feature_names)
        self.predict_from_surv = predict_from_surv
        self.risk_horizon = float(risk_horizon)

    def predict(self, X: Any) -> np.ndarray:
        X_df = as_dataframe(X, self.feature_names)
        x_mat = self.preprocessor.transform(X_df).to_numpy(dtype='float32')
        if self.predict_from_surv:
            surv_df = self.model.predict_surv_df(x_mat)
            return _risk_from_surv_df(surv_df, self.risk_horizon)
        preds = self.model.predict(x_mat)
        return np.asarray(preds, dtype=float).reshape(-1)

    def predict_survival_matrix(self, X: Any, eval_times: np.ndarray) -> np.ndarray:
        X_df = as_dataframe(X, self.feature_names)
        x_mat = self.preprocessor.transform(X_df).to_numpy(dtype='float32')
        surv_df = _clean_pycox_surv_df(self.model.predict_surv_df(x_mat))
        augmented_index = surv_df.index.union(pd.Index(eval_times)).sort_values()
        augmented = surv_df.reindex(augmented_index).ffill().fillna(1.0)
        return augmented.loc[eval_times].to_numpy().T

    def predict_survival_function(self, X: Any) -> List[SimpleStepFunction]:
        X_df = as_dataframe(X, self.feature_names)
        x_mat = self.preprocessor.transform(X_df).to_numpy(dtype='float32')
        surv_df = _clean_pycox_surv_df(self.model.predict_surv_df(x_mat))
        times = surv_df.index.to_numpy(dtype=float)
        return [SimpleStepFunction(times, surv_df.iloc[:, i].to_numpy(dtype=float)) for i in range(surv_df.shape[1])]

def setup_logging(log_path: str) -> None:
    os.makedirs(os.path.dirname(log_path), exist_ok=True)
    logging.basicConfig(level=logging.INFO, format='%(asctime)s | %(levelname)s | %(message)s', handlers=[logging.FileHandler(log_path, encoding='utf-8'), logging.StreamHandler(sys.stdout)])

def set_global_seed(seed: int) -> None:
    random.seed(seed)
    np.random.seed(seed)
    try:
        import torch
        requested_threads = max(1, int(os.environ.get('OMP_NUM_THREADS', '1')))
        torch.set_num_threads(requested_threads)
        try:
            torch.set_num_interop_threads(1)
        except RuntimeError:
            pass
        torch.manual_seed(seed)
        if torch.cuda.is_available():
            torch.cuda.manual_seed_all(seed)
        try:
            torch.use_deterministic_algorithms(True)
        except Exception:
            pass
    except Exception:
        pass

def ensure_dir(path: str) -> str:
    os.makedirs(path, exist_ok=True)
    return path

def save_json(obj: Any, path: str) -> None:
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(obj, f, indent=2, default=str)
CHECKPOINT_SCHEMA_VERSION = 5

def _canonical_json_text(obj: Any) -> str:
    return json.dumps(obj, sort_keys=True, separators=(',', ':'), ensure_ascii=False, allow_nan=False, default=str)

def _sha256_text(text_value: str) -> str:
    return hashlib.sha256(text_value.encode('utf-8')).hexdigest()

def file_sha256(path: str) -> str:
    digest = hashlib.sha256()
    with open(path, 'rb') as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b''):
            digest.update(chunk)
    return digest.hexdigest()

def dataframe_identity(df: pd.DataFrame) -> Dict[str, Any]:
    columns = [STUDY_ID_COL, TIME_COL, EVENT_COL, *PREDICTOR_COLUMNS]
    frame = df.loc[:, columns]
    csv_text = frame.to_csv(index=False, na_rep='<NA>', float_format='%.17g', lineterminator='\n')
    return {'n_rows': int(len(frame)), 'columns': columns, 'dtypes': {column: str(frame[column].dtype) for column in columns}, 'data_sha256': _sha256_text(csv_text)}

def current_script_sha256() -> str:
    script_path = os.path.abspath(__file__)
    if not os.path.exists(script_path):
        raise RuntimeError('The analysis script path could not be resolved for checkpoint validation.')
    return file_sha256(script_path)

def _temporary_path(final_path: str) -> Tuple[int, str]:
    directory = os.path.dirname(os.path.abspath(final_path))
    os.makedirs(directory, exist_ok=True)
    return tempfile.mkstemp(prefix=f'.{os.path.basename(final_path)}.', suffix='.tmp', dir=directory)

def atomic_write_json(obj: Any, final_path: str) -> None:
    fd, temporary_path = _temporary_path(final_path)
    try:
        with os.fdopen(fd, 'w', encoding='utf-8') as handle:
            json.dump(obj, handle, indent=2, sort_keys=True, default=str)
            handle.write('\n')
            handle.flush()
            os.fsync(handle.fileno())
        os.replace(temporary_path, final_path)
    except Exception:
        try:
            os.unlink(temporary_path)
        except FileNotFoundError:
            pass
        raise

def atomic_write_dataframe_csv(df: pd.DataFrame, final_path: str) -> None:
    fd, temporary_path = _temporary_path(final_path)
    try:
        with os.fdopen(fd, 'w', encoding='utf-8', newline='') as handle:
            df.to_csv(handle, index=False, float_format='%.17g')
            handle.flush()
            os.fsync(handle.fileno())
        os.replace(temporary_path, final_path)
    except Exception:
        try:
            os.unlink(temporary_path)
        except FileNotFoundError:
            pass
        raise

def prepare_checkpoint_manifest(checkpoint_dir: str, run_type: str, configuration: Dict[str, Any], resume: bool, allow_existing_validated: bool=False) -> str:
    os.makedirs(checkpoint_dir, exist_ok=True)
    manifest_path = os.path.join(checkpoint_dir, 'run_manifest.json')
    configuration_sha256 = _sha256_text(_canonical_json_text(configuration))
    expected_manifest = {'checkpoint_schema_version': CHECKPOINT_SCHEMA_VERSION, 'run_type': run_type, 'configuration_sha256': configuration_sha256, 'configuration': configuration}
    if os.path.exists(manifest_path):
        if not resume and (not allow_existing_validated):
            raise RuntimeError(f'Checkpoint manifest already exists at {manifest_path}. Use --resume only if this is the same run, or choose a new output directory.')
        with open(manifest_path, 'r', encoding='utf-8') as handle:
            existing_manifest = json.load(handle)
        if existing_manifest.get('checkpoint_schema_version') != CHECKPOINT_SCHEMA_VERSION or existing_manifest.get('run_type') != run_type or existing_manifest.get('configuration_sha256') != configuration_sha256 or (_canonical_json_text(existing_manifest.get('configuration')) != _canonical_json_text(configuration)):
            raise RuntimeError(f"Checkpoint manifest mismatch. Resume was stopped before loading any repeat. Existing hash={existing_manifest.get('configuration_sha256')}; requested hash={configuration_sha256}. Use the original configuration or a new output directory.")
        LOGGER.info('Validated checkpoint manifest: %s', manifest_path)
        return configuration_sha256
    remaining_files = [name for name in os.listdir(checkpoint_dir) if not name.endswith('.tmp')]
    if remaining_files:
        raise RuntimeError(f'Checkpoint directory {checkpoint_dir} contains files but no run manifest. Resume was stopped to avoid mixing runs. Use a new output directory.')
    atomic_write_json(expected_manifest, manifest_path)
    if resume:
        LOGGER.info('No prior checkpoint manifest was present; started a new resumable run at %s', checkpoint_dir)
    else:
        LOGGER.info('Created checkpoint manifest: %s', manifest_path)
    return configuration_sha256
FIT_REQUIRED_METRICS = ['harrell_c', 'uno_c_tau', 'brier_score_at_horizon', 'integrated_brier_score', 'calibration_slope', 'calibration_oe_ratio']

def _required_metric_columns(model_name: str) -> List[str]:
    if model_name == 'SVM':
        return ['harrell_c', 'uno_c_tau']
    return list(FIT_REQUIRED_METRICS)

def validate_metric_bounds(metric_name: str, value: Any, context: str) -> float:
    numeric = float(value)
    if not np.isfinite(numeric):
        raise ValueError(f'{context}: {metric_name} is non-finite.')
    if metric_name in {'harrell_c', 'uno_c_tau', 'brier_score_at_horizon', 'integrated_brier_score'} and (not 0.0 <= numeric <= 1.0):
        raise ValueError(f'{context}: {metric_name}={numeric} is outside [0, 1].')
    if metric_name == 'calibration_oe_ratio' and numeric <= 0.0:
        raise ValueError(f'{context}: O:E must be positive.')
    return numeric

def _validated_boolean_series(series: pd.Series, label: str) -> pd.Series:
    if pd.api.types.is_bool_dtype(series):
        return series.fillna(False).astype(bool)
    normalized = series.fillna('').astype(str).str.strip().str.lower()
    if not normalized.isin({'true', 'false', '1', '0'}).all():
        raise ValueError(f'{label} contained a value that was not a serialized boolean.')
    return normalized.isin({'true', '1'})
TUNING_SELECTION_COLUMNS = {'selection_score', 'selection_metric', 'selection_direction', 'model_selection_performed', 'validation_design', 'selection_window_start_months', 'selection_window_end_months', 'selection_grid_step_months', 'selection_grid_points', 'n_prespecified_configurations'}

def validate_tuning_selection_record(record: Any, model_name: str, context: str) -> None:
    missing = TUNING_SELECTION_COLUMNS.difference(record.index)
    if missing:
        raise ValueError(f'{context} lacks tuning fields: {sorted(missing)}.')
    expected = model_selection_metadata(model_name)
    for field in ['selection_metric', 'selection_direction', 'validation_design']:
        if str(record[field]) != str(expected[field]):
            raise ValueError(f'{context} has the wrong {field}.')
    selected = bool(_validated_boolean_series(pd.Series([record['model_selection_performed']]), f'{context} model_selection_performed').iloc[0])
    if selected != bool(expected['model_selection_performed']):
        raise ValueError(f'{context} has the wrong model-selection flag.')
    candidate_count = int(pd.to_numeric(pd.Series([record['n_prespecified_configurations']]), errors='raise').iloc[0])
    if candidate_count != int(expected['n_prespecified_configurations']):
        raise ValueError(f'{context} has the wrong candidate count.')
    selection_score = pd.to_numeric(pd.Series([record['selection_score']]), errors='coerce').iloc[0]
    if selected:
        if not np.isfinite(selection_score) or not 0.0 <= float(selection_score) <= 1.0:
            raise ValueError(f'{context} has an invalid selection score.')
    elif not pd.isna(selection_score):
        raise ValueError(f'{context} fixed model unexpectedly has a selection score.')
    window_fields = ['selection_window_start_months', 'selection_window_end_months', 'selection_grid_step_months', 'selection_grid_points']
    if expected['selection_metric'] == IBS_SELECTION_METRIC:
        for field in window_fields:
            observed = float(record[field])
            if not np.isclose(observed, float(expected[field]), rtol=0.0, atol=1e-12):
                raise ValueError(f'{context} has the wrong {field}.')
    elif not all((pd.isna(record[field]) for field in window_fields)):
        raise ValueError(f'{context} non-IBS model unexpectedly records an IBS window.')

def validate_model_imputation_tuning_frame(tuning_df: pd.DataFrame, repeat: int, expected_fit_rows: int) -> None:
    required = {'repeat', 'model', 'best_params_json', *TUNING_SELECTION_COLUMNS}
    if not required.issubset(tuning_df.columns):
        raise ValueError(f'Tuning checkpoint columns were missing: {sorted(required.difference(tuning_df.columns))}')
    expected_models = set(MODEL_ORDER)
    if not (pd.to_numeric(tuning_df['repeat'], errors='raise').astype(int) == repeat).all():
        raise ValueError('Tuning checkpoint rows contained the wrong fixed-split label.')
    if 'imputation' in tuning_df.columns:
        if expected_fit_rows % len(MODEL_ORDER):
            raise ValueError('Expected fit rows do not describe a complete tuning grid.')
        m_imputations = expected_fit_rows // len(MODEL_ORDER)
        expected_keys = {(repeat, imputation, model_name) for imputation in range(1, m_imputations + 1) for model_name in MODEL_ORDER}
        observed_keys = set(zip(pd.to_numeric(tuning_df['repeat'], errors='raise').astype(int), pd.to_numeric(tuning_df['imputation'], errors='raise').astype(int), tuning_df['model'].astype(str)))
        if len(tuning_df) != expected_fit_rows or observed_keys != expected_keys:
            raise ValueError('Tuning checkpoint did not contain the exact model-imputation grid.')
        if tuning_df.duplicated(['repeat', 'imputation', 'model']).any():
            raise ValueError('Tuning checkpoint contained duplicate keys.')
    else:
        if len(tuning_df) != len(MODEL_ORDER) or set(tuning_df['model'].astype(str)) != expected_models:
            raise ValueError('Fixed tuning checkpoint did not contain every model.')
        if tuning_df.duplicated(['repeat', 'model']).any():
            raise ValueError('Fixed tuning checkpoint contained duplicate keys.')
    for value in tuning_df['best_params_json'].astype(str):
        if not isinstance(json.loads(value), dict):
            raise ValueError('A tuning checkpoint best_params_json value was not an object.')
    for _, row in tuning_df.iterrows():
        validate_tuning_selection_record(row, str(row['model']), 'Tuning checkpoint')

def as_dataframe(x: Any, columns: Sequence[str]) -> pd.DataFrame:
    if isinstance(x, pd.DataFrame):
        return x.loc[:, list(columns)].copy()
    if isinstance(x, np.ndarray):
        return pd.DataFrame(x, columns=list(columns))
    return pd.DataFrame(np.asarray(x), columns=list(columns))

def format_mean_sd(series: pd.Series, decimals: int=1) -> str:
    return f'{series.mean():.{decimals}f} ({series.std(ddof=1):.{decimals}f})'

def format_median_iqr(series: pd.Series, decimals: int=1) -> str:
    q1, q3 = series.quantile([0.25, 0.75])
    return f'{series.median():.{decimals}f} [{q1:.{decimals}f}, {q3:.{decimals}f}]'

def format_n_pct(series: pd.Series, positive_value: int=1) -> str:
    n = int((series == positive_value).sum())
    pct = 100.0 * n / len(series)
    return f'{n} ({pct:.1f}%)'

def smd_continuous(x1: pd.Series, x2: pd.Series) -> float:
    x1 = x1.astype(float)
    x2 = x2.astype(float)
    v1 = x1.var(ddof=1)
    v2 = x2.var(ddof=1)
    pooled_sd = np.sqrt((v1 + v2) / 2.0)
    if pooled_sd == 0 or np.isnan(pooled_sd):
        return 0.0
    return float((x1.mean() - x2.mean()) / pooled_sd)

def smd_binary(x1: pd.Series, x2: pd.Series, positive_value: int=1) -> float:
    p1 = (x1 == positive_value).mean()
    p2 = (x2 == positive_value).mean()
    denom = np.sqrt((p1 * (1 - p1) + p2 * (1 - p2)) / 2.0)
    if denom == 0 or np.isnan(denom):
        return 0.0
    return float((p1 - p2) / denom)

def get_package_versions() -> Dict[str, Optional[str]]:
    packages = ['numpy', 'pandas', 'scikit-learn', 'joblib', 'scikit-survival', 'lifelines', 'xgboost', 'shap', 'matplotlib', 'scipy', 'torch', 'torchtuples', 'pycox', 'python-docx', 'Pillow', 'openpyxl']
    versions: Dict[str, Optional[str]] = {'python': platform.python_version()}
    for pkg in packages:
        try:
            versions[pkg] = importlib.metadata.version(pkg)
        except importlib.metadata.PackageNotFoundError:
            versions[pkg] = None
    return versions

def validate_runtime_and_data(csv_path: str) -> None:
    required_imports = {'numpy': 'numpy', 'pandas': 'pandas', 'scikit-learn': 'sklearn', 'joblib': 'joblib', 'scikit-survival': 'sksurv', 'lifelines': 'lifelines', 'xgboost': 'xgboost', 'shap': 'shap', 'matplotlib': 'matplotlib', 'scipy': 'scipy', 'torch': 'torch', 'torchtuples': 'torchtuples', 'pycox': 'pycox', 'python-docx': 'docx', 'Pillow': 'PIL', 'openpyxl': 'openpyxl'}
    required_distributions = list(required_imports)
    versions = get_package_versions()
    missing = [name for name in required_distributions if versions.get(name) is None]
    if missing:
        raise RuntimeError('Required Python packages are missing: ' + ', '.join(missing) + '. Install requirements.txt in the same Python environment before starting the analysis.')
    import_errors: List[str] = []
    for distribution, module_name in required_imports.items():
        try:
            importlib.import_module(module_name)
        except Exception as exc:
            import_errors.append(f'{distribution} ({type(exc).__name__}: {exc})')
    if import_errors:
        raise RuntimeError('Installed packages could not be imported: ' + '; '.join(import_errors))
    version_mismatches: List[str] = []
    if versions['python'] != EXPECTED_PYTHON_VERSION:
        version_mismatches.append(f"python: expected {EXPECTED_PYTHON_VERSION}, found {versions['python']}")
    version_mismatches.extend((f'{name}: expected {expected}, found {versions.get(name)}' for name, expected in EXPECTED_PACKAGE_VERSIONS.items() if versions.get(name) != expected))
    if version_mismatches:
        raise RuntimeError('The Python environment does not match the locked analysis environment. Do not start or resume the long analysis with different versions. Create/select the CPython 3.11.7 environment and install requirements.txt. Differences: ' + '; '.join(version_mismatches))
    source_row_count = int(len(pd.read_csv(csv_path, usecols=[STUDY_ID_COL])))
    df = load_dataset(csv_path)
    if df[STUDY_ID_COL].duplicated().any():
        duplicate_count = int(df[STUDY_ID_COL].duplicated(keep=False).sum())
        raise ValueError(f'The filtered analysis dataset contains {duplicate_count} rows with duplicated study IDs.')
    event_count = int(df[EVENT_COL].sum())
    expected_checks = {'source rows': (source_row_count, EXPECTED_SOURCE_ROWS), 'filtered participants': (len(df), EXPECTED_FILTERED_ROWS), 'ASCVD events': (event_count, EXPECTED_ASCVD_EVENTS), 'unique study IDs': (int(df[STUDY_ID_COL].nunique()), EXPECTED_UNIQUE_STUDY_IDS)}
    count_mismatches = [f'{label}: found {found}, expected {expected}' for label, (found, expected) in expected_checks.items() if found != expected]
    if count_mismatches:
        raise ValueError('The CSV does not match the validated study dataset: ' + '; '.join(count_mismatches))
    missingness = create_missingness_table(df)
    print('Runtime and data validation passed.')
    print(f"Python: {versions['python']}")
    print(f'Filtered participants: {len(df)}')
    print(f'ASCVD events: {event_count}')
    print(f'Unique study IDs: {df[STUDY_ID_COL].nunique()}')
    print(f'Predictors with missing values: {len(missingness)}')
    print('Detected package versions:')
    for name in required_distributions:
        print(f'  {name}: {versions[name]}')
    print('Locked scientific package pins, including joblib, match and are recorded in the run and checkpoint identities.')

def checkpoint_common_identity(df: pd.DataFrame) -> Dict[str, Any]:
    scientific_versions = get_package_versions().copy()
    scientific_versions.pop('platform', None)
    return {'dataset': dataframe_identity(df), 'analysis_script_sha256': current_script_sha256(), 'package_versions': scientific_versions, 'predictor_columns': list(PREDICTOR_COLUMNS), 'continuous_columns': list(CONTINUOUS_COLUMNS), 'binary_columns': list(BINARY_COLUMNS), 'model_order': list(MODEL_ORDER), 'deepsurv_batching': 'full_training_risk_set', 'model_selection_policy': {model_name: model_selection_metadata(model_name) for model_name in MODEL_ORDER}}

def repeated_split_checkpoint_configuration(df: pd.DataFrame, seed: int, n_splits: int, test_size: float, inner_folds: int, horizon_months: float, ibs_start_month: float, calibration_groups: int, n_jobs: int, impute_max_iter: int, calibration_bootstrap_reps: int) -> Dict[str, Any]:
    return {'identity': checkpoint_common_identity(df), 'settings': {'seed': int(seed), 'n_splits': int(n_splits), 'test_size': float(test_size), 'expected_training_participants_per_split': 992, 'expected_test_participants_per_split': 249, 'expected_training_events_per_split': 92, 'expected_test_events_per_split': 24, 'expected_held_out_prediction_rows': 49800, 'expected_participant_coverage_minimum': 21, 'expected_participant_coverage_maximum': 57, 'inner_folds': int(inner_folds), 'horizon_months': float(horizon_months), 'ibs_start_month': float(ibs_start_month), 'calibration_groups': int(calibration_groups), 'n_jobs': int(n_jobs), 'imputations_per_outer_split': 1, 'impute_max_iter': int(impute_max_iter), 'calibration_bootstrap_reps': int(calibration_bootstrap_reps), 'splitter': 'single_StratifiedShuffleSplit_stream_joint_event_status_and_sex', 'seed_scheme': 'splitter random_state=seed; split_base_seed=seed+100000*split_id; outer Bayesian imputer=split_base_seed; model/tuning=split_base_seed+1; classical inner imputer=split_base_seed+100+inner_fold; neural development imputer=split_base_seed+500', 'classical_tuning': 'fixed_CoxPH; configurable_absolute_risk_models_minimize_mean_IBS_12_120_over_three_folds; SVM_maximizes_mean_Harrell_C; imputer_refit_on_each_inner_training_fold', 'neural_tuning': 'prespecified_configurations_on_internal_stratified_80_20_validation; configuration_selected_by_minimum_IBS_12_120; DeepSurv_full_training_risk_set_batch; selected_configuration_refit_on_full_outer_training_split', 'selection_grid_months': list(range(12, 121)), 'imputation_timing': 'outer_evaluation_imputer_fit_on_full_outer_training_predictors; inner_selection_imputer_refit_on_each_inner_training_fold; neural_selection_imputer_fit_on_internal_development_predictors', 'performance_summary': 'mean_and_empirical_2.5th_97.5th_percentiles_across_200_split_estimates; descriptive_not_confidence_intervals; no_P_values', 'repeated_split_calibration': 'average_each_participant_predictions_only_across_splits_where_held_out; one_row_per_participant_for_OE_slope_grouped_and_flexible_calibration; never_stack_repeated_predictions_as_independent_observations; participant_bootstrap_bands_conditional_on_averaged_predictions'}}

def load_dataset(csv_path: str) -> pd.DataFrame:
    df = pd.read_csv(csv_path)
    df = df.copy()
    required_core = [STUDY_ID_COL, TIME_COL, EVENT_COL]
    missing_core = [column for column in required_core if column not in df.columns]
    if missing_core:
        raise ValueError(f'Dataset is missing required columns: {missing_core}')
    missing_predictors = [c for c in PREDICTOR_COLUMNS if c not in df.columns]
    if missing_predictors:
        raise ValueError(f'Missing predictor columns: {missing_predictors}')
    if df[STUDY_ID_COL].isna().any():
        raise ValueError(f'{STUDY_ID_COL!r} contains missing values.')
    df[TIME_COL] = pd.to_numeric(df[TIME_COL], errors='raise')
    event_values = pd.to_numeric(df[EVENT_COL], errors='raise')
    if event_values.isna().any() or not event_values.isin([0, 1]).all():
        invalid_values = sorted(event_values.loc[~event_values.isin([0, 1])].dropna().unique().tolist())
        raise ValueError(f'{EVENT_COL!r} must contain only non-missing 0 and 1 values; found {invalid_values}.')
    df[EVENT_COL] = event_values.astype(int)
    for column in PREDICTOR_COLUMNS:
        df[column] = pd.to_numeric(df[column], errors='raise')
    if not np.isfinite(df[TIME_COL].to_numpy(dtype=float)).all():
        raise ValueError(f'{TIME_COL!r} must contain finite numeric values.')
    finite_predictors = df[PREDICTOR_COLUMNS].replace([np.inf, -np.inf], np.nan)
    newly_nonfinite = finite_predictors.isna() & ~df[PREDICTOR_COLUMNS].isna()
    if newly_nonfinite.any().any():
        bad_columns = newly_nonfinite.columns[newly_nonfinite.any()].tolist()
        raise ValueError(f'Predictor columns contain infinite values: {bad_columns}')
    for column in BINARY_COLUMNS:
        observed = df[column].dropna()
        if not observed.isin([0, 1]).all():
            invalid_values = sorted(observed.loc[~observed.isin([0, 1])].unique().tolist())
            raise ValueError(f'{column!r} must contain only 0, 1, or missing values; found {invalid_values}.')
    before = len(df)
    df = df[df[TIME_COL] > 0].reset_index(drop=True)
    dropped = before - len(df)
    if dropped > 0:
        LOGGER.info('Dropped %d rows with non-positive follow-up time.', dropped)
    return df

def create_missingness_table(df: pd.DataFrame) -> pd.DataFrame:
    predictor_df = df.loc[:, PREDICTOR_COLUMNS]
    out = pd.DataFrame({'Variable': predictor_df.columns, 'n_missing': predictor_df.isna().sum().values, 'pct_missing': 100.0 * predictor_df.isna().mean().values, 'dtype': [str(predictor_df[c].dtype) for c in predictor_df.columns]})
    return out[out['n_missing'] > 0].sort_values(['pct_missing', 'Variable'], ascending=[False, True]).reset_index(drop=True)

def make_strata(df: pd.DataFrame) -> pd.Series:
    candidate = df[EVENT_COL].astype(str) + '_' + df[SEX_COL].astype(str)
    if candidate.value_counts().min() >= 2:
        return candidate
    return df[EVENT_COL].astype(str)

def make_repeated_split_strata(df: pd.DataFrame) -> pd.Series:
    candidate = df[EVENT_COL].astype(str) + '_' + df[SEX_COL].astype(str)
    if int(candidate.value_counts().min()) < 2:
        raise ValueError('At least one joint event/sex stratum has fewer than two records; stratified repeated 80/20 splitting is not feasible.')
    return candidate

def iter_repeated_splits(df: pd.DataFrame, seed: int, n_splits: int, test_size: float) -> Iterable[Tuple[int, np.ndarray, np.ndarray]]:
    if int(seed) != PRESPECIFIED_SPLITTER_SEED:
        raise ValueError(f'The repeated-split seed must be {PRESPECIFIED_SPLITTER_SEED}.')
    if int(n_splits) != PRESPECIFIED_N_SPLITS:
        raise ValueError(f'The analysis requires exactly {PRESPECIFIED_N_SPLITS} splits.')
    if not math.isclose(float(test_size), PRESPECIFIED_TEST_SIZE, rel_tol=0.0, abs_tol=1e-12):
        raise ValueError(f'The repeated-split test fraction must be {PRESPECIFIED_TEST_SIZE:.2f}.')
    strata = make_repeated_split_strata(df)
    splitter = StratifiedShuffleSplit(n_splits=int(n_splits), test_size=float(test_size), random_state=int(seed))
    seen_test_sets: set[Tuple[int, ...]] = set()
    for split_id, (train_idx, test_idx) in enumerate(splitter.split(df, strata), start=1):
        identity = tuple(sorted(np.asarray(test_idx, dtype=int).tolist()))
        if identity in seen_test_sets:
            raise RuntimeError(f'Repeated split {split_id} duplicates an earlier held-out set.')
        seen_test_sets.add(identity)
        if len(train_idx) != 992 or len(test_idx) != 249:
            raise RuntimeError(f'Split {split_id} does not contain 992 training and 249 test participants.')
        train_events = int(df.iloc[train_idx][EVENT_COL].sum())
        test_events = int(df.iloc[test_idx][EVENT_COL].sum())
        if (train_events, test_events) != (92, 24):
            raise RuntimeError(f'Split {split_id} has {train_events}/{test_events} train/test events; expected 92/24.')
        yield (split_id, np.asarray(train_idx, dtype=int), np.asarray(test_idx, dtype=int))
    if len(seen_test_sets) != PRESPECIFIED_N_SPLITS:
        raise RuntimeError('The repeated-split stream did not produce 200 unique held-out sets.')

def primary_split(df: pd.DataFrame, seed: int, test_size: float) -> Tuple[np.ndarray, np.ndarray]:
    splitter = StratifiedShuffleSplit(n_splits=1, test_size=test_size, random_state=seed)
    strata = make_strata(df)
    train_idx, test_idx = next(splitter.split(df, strata))
    return (train_idx, test_idx)

def get_surv_array(df: pd.DataFrame) -> Any:
    from sksurv.util import Surv
    try:
        return Surv.from_arrays(event=df[EVENT_COL].astype(bool).to_numpy(), time=df[TIME_COL].astype(float).to_numpy(), name_event='event', name_time='time')
    except TypeError:
        return Surv.from_arrays(event=df[EVENT_COL].astype(bool).to_numpy(), time=df[TIME_COL].astype(float).to_numpy())

def get_structured_field_names(y: Any) -> Tuple[str, str]:
    event_name, time_name = y.dtype.names
    return (event_name, time_name)

def harrell_c_index(y_true: Any, risk_scores: np.ndarray) -> float:
    from sksurv.metrics import concordance_index_censored
    event_name, time_name = get_structured_field_names(y_true)
    return float(concordance_index_censored(y_true[event_name], y_true[time_name], np.asarray(risk_scores, dtype=float))[0])

def uno_c_index(y_train: Any, y_test: Any, risk_scores: np.ndarray, tau: float) -> float:
    from sksurv.metrics import concordance_index_ipcw
    risk_scores = np.asarray(risk_scores, dtype=float)
    return float(concordance_index_ipcw(y_train, y_test, risk_scores, tau=tau)[0])

def get_eval_times(train_df: pd.DataFrame, test_df: pd.DataFrame, start_month: float, horizon_months: float) -> np.ndarray:
    del test_df
    start = float(start_month)
    horizon = float(horizon_months)
    if not (np.isfinite(start) and np.isfinite(horizon) and (0.0 < start < horizon)):
        raise ValueError('The IBS start and horizon must be finite with 0 < start < horizon.')
    if float(train_df[TIME_COL].max()) <= horizon:
        raise ValueError('Training follow-up must extend strictly beyond the prespecified Brier horizon.')
    times = np.arange(start, horizon + 1.0, 1.0, dtype=float)
    if times.size < 2 or not np.isclose(times[-1], horizon, rtol=0.0, atol=1e-10):
        raise ValueError('The prespecified one-month IBS grid must end exactly at the horizon.')
    return times

def stepwise_survival_matrix(source_times: np.ndarray, source_survival: np.ndarray, eval_times: np.ndarray) -> np.ndarray:
    times = np.asarray(source_times, dtype=float).reshape(-1)
    survival = np.asarray(source_survival, dtype=float)
    requested = np.asarray(eval_times, dtype=float).reshape(-1)
    if survival.ndim != 2 or survival.shape[1] != len(times):
        raise ValueError('The source survival matrix does not match its time grid.')
    if times.size == 0 or not np.isfinite(times).all() or np.any(np.diff(times) <= 0.0) or (not np.isfinite(survival).all()):
        raise ValueError('The source survival curves are invalid.')
    indices = np.searchsorted(times, requested, side='right') - 1
    out = np.ones((survival.shape[0], len(requested)), dtype=float)
    available = indices >= 0
    if np.any(available):
        out[:, available] = survival[:, indices[available]]
    return out

def predict_survival_matrix(fitted_model: Any, X: pd.DataFrame, eval_times: np.ndarray) -> np.ndarray:
    if hasattr(fitted_model, 'named_steps') and {'prep', 'model'}.issubset(fitted_model.named_steps):
        transformed = fitted_model.named_steps['prep'].transform(X)
        final_estimator = fitted_model.named_steps['model']
        if hasattr(final_estimator, 'predict_survival_matrix'):
            return np.asarray(final_estimator.predict_survival_matrix(transformed, eval_times), dtype=float)
        try:
            survival_at_native_times = np.asarray(final_estimator.predict_survival_function(transformed, return_array=True), dtype=float)
            native_times = np.asarray(final_estimator.unique_times_, dtype=float)
            return stepwise_survival_matrix(native_times, survival_at_native_times, eval_times)
        except (AttributeError, TypeError):
            pass
        surv_fns = final_estimator.predict_survival_function(transformed)
        return np.asarray([[fn(t) for t in eval_times] for fn in surv_fns], dtype=float)
    if hasattr(fitted_model, 'predict_survival_matrix'):
        return np.asarray(fitted_model.predict_survival_matrix(X, eval_times), dtype=float)
    surv_fns = fitted_model.predict_survival_function(X)
    return np.asarray([[fn(t) for t in eval_times] for fn in surv_fns], dtype=float)

def prepare_brier_scoring_context(y_train: Any, y_test: Any, eval_times: np.ndarray, horizon_months: float) -> Dict[str, Any]:
    from sksurv.nonparametric import CensoringDistributionEstimator
    times = np.asarray(eval_times, dtype=float).reshape(-1)
    if times.size < 2 or not np.isfinite(times).all() or np.any(np.diff(times) <= 0.0):
        raise ValueError('Brier evaluation times must be finite and strictly increasing.')
    event_name, time_name = get_structured_field_names(y_test)
    test_events = np.asarray(y_test[event_name], dtype=bool)
    test_times = np.asarray(y_test[time_name], dtype=float)
    if not np.isfinite(test_times).all() or np.any(test_times <= 0.0):
        raise ValueError('Held-out survival times must be positive and finite.')
    censoring = CensoringDistributionEstimator().fit(y_train)
    g_at_times = np.asarray(censoring.predict_proba(times), dtype=float)
    if g_at_times.shape != times.shape or not np.isfinite(g_at_times).all() or np.any(g_at_times <= 0.0):
        raise ValueError('Training censoring survival is non-positive on the Brier grid.')
    event_by_end = test_events & (test_times <= times[-1])
    g_at_event = np.ones(len(y_test), dtype=float)
    if np.any(event_by_end):
        queried = np.asarray(censoring.predict_proba(test_times[event_by_end]), dtype=float)
        if not np.isfinite(queried).all() or np.any(queried <= 0.0):
            raise ValueError('Training censoring survival is non-positive at a held-out event time.')
        g_at_event[event_by_end] = queried
    n_test = float(len(y_test))
    if n_test <= 0.0:
        raise ValueError('Brier scoring requires at least one held-out observation.')
    horizon_index = int(np.argmin(np.abs(times - float(horizon_months))))
    horizon_used = float(times[horizon_index])
    if not np.isclose(horizon_used, float(horizon_months), rtol=0.0, atol=1e-10):
        raise ValueError('The Brier grid did not contain the prespecified horizon.')
    cases = test_events[:, None] & (test_times[:, None] <= times[None, :])
    controls = test_times[:, None] > times[None, :]
    case_weights = cases.astype(float) / g_at_event[:, None]
    control_weights = controls.astype(float) / g_at_times[None, :]
    if not np.isfinite(case_weights).all() or not np.isfinite(control_weights).all() or np.any(case_weights < 0.0) or np.any(control_weights < 0.0):
        raise ValueError('Prepared Brier weights are invalid.')
    return {'times': times, 'n_test': int(len(y_test)), 'case_weights': case_weights, 'control_weights': control_weights, 'horizon_index': horizon_index, 'horizon_used': horizon_used}

def compute_brier_metrics_from_context(surv_matrix: np.ndarray, context: Dict[str, Any]) -> Dict[str, float]:
    times = np.asarray(context['times'], dtype=float)
    n_test = int(context['n_test'])
    survival = np.asarray(surv_matrix, dtype=float)
    if survival.shape != (n_test, len(times)):
        raise ValueError('The survival-probability matrix has the wrong Brier shape.')
    if not np.isfinite(survival).all() or np.any(survival < -1e-08) or np.any(survival > 1.0 + 1e-08):
        raise ValueError('Brier survival probabilities must be finite and lie in [0, 1].')
    survival = np.clip(survival, 0.0, 1.0)
    if np.any(np.diff(survival, axis=1) > 1e-06):
        raise ValueError('A predicted survival curve increased over time.')
    case_weights = np.asarray(context['case_weights'], dtype=float)
    control_weights = np.asarray(context['control_weights'], dtype=float)
    if case_weights.shape != survival.shape or control_weights.shape != survival.shape:
        raise ValueError('The prepared Brier weights have the wrong shape.')
    scores = np.mean(np.square(survival) * case_weights + np.square(1.0 - survival) * control_weights, axis=0)
    if not np.isfinite(scores).all():
        raise ValueError('A Graf Brier score was non-finite.')
    horizon_index = int(context['horizon_index'])
    horizon_used = float(context['horizon_used'])
    ibs = float(np.trapz(scores, times) / (times[-1] - times[0]))
    return {'brier_horizon_months': horizon_used, 'brier_score_at_horizon': float(scores[horizon_index]), 'integrated_brier_score': ibs}

def compute_brier_metrics(y_train: Any, y_test: Any, surv_matrix: np.ndarray, eval_times: np.ndarray, horizon_months: float) -> Dict[str, float]:
    context = prepare_brier_scoring_context(y_train=y_train, y_test=y_test, eval_times=eval_times, horizon_months=horizon_months)
    return compute_brier_metrics_from_context(surv_matrix, context)

def km_risk_with_greenwood_ci(durations: Sequence[float], events: Sequence[int], time_point: float, alpha: float=0.05) -> Dict[str, float]:
    durations_arr = np.asarray(durations, dtype=float)
    events_arr = np.asarray(events, dtype=int)
    if durations_arr.ndim != 1 or events_arr.ndim != 1 or len(durations_arr) != len(events_arr):
        raise ValueError('Durations and events must be one-dimensional arrays of equal length.')
    if len(durations_arr) == 0:
        raise ValueError('Kaplan-Meier risk cannot be estimated from an empty sample.')
    if not np.all(np.isfinite(durations_arr)) or not np.all(np.isin(events_arr, [0, 1])):
        raise ValueError('Durations must be finite and events must be coded 0/1.')
    survival = 1.0
    greenwood_sum = 0.0
    event_times = np.sort(np.unique(durations_arr[(events_arr == 1) & (durations_arr <= float(time_point))]))
    for event_time in event_times:
        n_at_risk = int(np.sum(durations_arr >= event_time))
        n_events = int(np.sum((durations_arr == event_time) & (events_arr == 1)))
        if n_events == 0:
            continue
        if n_at_risk <= n_events:
            survival = 0.0
            greenwood_sum = float('inf')
            break
        survival *= 1.0 - n_events / n_at_risk
        greenwood_sum += n_events / (n_at_risk * (n_at_risk - n_events))
    risk = float(1.0 - survival)
    if survival >= 1.0 or greenwood_sum == 0.0:
        risk_low, risk_high = (risk, risk)
    elif survival <= 0.0 or not np.isfinite(greenwood_sum):
        risk_low, risk_high = (1.0, 1.0)
    else:
        z_value = float(norm.ppf(1.0 - alpha / 2.0))
        log_minus_log_survival = float(np.log(-np.log(survival)))
        se_log_minus_log = float(np.sqrt(greenwood_sum) / abs(np.log(survival)))
        survival_low = float(np.exp(-np.exp(log_minus_log_survival + z_value * se_log_minus_log)))
        survival_high = float(np.exp(-np.exp(log_minus_log_survival - z_value * se_log_minus_log)))
        risk_low = float(1.0 - survival_high)
        risk_high = float(1.0 - survival_low)
    return {'risk': risk, 'risk_ci_low': float(np.clip(risk_low, 0.0, 1.0)), 'risk_ci_high': float(np.clip(risk_high, 0.0, 1.0)), 'survival': float(survival), 'greenwood_sum': float(greenwood_sum)}

def km_risk_at_time(durations: Sequence[float], events: Sequence[int], time_point: float) -> float:
    return km_risk_with_greenwood_ci(durations, events, time_point)['risk']

def calibration_oe_from_probs(test_df: pd.DataFrame, predicted_event_prob: np.ndarray, time_point: float) -> Dict[str, float]:
    probabilities = np.asarray(predicted_event_prob, dtype=float)
    if len(probabilities) != len(test_df) or not np.all(np.isfinite(probabilities)):
        raise ValueError('Predicted probabilities must be finite and match the test-set length.')
    if np.any((probabilities < -1e-08) | (probabilities > 1.0 + 1e-08)):
        raise ValueError('Predicted event probabilities must lie between 0 and 1.')
    probabilities = np.clip(probabilities, 0.0, 1.0)
    expected_mean_risk = float(np.mean(probabilities))
    if expected_mean_risk <= 0.0:
        raise ValueError('O:E cannot be calculated when mean predicted risk is zero.')
    km = km_risk_with_greenwood_ci(test_df[TIME_COL].to_numpy(dtype=float), test_df[EVENT_COL].to_numpy(dtype=int), time_point)
    observed_risk = float(km['risk'])
    n = int(len(test_df))
    return {'calibration_observed_km_risk': observed_risk, 'calibration_observed_km_risk_ci_low': float(km['risk_ci_low']), 'calibration_observed_km_risk_ci_high': float(km['risk_ci_high']), 'calibration_mean_predicted_risk': expected_mean_risk, 'calibration_observed_events_km': float(n * observed_risk), 'calibration_expected_events': float(np.sum(probabilities)), 'calibration_oe_ratio': float(observed_risk / expected_mean_risk), 'calibration_oe_ratio_ci_low': float(km['risk_ci_low'] / expected_mean_risk), 'calibration_oe_ratio_ci_high': float(km['risk_ci_high'] / expected_mean_risk)}

def grouped_calibration_table(test_df: pd.DataFrame, predicted_event_prob: np.ndarray, time_point: float, n_groups: int=5) -> pd.DataFrame:
    probabilities = np.asarray(predicted_event_prob, dtype=float)
    if len(probabilities) != len(test_df) or not np.all(np.isfinite(probabilities)):
        raise ValueError('Predicted probabilities must be finite and match the test-set length.')
    if np.any((probabilities < -1e-08) | (probabilities > 1.0 + 1e-08)):
        raise ValueError('Predicted event probabilities must lie between 0 and 1.')
    cal_df = pd.DataFrame({'predicted_event_prob': np.clip(probabilities, 0.0, 1.0), TIME_COL: test_df[TIME_COL].astype(float).values, EVENT_COL: test_df[EVENT_COL].astype(int).values})
    n_groups = max(2, min(n_groups, cal_df.shape[0]))
    cal_df['group'] = pd.qcut(cal_df['predicted_event_prob'].rank(method='first'), q=n_groups, labels=False, duplicates='drop')
    rows: List[Dict[str, Any]] = []
    for grp, grp_df in cal_df.groupby('group'):
        km = km_risk_with_greenwood_ci(grp_df[TIME_COL].values, grp_df[EVENT_COL].values, time_point=time_point)
        events_by_horizon = int(((grp_df[EVENT_COL] == 1) & (grp_df[TIME_COL] <= time_point)).sum())
        n_at_risk_at_horizon = int((grp_df[TIME_COL] >= time_point).sum())
        interval_estimable = bool(events_by_horizon > 0 and km['risk'] < 1.0)
        rows.append({'group': int(grp) + 1, 'risk_group_order': '1=lowest predicted risk; 5=highest predicted risk', 'n': int(len(grp_df)), 'events_by_horizon': events_by_horizon, 'n_at_risk_at_horizon': n_at_risk_at_horizon, 'mean_predicted_risk': float(grp_df['predicted_event_prob'].mean()), 'observed_km_risk': float(km['risk']), 'observed_km_risk_95ci_low': float(km['risk_ci_low']) if interval_estimable else np.nan, 'observed_km_risk_95ci_high': float(km['risk_ci_high']) if interval_estimable else np.nan, 'observed_km_risk_95ci_estimable': interval_estimable, 'observed_km_risk_95ci_method': 'Greenwood_log_log_pointwise' if interval_estimable else 'not_estimable_zero_events_or_boundary'})
    return pd.DataFrame(rows).sort_values('group').reset_index(drop=True)

def _censoring_survival_step(durations: np.ndarray, events: np.ndarray) -> Tuple[np.ndarray, np.ndarray]:
    censoring_times = np.sort(np.unique(durations[events == 0]))
    survival_after: List[float] = []
    survival = 1.0
    for censoring_time in censoring_times:
        n_at_risk = int(np.sum(durations >= censoring_time))
        tied_outcome_events = int(np.sum((durations == censoring_time) & (events == 1)))
        censoring_risk_set = n_at_risk - tied_outcome_events
        n_censored = int(np.sum((durations == censoring_time) & (events == 0)))
        if censoring_risk_set <= 0:
            continue
        survival *= max(0.0, 1.0 - n_censored / censoring_risk_set)
        survival_after.append(survival)
    return (censoring_times, np.asarray(survival_after, dtype=float))

def _step_survival_at(step_times: np.ndarray, survival_after: np.ndarray, query_time: float, left_limit: bool) -> float:
    side = 'left' if left_limit else 'right'
    idx = int(np.searchsorted(step_times, query_time, side=side) - 1)
    return 1.0 if idx < 0 else float(survival_after[idx])

def ipcw_binary_outcome_at_horizon(test_df: pd.DataFrame, time_point: float) -> Tuple[np.ndarray, np.ndarray]:
    durations = test_df[TIME_COL].to_numpy(dtype=float)
    events = test_df[EVENT_COL].to_numpy(dtype=int)
    censoring_times, censoring_survival = _censoring_survival_step(durations, events)
    outcome = np.zeros(len(test_df), dtype=int)
    weights = np.zeros(len(test_df), dtype=float)
    event_by_horizon = (events == 1) & (durations <= time_point)
    known_event_free = durations > time_point
    outcome[event_by_horizon] = 1
    for idx in np.flatnonzero(event_by_horizon):
        g_value = _step_survival_at(censoring_times, censoring_survival, float(durations[idx]), left_limit=True)
        if g_value <= 0.0:
            raise ValueError('The censoring survival estimate reached zero before an event time.')
        weights[idx] = 1.0 / g_value
    g_horizon = _step_survival_at(censoring_times, censoring_survival, float(time_point), left_limit=False)
    if g_horizon <= 0.0 and np.any(known_event_free):
        raise ValueError('The censoring survival estimate reached zero before the prediction horizon.')
    weights[known_event_free] = 1.0 / g_horizon
    positive = weights > 0.0
    if positive.sum() < 10 or np.unique(outcome[positive]).size < 2:
        raise ValueError('Too few informative observations or outcome classes for flexible calibration.')
    weights[positive] *= float(positive.sum()) / float(weights[positive].sum())
    return (outcome, weights)

def _fit_flexible_calibration_curve(test_df: pd.DataFrame, predicted_event_prob: np.ndarray, prediction_grid: np.ndarray, time_point: float) -> np.ndarray:
    probabilities = np.clip(np.asarray(predicted_event_prob, dtype=float), 1e-05, 1.0 - 1e-05)
    outcome, weights = ipcw_binary_outcome_at_horizon(test_df, time_point)
    informative = weights > 0.0
    x = np.log(probabilities[informative] / (1.0 - probabilities[informative])).reshape(-1, 1)
    grid = np.clip(np.asarray(prediction_grid, dtype=float), 1e-05, 1.0 - 1e-05)
    grid_x = np.log(grid / (1.0 - grid)).reshape(-1, 1)
    n_unique = int(np.unique(x).size)
    n_knots = min(4, max(2, n_unique - 1))
    calibration_model = Pipeline([('spline', SplineTransformer(n_knots=n_knots, degree=3, knots='quantile', include_bias=False)), ('logistic', LogisticRegression(C=10.0, solver='lbfgs', max_iter=2000, random_state=0))])
    calibration_model.fit(x, outcome[informative], logistic__sample_weight=weights[informative])
    return calibration_model.predict_proba(grid_x)[:, 1].astype(float)

def flexible_calibration_curve_with_bootstrap(test_df: pd.DataFrame, predicted_event_prob: np.ndarray, time_point: float, n_bootstrap: int, random_state: int, n_grid: int=101) -> pd.DataFrame:
    if n_bootstrap < 0:
        raise ValueError('n_bootstrap cannot be negative.')
    if n_grid < 2:
        raise ValueError('n_grid must be at least 2.')
    probabilities = np.asarray(predicted_event_prob, dtype=float)
    if len(probabilities) != len(test_df) or not np.all(np.isfinite(probabilities)):
        raise ValueError('Predicted probabilities must be finite and match the test-set length.')
    if np.any((probabilities < -1e-08) | (probabilities > 1.0 + 1e-08)):
        raise ValueError('Predicted event probabilities must lie between 0 and 1.')
    probabilities = np.clip(probabilities, 0.0, 1.0)
    grid_min = max(1e-05, float(np.min(probabilities)))
    grid_max = min(1.0 - 1e-05, float(np.max(probabilities)))
    if grid_max <= grid_min:
        raise ValueError('Flexible calibration requires variation in predicted risk.')
    prediction_grid = np.linspace(grid_min, grid_max, int(n_grid), dtype=float)
    fitted_curve = _fit_flexible_calibration_curve(test_df, probabilities, prediction_grid, time_point)
    bootstrap_curves: List[np.ndarray] = []
    rng = np.random.default_rng(random_state)
    for _ in range(int(n_bootstrap)):
        sampled = rng.integers(0, len(test_df), size=len(test_df))
        bootstrap_df = test_df.iloc[sampled].reset_index(drop=True)
        bootstrap_probabilities = probabilities[sampled]
        try:
            bootstrap_curves.append(_fit_flexible_calibration_curve(bootstrap_df, bootstrap_probabilities, prediction_grid, time_point))
        except Exception:
            continue
    if bootstrap_curves:
        bootstrap_matrix = np.vstack(bootstrap_curves)
        band_low, band_high = np.percentile(bootstrap_matrix, [2.5, 97.5], axis=0)
    else:
        band_low = np.full_like(prediction_grid, np.nan, dtype=float)
        band_high = np.full_like(prediction_grid, np.nan, dtype=float)
    return pd.DataFrame({'predicted_risk': prediction_grid, 'observed_risk_smoothed': fitted_curve, 'bootstrap_95ci_low': band_low, 'bootstrap_95ci_high': band_high, 'horizon_months': float(time_point), 'n_test': int(len(test_df)), 'n_bootstrap_requested': int(n_bootstrap), 'n_bootstrap_successful': int(len(bootstrap_curves)), 'band_type': 'pointwise_participant_bootstrap_95_percent' if n_bootstrap > 0 else 'not_computed'})

def plot_calibration_panel(ax: Any, calibration_df: pd.DataFrame, model_name: str, time_point: float, axis_max: float) -> None:
    ax.plot([0, axis_max], [0, axis_max], linestyle='--', linewidth=1, color='0.4')
    y_values = calibration_df['observed_km_risk'].to_numpy(dtype=float)
    y_error = np.vstack([y_values - calibration_df['observed_km_risk_95ci_low'].to_numpy(dtype=float), calibration_df['observed_km_risk_95ci_high'].to_numpy(dtype=float) - y_values])
    finite_interval = np.isfinite(y_error).all(axis=0)
    if finite_interval.any():
        ax.errorbar(calibration_df.loc[finite_interval, 'mean_predicted_risk'], y_values[finite_interval], yerr=np.maximum(y_error[:, finite_interval], 0.0), marker='o', linestyle='none', linewidth=1, capsize=2)
    if (~finite_interval).any():
        ax.scatter(calibration_df.loc[~finite_interval, 'mean_predicted_risk'], y_values[~finite_interval], marker='o')
    for _, row in calibration_df.iterrows():
        ax.annotate(str(int(row['group'])), (row['mean_predicted_risk'], row['observed_km_risk']), textcoords='offset points', xytext=(4, 4), fontsize=8)
    count_columns = {'n', 'events_by_horizon', 'n_at_risk_at_horizon'}
    if count_columns.issubset(calibration_df.columns):
        count_lines = ['Group  N/events/at risk']
        for _, row in calibration_df.sort_values('group').iterrows():
            count_lines.append(f"{int(row['group']):>2}   {int(row['n'])}/{int(row['events_by_horizon'])}/{int(row['n_at_risk_at_horizon'])}")
        ax.text(0.98, 0.02, '\n'.join(count_lines), transform=ax.transAxes, ha='right', va='bottom', fontsize=6.5, family='monospace', linespacing=1.05, bbox={'facecolor': 'white', 'edgecolor': '0.8', 'alpha': 0.88, 'pad': 1.5})
    ax.set_title(model_name, fontsize=11)
    ax.set_xlabel(f'Predicted {int(time_point)}-month risk', fontsize=9)
    ax.set_ylabel(f'Observed KM risk', fontsize=9)
    ax.set_xlim(0, axis_max)
    ax.set_ylim(0, axis_max)
    ax.set_aspect('equal', adjustable='box')
    ax.tick_params(axis='both', labelsize=8)

def calibration_slope_from_probs(test_df: pd.DataFrame, predicted_event_prob: np.ndarray, time_point: float) -> Dict[str, float]:
    from lifelines import CoxPHFitter
    probs = np.asarray(predicted_event_prob, dtype=float)
    probs = np.clip(probs, 1e-06, 1 - 1e-06)
    lp = np.log(-np.log(1.0 - probs))
    original_time = test_df[TIME_COL].to_numpy(dtype=float)
    original_event = test_df[EVENT_COL].to_numpy(dtype=int)
    cal_df = pd.DataFrame({TIME_COL: np.minimum(original_time, float(time_point)), EVENT_COL: ((original_event == 1) & (original_time <= float(time_point))).astype(int), 'lp': lp})
    cph = CoxPHFitter()
    cph.fit(cal_df, duration_col=TIME_COL, event_col=EVENT_COL)
    slope = float(cph.params_['lp'])
    ci_low, ci_high = map(float, cph.confidence_intervals_.loc['lp'].values.tolist())
    p_value = float(cph.summary.loc['lp', 'p'])
    return {'calibration_slope': slope, 'calibration_slope_ci_low': ci_low, 'calibration_slope_ci_high': ci_high, 'calibration_slope_p': p_value, 'calibration_slope_horizon_months': float(time_point)}

def format_mean_percentile_range(series: pd.Series, decimals: int=3) -> str:
    s = series.dropna()
    if s.empty:
        return 'NA'
    low, high = np.percentile(s.to_numpy(dtype=float), [2.5, 97.5])
    return f'{s.mean():.{decimals}f} ({low:.{decimals}f} to {high:.{decimals}f})'

def validate_split_performance_grid(perf_df: pd.DataFrame, n_splits: int) -> pd.DataFrame:
    required = {'split_id', 'model', 'n_test'}
    if not required.issubset(perf_df.columns):
        raise ValueError(f'Performance data lack split-summary fields: {sorted(required - set(perf_df.columns))}')
    estimates = perf_df.copy()
    expected_keys = {(model_name, split_id) for model_name in MODEL_ORDER for split_id in range(1, int(n_splits) + 1)}
    observed_keys = set(zip(estimates['model'].astype(str), estimates['split_id'].astype(int)))
    if observed_keys != expected_keys or estimates.duplicated(['model', 'split_id']).any():
        raise ValueError('Split-level performance estimates have an incomplete key grid.')
    if not (pd.to_numeric(estimates['n_test'], errors='raise').astype(int) == 249).all():
        raise ValueError('Every repeated 80/20 split must contain exactly 249 test participants.')
    return estimates

def summarise_performance(perf_df: pd.DataFrame, n_splits: int=PRESPECIFIED_N_SPLITS) -> Tuple[pd.DataFrame, pd.DataFrame]:
    required_identity = {'split_id', 'model'}
    if not required_identity.issubset(perf_df.columns):
        raise ValueError('Performance data must identify every row by split_id and model.')
    split_estimates = validate_split_performance_grid(perf_df, n_splits=n_splits)
    summary_rows = []
    formatted_rows = []
    metric_map = {'Harrell_C': 'harrell_c', 'Uno_C_tau': 'uno_c_tau', 'Brier_120m': 'brier_score_at_horizon', 'IBS_12_120m': 'integrated_brier_score'}
    for model_name, sub in split_estimates.groupby('model'):
        row = {'model': model_name, 'n_split_estimates': int(len(sub)), 'summary_unit': 'held_out_20_percent_split', 'range_type': 'empirical_2.5th_to_97.5th_percentile_descriptive_not_CI'}
        fmt_row = {'Model': model_name}
        for label, col in metric_map.items():
            values = sub[col].dropna() if col in sub else pd.Series(dtype=float)
            if values.empty:
                row[f'{col}_mean'] = np.nan
                row[f'{col}_empirical_2_5_percentile'] = np.nan
                row[f'{col}_empirical_97_5_percentile'] = np.nan
                fmt_row[label] = 'NA'
            else:
                row[f'{col}_mean'] = float(values.mean())
                low, high = np.percentile(values.to_numpy(dtype=float), [2.5, 97.5])
                row[f'{col}_empirical_2_5_percentile'] = float(low)
                row[f'{col}_empirical_97_5_percentile'] = float(high)
                fmt_row[label] = format_mean_percentile_range(values)
        summary_rows.append(row)
        formatted_rows.append(fmt_row)
    summary_df = pd.DataFrame(summary_rows).set_index('model').loc[MODEL_ORDER].reset_index()
    formatted_df = pd.DataFrame(formatted_rows).set_index('Model').loc[MODEL_ORDER].reset_index()
    return (summary_df, formatted_df)

def compare_models_vs_reference(perf_df: pd.DataFrame, reference_model: str='CoxPH') -> pd.DataFrame:
    required_identity = {'split_id', 'model'}
    if not required_identity.issubset(perf_df.columns):
        raise ValueError('Paired comparisons require split_id and model identifiers.')
    rows: List[Dict[str, Any]] = []
    metrics = ['harrell_c', 'uno_c_tau', 'brier_score_at_horizon', 'integrated_brier_score']
    models = [m for m in MODEL_ORDER if m != reference_model and m in perf_df['model'].unique()]
    for metric in metrics:
        if metric not in perf_df:
            continue
        wide = perf_df.pivot(index='split_id', columns='model', values=metric)
        if reference_model not in wide.columns:
            continue
        for model_name in models:
            if model_name not in wide.columns:
                continue
            split_diff = (wide[model_name] - wide[reference_model]).rename('split_difference').dropna()
            if split_diff.empty:
                continue
            low, high = np.percentile(split_diff.to_numpy(dtype=float), [2.5, 97.5])
            row = {'Model': model_name, 'Reference_model': reference_model, 'Metric': metric, 'Mean_difference': float(split_diff.mean()), 'n_paired_splits': int(split_diff.shape[0]), 'empirical_2_5_percentile': float(low), 'empirical_97_5_percentile': float(high), 'comparison_unit': 'identical_held_out_20_percent_split', 'range_type': 'descriptive_not_confidence_interval', 'direction_favoring_candidate': 'positive_difference' if metric in {'harrell_c', 'uno_c_tau'} else 'negative_difference', 'hypothesis_testing': 'not_performed', 'multiplicity_adjustment': 'not_applicable_no_hypothesis_tests'}
            rows.append(row)
    out = pd.DataFrame(rows)
    if not out.empty:
        metric_order = {m: i for i, m in enumerate(metrics)}
        out['_metric_order'] = out['Metric'].map(metric_order)
        out['_model_order'] = out['Model'].map({m: i for i, m in enumerate(MODEL_ORDER)})
        out = out.sort_values(['_metric_order', '_model_order']).drop(columns=['_metric_order', '_model_order']).reset_index(drop=True)
    return out

def build_paired_differences_by_split(perf_df: pd.DataFrame, reference_model: str='CoxPH') -> pd.DataFrame:
    rows: List[Dict[str, Any]] = []
    metrics = ['harrell_c', 'uno_c_tau', 'brier_score_at_horizon', 'integrated_brier_score']
    for metric in metrics:
        wide = perf_df.pivot(index='split_id', columns='model', values=metric)
        if reference_model not in wide.columns:
            raise ValueError('Reference model is absent from split-level performance data.')
        for model_name in [model for model in MODEL_ORDER if model != reference_model]:
            if model_name not in wide.columns:
                raise ValueError(f'{model_name} is absent from split-level performance data.')
            differences = (wide[model_name] - wide[reference_model]).dropna()
            if model_name == 'SVM' and metric in {'brier_score_at_horizon', 'integrated_brier_score'}:
                if not differences.empty:
                    raise ValueError('SVM unexpectedly has paired absolute-risk differences.')
                continue
            if len(differences) != PRESPECIFIED_N_SPLITS:
                raise ValueError(f'{model_name} {metric} lacks 200 paired split differences.')
            rows.extend(({'split_id': int(split_id), 'Model': model_name, 'Reference_model': reference_model, 'Metric': metric, 'Difference': float(value), 'direction_favoring_candidate': 'positive_difference' if metric in {'harrell_c', 'uno_c_tau'} else 'negative_difference'} for split_id, value in differences.items()))
    out = pd.DataFrame(rows)
    if len(out) != 5200 or out.duplicated(['split_id', 'Model', 'Metric']).any():
        raise RuntimeError('Split-level paired-difference audit table is incomplete.')
    return out

def summarise_tuning_records(tuning_df: pd.DataFrame) -> pd.DataFrame:
    if tuning_df.empty:
        return tuning_df
    normalized = tuning_df.copy()
    selected_configurations: List[str] = []
    refit_epochs: List[float] = []
    for value in normalized['best_params_json'].astype(str):
        parameters = json.loads(value)
        if not isinstance(parameters, dict):
            raise ValueError('A selected parameter record was not a JSON object.')
        epoch_value = parameters.pop('refit_epochs', None)
        selected_configurations.append(json.dumps(parameters, sort_keys=True))
        refit_epochs.append(float(epoch_value) if epoch_value is not None else float('nan'))
    normalized['selected_configuration_json'] = selected_configurations
    normalized['selected_refit_epochs'] = refit_epochs
    summary = normalized.groupby(['model', 'selected_configuration_json'], dropna=False).agg(n_times_selected=('model', 'size'), refit_epochs_median=('selected_refit_epochs', 'median'), refit_epochs_minimum=('selected_refit_epochs', 'min'), refit_epochs_maximum=('selected_refit_epochs', 'max')).reset_index().sort_values(['model', 'n_times_selected'], ascending=[True, False]).reset_index(drop=True)
    return summary

def build_predictor_imputation_matrix(df: pd.DataFrame) -> pd.DataFrame:
    return df.loc[:, PREDICTOR_COLUMNS].copy()

def postprocess_imputed_predictors(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    for col in BINARY_COLUMNS:
        vals = np.asarray(out[col], dtype=float)
        out[col] = (vals >= 0.5).astype(int)
    lower_bounds = {'age': 18.0, 'sbp': 40.0, 'dbp': 20.0, 'bmi': 10.0, 'chol': 0.1, 'hdl': 0.1, 'trig': 0.05, 'hba1c': 2.0, 'egfr': 1.0}
    for col, lb in lower_bounds.items():
        out[col] = np.clip(np.asarray(out[col], dtype=float), lb, None)
    return out

def fit_deployable_predictor_imputer(train_df: pd.DataFrame, random_state: int, max_iter: int) -> Tuple[IterativeImputer, pd.DataFrame]:
    train_matrix = build_predictor_imputation_matrix(train_df)
    imputer = IterativeImputer(estimator=BayesianRidge(), max_iter=max_iter, sample_posterior=True, random_state=random_state, initial_strategy='median', imputation_order='ascending')
    completed = pd.DataFrame(imputer.fit_transform(train_matrix), columns=PREDICTOR_COLUMNS, index=train_df.index)
    return (imputer, completed)

def transform_predictors_for_deployment(imputer: IterativeImputer, df: pd.DataFrame) -> pd.DataFrame:
    predictor_matrix = build_predictor_imputation_matrix(df)
    return pd.DataFrame(imputer.transform(predictor_matrix), columns=PREDICTOR_COLUMNS, index=df.index)

def generate_imputed_split_datasets(train_df: pd.DataFrame, test_df: pd.DataFrame, m: int, seed: int, max_iter: int) -> List[Dict[str, pd.DataFrame]]:
    out: List[Dict[str, pd.DataFrame]] = []
    for i in range(m):
        imputer, train_completed = fit_deployable_predictor_imputer(train_df=train_df, random_state=seed + i, max_iter=max_iter)
        test_completed = transform_predictors_for_deployment(imputer, test_df)
        train_pred = postprocess_imputed_predictors(train_completed)
        test_pred = postprocess_imputed_predictors(test_completed)
        imputed_train = pd.concat([train_df[[STUDY_ID_COL, TIME_COL, EVENT_COL]].reset_index(drop=True), train_pred.reset_index(drop=True)], axis=1)
        imputed_test = pd.concat([test_df[[STUDY_ID_COL, TIME_COL, EVENT_COL]].reset_index(drop=True), test_pred.reset_index(drop=True)], axis=1)
        imputed_combined = pd.concat([imputed_train, imputed_test], axis=0, ignore_index=True)
        out.append({'train': imputed_train, 'test': imputed_test, 'combined': imputed_combined})
    return out

def generate_nested_inner_imputed_splits(raw_outer_train_df: pd.DataFrame, inner_folds: int, split_seed: int, imputation_seed: int, max_iter: int) -> List[Tuple[pd.DataFrame, pd.DataFrame]]:
    if inner_folds < 2:
        raise ValueError('Nested model selection requires at least two inner folds.')
    splitter = StratifiedKFold(n_splits=inner_folds, shuffle=True, random_state=split_seed)
    strata = make_strata(raw_outer_train_df)
    completed_splits: List[Tuple[pd.DataFrame, pd.DataFrame]] = []
    held_out_ids: List[Any] = []
    for inner_fold, (inner_train_idx, inner_validation_idx) in enumerate(splitter.split(raw_outer_train_df, strata), start=1):
        completed = generate_imputed_split_datasets(raw_outer_train_df.iloc[inner_train_idx].reset_index(drop=True), raw_outer_train_df.iloc[inner_validation_idx].reset_index(drop=True), m=1, seed=imputation_seed + inner_fold, max_iter=max_iter)[0]
        inner_train = completed['train']
        inner_validation = completed['test']
        if inner_train[PREDICTOR_COLUMNS].isna().any().any() or inner_validation[PREDICTOR_COLUMNS].isna().any().any():
            raise RuntimeError(f'Inner fold {inner_fold} retained a missing predictor after imputation.')
        completed_splits.append((inner_train, inner_validation))
        held_out_ids.extend(inner_validation[STUDY_ID_COL].tolist())
    expected_ids = raw_outer_train_df[STUDY_ID_COL].tolist()
    if sorted(map(str, held_out_ids)) != sorted(map(str, expected_ids)):
        raise RuntimeError('Nested inner-fold validation did not hold out every outer-training participant exactly once.')
    return completed_splits

def generate_nested_neural_selection_split(raw_outer_train_df: pd.DataFrame, split_seed: int, imputation_seed: int, max_iter: int) -> Tuple[pd.DataFrame, pd.DataFrame]:
    splitter = StratifiedShuffleSplit(n_splits=1, test_size=0.2, random_state=split_seed)
    strata = make_strata(raw_outer_train_df)
    development_idx, validation_idx = next(splitter.split(raw_outer_train_df, strata))
    completed = generate_imputed_split_datasets(raw_outer_train_df.iloc[development_idx].reset_index(drop=True), raw_outer_train_df.iloc[validation_idx].reset_index(drop=True), m=1, seed=imputation_seed, max_iter=max_iter)[0]
    return (completed['train'], completed['test'])

def aggregate_table1_over_imputations(imputed_splits: List[Dict[str, pd.DataFrame]], raw_train_df: pd.DataFrame, raw_test_df: pd.DataFrame, raw_df: pd.DataFrame) -> pd.DataFrame:
    per_imp_rows: List[pd.DataFrame] = []
    for split in imputed_splits:
        df = split['combined']
        train_df = split['train']
        test_df = split['test']
        rows = []
        for col in PREDICTOR_COLUMNS:
            label = TABLE1_LABELS[col]
            if col in CONTINUOUS_COLUMNS:
                rows.append({'Variable': label, 'Overall_mean': float(df[col].mean()), 'Overall_sd': float(df[col].std(ddof=1)), 'Training_mean': float(train_df[col].mean()), 'Training_sd': float(train_df[col].std(ddof=1)), 'Testing_mean': float(test_df[col].mean()), 'Testing_sd': float(test_df[col].std(ddof=1)), 'SMD_train_vs_test': float(abs(smd_continuous(train_df[col], test_df[col]))), 'is_binary': 0})
            else:
                pos = POSITIVE_LEVELS.get(col, 1)
                rows.append({'Variable': label, 'Overall_n': float((df[col] == pos).sum()), 'Overall_pct': float(100.0 * (df[col] == pos).mean()), 'Training_n': float((train_df[col] == pos).sum()), 'Training_pct': float(100.0 * (train_df[col] == pos).mean()), 'Testing_n': float((test_df[col] == pos).sum()), 'Testing_pct': float(100.0 * (test_df[col] == pos).mean()), 'SMD_train_vs_test': float(abs(smd_binary(train_df[col], test_df[col], positive_value=pos))), 'is_binary': 1})
        per_imp_rows.append(pd.DataFrame(rows))
    combined = pd.concat(per_imp_rows, axis=0, ignore_index=True)
    out_rows = []
    for variable, sub in combined.groupby('Variable'):
        is_binary = int(sub['is_binary'].iloc[0])
        if is_binary == 0:
            out_rows.append({'Variable': variable, 'Overall': f"{sub['Overall_mean'].mean():.1f} ({sub['Overall_sd'].mean():.1f})", 'Training': f"{sub['Training_mean'].mean():.1f} ({sub['Training_sd'].mean():.1f})", 'Testing': f"{sub['Testing_mean'].mean():.1f} ({sub['Testing_sd'].mean():.1f})", 'SMD_train_vs_test': round(float(sub['SMD_train_vs_test'].mean()), 3)})
        else:
            out_rows.append({'Variable': variable, 'Overall': f"{int(round(sub['Overall_n'].mean()))} ({sub['Overall_pct'].mean():.1f}%)", 'Training': f"{int(round(sub['Training_n'].mean()))} ({sub['Training_pct'].mean():.1f}%)", 'Testing': f"{int(round(sub['Testing_n'].mean()))} ({sub['Testing_pct'].mean():.1f}%)", 'SMD_train_vs_test': round(float(sub['SMD_train_vs_test'].mean()), 3)})
    out_rows.extend([{'Variable': 'ASCVD events, n (%)', 'Overall': format_n_pct(raw_df[EVENT_COL], positive_value=1), 'Training': format_n_pct(raw_train_df[EVENT_COL], positive_value=1), 'Testing': format_n_pct(raw_test_df[EVENT_COL], positive_value=1), 'SMD_train_vs_test': round(abs(smd_binary(raw_train_df[EVENT_COL], raw_test_df[EVENT_COL], positive_value=1)), 3)}, {'Variable': 'Duration (months), median [IQR]', 'Overall': format_median_iqr(raw_df[TIME_COL]), 'Training': format_median_iqr(raw_train_df[TIME_COL]), 'Testing': format_median_iqr(raw_test_df[TIME_COL]), 'SMD_train_vs_test': ''}])
    out_df = pd.DataFrame(out_rows)
    order = [TABLE1_LABELS[c] for c in PREDICTOR_COLUMNS] + ['ASCVD events, n (%)', 'Duration (months), median [IQR]']
    out_df['_order'] = out_df['Variable'].map({v: i for i, v in enumerate(order)})
    out_df = out_df.sort_values('_order').drop(columns='_order').reset_index(drop=True)
    return out_df

def fit_lifelines_cox_and_ph_test(train_df: pd.DataFrame) -> Tuple[Any, pd.DataFrame]:
    from lifelines import CoxPHFitter
    from lifelines.statistics import proportional_hazard_test
    cox_df = train_df[PREDICTOR_COLUMNS + [TIME_COL, EVENT_COL]].copy()
    cph = CoxPHFitter()
    cph.fit(cox_df, duration_col=TIME_COL, event_col=EVENT_COL)
    summary = cph.summary.reset_index().rename(columns={'covariate': 'variable'})
    summary['Variable'] = summary['variable'].map(DISPLAY_NAMES).fillna(summary['variable'])
    ph_test = proportional_hazard_test(cph, cox_df, time_transform='rank')
    ph_df = ph_test.summary.reset_index().rename(columns={'index': 'variable'})
    ph_df = ph_df.rename(columns={'p': 'ph_test_p', 'test_statistic': 'ph_test_statistic'})
    merged = summary.merge(ph_df[['variable', 'ph_test_statistic', 'ph_test_p']], on='variable', how='left')
    out = merged[['Variable', 'variable', 'coef', 'se(coef)', 'z', 'p', 'ph_test_statistic', 'ph_test_p']].copy()
    out = out.rename(columns={'coef': 'beta', 'se(coef)': 'SE_beta', 'p': 'p_value'})
    return (cph, out)

def pool_cox_results(primary_imputed_trains: List[pd.DataFrame]) -> Tuple[pd.DataFrame, pd.DataFrame]:
    per_imp = []
    for i, train_df in enumerate(primary_imputed_trains):
        _, tbl = fit_lifelines_cox_and_ph_test(train_df)
        tbl = tbl.copy()
        tbl['imputation'] = i + 1
        per_imp.append(tbl)
    all_tbl = pd.concat(per_imp, axis=0, ignore_index=True)
    rows = []
    m = len(primary_imputed_trains)
    for variable, sub in all_tbl.groupby('variable'):
        q = sub['beta'].to_numpy(dtype=float)
        u = sub['SE_beta'].to_numpy(dtype=float) ** 2
        qbar = float(np.mean(q))
        ubar = float(np.mean(u))
        b = float(np.var(q, ddof=1)) if len(q) > 1 else 0.0
        total_var = ubar + (1.0 + 1.0 / max(m, 1)) * b
        se = float(np.sqrt(total_var))
        test_statistic = qbar / se if se > 0 else np.nan
        between_component = (1.0 + 1.0 / max(m, 1)) * b
        if m > 1 and between_component > 0.0:
            relative_increase = between_component / max(ubar, np.finfo(float).tiny)
            rubin_df = float((m - 1) * (1.0 + 1.0 / relative_increase) ** 2)
        else:
            rubin_df = float('inf')
        critical = float(stats.t.ppf(0.975, df=rubin_df)) if np.isfinite(rubin_df) else 1.959963984540054
        p = float(2.0 * stats.t.sf(abs(test_statistic), df=rubin_df)) if np.isfinite(test_statistic) and np.isfinite(rubin_df) else float(2.0 * norm.sf(abs(test_statistic))) if np.isfinite(test_statistic) else np.nan
        ci_low = qbar - critical * se
        ci_high = qbar + critical * se
        rows.append({'Variable': DISPLAY_NAMES.get(variable, variable), 'variable': variable, 'beta': qbar, 'HR': float(np.exp(qbar)), 'HR_95CI_lower': float(np.exp(ci_low)), 'HR_95CI_upper': float(np.exp(ci_high)), 'SE_beta': se, 'z': test_statistic, 't_statistic': test_statistic, 'rubin_df': rubin_df, 'p_value': p, 'ph_test_p_min': float(sub['ph_test_p'].min()), 'ph_test_p_median': float(sub['ph_test_p'].median()), 'ph_test_p_max': float(sub['ph_test_p'].max()), 'ph_test_n_below_0_05': int((sub['ph_test_p'] < 0.05).sum()), 'ph_test_n_imputations': int(len(sub))})
    pooled = pd.DataFrame(rows)
    pooled['_order'] = pooled['variable'].map({c: i for i, c in enumerate(PREDICTOR_COLUMNS)})
    pooled = pooled.sort_values('_order').drop(columns='_order').reset_index(drop=True)
    ph_by_imputation = all_tbl[['imputation', 'Variable', 'variable', 'ph_test_statistic', 'ph_test_p']].sort_values(['imputation', 'variable']).reset_index(drop=True)
    required_ph = ph_by_imputation[['ph_test_statistic', 'ph_test_p']].to_numpy(float)
    if not np.isfinite(required_ph).all():
        raise RuntimeError('A proportional-hazards diagnostic was non-finite.')
    return (pooled, ph_by_imputation)

def make_shap_outputs(model_name: str, fitted_model: Any, train_df: pd.DataFrame, test_df: pd.DataFrame, out_dir: str, make_force_plot: bool=False, max_background: int=100, max_explain: int=150) -> Optional[pd.DataFrame]:
    try:
        import shap
    except Exception as exc:
        LOGGER.warning('Skipping SHAP for %s because shap is unavailable: %s', model_name, exc)
        return None
    X_train = train_df[PREDICTOR_COLUMNS].copy()
    X_test = test_df[PREDICTOR_COLUMNS].copy()
    feature_names = list(X_test.columns)
    shap_feature_names = [SHAP_DISPLAY_NAMES.get(c, c) for c in feature_names]
    reverse_name_map = {SHAP_DISPLAY_NAMES.get(c, c): c for c in feature_names}
    background_plot = X_train.sample(n=min(max_background, len(X_train)), random_state=0).copy()
    background_plot.columns = shap_feature_names
    explain_df_plot = X_test.sample(n=min(max_explain, len(X_test)), random_state=0, replace=False).copy()
    explain_df_plot.columns = shap_feature_names

    def predict_fn(x: Any) -> np.ndarray:
        x_df = as_dataframe(x, shap_feature_names)
        x_df = x_df.rename(columns=reverse_name_map)
        x_df = x_df.loc[:, feature_names]
        preds = fitted_model.predict(x_df)
        return np.asarray(preds, dtype=float).reshape(-1)
    try:
        masker = shap.maskers.Independent(background_plot)
        explainer = shap.Explainer(predict_fn, masker=masker, feature_names=shap_feature_names, algorithm='permutation', seed=0)
        shap_values = explainer(explain_df_plot, max_evals=10 * (2 * len(shap_feature_names) + 1))
    except Exception as exc:
        LOGGER.warning('SHAP failed for %s: %s', model_name, exc)
        return None
    mean_abs = np.abs(shap_values.values).mean(axis=0)
    shap_table = pd.DataFrame({'variable': feature_names, 'Variable': shap_feature_names, 'mean_abs_shap': mean_abs}).sort_values('mean_abs_shap', ascending=False).reset_index(drop=True)
    summary_png = os.path.join(out_dir, f'shap_summary_{model_name}.png')
    try:
        plt.figure()
        shap.summary_plot(shap_values, explain_df_plot, max_display=len(shap_feature_names), show=False)
        ax = plt.gca()
        ax.tick_params(axis='y', labelsize=9)
        ax.tick_params(axis='x', labelsize=9)
        ax.set_xlabel('SHAP value (impact on model output)', fontsize=10)
        plt.tight_layout()
        plt.savefig(summary_png, dpi=400, bbox_inches='tight')
        plt.close()
    except Exception as exc:
        LOGGER.warning('SHAP summary plot failed for %s: %s', model_name, exc)
    if make_force_plot and len(explain_df_plot) > 0:
        try:
            force_png = os.path.join(out_dir, f'shap_force_{model_name}.png')
            shap.force_plot(shap_values.base_values[0], shap_values.values[0], explain_df_plot.iloc[0], matplotlib=True, show=False, contribution_threshold=0.05, text_rotation=45)
            plt.tight_layout()
            plt.savefig(force_png, dpi=400, bbox_inches='tight')
            plt.close()
        except Exception as exc:
            LOGGER.warning('SHAP force plot failed for %s: %s', model_name, exc)
    shap_table.to_csv(os.path.join(out_dir, f'shap_importance_{model_name}.csv'), index=False)
    return shap_table

def get_deepsurv_configs() -> List[Dict[str, Any]]:
    return [{'num_nodes': [32, 16], 'dropout': 0.1, 'lr': 0.01, 'batch_size': 'full_training_risk_set', 'epochs': 256}, {'num_nodes': [32, 16], 'dropout': 0.2, 'lr': 0.01, 'batch_size': 'full_training_risk_set', 'epochs': 256}, {'num_nodes': [32, 32], 'dropout': 0.1, 'lr': 0.01, 'batch_size': 'full_training_risk_set', 'epochs': 256}, {'num_nodes': [64, 32], 'dropout': 0.2, 'lr': 0.001, 'batch_size': 'full_training_risk_set', 'epochs': 256}, {'num_nodes': [64, 32], 'dropout': 0.1, 'lr': 0.001, 'batch_size': 'full_training_risk_set', 'epochs': 256}, {'num_nodes': [16, 16], 'dropout': 0.2, 'lr': 0.01, 'batch_size': 'full_training_risk_set', 'epochs': 256}]

def get_coxtime_configs() -> List[Dict[str, Any]]:
    return [{'num_nodes': [32, 16], 'dropout': 0.1, 'lr': 0.01, 'batch_size': 64, 'epochs': 256}, {'num_nodes': [32, 16], 'dropout': 0.2, 'lr': 0.01, 'batch_size': 64, 'epochs': 256}, {'num_nodes': [32, 32], 'dropout': 0.1, 'lr': 0.01, 'batch_size': 64, 'epochs': 256}, {'num_nodes': [64, 32], 'dropout': 0.2, 'lr': 0.001, 'batch_size': 64, 'epochs': 256}, {'num_nodes': [64, 32], 'dropout': 0.1, 'lr': 0.001, 'batch_size': 128, 'epochs': 256}, {'num_nodes': [16, 16], 'dropout': 0.2, 'lr': 0.01, 'batch_size': 128, 'epochs': 256}]

def _require_finite_neural_training_history(model: Any, context: str, require_validation: bool) -> pd.DataFrame:
    try:
        history = model.log.to_pandas()
    except Exception as exc:
        raise RuntimeError(f'{context} did not provide a readable training log.') from exc
    if history.empty:
        raise RuntimeError(f'{context} produced an empty training log.')
    required_columns = ['train_loss']
    if require_validation:
        required_columns.append('val_loss')
    missing = [column for column in required_columns if column not in history]
    if missing:
        raise RuntimeError(f'{context} training log was missing required columns: {missing}.')
    for column in required_columns:
        values = history[column].to_numpy(dtype=float)
        if values.size == 0 or not np.isfinite(values).all():
            raise RuntimeError(f'{context} produced a non-finite {column} history.')
    return history

def _best_epoch_from_training_log(model: Any, maximum_epochs: int) -> int:
    history = _require_finite_neural_training_history(model, context='Neural configuration selection', require_validation=True)
    values = history['val_loss'].to_numpy(dtype=float)
    minimum = float(np.min(values))
    tied_minima = np.flatnonzero(values == minimum)
    if tied_minima.size == 0:
        raise RuntimeError('Neural validation loss did not have a finite minimum.')
    return min(int(maximum_epochs), max(1, int(tied_minima[-1]) + 1))

def _deepsurv_full_batch_size(cfg: Dict[str, Any], events: np.ndarray, context: str) -> int:
    if cfg.get('batch_size') != 'full_training_risk_set':
        raise ValueError(f"{context} must use batch_size='full_training_risk_set'.")
    events_array = np.asarray(events).reshape(-1)
    if events_array.size == 0:
        raise ValueError(f'{context} data cannot be empty.')
    if int(np.count_nonzero(events_array)) == 0:
        raise ValueError(f'{context} requires at least one observed event.')
    return int(events_array.size)

def _require_finite_network_parameters(model: Any, context: str) -> None:
    import torch
    for parameter_name, parameter in model.net.named_parameters():
        if not bool(torch.isfinite(parameter.detach()).all().item()):
            raise RuntimeError(f'{context} produced a non-finite network parameter: {parameter_name}.')

def _require_finite_deepsurv_state(model: Any, prediction_matrices: Dict[str, np.ndarray], context: str) -> None:
    _require_finite_network_parameters(model, context)
    for matrix_name, matrix in prediction_matrices.items():
        native_predictions = np.asarray(model.predict(matrix)).reshape(-1)
        if native_predictions.shape != (len(matrix),) or not np.isfinite(native_predictions).all():
            raise RuntimeError(f'{context} produced non-finite raw log-risk predictions for {matrix_name}.')
        with np.errstate(over='ignore', under='ignore', invalid='ignore'):
            relative_risk = np.exp(native_predictions)
        if relative_risk.shape != (len(matrix),) or not np.isfinite(relative_risk).all() or np.any(relative_risk <= 0.0):
            raise RuntimeError(f'{context} produced invalid exponentiated risk for {matrix_name}.')

def _require_deepsurv_single_batch_execution(model: Any, context: str, require_validation: bool) -> None:
    training_batches = getattr(model, 'fit_info', {}).get('batches_per_epoch')
    if int(training_batches or 0) != 1:
        raise RuntimeError(f'{context} used {training_batches!r} training batches; expected one.')
    if require_validation:
        validation_loader = getattr(getattr(model, 'val_metrics', None), 'dataloader', None)
        if validation_loader is None or len(validation_loader) != 1:
            raise RuntimeError(f'{context} did not use one complete internal validation batch.')

def _require_finite_pycox_baseline(model: Any, context: str) -> None:
    baseline_values: Dict[str, np.ndarray] = {}
    for attribute in ['baseline_hazards_', 'baseline_cumulative_hazards_']:
        if not hasattr(model, attribute):
            raise RuntimeError(f'{context} did not create {attribute}.')
        values = np.asarray(getattr(model, attribute), dtype=float).reshape(-1)
        if values.size == 0 or not np.isfinite(values).all():
            raise RuntimeError(f'{context} produced a non-finite {attribute}.')
        if np.any(values < 0.0):
            raise RuntimeError(f'{context} produced a negative {attribute}.')
        if not np.any(values > 0.0):
            raise RuntimeError(f'{context} produced an all-zero {attribute} despite observed events.')
        if attribute == 'baseline_cumulative_hazards_' and np.any(np.diff(values) < -1e-12):
            raise RuntimeError(f'{context} produced a decreasing baseline cumulative hazard.')
        baseline_values[attribute] = values
    hazards = baseline_values['baseline_hazards_']
    cumulative = baseline_values['baseline_cumulative_hazards_']
    if hazards.shape != cumulative.shape or not np.allclose(cumulative, np.cumsum(hazards), rtol=1e-06, atol=1e-12):
        raise RuntimeError(f'{context} produced inconsistent baseline and cumulative hazards.')

def _ensure_pycox_coxtime_pandas_compatibility() -> None:
    if not hasattr(pd.Series, 'items'):
        raise RuntimeError('The locked pandas runtime does not provide Series.items().')
    if not hasattr(pd.Series, 'iteritems'):
        setattr(pd.Series, 'iteritems', pd.Series.items)
    if not callable(getattr(pd.Series, 'iteritems', None)):
        raise RuntimeError('Could not provide the pycox CoxTime pandas compatibility alias.')

def fit_deepsurv_fixed_config(train_df: pd.DataFrame, random_state: int, cfg: Dict[str, Any]) -> PyCoxAdapter:
    import torchtuples as tt
    from pycox.models import CoxPH as PyCoxCoxPH
    set_global_seed(random_state)
    feature_names = list(PREDICTOR_COLUMNS)
    prep = TabularSurvivalPreprocessor(CONTINUOUS_COLUMNS, BINARY_COLUMNS, scale_continuous=True)
    prep.fit(train_df[feature_names])
    x_full = prep.transform(train_df[feature_names]).to_numpy(dtype='float32')
    y_full = (train_df[TIME_COL].astype('float32').to_numpy(), train_df[EVENT_COL].astype('int64').to_numpy())
    full_batch_size = _deepsurv_full_batch_size(cfg, y_full[1], context='DeepSurv final refit')
    net = tt.practical.MLPVanilla(in_features=x_full.shape[1], num_nodes=cfg['num_nodes'], out_features=1, batch_norm=True, dropout=cfg['dropout'], output_bias=False)
    model = PyCoxCoxPH(net, tt.optim.Adam)
    model.optimizer.set_lr(cfg['lr'])
    model.fit(x_full, y_full, batch_size=full_batch_size, epochs=int(cfg.get('refit_epochs', cfg['epochs'])), callbacks=[], verbose=False, shuffle=True)
    _require_deepsurv_single_batch_execution(model, context='DeepSurv final refit', require_validation=False)
    _require_finite_neural_training_history(model, context='DeepSurv final refit', require_validation=False)
    _require_finite_deepsurv_state(model, {'full outer-training data': x_full}, context='DeepSurv final refit')
    model.compute_baseline_hazards()
    _require_finite_pycox_baseline(model, context='DeepSurv final refit')
    return PyCoxAdapter(model=model, preprocessor=prep, feature_names=feature_names)

def fit_deepsurv(train_df: pd.DataFrame, random_state: int, configs: Optional[List[Dict[str, Any]]]=None, selection_split: Optional[Tuple[pd.DataFrame, pd.DataFrame]]=None, require_all_candidates: bool=True, ibs_start_month: float=PRESPECIFIED_IBS_START_MONTH, horizon_months: float=PRESPECIFIED_HORIZON_MONTHS) -> Tuple[PyCoxAdapter, Dict[str, Any], float]:
    import torchtuples as tt
    from pycox.models import CoxPH as PyCoxCoxPH
    set_global_seed(random_state)
    feature_names = list(PREDICTOR_COLUMNS)
    if selection_split is None:
        splitter = StratifiedShuffleSplit(n_splits=1, test_size=0.2, random_state=random_state)
        strata = make_strata(train_df)
        tr_idx, val_idx = next(splitter.split(train_df, strata))
        tr_df = train_df.iloc[tr_idx].copy()
        val_df = train_df.iloc[val_idx].copy()
    else:
        tr_df, val_df = (frame.copy() for frame in selection_split)
        if tr_df.empty or val_df.empty:
            raise ValueError('DeepSurv selection data cannot be empty.')
    selection_prep = TabularSurvivalPreprocessor(CONTINUOUS_COLUMNS, BINARY_COLUMNS, scale_continuous=True)
    selection_prep.fit(tr_df[feature_names])
    x_tr = selection_prep.transform(tr_df[feature_names]).to_numpy(dtype='float32')
    x_val = selection_prep.transform(val_df[feature_names]).to_numpy(dtype='float32')
    y_tr = (tr_df[TIME_COL].astype('float32').to_numpy(), tr_df[EVENT_COL].astype('int64').to_numpy())
    y_val = (val_df[TIME_COL].astype('float32').to_numpy(), val_df[EVENT_COL].astype('int64').to_numpy())
    selection_times = get_eval_times(tr_df, val_df, start_month=ibs_start_month, horizon_months=horizon_months)
    selection_brier_context = prepare_brier_scoring_context(y_train=get_surv_array(tr_df), y_test=get_surv_array(val_df), eval_times=selection_times, horizon_months=horizon_months)
    best_score = np.inf
    best_cfg: Dict[str, Any] = {}
    best_refit_epochs = 0
    candidate_failures: List[str] = []
    candidate_configs = configs if configs is not None else get_deepsurv_configs()
    if not candidate_configs:
        raise ValueError('DeepSurv requires at least one candidate configuration.')
    for i, cfg in enumerate(candidate_configs, start=1):
        try:
            set_global_seed(random_state + i)
            net = tt.practical.MLPVanilla(in_features=x_tr.shape[1], num_nodes=cfg['num_nodes'], out_features=1, batch_norm=True, dropout=cfg['dropout'], output_bias=False)
            model = PyCoxCoxPH(net, tt.optim.Adam)
            model.optimizer.set_lr(cfg['lr'])
            callbacks = [tt.callbacks.EarlyStopping(metric='loss', dataset='val', minimize=True, patience=10, load_best=True)]
            full_batch_size = _deepsurv_full_batch_size(cfg, y_tr[1], context=f'DeepSurv candidate {i} development fit')
            validation_batch_size = _deepsurv_full_batch_size(cfg, y_val[1], context=f'DeepSurv candidate {i} validation loss')
            model.fit(x_tr, y_tr, batch_size=full_batch_size, epochs=cfg['epochs'], callbacks=callbacks, verbose=False, val_data=(x_val, y_val), val_batch_size=validation_batch_size, shuffle=True)
            _require_deepsurv_single_batch_execution(model, context=f'DeepSurv candidate {i}', require_validation=True)
            candidate_refit_epochs = _best_epoch_from_training_log(model, cfg['epochs'])
            _require_finite_deepsurv_state(model, {'development data': x_tr, 'internal validation data': x_val}, context=f'DeepSurv candidate {i}')
            model.compute_baseline_hazards()
            _require_finite_pycox_baseline(model, context=f'DeepSurv candidate {i}')
            validation_survival = _clean_pycox_surv_df(model.predict_surv_df(x_val))
            survival_matrix = stepwise_survival_matrix(validation_survival.index.to_numpy(dtype=float), validation_survival.to_numpy(dtype=float).T, selection_times)
            score = compute_brier_metrics_from_context(survival_matrix, selection_brier_context)['integrated_brier_score']
            if not np.isfinite(score) or not 0.0 <= score <= 1.0:
                raise RuntimeError('validation IBS was invalid')
            if score < best_score:
                best_score = score
                best_cfg = dict(cfg)
                best_refit_epochs = int(candidate_refit_epochs)
        except Exception as exc:
            candidate_failures.append(f'candidate {i}: {type(exc).__name__}: {exc}')
            LOGGER.warning('DeepSurv %s', candidate_failures[-1])
    if require_all_candidates and candidate_failures:
        raise RuntimeError('DeepSurv rejected one or more prespecified candidates. ' + ' | '.join(candidate_failures))
    if not best_cfg:
        raise RuntimeError('DeepSurv tuning did not produce a valid model. ' + ' | '.join(candidate_failures))
    best_cfg['refit_epochs'] = int(best_refit_epochs)
    final_model = fit_deepsurv_fixed_config(train_df, random_state + 10000, best_cfg)
    return (final_model, best_cfg, float(best_score))

def fit_coxtime_fixed_config(train_df: pd.DataFrame, random_state: int, cfg: Dict[str, Any]) -> PyCoxAdapter:
    import torchtuples as tt
    from pycox.models import CoxTime
    from pycox.models.cox_time import MLPVanillaCoxTime
    _ensure_pycox_coxtime_pandas_compatibility()
    set_global_seed(random_state)
    feature_names = list(PREDICTOR_COLUMNS)
    prep = TabularSurvivalPreprocessor(CONTINUOUS_COLUMNS, BINARY_COLUMNS, scale_continuous=True)
    prep.fit(train_df[feature_names])
    x_full = prep.transform(train_df[feature_names]).to_numpy(dtype='float32')
    durations_full = train_df[TIME_COL].astype('float32').to_numpy()
    events_full = train_df[EVENT_COL].astype('float32').to_numpy()
    labtrans = CoxTime.label_transform()
    y_full = labtrans.fit_transform(durations_full, events_full)
    net = MLPVanillaCoxTime(in_features=x_full.shape[1], num_nodes=cfg['num_nodes'], batch_norm=True, dropout=cfg['dropout'])
    model = CoxTime(net, tt.optim.Adam, labtrans=labtrans)
    model.optimizer.set_lr(cfg['lr'])
    model.fit(x_full, y_full, batch_size=cfg['batch_size'], epochs=int(cfg.get('refit_epochs', cfg['epochs'])), callbacks=[], verbose=False)
    _require_finite_neural_training_history(model, context='CoxTime final refit', require_validation=False)
    _require_finite_network_parameters(model, context='CoxTime final refit')
    model.compute_baseline_hazards()
    _require_finite_pycox_baseline(model, context='CoxTime final refit')
    return PyCoxAdapter(model=model, preprocessor=prep, feature_names=feature_names, predict_from_surv=True, risk_horizon=120.0)

def fit_coxtime(train_df: pd.DataFrame, random_state: int, configs: Optional[List[Dict[str, Any]]]=None, selection_split: Optional[Tuple[pd.DataFrame, pd.DataFrame]]=None, ibs_start_month: float=PRESPECIFIED_IBS_START_MONTH, horizon_months: float=PRESPECIFIED_HORIZON_MONTHS, require_all_candidates: bool=True) -> Tuple[PyCoxAdapter, Dict[str, Any], float]:
    import torchtuples as tt
    from pycox.models import CoxTime
    from pycox.models.cox_time import MLPVanillaCoxTime
    _ensure_pycox_coxtime_pandas_compatibility()
    set_global_seed(random_state)
    feature_names = list(PREDICTOR_COLUMNS)
    if selection_split is None:
        splitter = StratifiedShuffleSplit(n_splits=1, test_size=0.2, random_state=random_state)
        strata = make_strata(train_df)
        tr_idx, val_idx = next(splitter.split(train_df, strata))
        tr_df = train_df.iloc[tr_idx].copy()
        val_df = train_df.iloc[val_idx].copy()
    else:
        tr_df, val_df = (frame.copy() for frame in selection_split)
        if tr_df.empty or val_df.empty:
            raise ValueError('CoxTime selection data cannot be empty.')
    selection_prep = TabularSurvivalPreprocessor(CONTINUOUS_COLUMNS, BINARY_COLUMNS, scale_continuous=True)
    selection_prep.fit(tr_df[feature_names])
    x_tr = selection_prep.transform(tr_df[feature_names]).to_numpy(dtype='float32')
    x_val = selection_prep.transform(val_df[feature_names]).to_numpy(dtype='float32')
    durations_tr = tr_df[TIME_COL].astype('float32').to_numpy()
    events_tr = tr_df[EVENT_COL].astype('float32').to_numpy()
    durations_val = val_df[TIME_COL].astype('float32').to_numpy()
    events_val = val_df[EVENT_COL].astype('float32').to_numpy()
    selection_times = get_eval_times(tr_df, val_df, start_month=ibs_start_month, horizon_months=horizon_months)
    selection_brier_context = prepare_brier_scoring_context(y_train=get_surv_array(tr_df), y_test=get_surv_array(val_df), eval_times=selection_times, horizon_months=horizon_months)
    best_score = np.inf
    best_cfg: Dict[str, Any] = {}
    best_refit_epochs = 0
    candidate_failures: List[str] = []
    candidate_configs = configs if configs is not None else get_coxtime_configs()
    if not candidate_configs:
        raise ValueError('CoxTime requires at least one candidate configuration.')
    for i, cfg in enumerate(candidate_configs, start=1):
        try:
            set_global_seed(random_state + 100 + i)
            labtrans = CoxTime.label_transform()
            y_tr = labtrans.fit_transform(durations_tr, events_tr)
            y_val = labtrans.transform(durations_val, events_val)
            net = MLPVanillaCoxTime(in_features=x_tr.shape[1], num_nodes=cfg['num_nodes'], batch_norm=True, dropout=cfg['dropout'])
            model = CoxTime(net, tt.optim.Adam, labtrans=labtrans)
            model.optimizer.set_lr(cfg['lr'])
            callbacks = [tt.callbacks.EarlyStopping(metric='loss', dataset='val', minimize=True, patience=10, load_best=True)]
            val_data = tt.tuplefy(x_val, y_val)
            model.fit(x_tr, y_tr, batch_size=cfg['batch_size'], epochs=cfg['epochs'], callbacks=callbacks, verbose=False, val_data=val_data, val_batch_size=cfg['batch_size'])
            _require_finite_neural_training_history(model, context=f'CoxTime candidate {i}', require_validation=True)
            candidate_refit_epochs = _best_epoch_from_training_log(model, cfg['epochs'])
            _require_finite_network_parameters(model, context=f'CoxTime candidate {i}')
            model.compute_baseline_hazards()
            _require_finite_pycox_baseline(model, context=f'CoxTime candidate {i}')
            validation_survival = _clean_pycox_surv_df(model.predict_surv_df(x_val))
            survival_matrix = stepwise_survival_matrix(validation_survival.index.to_numpy(dtype=float), validation_survival.to_numpy(dtype=float).T, selection_times)
            score = compute_brier_metrics_from_context(survival_matrix, selection_brier_context)['integrated_brier_score']
            if not np.isfinite(score) or not 0.0 <= score <= 1.0:
                raise RuntimeError('validation IBS was invalid')
            if score < best_score:
                best_score = score
                best_cfg = dict(cfg)
                best_refit_epochs = int(candidate_refit_epochs)
        except Exception as exc:
            candidate_failures.append(f'candidate {i}: {type(exc).__name__}: {exc}')
            LOGGER.warning('CoxTime %s', candidate_failures[-1])
    if require_all_candidates and candidate_failures:
        raise RuntimeError('CoxTime rejected one or more prespecified candidates. ' + ' | '.join(candidate_failures))
    if not best_cfg:
        raise RuntimeError('CoxTime tuning did not produce a valid model. ' + ' | '.join(candidate_failures))
    best_cfg['refit_epochs'] = int(best_refit_epochs)
    final_model = fit_coxtime_fixed_config(train_df, random_state + 20000, best_cfg)
    return (final_model, best_cfg, float(best_score))

def get_classical_model_specs(random_state: int, n_jobs: int) -> Dict[str, Dict[str, Any]]:
    from sksurv.ensemble import GradientBoostingSurvivalAnalysis, RandomSurvivalForest
    from sksurv.linear_model import CoxPHSurvivalAnalysis
    from sksurv.svm import FastSurvivalSVM
    del n_jobs

    def singleton_configurations(configurations: Sequence[Dict[str, Any]]) -> List[Dict[str, List[Any]]]:
        return [{parameter: [value] for parameter, value in configuration.items()} for configuration in configurations]
    return {'CoxPH': {'pipeline': Pipeline([('prep', TabularSurvivalPreprocessor(CONTINUOUS_COLUMNS, BINARY_COLUMNS, scale_continuous=False)), ('model', CoxPHSurvivalAnalysis())]), 'param_grid': singleton_configurations([{'model__ties': 'efron', 'model__tol': 1e-07, 'model__n_iter': 300}]), 'supports_survival_probs': True}, 'ElasticNetCox': {'pipeline': Pipeline([('prep', TabularSurvivalPreprocessor(CONTINUOUS_COLUMNS, BINARY_COLUMNS, scale_continuous=True)), ('model', LifelinesElasticNetCox())]), 'param_grid': singleton_configurations([{'model__penalizer': penalizer, 'model__l1_ratio': l1_ratio} for penalizer, l1_ratio in [(0.001, 0.0), (0.01, 0.0), (0.1, 0.0), (1.0, 0.0), (0.01, 0.5), (0.1, 0.5), (0.01, 1.0), (0.1, 1.0)]]), 'supports_survival_probs': True}, 'RSF': {'pipeline': Pipeline([('prep', TabularSurvivalPreprocessor(CONTINUOUS_COLUMNS, BINARY_COLUMNS, scale_continuous=False)), ('model', RandomSurvivalForest(random_state=random_state, n_jobs=1))]), 'param_grid': singleton_configurations([{'model__n_estimators': 300, 'model__min_samples_split': min_split, 'model__min_samples_leaf': min_leaf, 'model__max_features': max_features, 'model__max_depth': max_depth} for min_split, min_leaf, max_features, max_depth in [(10, 5, 'sqrt', None), (20, 10, 'sqrt', None), (40, 20, 'sqrt', None), (10, 5, 0.5, None), (20, 10, 0.5, None), (20, 10, 0.5, 5)]]), 'supports_survival_probs': True}, 'GBSA': {'pipeline': Pipeline([('prep', TabularSurvivalPreprocessor(CONTINUOUS_COLUMNS, BINARY_COLUMNS, scale_continuous=False)), ('model', GradientBoostingSurvivalAnalysis(random_state=random_state))]), 'param_grid': singleton_configurations([{'model__n_estimators': n_estimators, 'model__learning_rate': learning_rate, 'model__max_depth': max_depth, 'model__min_samples_leaf': min_leaf, 'model__subsample': subsample} for n_estimators, learning_rate, max_depth, min_leaf, subsample in [(100, 0.05, 1, 10, 1.0), (200, 0.03, 1, 10, 0.7), (100, 0.1, 1, 10, 0.7), (100, 0.05, 2, 10, 1.0), (200, 0.03, 2, 10, 0.7), (100, 0.1, 2, 20, 0.7)]]), 'supports_survival_probs': True}, 'XGBoost': {'pipeline': Pipeline([('prep', TabularSurvivalPreprocessor(CONTINUOUS_COLUMNS, BINARY_COLUMNS, scale_continuous=False)), ('model', XGBoostCoxWrapper(random_state=random_state, n_jobs=1))]), 'param_grid': singleton_configurations([{'model__n_estimators': n_estimators, 'model__learning_rate': learning_rate, 'model__max_depth': max_depth, 'model__subsample': subsample, 'model__colsample_bytree': colsample, 'model__min_child_weight': min_child_weight, 'model__reg_lambda': reg_lambda} for n_estimators, learning_rate, max_depth, subsample, colsample, min_child_weight, reg_lambda in [(100, 0.05, 1, 1.0, 1.0, 1.0, 1.0), (200, 0.03, 2, 1.0, 1.0, 1.0, 1.0), (200, 0.03, 2, 0.8, 0.8, 1.0, 1.0), (300, 0.01, 2, 0.8, 0.8, 1.0, 5.0), (100, 0.05, 2, 0.8, 0.8, 5.0, 5.0), (100, 0.05, 3, 0.8, 0.8, 5.0, 5.0), (200, 0.03, 3, 0.8, 0.8, 5.0, 5.0), (100, 0.1, 2, 0.8, 0.8, 5.0, 5.0)]]), 'supports_survival_probs': True}, 'SVM': {'pipeline': Pipeline([('prep', TabularSurvivalPreprocessor(CONTINUOUS_COLUMNS, BINARY_COLUMNS, scale_continuous=True)), ('model', FastSurvivalSVM(random_state=random_state))]), 'param_grid': singleton_configurations([{'model__alpha': alpha, 'model__optimizer': 'avltree', 'model__max_iter': 500, 'model__rank_ratio': 1.0} for alpha in [0.01, 0.1, 1.0, 10.0]]), 'supports_survival_probs': False}}

def save_search_spaces(out_path: str, specs: Dict[str, Dict[str, Any]]) -> None:
    expected_counts = {model_name: PRESPECIFIED_CANDIDATE_COUNTS[model_name] for model_name in ['CoxPH', 'ElasticNetCox', 'RSF', 'GBSA', 'XGBoost', 'SVM']}
    observed_counts = {model_name: len(list(ParameterGrid(specs[model_name]['param_grid']))) for model_name in expected_counts}
    if observed_counts != expected_counts or sum(observed_counts.values()) != 33:
        raise RuntimeError(f'The locked 33-configuration classical search changed: {observed_counts}.')
    payload = {}
    for model_name, spec in specs.items():
        payload[model_name] = {'description': MODEL_DESCRIPTIONS.get(model_name), 'param_grid': spec['param_grid'], 'n_prespecified_configurations': observed_counts[model_name], 'supports_survival_probs': spec['supports_survival_probs'], **model_selection_metadata(model_name)}
    payload['DeepSurv'] = {'description': MODEL_DESCRIPTIONS['DeepSurv'], 'configs': get_deepsurv_configs(), 'supports_survival_probs': True, **model_selection_metadata('DeepSurv')}
    payload['CoxTime'] = {'description': MODEL_DESCRIPTIONS['CoxTime'], 'configs': get_coxtime_configs(), 'supports_survival_probs': True, **model_selection_metadata('CoxTime')}
    save_json(payload, out_path)

def fit_gridsearched_model(model_name: str, train_df: pd.DataFrame, inner_folds: int, random_state: int, n_jobs: int, specs: Dict[str, Dict[str, Any]], inner_completed_splits: Optional[Sequence[Tuple[pd.DataFrame, pd.DataFrame]]]=None, ibs_start_month: float=PRESPECIFIED_IBS_START_MONTH, horizon_months: float=PRESPECIFIED_HORIZON_MONTHS) -> Tuple[Any, Dict[str, Any], Optional[float]]:
    del random_state
    if model_name not in specs:
        raise ValueError(f'Unknown classical model: {model_name!r}.')
    candidates = list(ParameterGrid(specs[model_name]['param_grid']))
    if not candidates:
        raise ValueError(f'{model_name} has no prespecified configurations.')
    if model_name in FIXED_SPECIFICATION_MODELS:
        if len(candidates) != 1:
            raise ValueError('CoxPH must have exactly one fixed specification.')
        fixed_parameters = dict(candidates[0])
        fixed_estimator = clone(specs[model_name]['pipeline'])
        fixed_estimator.set_params(**fixed_parameters)
        fixed_estimator.fit(train_df[PREDICTOR_COLUMNS], get_surv_array(train_df))
        return (fixed_estimator, fixed_parameters, None)
    if inner_completed_splits is None:
        raise ValueError(f'{model_name} requires completed inner folds so imputation and IPCW estimation remain training-only.')
    if len(inner_completed_splits) != inner_folds:
        raise ValueError(f'{model_name} received {len(inner_completed_splits)} completed inner folds; expected {inner_folds}.')
    scoring_contexts: List[Optional[Dict[str, Any]]] = []
    evaluation_times: List[Optional[np.ndarray]] = []
    for inner_train, inner_validation in inner_completed_splits:
        if model_name in IBS_SELECTED_MODELS:
            times = get_eval_times(inner_train, inner_validation, start_month=ibs_start_month, horizon_months=horizon_months)
            context = prepare_brier_scoring_context(y_train=get_surv_array(inner_train), y_test=get_surv_array(inner_validation), eval_times=times, horizon_months=horizon_months)
            evaluation_times.append(times)
            scoring_contexts.append(context)
        elif model_name in HARRELL_SELECTED_MODELS:
            evaluation_times.append(None)
            scoring_contexts.append(None)
        else:
            raise ValueError(f'No selection rule is defined for {model_name!r}.')

    def evaluate_candidate(candidate_index: int, parameters: Dict[str, Any]) -> Tuple[int, Dict[str, Any], float, List[str]]:
        fold_scores: List[float] = []
        failures: List[str] = []
        for inner_fold, (inner_train, inner_validation) in enumerate(inner_completed_splits, start=1):
            try:
                estimator = clone(specs[model_name]['pipeline'])
                estimator.set_params(**parameters)
                estimator.fit(inner_train[PREDICTOR_COLUMNS], get_surv_array(inner_train))
                if model_name in IBS_SELECTED_MODELS:
                    times = evaluation_times[inner_fold - 1]
                    context = scoring_contexts[inner_fold - 1]
                    if times is None or context is None:
                        raise RuntimeError('IBS scoring context was not prepared.')
                    survival_matrix = predict_survival_matrix(estimator, inner_validation[PREDICTOR_COLUMNS], times)
                    score = compute_brier_metrics_from_context(survival_matrix, context)['integrated_brier_score']
                    if not np.isfinite(score) or not 0.0 <= score <= 1.0:
                        raise RuntimeError('validation IBS was invalid')
                else:
                    risk = np.asarray(estimator.predict(inner_validation[PREDICTOR_COLUMNS]), dtype=float)
                    score = harrell_c_index(get_surv_array(inner_validation), risk)
                    if not np.isfinite(score) or not 0.0 <= score <= 1.0:
                        raise RuntimeError('validation Harrell C was invalid')
                fold_scores.append(float(score))
            except Exception as exc:
                failures.append(f'candidate {candidate_index}, inner fold {inner_fold}: {type(exc).__name__}: {exc}')
        mean_score = float(np.mean(fold_scores)) if len(fold_scores) == inner_folds else float('nan')
        return (candidate_index, dict(parameters), mean_score, failures)
    evaluated = Parallel(n_jobs=n_jobs, prefer='processes')((delayed(evaluate_candidate)(candidate_index, parameters) for candidate_index, parameters in enumerate(candidates)))
    valid = [record for record in evaluated if np.isfinite(record[2])]
    if len(valid) != len(candidates):
        failure_text = ' | '.join((failure for _, _, _, failures in evaluated for failure in failures))
        raise RuntimeError(f'Not every prespecified hyperparameter candidate completed all inner folds for {model_name}. ' + failure_text)
    if model_name in IBS_SELECTED_MODELS:
        best_index, best_params, best_score, _ = min(valid, key=lambda record: (record[2], record[0]))
    else:
        best_index, best_params, best_score, _ = max(valid, key=lambda record: (record[2], -record[0]))
    del best_index
    final_estimator = clone(specs[model_name]['pipeline'])
    final_estimator.set_params(**best_params)
    final_estimator.fit(train_df[PREDICTOR_COLUMNS], get_surv_array(train_df))
    return (final_estimator, best_params, float(best_score))

def evaluate_model(model_name: str, fitted_model: Any, train_df: pd.DataFrame, test_df: pd.DataFrame, horizon_months: float, ibs_start_month: float, calibration_groups: int, compute_auxiliary_calibration: bool=True) -> Tuple[Dict[str, Any], Optional[pd.DataFrame], Optional[np.ndarray]]:
    y_train = get_surv_array(train_df)
    y_test = get_surv_array(test_df)
    X_test = test_df[PREDICTOR_COLUMNS].copy()
    tau = min(float(horizon_months), np.nextafter(float(train_df[TIME_COL].max()), -np.inf))
    uno_tau = min(np.nextafter(tau, np.inf), np.nextafter(float(train_df[TIME_COL].max()), -np.inf))
    metrics: Dict[str, Any] = {'harrell_c': np.nan, 'uno_c_tau': np.nan, 'tau_months': tau, 'brier_score_at_horizon': np.nan, 'brier_horizon_months': np.nan, 'integrated_brier_score': np.nan, 'calibration_slope': np.nan, 'calibration_slope_ci_low': np.nan, 'calibration_slope_ci_high': np.nan, 'calibration_slope_p': np.nan, 'calibration_observed_km_risk': np.nan, 'calibration_observed_km_risk_ci_low': np.nan, 'calibration_observed_km_risk_ci_high': np.nan, 'calibration_mean_predicted_risk': np.nan, 'calibration_observed_events_km': np.nan, 'calibration_expected_events': np.nan, 'calibration_oe_ratio': np.nan, 'calibration_oe_ratio_ci_low': np.nan, 'calibration_oe_ratio_ci_high': np.nan}
    metric_errors: Dict[str, str] = {}
    try:
        risk_test = np.asarray(fitted_model.predict(X_test), dtype=float)
    except Exception as exc:
        risk_test = np.full(len(test_df), np.nan, dtype=float)
        metric_errors['risk_prediction'] = f'{type(exc).__name__}: {exc}'
    else:
        try:
            metrics['harrell_c'] = harrell_c_index(y_test, risk_test)
        except Exception as exc:
            metric_errors['harrell_c'] = f'{type(exc).__name__}: {exc}'
        try:
            metrics['uno_c_tau'] = uno_c_index(y_train, y_test, risk_test, tau=uno_tau)
            if not np.isfinite(metrics['uno_c_tau']):
                metric_errors['uno_c_tau'] = 'Metric returned a non-finite value.'
        except Exception as exc:
            metric_errors['uno_c_tau'] = f'{type(exc).__name__}: {exc}'
    calibration_df: Optional[pd.DataFrame] = None
    pred_event_prob: Optional[np.ndarray] = None
    if model_name != 'SVM':
        try:
            eval_times = get_eval_times(train_df, test_df, start_month=ibs_start_month, horizon_months=horizon_months)
            surv_matrix = predict_survival_matrix(fitted_model, X_test, eval_times)
            idx = int(np.argmin(np.abs(eval_times - horizon_months)))
            horizon_used = float(eval_times[idx])
            metrics['brier_horizon_months'] = horizon_used
            pred_event_prob = 1.0 - surv_matrix[:, idx]
        except Exception as exc:
            eval_times = np.asarray([], dtype=float)
            surv_matrix = np.empty((len(test_df), 0), dtype=float)
            horizon_used = float('nan')
            metric_errors['survival_probability_prediction'] = f'{type(exc).__name__}: {exc}'
        if eval_times.size > 0:
            try:
                metrics.update(compute_brier_metrics(y_train, y_test, surv_matrix, eval_times, horizon_months=horizon_months))
            except Exception as exc:
                error_text = f'{type(exc).__name__}: {exc}'
                metric_errors['brier_score_at_horizon'] = error_text
                metric_errors['integrated_brier_score'] = error_text
            if compute_auxiliary_calibration:
                try:
                    calibration_df = grouped_calibration_table(test_df, pred_event_prob, time_point=horizon_used, n_groups=calibration_groups)
                except Exception as exc:
                    LOGGER.warning('%s auxiliary per-fit grouped calibration could not be created: %s: %s', model_name, type(exc).__name__, exc)
                try:
                    metrics.update(calibration_oe_from_probs(test_df, pred_event_prob, time_point=horizon_used))
                except Exception as exc:
                    metric_errors['calibration_oe_ratio'] = f'{type(exc).__name__}: {exc}'
                try:
                    metrics.update(calibration_slope_from_probs(test_df, pred_event_prob, time_point=horizon_used))
                except Exception as exc:
                    metric_errors['calibration_slope'] = f'{type(exc).__name__}: {exc}'
    expected_metrics = ['harrell_c', 'uno_c_tau', 'brier_score_at_horizon', 'integrated_brier_score', 'calibration_slope', 'calibration_oe_ratio']
    if model_name == 'SVM':
        expected_metrics = ['harrell_c', 'uno_c_tau']
    elif not compute_auxiliary_calibration:
        expected_metrics = ['harrell_c', 'uno_c_tau', 'brier_score_at_horizon', 'integrated_brier_score']
    for name in expected_metrics:
        if not np.isfinite(metrics[name]) and name not in metric_errors:
            metric_errors[name] = 'Metric returned a non-finite value.'
    available = sum((bool(np.isfinite(metrics[name])) for name in expected_metrics))
    metrics['n_metrics_expected'] = len(expected_metrics)
    metrics['n_metrics_available'] = available
    metrics['metrics_complete'] = bool(available == len(expected_metrics))
    metrics['metric_errors_json'] = json.dumps(metric_errors, sort_keys=True)
    if metric_errors:
        LOGGER.warning('%s metric completeness %d/%d; errors=%s', model_name, available, len(expected_metrics), metric_errors)
    return (metrics, calibration_df, pred_event_prob)

def require_complete_metrics(model_name: str, metrics: Dict[str, Any], context: str) -> None:
    if bool(metrics.get('metrics_complete', False)):
        return
    raise RuntimeError(f"{context}: {model_name} produced {metrics.get('n_metrics_available', 0)}/{metrics.get('n_metrics_expected', '?')} required metrics. Details: {metrics.get('metric_errors_json', '{}')}")

def metric_completeness_summary(metrics_df: pd.DataFrame) -> pd.DataFrame:
    metric_columns = ['harrell_c', 'uno_c_tau', 'brier_score_at_horizon', 'integrated_brier_score', 'calibration_slope', 'calibration_oe_ratio']
    rows: List[Dict[str, Any]] = []
    for model_name, sub in metrics_df.groupby('model'):
        row: Dict[str, Any] = {'model': model_name, 'n_model_fits': int(len(sub))}
        for metric in metric_columns:
            row[f'{metric}_n_available'] = int(sub.get(metric, pd.Series(dtype=float)).notna().sum())
        if 'metrics_complete' in sub:
            complete_values = sub['metrics_complete'].fillna(False)
            if complete_values.dtype != bool:
                complete_values = complete_values.astype(str).str.lower().isin(['true', '1'])
            row['n_complete_fits'] = int(complete_values.astype(bool).sum())
        else:
            row['n_complete_fits'] = 0
        if 'metric_errors_json' in sub:
            row['n_fits_with_logged_metric_error'] = int(sub['metric_errors_json'].fillna('{}').astype(str).ne('{}').sum())
        else:
            row['n_fits_with_logged_metric_error'] = int(len(sub))
        rows.append(row)
    return pd.DataFrame(rows)

def failed_model_metric_record(exc: Exception) -> Dict[str, Any]:
    return {'harrell_c': np.nan, 'uno_c_tau': np.nan, 'tau_months': np.nan, 'brier_score_at_horizon': np.nan, 'brier_horizon_months': np.nan, 'integrated_brier_score': np.nan, 'calibration_slope': np.nan, 'calibration_slope_ci_low': np.nan, 'calibration_slope_ci_high': np.nan, 'calibration_slope_p': np.nan, 'calibration_observed_km_risk': np.nan, 'calibration_observed_km_risk_ci_low': np.nan, 'calibration_observed_km_risk_ci_high': np.nan, 'calibration_mean_predicted_risk': np.nan, 'calibration_observed_events_km': np.nan, 'calibration_expected_events': np.nan, 'calibration_oe_ratio': np.nan, 'calibration_oe_ratio_ci_low': np.nan, 'calibration_oe_ratio_ci_high': np.nan, 'n_metrics_expected': np.nan, 'n_metrics_available': 0, 'metrics_complete': False, 'metric_errors_json': json.dumps({'model_fit': f'{type(exc).__name__}: {exc}'}, sort_keys=True)}

def combine_shap_summary_images(image_paths: List[str], titles: List[str], out_path: str, ncols: int=2) -> None:
    valid = [(p, t) for p, t in zip(image_paths, titles) if os.path.exists(p)]
    if not valid:
        return
    n = len(valid)
    ncols = min(ncols, n)
    nrows = int(math.ceil(n / ncols))
    fig, axes = plt.subplots(nrows, ncols, figsize=(11.5, 4.2 * nrows))
    axes = np.asarray(axes).reshape(-1)
    for ax in axes:
        ax.axis('off')
    for panel_index, (ax, (img_path, title)) in enumerate(zip(axes, valid)):
        ax.imshow(mpimg.imread(img_path))
        ax.set_title(title, fontsize=11)
        ax.text(0.01, 0.99, chr(ord('A') + panel_index), transform=ax.transAxes, ha='left', va='top', fontsize=13, fontweight='bold', bbox={'facecolor': 'white', 'edgecolor': 'none', 'alpha': 0.85, 'pad': 1.5})
        ax.axis('off')
    plt.tight_layout()
    plt.savefig(out_path, dpi=300, bbox_inches='tight')
    plt.close()

def combine_force_plots(image_paths: List[str], titles: List[str], out_path: str) -> None:
    valid = [(p, t) for p, t in zip(image_paths, titles) if os.path.exists(p)]
    if not valid:
        return
    fig, axes = plt.subplots(1, len(valid), figsize=(7 * len(valid), 5))
    axes = np.atleast_1d(axes)
    for panel_index, (ax, (img_path, title)) in enumerate(zip(axes, valid)):
        ax.imshow(mpimg.imread(img_path))
        ax.set_title(title, fontsize=11)
        ax.text(0.01, 0.99, chr(ord('A') + panel_index), transform=ax.transAxes, ha='left', va='top', fontsize=13, fontweight='bold', bbox={'facecolor': 'white', 'edgecolor': 'none', 'alpha': 0.85, 'pad': 1.5})
        ax.axis('off')
    plt.tight_layout()
    plt.savefig(out_path, dpi=300, bbox_inches='tight')
    plt.close()

def make_combined_calibration_figure(calibration_tables: Dict[str, pd.DataFrame], out_path: str, horizon_months: float, source_note: Optional[str]=None) -> None:
    del source_note
    fig, axes = plt.subplots(2, 4, figsize=(14.5, 7.2))
    axes = axes.ravel()
    models = ['CoxPH', 'ElasticNetCox', 'RSF', 'GBSA', 'DeepSurv', 'CoxTime', 'XGBoost', 'SVM']
    finite_maxima = [float(np.nanmax(table[['mean_predicted_risk', 'observed_km_risk', 'observed_km_risk_95ci_high']].to_numpy(dtype=float))) for table in calibration_tables.values() if not table.empty]
    finite_maxima = [value for value in finite_maxima if np.isfinite(value)]
    raw_max = max(finite_maxima, default=0.05) * 1.1
    axis_max = min(1.0, max(0.05, math.ceil(raw_max / 0.05) * 0.05))
    for panel_index, (ax, model_name) in enumerate(zip(axes, models)):
        panel_title = f"{chr(ord('A') + panel_index)}. {model_name}"
        if model_name in calibration_tables:
            plot_calibration_panel(ax, calibration_tables[model_name], panel_title, time_point=horizon_months, axis_max=axis_max)
        else:
            ax.axis('off')
            ax.text(0.5, 0.5, f'{panel_title}\nNot available\n(no survival probabilities)', ha='center', va='center', fontsize=11)
    plt.tight_layout(w_pad=1.0, h_pad=1.0)
    _save_calibration_composite(fig, out_path, dpi=600)
    plt.close()

def make_combined_flexible_calibration_figure(calibration_tables: Dict[str, pd.DataFrame], out_path: str, horizon_months: float, source_note: Optional[str]=None) -> None:
    del source_note
    fig, axes = plt.subplots(2, 4, figsize=(14.5, 7.2))
    axes = axes.ravel()
    models = ['CoxPH', 'ElasticNetCox', 'RSF', 'GBSA', 'DeepSurv', 'CoxTime', 'XGBoost', 'SVM']
    maxima: List[float] = []
    for table in calibration_tables.values():
        for column in ['predicted_risk', 'observed_risk_smoothed', 'bootstrap_95ci_high']:
            if column in table and table[column].notna().any():
                maxima.append(float(np.nanmax(table[column].to_numpy(dtype=float))))
    raw_max = max(maxima, default=0.05) * 1.1
    axis_max = min(1.0, max(0.05, math.ceil(raw_max / 0.05) * 0.05))
    for panel_index, (ax, model_name) in enumerate(zip(axes, models)):
        panel_title = f"{chr(ord('A') + panel_index)}. {model_name}"
        table = calibration_tables.get(model_name)
        if table is None or table.empty:
            ax.axis('off')
            ax.text(0.5, 0.5, f'{panel_title}\nNot available\n(no survival probabilities)', ha='center', va='center', fontsize=11)
            continue
        ax.plot([0, axis_max], [0, axis_max], linestyle='--', linewidth=1, color='0.4')
        ax.plot(table['predicted_risk'], table['observed_risk_smoothed'], linewidth=1.5, color='#1f77b4')
        if table['bootstrap_95ci_low'].notna().any():
            ax.fill_between(table['predicted_risk'], table['bootstrap_95ci_low'], table['bootstrap_95ci_high'], color='#1f77b4', alpha=0.2, linewidth=0)
        ax.set_title(panel_title, fontsize=11)
        ax.set_xlabel(f'Predicted {int(horizon_months)}-month risk', fontsize=9)
        ax.set_ylabel('Smoothed observed risk', fontsize=9)
        ax.set_xlim(0, axis_max)
        ax.set_ylim(0, axis_max)
        ax.set_aspect('equal', adjustable='box')
        ax.tick_params(axis='both', labelsize=8)
    plt.tight_layout(w_pad=1.0, h_pad=1.0)
    _save_calibration_composite(fig, out_path, dpi=600)
    plt.close()

def _save_calibration_composite(fig: Any, out_path: str, dpi: int=600) -> None:
    if getattr(fig, '_suptitle', None) is not None or list(getattr(fig, 'texts', [])):
        raise RuntimeError('Calibration composites must not contain an overall title or figure-level footer.')
    root, extension = os.path.splitext(out_path)
    png_path = out_path if extension.lower() == '.png' else root + '.png'
    fig.savefig(png_path, dpi=int(dpi), bbox_inches='tight', facecolor='white')
    tiff_path = root + '.tiff'
    with Image.open(png_path) as image:
        image.convert('RGB').save(tiff_path, format='TIFF', compression='tiff_lzw', dpi=(int(dpi), int(dpi)))

def _primary_imputation_checkpoint_paths(checkpoint_dir: str, imputation: int) -> Dict[str, str]:
    directory = ensure_dir(os.path.join(checkpoint_dir, 'imputations'))
    prefix = f'imputation_{imputation:04d}'
    return {'metrics': os.path.join(directory, f'{prefix}_metrics.csv'), 'tuning': os.path.join(directory, f'{prefix}_tuning.csv'), 'predictions': os.path.join(directory, f'{prefix}_predictions.csv'), 'marker': os.path.join(directory, f'{prefix}_complete.json')}

def validate_primary_imputation_checkpoint_frames(metrics_df: pd.DataFrame, tuning_df: pd.DataFrame, predictions_df: pd.DataFrame, imputation: int, expected_test_ids: pd.Series) -> None:
    required_metrics = {'imputation', 'model', 'metrics_complete', 'metric_errors_json', *FIT_REQUIRED_METRICS}
    if not required_metrics.issubset(metrics_df.columns):
        raise ValueError('Fixed-split imputation checkpoint metric columns were incomplete.')
    if len(metrics_df) != len(MODEL_ORDER) or set(metrics_df['model'].astype(str)) != set(MODEL_ORDER) or metrics_df['model'].astype(str).duplicated().any() or (set(pd.to_numeric(metrics_df['imputation'], errors='raise').astype(int)) != {int(imputation)}):
        raise ValueError('Fixed-split imputation checkpoint metric keys were incomplete.')
    if not _validated_boolean_series(metrics_df['metrics_complete'], 'fixed-split imputation metrics_complete').all():
        raise ValueError('Fixed-split imputation checkpoint included an incomplete model fit.')
    for value in metrics_df['metric_errors_json'].fillna('{}').astype(str):
        parsed = json.loads(value)
        if not isinstance(parsed, dict) or parsed:
            raise ValueError('Fixed-split imputation checkpoint included a metric error.')
    for model_name, rows in metrics_df.groupby('model'):
        for column in _required_metric_columns(str(model_name)):
            values = pd.to_numeric(rows[column], errors='coerce').to_numpy(float)
            if not np.isfinite(values).all():
                raise ValueError(f'Fixed-split imputation checkpoint has a non-finite {column} for {model_name}.')
            for value in values:
                validate_metric_bounds(column, value, f'Fixed-split imputation {imputation} {model_name}')
        if not (pd.to_numeric(rows['tau_months'], errors='raise') == PRESPECIFIED_HORIZON_MONTHS).all():
            raise ValueError('Fixed-split checkpoint has the wrong Uno horizon.')
        if str(model_name) != 'SVM' and (not (pd.to_numeric(rows['brier_horizon_months'], errors='raise') == PRESPECIFIED_HORIZON_MONTHS).all()):
            raise ValueError('Fixed-split checkpoint has the wrong Brier horizon.')
    required_tuning = {'split', 'repeat', 'imputation', 'model', 'tuning_method', 'best_params_json', *TUNING_SELECTION_COLUMNS}
    if not required_tuning.issubset(tuning_df.columns):
        raise ValueError('Fixed-split imputation tuning checkpoint columns were incomplete.')
    if len(tuning_df) != len(MODEL_ORDER) or set(tuning_df['model'].astype(str)) != set(MODEL_ORDER) or tuning_df['model'].astype(str).duplicated().any() or (set(pd.to_numeric(tuning_df['repeat'], errors='raise').astype(int)) != {0}) or (set(pd.to_numeric(tuning_df['imputation'], errors='raise').astype(int)) != {int(imputation)}):
        raise ValueError('Fixed-split imputation tuning checkpoint keys were incomplete.')
    for value in tuning_df['best_params_json'].astype(str):
        parsed = json.loads(value)
        if not isinstance(parsed, dict):
            raise ValueError('Fixed-split imputation best_params_json was not an object.')
    for _, row in tuning_df.iterrows():
        model_name = str(row['model'])
        if str(row['tuning_method']) != main_tuning_method(model_name):
            raise ValueError(f'Fixed-split {model_name} tuning method is wrong.')
        validate_tuning_selection_record(row, model_name, f'Fixed-split imputation {imputation} {model_name}')
    expected_prediction_columns = [STUDY_ID_COL, *CALIBRATION_MODELS]
    if list(predictions_df.columns) != expected_prediction_columns:
        raise ValueError('Fixed-split imputation prediction checkpoint columns were incomplete.')
    if len(predictions_df) != len(expected_test_ids):
        raise ValueError('Fixed-split imputation prediction checkpoint has the wrong row count.')
    if not np.array_equal(predictions_df[STUDY_ID_COL].to_numpy(), expected_test_ids.to_numpy()):
        raise ValueError('Fixed-split imputation prediction checkpoint study IDs changed order.')
    probabilities = predictions_df[CALIBRATION_MODELS].apply(pd.to_numeric, errors='raise').to_numpy(float)
    if not np.isfinite(probabilities).all() or np.any((probabilities < -1e-08) | (probabilities > 1.0 + 1e-08)):
        raise ValueError('Fixed-split imputation prediction checkpoint has invalid probabilities.')

def write_primary_imputation_checkpoint(checkpoint_dir: str, manifest_sha256: str, imputation: int, metrics_df: pd.DataFrame, tuning_df: pd.DataFrame, predictions_df: pd.DataFrame, expected_test_ids: pd.Series) -> None:
    validate_primary_imputation_checkpoint_frames(metrics_df, tuning_df, predictions_df, imputation=imputation, expected_test_ids=expected_test_ids)
    paths = _primary_imputation_checkpoint_paths(checkpoint_dir, imputation)
    atomic_write_dataframe_csv(metrics_df, paths['metrics'])
    atomic_write_dataframe_csv(tuning_df, paths['tuning'])
    atomic_write_dataframe_csv(predictions_df, paths['predictions'])
    files: Dict[str, Dict[str, Any]] = {}
    for key in ['metrics', 'tuning', 'predictions']:
        frame = {'metrics': metrics_df, 'tuning': tuning_df, 'predictions': predictions_df}[key]
        files[key] = {'name': os.path.basename(paths[key]), 'sha256': file_sha256(paths[key]), 'n_rows': int(len(frame))}
    atomic_write_json({'checkpoint_schema_version': CHECKPOINT_SCHEMA_VERSION, 'complete': True, 'manifest_sha256': manifest_sha256, 'imputation': int(imputation), 'files': files}, paths['marker'])

def load_primary_imputation_checkpoint(checkpoint_dir: str, manifest_sha256: str, imputation: int, expected_test_ids: pd.Series) -> Optional[Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]]:
    paths = _primary_imputation_checkpoint_paths(checkpoint_dir, imputation)
    if not os.path.exists(paths['marker']):
        return None
    try:
        with open(paths['marker'], 'r', encoding='utf-8') as handle:
            marker = json.load(handle)
        if marker.get('checkpoint_schema_version') != CHECKPOINT_SCHEMA_VERSION or marker.get('complete') is not True or marker.get('manifest_sha256') != manifest_sha256 or (int(marker.get('imputation', -1)) != int(imputation)):
            raise ValueError('Fixed-split imputation completion marker did not validate.')
        loaded: Dict[str, pd.DataFrame] = {}
        for key in ['metrics', 'tuning', 'predictions']:
            record = marker['files'][key]
            if record.get('name') != os.path.basename(paths[key]):
                raise ValueError('Fixed-split imputation checkpoint filename changed.')
            if not os.path.exists(paths[key]) or file_sha256(paths[key]) != record.get('sha256'):
                raise ValueError('Fixed-split imputation checkpoint file was missing or changed.')
            frame = pd.read_csv(paths[key], float_precision='round_trip')
            if len(frame) != int(record.get('n_rows', -1)):
                raise ValueError('Fixed-split imputation checkpoint row count changed.')
            loaded[key] = frame
        validate_primary_imputation_checkpoint_frames(loaded['metrics'], loaded['tuning'], loaded['predictions'], imputation=imputation, expected_test_ids=expected_test_ids)
        return (loaded['metrics'], loaded['tuning'], loaded['predictions'])
    except Exception as exc:
        LOGGER.warning('Fixed-split imputation %d checkpoint was invalid and will be recomputed: %s', imputation, exc)
        return None

def refit_primary_shap_reference_models(first_imputed_split: Dict[str, pd.DataFrame], first_imputation_tuning: pd.DataFrame, seed: int, n_jobs: int) -> Tuple[Dict[str, Any], pd.DataFrame, pd.DataFrame]:
    train_df = first_imputed_split['train']
    test_df = first_imputed_split['test']
    split_seed = seed + 1000 + 1
    parameter_map = {str(row['model']): json.loads(str(row['best_params_json'])) for _, row in first_imputation_tuning.iterrows()}
    models: Dict[str, Any] = {}
    specs = get_classical_model_specs(random_state=split_seed, n_jobs=n_jobs)
    x_train = train_df[PREDICTOR_COLUMNS].copy()
    y_train = get_surv_array(train_df)
    for model_name in ['CoxPH', 'ElasticNetCox', 'RSF', 'GBSA', 'XGBoost', 'SVM']:
        fitted = clone(specs[model_name]['pipeline'])
        fitted.set_params(**parameter_map[model_name])
        fitted.fit(x_train, y_train)
        models[model_name] = fitted
    models['DeepSurv'] = fit_deepsurv_fixed_config(train_df, split_seed + 10000, parameter_map['DeepSurv'])
    models['CoxTime'] = fit_coxtime_fixed_config(train_df, split_seed + 20000, parameter_map['CoxTime'])
    return (models, train_df.copy(), test_df.copy())

def run_primary_split_analysis(df: pd.DataFrame, out_dir: str, seed: int, test_size: float, inner_folds: int, horizon_months: float, ibs_start_month: float, calibration_groups: int, n_jobs: int, m_imputations: int, impute_max_iter: int, make_shap: bool, shap_max_background: int, shap_max_explain: int, calibration_bootstrap_reps: int, resume: bool, checkpoint_dir: str, manifest_sha256: str, deadline_monotonic: Optional[float]=None) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    LOGGER.info('Running the prespecified fixed 80/20 analysis (master seed %d).', seed)
    tables_dir = ensure_dir(os.path.join(out_dir, 'tables'))
    figs_dir = ensure_dir(os.path.join(out_dir, 'figures'))
    meta_dir = ensure_dir(os.path.join(out_dir, 'metadata'))
    train_idx, test_idx = primary_split(df, seed=seed, test_size=test_size)
    raw_train_df = df.iloc[train_idx].copy().reset_index(drop=True)
    raw_test_df = df.iloc[test_idx].copy().reset_index(drop=True)
    primary_seed = seed + 1000
    imputed_splits = generate_imputed_split_datasets(raw_train_df, raw_test_df, m=m_imputations, seed=primary_seed, max_iter=impute_max_iter)
    table1 = aggregate_table1_over_imputations(imputed_splits, raw_train_df, raw_test_df, df)
    table1.to_csv(os.path.join(tables_dir, 'table1_fixed_split.csv'), index=False)
    pooled_cox_table, ph_by_imputation = pool_cox_results([x['train'] for x in imputed_splits])
    pooled_cox_table.to_csv(os.path.join(tables_dir, 'table2_cox_hazard_ratios_fixed_split.csv'), index=False)
    ph_by_imputation.to_csv(os.path.join(tables_dir, 'proportional_hazards_tests_by_imputation.csv'), index=False)
    model_specs = get_classical_model_specs(random_state=seed, n_jobs=n_jobs)
    save_search_spaces(os.path.join(meta_dir, 'hyperparameter_search_spaces.json'), model_specs)
    primary_metrics_records: List[Dict[str, Any]] = []
    tuning_rows: List[Dict[str, Any]] = []
    calibration_pred_store: Dict[str, List[np.ndarray]] = {m: [] for m in CALIBRATION_MODELS}
    shap_reference_models: Dict[str, Any] = {}
    shap_reference_train: Optional[pd.DataFrame] = None
    shap_reference_test: Optional[pd.DataFrame] = None
    first_imputation_tuning: Optional[pd.DataFrame] = None
    for imp_idx, split in enumerate(imputed_splits, start=1):
        LOGGER.info('Fixed-split imputation %d / %d', imp_idx, m_imputations)
        train_df = split['train']
        test_df = split['test']
        loaded_imputation = None
        if resume:
            loaded_imputation = load_primary_imputation_checkpoint(checkpoint_dir=checkpoint_dir, manifest_sha256=manifest_sha256, imputation=imp_idx, expected_test_ids=raw_test_df[STUDY_ID_COL])
        if loaded_imputation is not None:
            metrics_checkpoint, tuning_checkpoint, prediction_checkpoint = loaded_imputation
            primary_metrics_records.extend(metrics_checkpoint.to_dict('records'))
            tuning_rows.extend(tuning_checkpoint.to_dict('records'))
            for model_name in CALIBRATION_MODELS:
                calibration_pred_store[model_name].append(prediction_checkpoint[model_name].to_numpy(dtype=float))
            if imp_idx == 1:
                first_imputation_tuning = tuning_checkpoint.copy()
            LOGGER.info('Resumed fixed-split imputation %d / %d from a validated checkpoint.', imp_idx, m_imputations)
            check_session_deadline(deadline_monotonic, f'fixed-split imputation {imp_idx}')
            continue
        split_seed = primary_seed + imp_idx
        specs = get_classical_model_specs(random_state=split_seed, n_jobs=n_jobs)
        imputation_metrics: List[Dict[str, Any]] = []
        imputation_tuning: List[Dict[str, Any]] = []
        imputation_predictions: Dict[str, np.ndarray] = {}
        inner_completed_splits: Optional[List[Tuple[pd.DataFrame, pd.DataFrame]]] = None
        neural_selection_split: Optional[Tuple[pd.DataFrame, pd.DataFrame]] = None
        for model_name in ['CoxPH', 'ElasticNetCox', 'RSF', 'GBSA', 'XGBoost', 'SVM']:
            partial = load_outer_model_checkpoint(checkpoint_dir=checkpoint_dir, manifest_sha256=manifest_sha256, repetition=0, fold=imp_idx, model_name=model_name, expected_test_ids=raw_test_df[STUDY_ID_COL]) if resume else None
            if partial is not None:
                model_metrics_df, model_tuning_df, model_predictions_df = partial
                imputation_metrics.extend(model_metrics_df.to_dict('records'))
                imputation_tuning.extend(model_tuning_df.to_dict('records'))
                if model_name != 'SVM':
                    imputation_predictions[model_name] = model_predictions_df['predicted_event_probability'].to_numpy(dtype=float)
                LOGGER.info('  resumed fixed-split imputation %d model %s', imp_idx, model_name)
                continue
            check_session_deadline(deadline_monotonic, f'fixed-split imputation {imp_idx} before model {model_name}')
            if model_name != 'CoxPH' and inner_completed_splits is None:
                inner_completed_splits = generate_nested_inner_imputed_splits(raw_outer_train_df=raw_train_df, inner_folds=inner_folds, split_seed=split_seed, imputation_seed=primary_seed + 100000 + 1000 * imp_idx, max_iter=impute_max_iter)
            best_estimator, best_params, selection_score = fit_gridsearched_model(model_name, train_df, inner_folds=inner_folds, random_state=split_seed, n_jobs=n_jobs, specs=specs, inner_completed_splits=inner_completed_splits, ibs_start_month=ibs_start_month, horizon_months=horizon_months)
            metrics, _, pred_event_prob = evaluate_model(model_name, best_estimator, train_df, test_df, horizon_months, ibs_start_month, calibration_groups)
            require_complete_metrics(model_name, metrics, f'Fixed-split imputation {imp_idx}')
            model_metrics_df = pd.DataFrame([{'repetition': 0, 'fold': imp_idx, 'imputation': imp_idx, 'n_test': int(len(test_df)), 'model': model_name, **metrics}])
            model_tuning_df = pd.DataFrame([{'split': 'fixed', 'repeat': 0, 'repetition': 0, 'fold': imp_idx, 'imputation': imp_idx, 'n_test': int(len(test_df)), 'model': model_name, 'tuning_method': main_tuning_method(model_name), **tuning_output_metadata(model_name, selection_score), 'best_params_json': json.dumps(best_params, sort_keys=True)}])
            model_probabilities = np.full(len(test_df), np.nan, dtype=float) if model_name == 'SVM' else np.asarray(pred_event_prob, dtype=float)
            model_predictions_df = pd.DataFrame({STUDY_ID_COL: raw_test_df[STUDY_ID_COL].to_numpy(), 'predicted_event_probability': model_probabilities})
            write_outer_model_checkpoint(checkpoint_dir=checkpoint_dir, manifest_sha256=manifest_sha256, repetition=0, fold=imp_idx, model_name=model_name, metrics_df=model_metrics_df, tuning_df=model_tuning_df, predictions_df=model_predictions_df, expected_test_ids=raw_test_df[STUDY_ID_COL])
            imputation_metrics.extend(model_metrics_df.to_dict('records'))
            imputation_tuning.extend(model_tuning_df.to_dict('records'))
            if model_name != 'SVM':
                imputation_predictions[model_name] = model_probabilities
            if imp_idx == 1 and model_name in SHAP_MODELS:
                shap_reference_models[model_name] = best_estimator
            check_session_deadline(deadline_monotonic, f'fixed-split imputation {imp_idx} model {model_name}')
        for model_name, fit_function in [('DeepSurv', fit_deepsurv), ('CoxTime', fit_coxtime)]:
            partial = load_outer_model_checkpoint(checkpoint_dir=checkpoint_dir, manifest_sha256=manifest_sha256, repetition=0, fold=imp_idx, model_name=model_name, expected_test_ids=raw_test_df[STUDY_ID_COL]) if resume else None
            if partial is not None:
                model_metrics_df, model_tuning_df, model_predictions_df = partial
                imputation_metrics.extend(model_metrics_df.to_dict('records'))
                imputation_tuning.extend(model_tuning_df.to_dict('records'))
                imputation_predictions[model_name] = model_predictions_df['predicted_event_probability'].to_numpy(dtype=float)
                LOGGER.info('  resumed fixed-split imputation %d model %s', imp_idx, model_name)
                continue
            check_session_deadline(deadline_monotonic, f'fixed-split imputation {imp_idx} before model {model_name}')
            if neural_selection_split is None:
                neural_selection_split = generate_nested_neural_selection_split(raw_outer_train_df=raw_train_df, split_seed=split_seed, imputation_seed=primary_seed + 200000 + 1000 * imp_idx, max_iter=impute_max_iter)
            fitted_model, best_params, selection_score = fit_function(train_df, random_state=split_seed, selection_split=neural_selection_split, require_all_candidates=True, ibs_start_month=ibs_start_month, horizon_months=horizon_months)
            metrics, _, pred_event_prob = evaluate_model(model_name, fitted_model, train_df, test_df, horizon_months, ibs_start_month, calibration_groups)
            require_complete_metrics(model_name, metrics, f'Fixed-split imputation {imp_idx}')
            model_metrics_df = pd.DataFrame([{'repetition': 0, 'fold': imp_idx, 'imputation': imp_idx, 'n_test': int(len(test_df)), 'model': model_name, **metrics}])
            model_tuning_df = pd.DataFrame([{'split': 'fixed', 'repeat': 0, 'repetition': 0, 'fold': imp_idx, 'imputation': imp_idx, 'model': model_name, 'tuning_method': main_tuning_method(model_name), **tuning_output_metadata(model_name, selection_score), 'best_params_json': json.dumps(best_params, sort_keys=True)}])
            model_probabilities = np.asarray(pred_event_prob, dtype=float)
            model_predictions_df = pd.DataFrame({STUDY_ID_COL: raw_test_df[STUDY_ID_COL].to_numpy(), 'predicted_event_probability': model_probabilities})
            write_outer_model_checkpoint(checkpoint_dir=checkpoint_dir, manifest_sha256=manifest_sha256, repetition=0, fold=imp_idx, model_name=model_name, metrics_df=model_metrics_df, tuning_df=model_tuning_df, predictions_df=model_predictions_df, expected_test_ids=raw_test_df[STUDY_ID_COL])
            imputation_metrics.extend(model_metrics_df.to_dict('records'))
            imputation_tuning.extend(model_tuning_df.to_dict('records'))
            imputation_predictions[model_name] = model_probabilities
            if imp_idx == 1:
                shap_reference_models[model_name] = fitted_model
                if model_name == 'CoxTime':
                    shap_reference_train = train_df.copy()
                    shap_reference_test = test_df.copy()
            if model_name != 'CoxTime':
                check_session_deadline(deadline_monotonic, f'fixed-split imputation {imp_idx} model {model_name}')
        missing_predictions = sorted(set(CALIBRATION_MODELS).difference(imputation_predictions))
        if missing_predictions:
            raise RuntimeError(f'Fixed-split imputation {imp_idx} did not produce absolute-risk predictions for: {missing_predictions}.')
        metrics_checkpoint = pd.DataFrame(imputation_metrics)
        tuning_checkpoint = pd.DataFrame(imputation_tuning)
        prediction_checkpoint = pd.DataFrame({STUDY_ID_COL: raw_test_df[STUDY_ID_COL].to_numpy(), **{model_name: imputation_predictions[model_name] for model_name in CALIBRATION_MODELS}})
        write_primary_imputation_checkpoint(checkpoint_dir=checkpoint_dir, manifest_sha256=manifest_sha256, imputation=imp_idx, metrics_df=metrics_checkpoint, tuning_df=tuning_checkpoint, predictions_df=prediction_checkpoint, expected_test_ids=raw_test_df[STUDY_ID_COL])
        primary_metrics_records.extend(metrics_checkpoint.to_dict('records'))
        tuning_rows.extend(tuning_checkpoint.to_dict('records'))
        for model_name in CALIBRATION_MODELS:
            calibration_pred_store[model_name].append(prediction_checkpoint[model_name].to_numpy(dtype=float))
        if imp_idx == 1:
            first_imputation_tuning = tuning_checkpoint.copy()
        LOGGER.info('Checkpointed fixed-split imputation %d / %d.', imp_idx, m_imputations)
        check_session_deadline(deadline_monotonic, f'fixed-split imputation {imp_idx}')
    if make_shap and (set(shap_reference_models) != set(SHAP_MODELS) or shap_reference_train is None or shap_reference_test is None):
        if first_imputation_tuning is None:
            raise RuntimeError('Fixed-split imputation 1 tuning was unavailable for SHAP refitting.')
        LOGGER.info('Refitting imputation-1 model configurations for SHAP without repeating tuning.')
        shap_reference_models, shap_reference_train, shap_reference_test = refit_primary_shap_reference_models(first_imputed_split=imputed_splits[0], first_imputation_tuning=first_imputation_tuning, seed=seed, n_jobs=n_jobs)
    primary_metrics_long = pd.DataFrame(primary_metrics_records)
    primary_metrics_long.to_csv(os.path.join(tables_dir, 'supplementary_table_s5_fixed_split_by_imputation.csv'), index=False)
    metric_completeness_summary(primary_metrics_long).to_csv(os.path.join(meta_dir, 'metric_completeness_fixed_split.csv'), index=False)
    avg_metrics_rows = []
    for model_name, sub in primary_metrics_long.groupby('model'):
        row = {'Model': model_name}
        for col in ['harrell_c', 'uno_c_tau', 'brier_score_at_horizon', 'integrated_brier_score', 'calibration_slope', 'calibration_oe_ratio']:
            vals = sub[col].dropna()
            row[col] = float(vals.mean()) if not vals.empty else np.nan
        avg_metrics_rows.append(row)
    primary_metrics_df = pd.DataFrame(avg_metrics_rows).set_index('Model').loc[MODEL_ORDER].reset_index()
    primary_metrics_df = primary_metrics_df.rename(columns={'harrell_c': 'Harrell_C', 'uno_c_tau': 'Uno_C_tau', 'brier_score_at_horizon': 'Brier_120m', 'integrated_brier_score': 'IBS_12_120m', 'calibration_slope': 'Calibration_slope_120m', 'calibration_oe_ratio': 'O_E_120m'})
    calibration_tables: Dict[str, pd.DataFrame] = {}
    flexible_calibration_tables: Dict[str, pd.DataFrame] = {}
    absolute_calibration_rows: List[Dict[str, Any]] = []
    bootstrap_completeness_rows: List[Dict[str, Any]] = []
    grouped_completeness_rows: List[Dict[str, Any]] = []
    for model_name, pred_list in calibration_pred_store.items():
        if len(pred_list) == 0:
            continue
        mean_pred = np.mean(np.vstack(pred_list), axis=0)
        try:
            oe_metrics = calibration_oe_from_probs(raw_test_df, mean_pred, time_point=horizon_months)
            absolute_calibration_rows.append({'Model': model_name, **oe_metrics})
            primary_metrics_df.loc[primary_metrics_df['Model'] == model_name, 'O_E_120m'] = oe_metrics['calibration_oe_ratio']
        except Exception as exc:
            LOGGER.exception('Aggregate O:E calibration failed for %s: %s', model_name, exc)
            absolute_calibration_rows.append({'Model': model_name, 'calibration_observed_km_risk': np.nan, 'calibration_observed_km_risk_ci_low': np.nan, 'calibration_observed_km_risk_ci_high': np.nan, 'calibration_mean_predicted_risk': np.nan, 'calibration_observed_events_km': np.nan, 'calibration_expected_events': np.nan, 'calibration_oe_ratio': np.nan, 'calibration_oe_ratio_ci_low': np.nan, 'calibration_oe_ratio_ci_high': np.nan, 'reason_unavailable': f'{type(exc).__name__}: {exc}'})
        try:
            calibration_df = grouped_calibration_table(raw_test_df, mean_pred, time_point=horizon_months, n_groups=calibration_groups)
            calibration_df.insert(0, 'model', model_name)
            calibration_df.insert(1, 'horizon_months', float(horizon_months))
            calibration_tables[model_name] = calibration_df
            calibration_df.to_csv(os.path.join(tables_dir, f'calibration_points_fixed_{model_name}.csv'), index=False)
            grouped_completeness_rows.append({'model': model_name, 'n_groups_requested': int(calibration_groups), 'n_groups_produced': int(len(calibration_df)), 'complete': bool(len(calibration_df) == calibration_groups), 'error': ''})
        except Exception as exc:
            LOGGER.exception('Grouped calibration failed for %s: %s', model_name, exc)
            grouped_completeness_rows.append({'model': model_name, 'n_groups_requested': int(calibration_groups), 'n_groups_produced': 0, 'complete': False, 'error': f'{type(exc).__name__}: {exc}'})
        try:
            flexible_df = flexible_calibration_curve_with_bootstrap(raw_test_df, mean_pred, time_point=horizon_months, n_bootstrap=calibration_bootstrap_reps, random_state=seed + 50000 + MODEL_ORDER.index(model_name))
            flexible_df.insert(0, 'model', model_name)
            flexible_calibration_tables[model_name] = flexible_df
            flexible_df.to_csv(os.path.join(tables_dir, f'flexible_calibration_fixed_{model_name}.csv'), index=False)
            n_success = int(flexible_df['n_bootstrap_successful'].iloc[0])
            minimum_successful = int(math.ceil(0.9 * calibration_bootstrap_reps))
            bootstrap_completeness_rows.append({'model': model_name, 'n_bootstrap_requested': int(calibration_bootstrap_reps), 'n_bootstrap_successful': n_success, 'minimum_successful_required': minimum_successful, 'complete': bool(n_success >= minimum_successful), 'error': ''})
            if n_success < calibration_bootstrap_reps:
                LOGGER.warning('%s flexible calibration completed %d/%d bootstrap resamples.', model_name, n_success, calibration_bootstrap_reps)
        except Exception as exc:
            LOGGER.exception('Flexible calibration failed for %s: %s', model_name, exc)
            bootstrap_completeness_rows.append({'model': model_name, 'n_bootstrap_requested': int(calibration_bootstrap_reps), 'n_bootstrap_successful': 0, 'complete': False, 'error': f'{type(exc).__name__}: {exc}'})
    completed_absolute_models = {row['Model'] for row in absolute_calibration_rows}
    for missing_model in [m for m in CALIBRATION_MODELS if m not in completed_absolute_models]:
        absolute_calibration_rows.append({'Model': missing_model, 'calibration_observed_km_risk': np.nan, 'calibration_observed_km_risk_ci_low': np.nan, 'calibration_observed_km_risk_ci_high': np.nan, 'calibration_mean_predicted_risk': np.nan, 'calibration_observed_events_km': np.nan, 'calibration_expected_events': np.nan, 'calibration_oe_ratio': np.nan, 'calibration_oe_ratio_ci_low': np.nan, 'calibration_oe_ratio_ci_high': np.nan, 'reason_unavailable': 'No valid absolute-risk predictions were produced.'})
        bootstrap_completeness_rows.append({'model': missing_model, 'n_bootstrap_requested': int(calibration_bootstrap_reps), 'n_bootstrap_successful': 0, 'complete': False, 'error': 'No valid absolute-risk predictions were produced.'})
        grouped_completeness_rows.append({'model': missing_model, 'n_groups_requested': int(calibration_groups), 'n_groups_produced': 0, 'complete': False, 'error': 'No valid absolute-risk predictions were produced.'})
    absolute_calibration_rows.append({'Model': 'SVM', 'calibration_observed_km_risk': np.nan, 'calibration_observed_km_risk_ci_low': np.nan, 'calibration_observed_km_risk_ci_high': np.nan, 'calibration_mean_predicted_risk': np.nan, 'calibration_observed_events_km': np.nan, 'calibration_expected_events': np.nan, 'calibration_oe_ratio': np.nan, 'calibration_oe_ratio_ci_low': np.nan, 'calibration_oe_ratio_ci_high': np.nan, 'reason_unavailable': 'Survival SVM provides risk scores but not survival probabilities.'})
    bootstrap_completeness_rows.append({'model': 'SVM', 'n_bootstrap_requested': 0, 'n_bootstrap_successful': 0, 'complete': False, 'error': 'Not applicable: Survival SVM does not provide survival probabilities.'})
    grouped_completeness_rows.append({'model': 'SVM', 'n_groups_requested': 0, 'n_groups_produced': 0, 'complete': False, 'error': 'Not applicable: Survival SVM does not provide survival probabilities.'})
    absolute_calibration_df = pd.DataFrame(absolute_calibration_rows)
    absolute_calibration_df = absolute_calibration_df.set_index('Model').loc[MODEL_ORDER].reset_index()
    expected_absolute_models = set(CALIBRATION_MODELS)
    if set(calibration_tables) != expected_absolute_models:
        missing = sorted(expected_absolute_models.difference(calibration_tables))
        raise RuntimeError(f'Grouped fixed-split calibration was incomplete for: {missing}')
    if set(flexible_calibration_tables) != expected_absolute_models:
        missing = sorted(expected_absolute_models.difference(flexible_calibration_tables))
        raise RuntimeError(f'Flexible fixed-split calibration was incomplete for: {missing}')
    absolute_non_svm = absolute_calibration_df.loc[absolute_calibration_df['Model'].isin(CALIBRATION_MODELS)]
    required_absolute_columns = ['calibration_observed_km_risk', 'calibration_mean_predicted_risk', 'calibration_oe_ratio']
    if not np.isfinite(absolute_non_svm[required_absolute_columns].to_numpy(dtype=float)).all():
        raise RuntimeError('Fixed-split absolute calibration contained a non-finite required value.')
    for row in grouped_completeness_rows:
        if row['model'] != 'SVM' and (not bool(row['complete'])):
            raise RuntimeError(f"Grouped calibration was incomplete for {row['model']}.")
    for row in bootstrap_completeness_rows:
        if row['model'] == 'SVM':
            continue
        requested = int(row['n_bootstrap_requested'])
        successful = int(row['n_bootstrap_successful'])
        minimum_successful = int(math.ceil(0.9 * requested))
        if requested < 1 or successful < minimum_successful or str(row['error']).strip():
            raise RuntimeError(f"Flexible calibration bootstrap was incomplete for {row['model']}: {successful}/{requested}; error={row['error']!r}.")
    absolute_calibration_df.to_csv(os.path.join(tables_dir, 'absolute_calibration_fixed_120m.csv'), index=False)
    pd.DataFrame(bootstrap_completeness_rows).to_csv(os.path.join(meta_dir, 'flexible_calibration_bootstrap_completeness.csv'), index=False)
    pd.DataFrame(grouped_completeness_rows).to_csv(os.path.join(meta_dir, 'grouped_calibration_completeness.csv'), index=False)
    if calibration_tables:
        pd.concat(calibration_tables.values(), ignore_index=True).to_csv(os.path.join(tables_dir, 'grouped_calibration_fixed_all_models.csv'), index=False)
    if flexible_calibration_tables:
        pd.concat(flexible_calibration_tables.values(), ignore_index=True).to_csv(os.path.join(tables_dir, 'flexible_calibration_fixed_all_models.csv'), index=False)
    primary_metrics_df.to_csv(os.path.join(tables_dir, 'supplementary_table_s5_fixed_split_metrics.csv'), index=False)
    fixed_source_note = 'Illustrative fixed 80/20 test set; predictions are averaged across 10 fixed-split imputations.'
    make_combined_calibration_figure(calibration_tables, os.path.join(figs_dir, 'supplementary_fixed_split_grouped_calibration_all_models.png'), horizon_months=horizon_months, source_note=fixed_source_note)
    make_combined_flexible_calibration_figure(flexible_calibration_tables, os.path.join(figs_dir, 'supplementary_fixed_split_flexible_calibration_all_models.png'), horizon_months=horizon_months, source_note=fixed_source_note + ' Bootstrap bands are conditional on these averaged predictions.')
    shap_tables: List[pd.DataFrame] = []
    if make_shap and shap_reference_train is not None and (shap_reference_test is not None):
        for model_name in SHAP_MODELS:
            fitted_model = shap_reference_models.get(model_name)
            if fitted_model is None:
                continue
            shap_table = make_shap_outputs(model_name=model_name, fitted_model=fitted_model, train_df=shap_reference_train, test_df=shap_reference_test, out_dir=figs_dir, make_force_plot=model_name in ['CoxPH', 'DeepSurv'], max_background=shap_max_background, max_explain=shap_max_explain)
            if shap_table is not None:
                shap_table.insert(0, 'Model', model_name)
                shap_tables.append(shap_table)
        shap_pngs = [os.path.join(figs_dir, f'shap_summary_{m}.png') for m in SHAP_MODELS]
        combine_shap_summary_images(shap_pngs[:4], SHAP_MODELS[:4], os.path.join(figs_dir, 'figure2_shap_summary_models_a_to_d.png'), ncols=2)
        combine_shap_summary_images(shap_pngs[4:], SHAP_MODELS[4:], os.path.join(figs_dir, 'figure3_shap_summary_models_e_to_h.png'), ncols=2)
        combine_force_plots([os.path.join(figs_dir, 'shap_force_CoxPH.png'), os.path.join(figs_dir, 'shap_force_DeepSurv.png')], ['CoxPH', 'DeepSurv'], os.path.join(figs_dir, 'supplementary_figure_s1_local_shap_force.png'))
    if make_shap:
        completed_shap_models = {str(table['Model'].iloc[0]) for table in shap_tables if not table.empty and 'Model' in table.columns}
        if completed_shap_models != set(SHAP_MODELS):
            missing = sorted(set(SHAP_MODELS).difference(completed_shap_models))
            raise RuntimeError(f'Fixed-split SHAP outputs were incomplete for: {missing}')
        required_shap_images = [*(os.path.join(figs_dir, f'shap_summary_{model_name}.png') for model_name in SHAP_MODELS), os.path.join(figs_dir, 'shap_force_CoxPH.png'), os.path.join(figs_dir, 'shap_force_DeepSurv.png'), os.path.join(figs_dir, 'figure2_shap_summary_models_a_to_d.png'), os.path.join(figs_dir, 'figure3_shap_summary_models_e_to_h.png'), os.path.join(figs_dir, 'supplementary_figure_s1_local_shap_force.png')]
        missing_images = [path for path in required_shap_images if not os.path.exists(path)]
        if missing_images:
            raise RuntimeError(f'Fixed-split SHAP image outputs were missing: {missing_images}')
        pd.concat(shap_tables, axis=0, ignore_index=True).to_csv(os.path.join(tables_dir, 'shap_importance_fixed_split_all_models.csv'), index=False)
    tuning_df = pd.DataFrame(tuning_rows)
    tuning_df.to_csv(os.path.join(meta_dir, 'tuning_records_fixed_split.csv'), index=False)
    return (table1, pooled_cox_table, primary_metrics_df, tuning_df)

def _outer_model_checkpoint_paths(checkpoint_dir: str, repetition: int, fold: int, model_name: str) -> Dict[str, str]:
    model_dir = ensure_dir(os.path.join(checkpoint_dir, 'partial_models'))
    prefix = f'repetition_{repetition:02d}_fold_{fold:02d}_{model_name}'
    return {'metrics': os.path.join(model_dir, f'{prefix}_metrics.csv'), 'tuning': os.path.join(model_dir, f'{prefix}_tuning.csv'), 'predictions': os.path.join(model_dir, f'{prefix}_predictions.csv'), 'marker': os.path.join(model_dir, f'{prefix}_complete.json')}
OUTER_FOLD_REQUIRED_METRICS = ['harrell_c', 'uno_c_tau', 'brier_score_at_horizon', 'integrated_brier_score']

def _outer_required_metrics(model_name: str) -> List[str]:
    if model_name == 'SVM':
        return ['harrell_c', 'uno_c_tau']
    return list(OUTER_FOLD_REQUIRED_METRICS)

def require_complete_outer_fold_metrics(model_name: str, metrics: Dict[str, Any], predicted_event_probability: Optional[np.ndarray], expected_n: int, context: str) -> None:
    missing = [name for name in _outer_required_metrics(model_name) if not np.isfinite(float(metrics.get(name, np.nan)))]
    if missing:
        raise RuntimeError(f'{context}: {model_name} has non-finite required metrics: {missing}.')
    if model_name == 'SVM':
        if predicted_event_probability is not None:
            raise RuntimeError(f'{context}: SVM unexpectedly produced absolute-risk predictions.')
        return
    if predicted_event_probability is None:
        raise RuntimeError(f'{context}: {model_name} did not produce absolute-risk predictions.')
    probabilities = np.asarray(predicted_event_probability, dtype=float).reshape(-1)
    if probabilities.shape != (int(expected_n),) or not np.isfinite(probabilities).all() or np.any((probabilities < -1e-08) | (probabilities > 1.0 + 1e-08)):
        raise RuntimeError(f'{context}: {model_name} produced invalid absolute-risk predictions.')

def validate_outer_model_checkpoint_frames(metrics_df: pd.DataFrame, tuning_df: pd.DataFrame, predictions_df: pd.DataFrame, repetition: int, fold: int, model_name: str, expected_test_ids: Sequence[Any]) -> None:
    if len(metrics_df) != 1 or len(tuning_df) != 1:
        raise ValueError('A partial model checkpoint must have one metric and tuning row.')
    for frame_name, frame in [('metrics', metrics_df), ('tuning', tuning_df)]:
        required = {'repetition', 'fold', 'model'}
        if not required.issubset(frame.columns):
            raise ValueError(f'Partial {frame_name} checkpoint lacks identity columns.')
        row = frame.iloc[0]
        if int(row['repetition']) != repetition or int(row['fold']) != fold or str(row['model']) != model_name:
            raise ValueError(f'Partial {frame_name} checkpoint identity changed.')
    if 'n_test' not in metrics_df.columns or int(metrics_df.iloc[0]['n_test']) != len(expected_test_ids):
        raise ValueError('Partial model checkpoint has the wrong held-out sample size.')
    for metric in _outer_required_metrics(model_name):
        if metric not in metrics_df.columns or not np.isfinite(float(metrics_df.iloc[0][metric])):
            raise ValueError(f'Partial model checkpoint has non-finite {metric}.')
        validate_metric_bounds(metric, metrics_df.iloc[0][metric], f'Partial repetition {repetition} fold {fold} {model_name}')
    if float(metrics_df.iloc[0].get('tau_months', np.nan)) != PRESPECIFIED_HORIZON_MONTHS:
        raise ValueError('Partial model checkpoint has the wrong Uno horizon.')
    if model_name != 'SVM' and float(metrics_df.iloc[0].get('brier_horizon_months', np.nan)) != PRESPECIFIED_HORIZON_MONTHS:
        raise ValueError('Partial model checkpoint has the wrong Brier horizon.')
    if not {'tuning_method', 'best_params_json', *TUNING_SELECTION_COLUMNS}.issubset(tuning_df.columns):
        raise ValueError('Partial tuning checkpoint lacks required fields.')
    if not isinstance(json.loads(str(tuning_df.iloc[0]['best_params_json'])), dict):
        raise ValueError('Partial tuning parameters are not a JSON object.')
    expected_method = main_tuning_method(model_name)
    if str(tuning_df.iloc[0]['tuning_method']) != expected_method:
        raise ValueError('Partial model tuning method is inconsistent with the model.')
    validate_tuning_selection_record(tuning_df.iloc[0], model_name, f'Partial fixed-split {model_name}')
    required_predictions = {STUDY_ID_COL, 'predicted_event_probability'}
    if not required_predictions.issubset(predictions_df.columns):
        raise ValueError('Partial model predictions lack required columns.')
    expected_ids = pd.Series(expected_test_ids).reset_index(drop=True)
    observed_ids = predictions_df[STUDY_ID_COL].reset_index(drop=True)
    if len(predictions_df) != len(expected_ids) or not np.array_equal(observed_ids.astype(str).to_numpy(), expected_ids.astype(str).to_numpy()):
        raise ValueError('Partial model prediction study IDs changed order or membership.')
    probabilities = pd.to_numeric(predictions_df['predicted_event_probability'], errors='coerce').to_numpy(float)
    if model_name == 'SVM':
        if not np.isnan(probabilities).all():
            raise ValueError('SVM partial checkpoint must use the no-absolute-risk sentinel.')
    elif not np.isfinite(probabilities).all() or np.any((probabilities < -1e-08) | (probabilities > 1.0 + 1e-08)):
        raise ValueError('Partial model checkpoint has invalid probabilities.')

def write_outer_model_checkpoint(checkpoint_dir: str, manifest_sha256: str, repetition: int, fold: int, model_name: str, metrics_df: pd.DataFrame, tuning_df: pd.DataFrame, predictions_df: pd.DataFrame, expected_test_ids: Sequence[Any]) -> None:
    validate_outer_model_checkpoint_frames(metrics_df, tuning_df, predictions_df, repetition, fold, model_name, expected_test_ids)
    paths = _outer_model_checkpoint_paths(checkpoint_dir, repetition, fold, model_name)
    frames = {'metrics': metrics_df, 'tuning': tuning_df, 'predictions': predictions_df}
    for key, frame in frames.items():
        atomic_write_dataframe_csv(frame, paths[key])
    marker = {'checkpoint_schema_version': CHECKPOINT_SCHEMA_VERSION, 'complete': True, 'manifest_sha256': manifest_sha256, 'repetition': int(repetition), 'fold': int(fold), 'model': model_name, 'absolute_risk_available': bool(model_name != 'SVM'), 'files': {key: {'name': os.path.basename(paths[key]), 'sha256': file_sha256(paths[key]), 'n_rows': int(len(frame))} for key, frame in frames.items()}}
    atomic_write_json(marker, paths['marker'])

def load_outer_model_checkpoint(checkpoint_dir: str, manifest_sha256: str, repetition: int, fold: int, model_name: str, expected_test_ids: Sequence[Any]) -> Optional[Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]]:
    paths = _outer_model_checkpoint_paths(checkpoint_dir, repetition, fold, model_name)
    if not os.path.isfile(paths['marker']):
        return None
    try:
        with open(paths['marker'], 'r', encoding='utf-8') as handle:
            marker = json.load(handle)
        if marker.get('checkpoint_schema_version') != CHECKPOINT_SCHEMA_VERSION or marker.get('complete') is not True or marker.get('manifest_sha256') != manifest_sha256 or (int(marker.get('repetition', -1)) != repetition) or (int(marker.get('fold', -1)) != fold) or (marker.get('model') != model_name) or (bool(marker.get('absolute_risk_available')) != (model_name != 'SVM')):
            raise ValueError('Partial model completion marker did not validate.')
        loaded: Dict[str, pd.DataFrame] = {}
        for key in ['metrics', 'tuning', 'predictions']:
            meta = marker['files'][key]
            if meta.get('name') != os.path.basename(paths[key]):
                raise ValueError(f'Partial model {key} filename changed.')
            if not os.path.isfile(paths[key]) or file_sha256(paths[key]) != meta.get('sha256'):
                raise ValueError(f'Partial model {key} file is missing or changed.')
            frame = pd.read_csv(paths[key], float_precision='round_trip')
            if len(frame) != int(meta.get('n_rows', -1)):
                raise ValueError(f'Partial model {key} row count changed.')
            loaded[key] = frame
        validate_outer_model_checkpoint_frames(loaded['metrics'], loaded['tuning'], loaded['predictions'], repetition, fold, model_name, expected_test_ids)
        return (loaded['metrics'], loaded['tuning'], loaded['predictions'])
    except Exception as exc:
        LOGGER.warning('Repetition %d fold %d partial %s checkpoint is invalid and will be refit: %s', repetition, fold, model_name, exc)
        return None
REPEATED_SPLIT_CHECKPOINT_KIND = 'repeated_200_split_v2_ibs_tuning'

def _repeated_split_checkpoint_paths(checkpoint_dir: str, split_id: int) -> Dict[str, str]:
    prefix = f'split_{int(split_id):03d}'
    return {'metrics': os.path.join(checkpoint_dir, f'{prefix}_metrics.csv'), 'tuning': os.path.join(checkpoint_dir, f'{prefix}_tuning.csv'), 'predictions': os.path.join(checkpoint_dir, f'{prefix}_held_out_predictions.csv'), 'marker': os.path.join(checkpoint_dir, f'{prefix}_complete.json')}

def _repeated_split_model_checkpoint_paths(checkpoint_dir: str, split_id: int, model_name: str) -> Dict[str, str]:
    model_dir = ensure_dir(os.path.join(checkpoint_dir, 'partial_models'))
    prefix = f'split_{int(split_id):03d}_{model_name}'
    return {'metrics': os.path.join(model_dir, f'{prefix}_metrics.csv'), 'tuning': os.path.join(model_dir, f'{prefix}_tuning.csv'), 'predictions': os.path.join(model_dir, f'{prefix}_predictions.csv'), 'marker': os.path.join(model_dir, f'{prefix}_complete.json')}

def validate_repeated_split_checkpoint_frames(metrics_df: pd.DataFrame, tuning_df: pd.DataFrame, predictions_df: pd.DataFrame, split_id: int, expected_test_ids: Sequence[Any]) -> None:
    if not 1 <= int(split_id) <= PRESPECIFIED_N_SPLITS:
        raise ValueError('Repeated split_id is outside 1-200.')
    metric_required = {'split_id', 'model', 'n_train', 'n_test', 'train_events', 'test_events', *OUTER_FOLD_REQUIRED_METRICS}
    tuning_required = {'split_id', 'model', 'tuning_method', 'best_params_json', *TUNING_SELECTION_COLUMNS, 'split_base_seed', 'outer_imputation_seed', 'model_seed', 'classical_inner_imputation_seed_base', 'neural_development_imputation_seed'}
    prediction_required = {'split_id', STUDY_ID_COL, *CALIBRATION_MODELS}
    for frame_name, frame, required in [('metrics', metrics_df, metric_required), ('tuning', tuning_df, tuning_required), ('predictions', predictions_df, prediction_required)]:
        if not required.issubset(frame.columns):
            raise ValueError(f'Repeated-split {frame_name} lack columns: {sorted(required - set(frame.columns))}')
        if not (pd.to_numeric(frame['split_id'], errors='raise').astype(int) == int(split_id)).all():
            raise ValueError(f'Repeated-split {frame_name} have the wrong split_id.')
    if len(metrics_df) != len(MODEL_ORDER) or len(tuning_df) != len(MODEL_ORDER):
        raise ValueError('Repeated-split metrics and tuning must contain one row per model.')
    if set(metrics_df['model'].astype(str)) != set(MODEL_ORDER):
        raise ValueError('Repeated-split metric model set is incomplete.')
    if set(tuning_df['model'].astype(str)) != set(MODEL_ORDER):
        raise ValueError('Repeated-split tuning model set is incomplete.')
    if metrics_df.duplicated(['split_id', 'model']).any():
        raise ValueError('Repeated-split metrics contain duplicate model keys.')
    if tuning_df.duplicated(['split_id', 'model']).any():
        raise ValueError('Repeated-split tuning contains duplicate model keys.')
    expected_ids = pd.Series(expected_test_ids).reset_index(drop=True)
    observed_ids = predictions_df[STUDY_ID_COL].reset_index(drop=True)
    if len(predictions_df) != len(expected_ids) or not np.array_equal(observed_ids.astype(str).to_numpy(), expected_ids.astype(str).to_numpy()):
        raise ValueError('Repeated-split prediction IDs changed order or membership.')
    if observed_ids.duplicated().any():
        raise ValueError('Repeated-split predictions contain a duplicate study ID.')
    if not (pd.to_numeric(metrics_df['n_test'], errors='raise').astype(int) == len(expected_ids)).all():
        raise ValueError('Repeated-split metrics contain the wrong test-set size.')
    for column, expected in [('n_train', 992), ('n_test', 249), ('train_events', 92), ('test_events', 24)]:
        if not (pd.to_numeric(metrics_df[column], errors='raise').astype(int) == expected).all():
            raise ValueError(f'Repeated-split metrics contain the wrong {column}.')
    for model_name, rows in metrics_df.groupby('model', sort=False):
        for metric in _outer_required_metrics(str(model_name)):
            value = float(rows.iloc[0][metric])
            if not np.isfinite(value):
                raise ValueError(f'Repeated-split checkpoint has non-finite {metric} for {model_name}.')
            validate_metric_bounds(metric, value, f'Split {split_id} {model_name}')
        row = rows.iloc[0]
        if float(row.get('tau_months', np.nan)) != PRESPECIFIED_HORIZON_MONTHS:
            raise ValueError('Repeated-split Uno horizon is not exactly 120 months.')
        if str(model_name) != 'SVM' and float(row.get('brier_horizon_months', np.nan)) != PRESPECIFIED_HORIZON_MONTHS:
            raise ValueError('Repeated-split Brier horizon is not exactly 120 months.')
    for model_name in CALIBRATION_MODELS:
        probabilities = pd.to_numeric(predictions_df[model_name], errors='coerce').to_numpy(float)
        if not np.isfinite(probabilities).all() or np.any((probabilities < -1e-08) | (probabilities > 1.0 + 1e-08)):
            raise ValueError(f'Repeated-split predictions are invalid for {model_name}.')
    for _, tuning_row in tuning_df.iterrows():
        model_name = str(tuning_row['model'])
        if str(tuning_row['tuning_method']) != main_tuning_method(model_name):
            raise ValueError(f'{model_name} records the wrong tuning method.')
        validate_tuning_selection_record(tuning_row, model_name, f'Repeated split {split_id} {model_name}')
    split_base_seed = PRESPECIFIED_SPLITTER_SEED + 100000 * int(split_id)
    if not (pd.to_numeric(tuning_df['split_base_seed'], errors='raise').astype(int) == split_base_seed).all():
        raise ValueError('Repeated-split tuning has the wrong split_base_seed.')
    if not (pd.to_numeric(tuning_df['outer_imputation_seed'], errors='raise').astype(int) == split_base_seed).all():
        raise ValueError('Repeated-split tuning has the wrong outer imputation seed.')
    if not (pd.to_numeric(tuning_df['model_seed'], errors='raise').astype(int) == split_base_seed + 1).all():
        raise ValueError('Repeated-split tuning has the wrong model seed.')
    if not (pd.to_numeric(tuning_df['classical_inner_imputation_seed_base'], errors='raise').astype(int) == split_base_seed + 100).all():
        raise ValueError('Repeated-split tuning has the wrong classical inner imputation seed base.')
    if not (pd.to_numeric(tuning_df['neural_development_imputation_seed'], errors='raise').astype(int) == split_base_seed + 500).all():
        raise ValueError('Repeated-split tuning has the wrong neural development imputation seed.')
    for value in tuning_df['best_params_json'].astype(str):
        if not isinstance(json.loads(value), dict):
            raise ValueError('Repeated-split best_params_json is not an object.')

def validate_repeated_split_model_checkpoint_frames(metrics_df: pd.DataFrame, tuning_df: pd.DataFrame, predictions_df: pd.DataFrame, split_id: int, model_name: str, expected_test_ids: Sequence[Any]) -> None:
    if len(metrics_df) != 1 or len(tuning_df) != 1:
        raise ValueError('A repeated-split partial model checkpoint needs one metric and tuning row.')
    for frame_name, frame in [('metrics', metrics_df), ('tuning', tuning_df)]:
        if not {'split_id', 'model'}.issubset(frame.columns):
            raise ValueError(f'Partial repeated-split {frame_name} lack identity columns.')
        row = frame.iloc[0]
        if int(row['split_id']) != int(split_id) or str(row['model']) != model_name:
            raise ValueError(f'Partial repeated-split {frame_name} identity changed.')
    metric_row = metrics_df.iloc[0]
    expected_counts = {'n_train': 992, 'n_test': 249, 'train_events': 92, 'test_events': 24}
    for column, expected in expected_counts.items():
        if int(metric_row.get(column, -1)) != expected:
            raise ValueError(f'Partial repeated-split model has the wrong {column}.')
    if len(expected_test_ids) != 249:
        raise ValueError('Partial repeated-split model did not receive 249 expected test IDs.')
    for metric in _outer_required_metrics(model_name):
        value = float(metrics_df.iloc[0].get(metric, np.nan))
        if not np.isfinite(value):
            raise ValueError(f'Partial repeated-split model has non-finite {metric}.')
        validate_metric_bounds(metric, value, f'Partial split {split_id} {model_name}')
    if float(metric_row.get('tau_months', np.nan)) != PRESPECIFIED_HORIZON_MONTHS:
        raise ValueError('Partial repeated-split model has the wrong Uno horizon.')
    if model_name != 'SVM' and float(metric_row.get('brier_horizon_months', np.nan)) != PRESPECIFIED_HORIZON_MONTHS:
        raise ValueError('Partial repeated-split model has the wrong Brier horizon.')
    expected_method = main_tuning_method(model_name)
    if str(tuning_df.iloc[0].get('tuning_method', '')) != expected_method:
        raise ValueError('Partial repeated-split tuning method is inconsistent.')
    validate_tuning_selection_record(tuning_df.iloc[0], model_name, f'Partial repeated split {split_id} {model_name}')
    split_base_seed = PRESPECIFIED_SPLITTER_SEED + 100000 * int(split_id)
    expected_seeds = {'split_base_seed': split_base_seed, 'outer_imputation_seed': split_base_seed, 'model_seed': split_base_seed + 1, 'classical_inner_imputation_seed_base': split_base_seed + 100, 'neural_development_imputation_seed': split_base_seed + 500}
    for column, expected in expected_seeds.items():
        if int(tuning_df.iloc[0].get(column, -1)) != expected:
            raise ValueError(f'Partial repeated-split tuning has the wrong {column}.')
    if not isinstance(json.loads(str(tuning_df.iloc[0].get('best_params_json', ''))), dict):
        raise ValueError('Partial repeated-split tuning parameters are not an object.')
    if not {STUDY_ID_COL, 'predicted_event_probability'}.issubset(predictions_df.columns):
        raise ValueError('Partial repeated-split predictions lack required columns.')
    expected_ids = pd.Series(expected_test_ids).reset_index(drop=True)
    observed_ids = predictions_df[STUDY_ID_COL].reset_index(drop=True)
    if len(predictions_df) != len(expected_ids) or not np.array_equal(observed_ids.astype(str).to_numpy(), expected_ids.astype(str).to_numpy()):
        raise ValueError('Partial repeated-split prediction IDs changed order or membership.')
    probabilities = pd.to_numeric(predictions_df['predicted_event_probability'], errors='coerce').to_numpy(float)
    if model_name == 'SVM':
        if not np.isnan(probabilities).all():
            raise ValueError('SVM partial repeated-split checkpoint must use NaN probabilities.')
    elif not np.isfinite(probabilities).all() or np.any((probabilities < -1e-08) | (probabilities > 1.0 + 1e-08)):
        raise ValueError('Partial repeated-split checkpoint has invalid probabilities.')

def write_repeated_split_model_checkpoint(checkpoint_dir: str, manifest_sha256: str, split_id: int, model_name: str, metrics_df: pd.DataFrame, tuning_df: pd.DataFrame, predictions_df: pd.DataFrame, expected_test_ids: Sequence[Any]) -> None:
    validate_repeated_split_model_checkpoint_frames(metrics_df, tuning_df, predictions_df, split_id, model_name, expected_test_ids)
    paths = _repeated_split_model_checkpoint_paths(checkpoint_dir, split_id, model_name)
    frames = {'metrics': metrics_df, 'tuning': tuning_df, 'predictions': predictions_df}
    for key, frame in frames.items():
        atomic_write_dataframe_csv(frame, paths[key])
    marker = {'checkpoint_schema_version': CHECKPOINT_SCHEMA_VERSION, 'checkpoint_kind': REPEATED_SPLIT_CHECKPOINT_KIND, 'complete': True, 'manifest_sha256': manifest_sha256, 'split_id': int(split_id), 'model': model_name, 'files': {key: {'name': os.path.basename(paths[key]), 'sha256': file_sha256(paths[key]), 'n_rows': int(len(frame))} for key, frame in frames.items()}}
    atomic_write_json(marker, paths['marker'])

def load_repeated_split_model_checkpoint(checkpoint_dir: str, manifest_sha256: str, split_id: int, model_name: str, expected_test_ids: Sequence[Any]) -> Optional[Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]]:
    paths = _repeated_split_model_checkpoint_paths(checkpoint_dir, split_id, model_name)
    if not os.path.isfile(paths['marker']):
        return None
    try:
        with open(paths['marker'], 'r', encoding='utf-8') as handle:
            marker = json.load(handle)
        if marker.get('checkpoint_schema_version') != CHECKPOINT_SCHEMA_VERSION or marker.get('checkpoint_kind') != REPEATED_SPLIT_CHECKPOINT_KIND or marker.get('complete') is not True or (marker.get('manifest_sha256') != manifest_sha256) or (int(marker.get('split_id', -1)) != int(split_id)) or (marker.get('model') != model_name):
            raise ValueError('Partial repeated-split marker did not validate.')
        loaded: Dict[str, pd.DataFrame] = {}
        for key in ['metrics', 'tuning', 'predictions']:
            metadata = marker['files'][key]
            if metadata.get('name') != os.path.basename(paths[key]):
                raise ValueError(f'Partial repeated-split {key} filename changed.')
            if not os.path.isfile(paths[key]) or file_sha256(paths[key]) != metadata.get('sha256'):
                raise ValueError(f'Partial repeated-split {key} file is missing or changed.')
            frame = pd.read_csv(paths[key], float_precision='round_trip')
            if len(frame) != int(metadata.get('n_rows', -1)):
                raise ValueError(f'Partial repeated-split {key} row count changed.')
            loaded[key] = frame
        validate_repeated_split_model_checkpoint_frames(loaded['metrics'], loaded['tuning'], loaded['predictions'], split_id, model_name, expected_test_ids)
        return (loaded['metrics'], loaded['tuning'], loaded['predictions'])
    except Exception as exc:
        LOGGER.warning('Split %d partial %s checkpoint is invalid and will be refit: %s', split_id, model_name, exc)
        return None

def write_repeated_split_checkpoint(checkpoint_dir: str, manifest_sha256: str, split_id: int, metrics_df: pd.DataFrame, tuning_df: pd.DataFrame, predictions_df: pd.DataFrame, expected_test_ids: Sequence[Any]) -> None:
    validate_repeated_split_checkpoint_frames(metrics_df, tuning_df, predictions_df, split_id, expected_test_ids)
    paths = _repeated_split_checkpoint_paths(checkpoint_dir, split_id)
    frames = {'metrics': metrics_df, 'tuning': tuning_df, 'predictions': predictions_df}
    for key, frame in frames.items():
        atomic_write_dataframe_csv(frame, paths[key])
    marker = {'checkpoint_schema_version': CHECKPOINT_SCHEMA_VERSION, 'checkpoint_kind': REPEATED_SPLIT_CHECKPOINT_KIND, 'complete': True, 'manifest_sha256': manifest_sha256, 'split_id': int(split_id), 'files': {key: {'name': os.path.basename(paths[key]), 'sha256': file_sha256(paths[key]), 'n_rows': int(len(frame))} for key, frame in frames.items()}}
    atomic_write_json(marker, paths['marker'])

def load_repeated_split_checkpoint(checkpoint_dir: str, manifest_sha256: str, split_id: int, expected_test_ids: Sequence[Any]) -> Optional[Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]]:
    paths = _repeated_split_checkpoint_paths(checkpoint_dir, split_id)
    if not os.path.isfile(paths['marker']):
        return None
    try:
        with open(paths['marker'], 'r', encoding='utf-8') as handle:
            marker = json.load(handle)
        if marker.get('checkpoint_schema_version') != CHECKPOINT_SCHEMA_VERSION or marker.get('checkpoint_kind') != REPEATED_SPLIT_CHECKPOINT_KIND or marker.get('complete') is not True or (marker.get('manifest_sha256') != manifest_sha256) or (int(marker.get('split_id', -1)) != int(split_id)):
            raise ValueError('Complete repeated-split marker did not validate.')
        loaded: Dict[str, pd.DataFrame] = {}
        for key in ['metrics', 'tuning', 'predictions']:
            metadata = marker['files'][key]
            if metadata.get('name') != os.path.basename(paths[key]):
                raise ValueError(f'Complete repeated-split {key} filename changed.')
            if not os.path.isfile(paths[key]) or file_sha256(paths[key]) != metadata.get('sha256'):
                raise ValueError(f'Complete repeated-split {key} file is missing or changed.')
            frame = pd.read_csv(paths[key], float_precision='round_trip')
            if len(frame) != int(metadata.get('n_rows', -1)):
                raise ValueError(f'Complete repeated-split {key} row count changed.')
            loaded[key] = frame
        validate_repeated_split_checkpoint_frames(loaded['metrics'], loaded['tuning'], loaded['predictions'], split_id, expected_test_ids)
        return (loaded['metrics'], loaded['tuning'], loaded['predictions'])
    except Exception as exc:
        LOGGER.warning('Split %d checkpoint is invalid and will be recomputed: %s', split_id, exc)
        return None

def validate_repeated_split_map(split_map: pd.DataFrame, df: pd.DataFrame, seed: int, n_splits: int) -> None:
    expected_columns = ['split_id', STUDY_ID_COL, 'role', 'split_seed', 'test_order']
    if list(split_map.columns) != expected_columns:
        raise ValueError(f'Repeated-split map columns must be exactly {expected_columns}.')
    if len(split_map) != int(n_splits) * 249:
        raise ValueError('Repeated-split map must contain exactly 49,800 test rows.')
    if set(split_map['role'].astype(str)) != {'test'}:
        raise ValueError('Repeated-split map may contain test rows only.')
    if not (pd.to_numeric(split_map['split_seed'], errors='raise').astype(int) == int(seed)).all():
        raise ValueError('Repeated-split map does not record the master splitter seed consistently.')
    if set(pd.to_numeric(split_map['split_id'], errors='raise').astype(int)) != set(range(1, int(n_splits) + 1)):
        raise ValueError('Repeated-split map does not contain split_id 1 through 200.')
    expected_ids = set(df[STUDY_ID_COL].tolist())
    coverage: Dict[Any, int] = {study_id: 0 for study_id in expected_ids}
    test_sets: set[Tuple[str, ...]] = set()
    for split_id, rows in split_map.groupby('split_id', sort=True):
        if len(rows) != 249 or rows[STUDY_ID_COL].nunique() != 249:
            raise ValueError(f'Split {split_id} must contain 249 unique held-out IDs.')
        order = pd.to_numeric(rows['test_order'], errors='raise').astype(int).tolist()
        if order != list(range(1, 250)):
            raise ValueError(f'Split {split_id} test_order must preserve positions 1 through 249.')
        ids = rows[STUDY_ID_COL].tolist()
        if not set(ids).issubset(expected_ids):
            raise ValueError(f'Split {split_id} contains an unknown study ID.')
        identity = tuple(sorted(map(str, ids)))
        if identity in test_sets:
            raise ValueError(f'Split {split_id} duplicates an earlier held-out set.')
        test_sets.add(identity)
        for study_id in ids:
            coverage[study_id] += 1
    coverage_values = np.asarray(list(coverage.values()), dtype=int)
    if len(coverage) != EXPECTED_FILTERED_ROWS or np.any(coverage_values < 1):
        raise ValueError('Repeated-split map does not hold out every participant at least once.')
    if int(coverage_values.min()) != 21 or int(coverage_values.max()) != 57:
        raise ValueError('Repeated-split coverage changed from the validated 21-57 appearances per participant.')

def build_repeated_split_map(df: pd.DataFrame, seed: int, n_splits: int, test_size: float) -> pd.DataFrame:
    rows: List[Dict[str, Any]] = []
    first_test_idx: Optional[np.ndarray] = None
    for split_id, train_idx, test_idx in iter_repeated_splits(df, seed=seed, n_splits=n_splits, test_size=test_size):
        if np.intersect1d(train_idx, test_idx).size or len(train_idx) + len(test_idx) != len(df):
            raise ValueError(f'Split {split_id} does not cleanly partition the cohort.')
        if split_id == 1:
            first_test_idx = test_idx.copy()
        rows.extend(({'split_id': int(split_id), STUDY_ID_COL: study_id, 'role': 'test', 'split_seed': int(seed), 'test_order': int(order)} for order, study_id in enumerate(df.iloc[test_idx][STUDY_ID_COL].tolist(), start=1)))
    fixed_train_idx, fixed_test_idx = primary_split(df, seed=seed, test_size=test_size)
    del fixed_train_idx
    if first_test_idx is None or not np.array_equal(first_test_idx, fixed_test_idx):
        raise RuntimeError('Repeated split 1 does not match the fixed supplementary 80/20 split.')
    split_map = pd.DataFrame(rows, columns=['split_id', STUDY_ID_COL, 'role', 'split_seed', 'test_order'])
    validate_repeated_split_map(split_map, df=df, seed=seed, n_splits=n_splits)
    return split_map

def validate_repeated_held_out_predictions(predictions: pd.DataFrame, split_map: pd.DataFrame, df: pd.DataFrame, n_splits: int) -> None:
    required = {'split_id', STUDY_ID_COL, *CALIBRATION_MODELS}
    if not required.issubset(predictions.columns):
        raise ValueError('Repeated held-out predictions lack required columns.')
    if len(predictions) != int(n_splits) * 249:
        raise ValueError('Repeated held-out predictions must contain exactly 49,800 rows.')
    if predictions.duplicated(['split_id', STUDY_ID_COL]).any():
        raise ValueError('Repeated held-out predictions contain a duplicate split-participant key.')
    expected_keys = split_map[['split_id', STUDY_ID_COL]].reset_index(drop=True)
    observed_keys = predictions[['split_id', STUDY_ID_COL]].reset_index(drop=True)
    if not (np.array_equal(pd.to_numeric(observed_keys['split_id'], errors='raise').astype(int).to_numpy(), pd.to_numeric(expected_keys['split_id'], errors='raise').astype(int).to_numpy()) and np.array_equal(observed_keys[STUDY_ID_COL].astype(str).to_numpy(), expected_keys[STUDY_ID_COL].astype(str).to_numpy())):
        raise ValueError('Repeated held-out predictions do not preserve the split-map row order.')
    expected_ids = set(df[STUDY_ID_COL].tolist())
    if set(predictions[STUDY_ID_COL].tolist()) != expected_ids:
        raise ValueError('Repeated held-out predictions do not cover all 1,241 participants.')
    for model_name in CALIBRATION_MODELS:
        probabilities = pd.to_numeric(predictions[model_name], errors='coerce').to_numpy(float)
        if not np.isfinite(probabilities).all() or np.any((probabilities < -1e-08) | (probabilities > 1.0 + 1e-08)):
            raise ValueError(f'Repeated held-out predictions are invalid for {model_name}.')

def create_repeated_split_calibration_outputs(df: pd.DataFrame, held_out_predictions: pd.DataFrame, split_map: pd.DataFrame, n_splits: int, horizon_months: float, calibration_groups: int, bootstrap_reps: int, out_dir: str, seed: int) -> pd.DataFrame:
    validate_repeated_split_map(split_map, df=df, seed=seed, n_splits=n_splits)
    validate_repeated_held_out_predictions(held_out_predictions, split_map=split_map, df=df, n_splits=n_splits)
    tables_dir = ensure_dir(os.path.join(out_dir, 'tables'))
    meta_dir = ensure_dir(os.path.join(out_dir, 'metadata'))
    figures_dir = ensure_dir(os.path.join(out_dir, 'figures'))
    coverage = held_out_predictions.groupby(STUDY_ID_COL, as_index=False).size().rename(columns={'size': 'n_held_out_appearances'})
    if len(coverage) != EXPECTED_FILTERED_ROWS or int(coverage['n_held_out_appearances'].min()) != 21 or int(coverage['n_held_out_appearances'].max()) != 57:
        raise RuntimeError('Repeated-split calibration coverage must include all 1,241 participants 21-57 times.')
    atomic_write_dataframe_csv(coverage, os.path.join(meta_dir, 'repeated_200_held_out_coverage.csv'))
    mean_predictions = held_out_predictions.groupby(STUDY_ID_COL, as_index=False)[CALIBRATION_MODELS].mean()
    calibration_data = df[[STUDY_ID_COL, TIME_COL, EVENT_COL]].merge(coverage, on=STUDY_ID_COL, how='left', validate='one_to_one').merge(mean_predictions, on=STUDY_ID_COL, how='left', validate='one_to_one')
    if len(calibration_data) != EXPECTED_FILTERED_ROWS or calibration_data.isna().any().any():
        raise RuntimeError('Participant-level repeated-split calibration data are incomplete.')
    atomic_write_dataframe_csv(calibration_data, os.path.join(tables_dir, 'repeated_200_mean_held_out_risk_per_participant.csv'))
    grouped_tables: Dict[str, pd.DataFrame] = {}
    flexible_tables: Dict[str, pd.DataFrame] = {}
    summary_rows: List[Dict[str, Any]] = []
    grouped_rows: List[pd.DataFrame] = []
    flexible_rows: List[pd.DataFrame] = []
    for model_name in CALIBRATION_MODELS:
        probabilities = calibration_data[model_name].to_numpy(dtype=float)
        oe = calibration_oe_from_probs(calibration_data, probabilities, time_point=horizon_months)
        slope = calibration_slope_from_probs(calibration_data, probabilities, time_point=horizon_months)
        if not np.isfinite([float(oe['calibration_oe_ratio']), float(slope['calibration_slope'])]).all():
            raise RuntimeError(f'Participant-level calibration is incomplete for {model_name}.')
        summary_rows.append({'model': model_name, 'n_participants': int(len(calibration_data)), 'total_held_out_prediction_rows': int(len(held_out_predictions)), 'minimum_held_out_appearances': int(coverage['n_held_out_appearances'].min()), 'maximum_held_out_appearances': int(coverage['n_held_out_appearances'].max()), 'calibration_oe_ratio': float(oe['calibration_oe_ratio']), 'calibration_slope': float(slope['calibration_slope']), 'observed_km_risk': float(oe['calibration_observed_km_risk']), 'mean_predicted_risk': float(oe['calibration_mean_predicted_risk']), 'calibration_unit': 'one_averaged_held_out_prediction_per_participant', 'calibration_estimand': 'cross_fitted_repeated_split_ensemble_not_one_final_deployable_model', 'uncertainty_scope': 'conditional_participant_bootstrap_does_not_include_model_development_or_split_stream_selection'})
        grouped = grouped_calibration_table(calibration_data, probabilities, time_point=horizon_months, n_groups=calibration_groups)
        if len(grouped) != calibration_groups:
            raise RuntimeError(f'Participant-level grouped calibration is incomplete for {model_name}.')
        grouped_tables[model_name] = grouped
        grouped_saved = grouped.copy()
        grouped_saved.insert(0, 'model', model_name)
        grouped_rows.append(grouped_saved)
        flexible = flexible_calibration_curve_with_bootstrap(calibration_data, probabilities, time_point=horizon_months, n_bootstrap=bootstrap_reps, random_state=seed + 700000 + MODEL_ORDER.index(model_name))
        minimum_successful = int(math.ceil(0.9 * bootstrap_reps))
        if int(flexible['n_bootstrap_successful'].iloc[0]) < minimum_successful:
            raise RuntimeError(f'Participant-level flexible calibration bootstrap is incomplete for {model_name}.')
        flexible_tables[model_name] = flexible
        flexible_saved = flexible.copy()
        flexible_saved.insert(0, 'model', model_name)
        flexible_rows.append(flexible_saved)
    calibration_summary = pd.DataFrame(summary_rows)
    atomic_write_dataframe_csv(calibration_summary, os.path.join(tables_dir, 'repeated_200_participant_level_calibration_summary.csv'))
    atomic_write_dataframe_csv(pd.concat(grouped_rows, ignore_index=True), os.path.join(tables_dir, 'repeated_200_grouped_calibration.csv'))
    atomic_write_dataframe_csv(pd.concat(flexible_rows, ignore_index=True), os.path.join(tables_dir, 'repeated_200_flexible_calibration.csv'))
    grouped_path = os.path.join(figures_dir, 'repeated_200_grouped_calibration.png')
    flexible_path = os.path.join(figures_dir, 'repeated_200_flexible_calibration.png')
    make_combined_calibration_figure(grouped_tables, grouped_path, horizon_months=horizon_months)
    make_combined_flexible_calibration_figure(flexible_tables, flexible_path, horizon_months=horizon_months)
    for path in [grouped_path, flexible_path, os.path.splitext(grouped_path)[0] + '.tiff', os.path.splitext(flexible_path)[0] + '.tiff']:
        if not os.path.isfile(path) or os.path.getsize(path) == 0:
            raise RuntimeError(f'Repeated-split calibration figure was not created: {path}')
    return calibration_summary

def add_participant_calibration_to_performance_tables(summary_df: pd.DataFrame, formatted_df: pd.DataFrame, calibration_summary: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame]:
    summary = summary_df.copy()
    formatted = formatted_df.copy()
    summary['calibration_summary_unit'] = 'one_averaged_held_out_prediction_per_participant'
    summary['calibration_slope'] = np.nan
    summary['calibration_oe_ratio'] = np.nan
    formatted['Calibration_slope_120m'] = 'NA'
    formatted['O_E_120m'] = 'NA'
    if len(calibration_summary) != len(CALIBRATION_MODELS) or calibration_summary.duplicated(['model']).any() or set(calibration_summary['model'].astype(str)) != set(CALIBRATION_MODELS):
        raise RuntimeError('Participant-level calibration summary has an incomplete model set.')
    indexed = calibration_summary.set_index('model')
    for model_name in MODEL_ORDER:
        summary_mask = summary['model'] == model_name
        formatted_mask = formatted['Model'] == model_name
        if model_name == 'SVM':
            continue
        slope = float(indexed.loc[model_name, 'calibration_slope'])
        oe = float(indexed.loc[model_name, 'calibration_oe_ratio'])
        summary.loc[summary_mask, 'calibration_slope'] = slope
        summary.loc[summary_mask, 'calibration_oe_ratio'] = oe
        formatted.loc[formatted_mask, 'Calibration_slope_120m'] = f'{slope:.3f}'
        formatted.loc[formatted_mask, 'O_E_120m'] = f'{oe:.3f}'
    return (summary, formatted)

def run_repeated_200_splits(df: pd.DataFrame, out_dir: str, seed: int, n_splits: int, test_size: float, inner_folds: int, horizon_months: float, ibs_start_month: float, calibration_groups: int, n_jobs: int, impute_max_iter: int, calibration_bootstrap_reps: int, resume: bool, prepared_manifest_sha256: Optional[str]=None, deadline_monotonic: Optional[float]=None) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    if int(seed) != PRESPECIFIED_SPLITTER_SEED:
        raise ValueError(f'The analysis seed must be {PRESPECIFIED_SPLITTER_SEED}.')
    if int(n_splits) != PRESPECIFIED_N_SPLITS:
        raise ValueError(f'The analysis requires exactly {PRESPECIFIED_N_SPLITS} splits.')
    if not math.isclose(float(test_size), PRESPECIFIED_TEST_SIZE, rel_tol=0.0, abs_tol=1e-12):
        raise ValueError(f'The repeated-split test fraction must be {PRESPECIFIED_TEST_SIZE:.2f}.')
    if inner_folds != 3:
        raise ValueError('Classical-model tuning requires exactly three inner folds.')
    tables_dir = ensure_dir(os.path.join(out_dir, 'tables'))
    meta_dir = ensure_dir(os.path.join(out_dir, 'metadata'))
    checkpoint_dir = ensure_dir(os.path.join(out_dir, 'checkpoints', 'repeated_200_splits_ibs_v1'))
    configuration = repeated_split_checkpoint_configuration(df=df, seed=seed, n_splits=n_splits, test_size=test_size, inner_folds=inner_folds, horizon_months=horizon_months, ibs_start_month=ibs_start_month, calibration_groups=calibration_groups, n_jobs=n_jobs, impute_max_iter=impute_max_iter, calibration_bootstrap_reps=calibration_bootstrap_reps)
    manifest_sha256 = prepared_manifest_sha256 if prepared_manifest_sha256 is not None else prepare_checkpoint_manifest(checkpoint_dir, run_type='repeated_200_splits_ibs_tuning', configuration=configuration, resume=resume)
    split_map = build_repeated_split_map(df, seed=seed, n_splits=n_splits, test_size=test_size)
    atomic_write_dataframe_csv(split_map, os.path.join(meta_dir, 'repeated_200_split_map.csv'))
    metrics_records: List[Dict[str, Any]] = []
    tuning_records: List[Dict[str, Any]] = []
    prediction_frames: List[pd.DataFrame] = []
    completed_splits = 0
    stage_start = time.monotonic()
    for split_id, train_idx, test_idx in iter_repeated_splits(df, seed=seed, n_splits=n_splits, test_size=test_size):
        raw_train_df = df.iloc[train_idx].copy().reset_index(drop=True)
        raw_test_df = df.iloc[test_idx].copy().reset_index(drop=True)
        expected_test_ids = raw_test_df[STUDY_ID_COL].reset_index(drop=True)
        if resume:
            loaded = load_repeated_split_checkpoint(checkpoint_dir=checkpoint_dir, manifest_sha256=manifest_sha256, split_id=split_id, expected_test_ids=expected_test_ids)
            if loaded is not None:
                split_metrics, split_tuning, split_predictions = loaded
                metrics_records.extend(split_metrics.to_dict('records'))
                tuning_records.extend(split_tuning.to_dict('records'))
                prediction_frames.append(split_predictions)
                completed_splits += 1
                LOGGER.info('Resumed repeated split %d/%d (%d complete).', split_id, n_splits, completed_splits)
                check_session_deadline(deadline_monotonic, f'repeated split {split_id}')
                continue
        LOGGER.info('Running repeated split %d/%d.', split_id, n_splits)
        split_base_seed = seed + 100000 * split_id
        completed = generate_imputed_split_datasets(raw_train_df, raw_test_df, m=1, seed=split_base_seed, max_iter=impute_max_iter)[0]
        train_df = completed['train']
        test_df = completed['test']
        fit_seed = split_base_seed + 1
        specs = get_classical_model_specs(random_state=fit_seed, n_jobs=n_jobs)
        inner_completed_splits: Optional[List[Tuple[pd.DataFrame, pd.DataFrame]]] = None
        neural_selection_split: Optional[Tuple[pd.DataFrame, pd.DataFrame]] = None
        split_metric_rows: List[Dict[str, Any]] = []
        split_tuning_rows: List[Dict[str, Any]] = []
        absolute_predictions: Dict[str, np.ndarray] = {}
        classical_models = ['CoxPH', 'ElasticNetCox', 'RSF', 'GBSA', 'XGBoost', 'SVM']
        for model_name in classical_models:
            partial = load_repeated_split_model_checkpoint(checkpoint_dir=checkpoint_dir, manifest_sha256=manifest_sha256, split_id=split_id, model_name=model_name, expected_test_ids=expected_test_ids) if resume else None
            if partial is not None:
                model_metrics_df, model_tuning_df, model_predictions_df = partial
                split_metric_rows.extend(model_metrics_df.to_dict('records'))
                split_tuning_rows.extend(model_tuning_df.to_dict('records'))
                if model_name != 'SVM':
                    absolute_predictions[model_name] = model_predictions_df['predicted_event_probability'].to_numpy(dtype=float)
                LOGGER.info('  resumed split %d model %s', split_id, model_name)
                continue
            check_session_deadline(deadline_monotonic, f'repeated split {split_id} before model {model_name}')
            try:
                if model_name != 'CoxPH' and inner_completed_splits is None:
                    inner_completed_splits = generate_nested_inner_imputed_splits(raw_outer_train_df=raw_train_df, inner_folds=inner_folds, split_seed=fit_seed, imputation_seed=split_base_seed + 100, max_iter=impute_max_iter)
                fitted, best_params, selection_score = fit_gridsearched_model(model_name, train_df, inner_folds=inner_folds, random_state=fit_seed, n_jobs=n_jobs, specs=specs, inner_completed_splits=inner_completed_splits, ibs_start_month=ibs_start_month, horizon_months=horizon_months)
                metrics, _, predicted_event_probability = evaluate_model(model_name, fitted, train_df, test_df, horizon_months, ibs_start_month, calibration_groups, compute_auxiliary_calibration=False)
                require_complete_outer_fold_metrics(model_name, metrics, predicted_event_probability, expected_n=len(test_df), context=f'Repeated split {split_id}')
            except Exception as exc:
                LOGGER.exception('Repeated split %d failed for %s: %s', split_id, model_name, exc)
                raise RuntimeError(f'Repeated split {split_id} failed for {model_name}. Earlier model checkpoints in this split remain reusable.') from exc
            model_metrics_df = pd.DataFrame([{'split_id': int(split_id), 'n_train': int(len(train_df)), 'n_test': int(len(test_df)), 'train_events': int(raw_train_df[EVENT_COL].sum()), 'test_events': int(raw_test_df[EVENT_COL].sum()), 'model': model_name, **metrics}])
            model_tuning_df = pd.DataFrame([{'split': 'repeated_200_stratified_80_20', 'split_id': int(split_id), 'model': model_name, 'tuning_method': main_tuning_method(model_name), **tuning_output_metadata(model_name, selection_score), 'best_params_json': json.dumps(best_params, sort_keys=True), 'split_base_seed': int(split_base_seed), 'outer_imputation_seed': int(split_base_seed), 'model_seed': int(fit_seed), 'classical_inner_imputation_seed_base': int(split_base_seed + 100), 'neural_development_imputation_seed': int(split_base_seed + 500)}])
            model_probabilities = np.full(len(test_df), np.nan, dtype=float) if model_name == 'SVM' else np.asarray(predicted_event_probability, dtype=float)
            model_predictions_df = pd.DataFrame({STUDY_ID_COL: expected_test_ids.to_numpy(), 'predicted_event_probability': model_probabilities})
            write_repeated_split_model_checkpoint(checkpoint_dir=checkpoint_dir, manifest_sha256=manifest_sha256, split_id=split_id, model_name=model_name, metrics_df=model_metrics_df, tuning_df=model_tuning_df, predictions_df=model_predictions_df, expected_test_ids=expected_test_ids)
            split_metric_rows.extend(model_metrics_df.to_dict('records'))
            split_tuning_rows.extend(model_tuning_df.to_dict('records'))
            if model_name != 'SVM':
                absolute_predictions[model_name] = model_probabilities
            LOGGER.info('  checkpointed split %d model %s', split_id, model_name)
            check_session_deadline(deadline_monotonic, f'repeated split {split_id} model {model_name}')
        for model_name, fit_function in [('DeepSurv', fit_deepsurv), ('CoxTime', fit_coxtime)]:
            partial = load_repeated_split_model_checkpoint(checkpoint_dir=checkpoint_dir, manifest_sha256=manifest_sha256, split_id=split_id, model_name=model_name, expected_test_ids=expected_test_ids) if resume else None
            if partial is not None:
                model_metrics_df, model_tuning_df, model_predictions_df = partial
                split_metric_rows.extend(model_metrics_df.to_dict('records'))
                split_tuning_rows.extend(model_tuning_df.to_dict('records'))
                absolute_predictions[model_name] = model_predictions_df['predicted_event_probability'].to_numpy(dtype=float)
                LOGGER.info('  resumed split %d model %s', split_id, model_name)
                continue
            check_session_deadline(deadline_monotonic, f'repeated split {split_id} before model {model_name}')
            try:
                if neural_selection_split is None:
                    neural_selection_split = generate_nested_neural_selection_split(raw_outer_train_df=raw_train_df, split_seed=fit_seed, imputation_seed=split_base_seed + 500, max_iter=impute_max_iter)
                fitted, best_params, selection_score = fit_function(train_df, random_state=fit_seed, selection_split=neural_selection_split, require_all_candidates=True, ibs_start_month=ibs_start_month, horizon_months=horizon_months)
                metrics, _, predicted_event_probability = evaluate_model(model_name, fitted, train_df, test_df, horizon_months, ibs_start_month, calibration_groups, compute_auxiliary_calibration=False)
                require_complete_outer_fold_metrics(model_name, metrics, predicted_event_probability, expected_n=len(test_df), context=f'Repeated split {split_id}')
            except Exception as exc:
                LOGGER.exception('Repeated split %d failed for %s: %s', split_id, model_name, exc)
                raise RuntimeError(f'Repeated split {split_id} failed for {model_name}. Earlier model checkpoints in this split remain reusable.') from exc
            model_metrics_df = pd.DataFrame([{'split_id': int(split_id), 'n_train': int(len(train_df)), 'n_test': int(len(test_df)), 'train_events': int(raw_train_df[EVENT_COL].sum()), 'test_events': int(raw_test_df[EVENT_COL].sum()), 'model': model_name, **metrics}])
            model_tuning_df = pd.DataFrame([{'split': 'repeated_200_stratified_80_20', 'split_id': int(split_id), 'model': model_name, 'tuning_method': main_tuning_method(model_name), **tuning_output_metadata(model_name, selection_score), 'best_params_json': json.dumps(best_params, sort_keys=True), 'split_base_seed': int(split_base_seed), 'outer_imputation_seed': int(split_base_seed), 'model_seed': int(fit_seed), 'classical_inner_imputation_seed_base': int(split_base_seed + 100), 'neural_development_imputation_seed': int(split_base_seed + 500)}])
            model_probabilities = np.asarray(predicted_event_probability, dtype=float)
            model_predictions_df = pd.DataFrame({STUDY_ID_COL: expected_test_ids.to_numpy(), 'predicted_event_probability': model_probabilities})
            write_repeated_split_model_checkpoint(checkpoint_dir=checkpoint_dir, manifest_sha256=manifest_sha256, split_id=split_id, model_name=model_name, metrics_df=model_metrics_df, tuning_df=model_tuning_df, predictions_df=model_predictions_df, expected_test_ids=expected_test_ids)
            split_metric_rows.extend(model_metrics_df.to_dict('records'))
            split_tuning_rows.extend(model_tuning_df.to_dict('records'))
            absolute_predictions[model_name] = model_probabilities
            LOGGER.info('  checkpointed split %d model %s', split_id, model_name)
            if model_name != 'CoxTime':
                check_session_deadline(deadline_monotonic, f'repeated split {split_id} model {model_name}')
        missing_absolute = sorted(set(CALIBRATION_MODELS).difference(absolute_predictions))
        if missing_absolute:
            raise RuntimeError(f'Repeated split {split_id} lacks probabilities for {missing_absolute}.')
        split_metrics = pd.DataFrame(split_metric_rows)
        split_tuning = pd.DataFrame(split_tuning_rows)
        split_predictions = pd.DataFrame({'split_id': int(split_id), STUDY_ID_COL: expected_test_ids.to_numpy(), **{model_name: absolute_predictions[model_name] for model_name in CALIBRATION_MODELS}})
        write_repeated_split_checkpoint(checkpoint_dir=checkpoint_dir, manifest_sha256=manifest_sha256, split_id=split_id, metrics_df=split_metrics, tuning_df=split_tuning, predictions_df=split_predictions, expected_test_ids=expected_test_ids)
        metrics_records.extend(split_metrics.to_dict('records'))
        tuning_records.extend(split_tuning.to_dict('records'))
        prediction_frames.append(split_predictions)
        completed_splits += 1
        elapsed = max(time.monotonic() - stage_start, 1e-09)
        LOGGER.info('Checkpointed repeated split %d/%d. Approximate remaining time in this process: %.1f hours.', split_id, n_splits, elapsed / max(completed_splits, 1) * (n_splits - completed_splits) / 3600.0)
        check_session_deadline(deadline_monotonic, f'repeated split {split_id}')
    performance_df = pd.DataFrame(metrics_records)
    published_columns = ['split_id', 'n_train', 'n_test', 'train_events', 'test_events', 'model', 'harrell_c', 'uno_c_tau', 'tau_months', 'brier_score_at_horizon', 'brier_horizon_months', 'integrated_brier_score']
    performance_df = performance_df.loc[:, published_columns].copy()
    for column in ['split_id', 'n_train', 'n_test', 'train_events', 'test_events']:
        performance_df[column] = pd.to_numeric(performance_df[column], errors='raise').astype(int)
    for column in ['harrell_c', 'uno_c_tau', 'tau_months', 'brier_score_at_horizon', 'brier_horizon_months', 'integrated_brier_score']:
        performance_df[column] = pd.to_numeric(performance_df[column], errors='coerce').astype(float)
    tuning_df = pd.DataFrame(tuning_records)
    for column in ['split_id', 'split_base_seed', 'outer_imputation_seed', 'model_seed', 'classical_inner_imputation_seed_base', 'neural_development_imputation_seed']:
        tuning_df[column] = pd.to_numeric(tuning_df[column], errors='raise').astype(int)
    tuning_df['selection_score'] = pd.to_numeric(tuning_df['selection_score'], errors='coerce').astype(float)
    held_out_predictions = pd.concat(prediction_frames, ignore_index=True)
    held_out_predictions['split_id'] = pd.to_numeric(held_out_predictions['split_id'], errors='raise').astype(int)
    for model_name in CALIBRATION_MODELS:
        held_out_predictions[model_name] = pd.to_numeric(held_out_predictions[model_name], errors='raise').astype(float)
    model_order_lookup = {model_name: index for index, model_name in enumerate(MODEL_ORDER)}
    performance_df['_model_order'] = performance_df['model'].map(model_order_lookup)
    performance_df = performance_df.sort_values(['split_id', '_model_order']).drop(columns='_model_order').reset_index(drop=True)
    tuning_df['_model_order'] = tuning_df['model'].map(model_order_lookup)
    tuning_df = tuning_df.sort_values(['split_id', '_model_order']).drop(columns='_model_order').reset_index(drop=True)
    expected_rows = int(n_splits) * len(MODEL_ORDER)
    if len(performance_df) != expected_rows or len(tuning_df) != expected_rows:
        raise RuntimeError('Repeated 200-split metrics or tuning grid is incomplete.')
    if performance_df.duplicated(['split_id', 'model']).any():
        raise RuntimeError('Repeated 200-split performance grid contains duplicate keys.')
    if tuning_df.duplicated(['split_id', 'model']).any():
        raise RuntimeError('Repeated 200-split tuning grid contains duplicate keys.')
    validate_split_performance_grid(performance_df, n_splits=n_splits)
    validate_repeated_held_out_predictions(held_out_predictions, split_map=split_map, df=df, n_splits=n_splits)
    atomic_write_dataframe_csv(performance_df, os.path.join(tables_dir, 'repeated_200_performance_by_split.csv'))
    atomic_write_dataframe_csv(held_out_predictions, os.path.join(tables_dir, 'repeated_200_held_out_predictions.csv'))
    atomic_write_dataframe_csv(tuning_df, os.path.join(tables_dir, 'repeated_200_tuning_by_split.csv'))
    atomic_write_dataframe_csv(summarise_tuning_records(tuning_df), os.path.join(meta_dir, 'repeated_200_tuning_summary.csv'))
    completeness_rows = []
    for model_name, rows in performance_df.groupby('model', sort=False):
        required_metrics = _outer_required_metrics(str(model_name))
        completeness_rows.append({'model': model_name, 'n_split_rows': int(len(rows)), 'required_metrics': ','.join(required_metrics), 'complete': bool(np.isfinite(rows[required_metrics].apply(pd.to_numeric, errors='coerce').to_numpy(float)).all())})
    completeness_df = pd.DataFrame(completeness_rows)
    if not completeness_df['complete'].all():
        raise RuntimeError('Repeated 200-split required-metric validation failed.')
    atomic_write_dataframe_csv(completeness_df, os.path.join(meta_dir, 'repeated_200_metric_completeness.csv'))
    calibration_summary = create_repeated_split_calibration_outputs(df=df, held_out_predictions=held_out_predictions, split_map=split_map, n_splits=n_splits, horizon_months=horizon_months, calibration_groups=calibration_groups, bootstrap_reps=calibration_bootstrap_reps, out_dir=out_dir, seed=seed)
    summary_df, table3_df = summarise_performance(performance_df, n_splits=n_splits)
    summary_df, table3_df = add_participant_calibration_to_performance_tables(summary_df, table3_df, calibration_summary)
    atomic_write_dataframe_csv(summary_df, os.path.join(tables_dir, 'repeated_200_performance_summary.csv'))
    atomic_write_dataframe_csv(table3_df, os.path.join(tables_dir, 'repeated_200_table3_formatted.csv'))
    paired_by_split = build_paired_differences_by_split(performance_df, reference_model='CoxPH')
    atomic_write_dataframe_csv(paired_by_split, os.path.join(tables_dir, 'repeated_200_paired_differences_by_split.csv'))
    paired_df = compare_models_vs_reference(performance_df, reference_model='CoxPH')
    expected_paired_keys = {(model_name, metric) for model_name in MODEL_ORDER if model_name != 'CoxPH' for metric in ['harrell_c', 'uno_c_tau', 'brier_score_at_horizon', 'integrated_brier_score'] if not (model_name == 'SVM' and metric in {'brier_score_at_horizon', 'integrated_brier_score'})}
    observed_paired_keys = set(zip(paired_df['Model'], paired_df['Metric']))
    if len(paired_df) != 26 or observed_paired_keys != expected_paired_keys or paired_df.duplicated(['Model', 'Metric']).any() or (not (paired_df['n_paired_splits'].astype(int) == n_splits).all()):
        raise RuntimeError('Paired differences versus CoxPH are incomplete.')
    atomic_write_dataframe_csv(paired_df, os.path.join(tables_dir, 'repeated_200_paired_differences_vs_coxph_summary.csv'))
    inference_note = 'The 200 repeated 80/20 splits overlap and are not independent cohorts. Means and empirical 2.5th-97.5th percentile ranges describe variation across the prespecified split stream; they are not confidence intervals. Model differences are paired within identical held-out sets. No P values or formal hypothesis tests were calculated, so no multiplicity adjustment was applied. Calibration uses one prediction per participant, averaged only across splits in which that participant was held out; the 49,800 prediction rows were not stacked as independent observations. Flexible calibration bands are conditional participant-bootstrap bands; they do not include model-development or split-stream-selection uncertainty. The 200 repeated stratified 80/20 splits are the primary internal-validation analysis in this release.'
    inference_note_path = os.path.join(meta_dir, 'repeated_200_inference_and_multiplicity_note.txt')
    fd, temporary_path = _temporary_path(inference_note_path)
    try:
        with os.fdopen(fd, 'w', encoding='utf-8') as handle:
            handle.write(inference_note + '\n')
            handle.flush()
            os.fsync(handle.fileno())
        os.replace(temporary_path, inference_note_path)
    except Exception:
        try:
            os.unlink(temporary_path)
        except FileNotFoundError:
            pass
        raise
    return (performance_df, table3_df, paired_df)

def build_fixed_vs_repeated_comparison(fixed_metrics: pd.DataFrame, repeated_performance: pd.DataFrame, n_splits: int) -> pd.DataFrame:
    repeated_summary, _ = summarise_performance(repeated_performance, n_splits=n_splits)
    repeated_index = repeated_summary.set_index('model')
    metric_map = {'Harrell_C': 'harrell_c', 'Uno_C_tau': 'uno_c_tau', 'Brier_120m': 'brier_score_at_horizon', 'IBS_12_120m': 'integrated_brier_score'}
    rows: List[Dict[str, Any]] = []
    fixed_index = fixed_metrics.set_index('Model')
    for model_name in MODEL_ORDER:
        for fixed_column, repeated_column in metric_map.items():
            fixed_value = float(fixed_index.loc[model_name, fixed_column])
            repeated_mean = float(repeated_index.loc[model_name, f'{repeated_column}_mean'])
            model_split_rows = repeated_performance.loc[repeated_performance['model'].eq(model_name), ['split_id', repeated_column]].copy()
            split_1_series = pd.to_numeric(model_split_rows.loc[pd.to_numeric(model_split_rows['split_id'], errors='raise').astype(int).eq(1), repeated_column], errors='coerce')
            if len(split_1_series) != 1:
                raise RuntimeError(f'Repeated split 1 is missing or duplicated for {model_name} {repeated_column}.')
            repeated_split_1 = float(split_1_series.iloc[0])
            all_repeated_values = pd.to_numeric(model_split_rows[repeated_column], errors='coerce').to_numpy(float)
            finite_repeated = all_repeated_values[np.isfinite(all_repeated_values)]
            if np.isfinite(repeated_split_1) and finite_repeated.size:
                split_1_percentile_rank = float(100.0 * np.mean(finite_repeated <= repeated_split_1))
            else:
                split_1_percentile_rank = np.nan
            rows.append({'Model': model_name, 'Metric': repeated_column, 'Fixed_80_20_value': fixed_value, 'Repeated_method_split_1_value': repeated_split_1, 'Repeated_200_mean': repeated_mean, 'Repeated_minus_fixed': repeated_mean - fixed_value, 'Fixed_minus_repeated_method_split_1': fixed_value - repeated_split_1 if np.isfinite(fixed_value) and np.isfinite(repeated_split_1) else np.nan, 'Repeated_mean_minus_repeated_method_split_1': repeated_mean - repeated_split_1 if np.isfinite(repeated_mean) and np.isfinite(repeated_split_1) else np.nan, 'Repeated_method_split_1_percentile_rank': split_1_percentile_rank, 'Repeated_empirical_2_5_percentile': float(repeated_index.loc[model_name, f'{repeated_column}_empirical_2_5_percentile']), 'Repeated_empirical_97_5_percentile': float(repeated_index.loc[model_name, f'{repeated_column}_empirical_97_5_percentile']), 'Interpretation': 'descriptive_only_fixed_split_is_split_1_and_does_not_add_an_independent_comparison', 'Procedure_note': 'Fixed and repeated split 1 have identical participant membership, but the fixed value averages 10 imputations and uses the fixed-stage seed/refitting procedure; their difference is not an allocation effect. Repeated mean minus repeated-method split 1 isolates how split 1 compares with the same repeated-analysis procedure across the stream.'})
    return pd.DataFrame(rows)

def build_supplementary_tuning_table(primary_tuning_df: pd.DataFrame) -> pd.DataFrame:
    rows = [{'Model': 'CoxPH', 'Implementation and role': 'CoxPHSurvivalAnalysis (semiparametric benchmark)', 'Prespecified tuning strategy': 'One numerical specification: Efron ties, tol = 1e-7, n_iter = 300'}, {'Model': 'ElasticNetCox', 'Implementation and role': 'Penalized CoxPHFitter (regularized Cox benchmark)', 'Prespecified tuning strategy': 'Eight prespecified (penalizer, l1_ratio) pairs spanning ridge, mixed and LASSO penalties; continuous predictors standardized within each training split'}, {'Model': 'RSF', 'Implementation and role': 'RandomSurvivalForest (tree-based survival ensemble)', 'Prespecified tuning strategy': 'Six prespecified configurations; 300 trees with sparse coverage of node size, feature fraction and depth; candidate fits are single-threaded'}, {'Model': 'GBSA', 'Implementation and role': 'GradientBoostingSurvivalAnalysis (boosted survival-tree model)', 'Prespecified tuning strategy': 'Six prespecified configurations spanning 100-200 trees, learning rate 0.03-0.10, depth 1-2, node size and subsampling'}, {'Model': 'DeepSurv', 'Implementation and role': 'DeepSurv (neural-network extension of Cox regression; continuous predictors standardized within each training split)', 'Prespecified tuning strategy': 'Six prespecified multilayer perceptron configurations were evaluated: (1) num_nodes = [32,16], dropout = 0.1, lr = 1e-2, max_epochs = 256; (2) num_nodes = [32,16], dropout = 0.2, lr = 1e-2, max_epochs = 256; (3) num_nodes = [32,32], dropout = 0.1, lr = 1e-2, max_epochs = 256; (4) num_nodes = [64,32], dropout = 0.2, lr = 1e-3, max_epochs = 256; (5) num_nodes = [64,32], dropout = 0.1, lr = 1e-3, max_epochs = 256; and (6) num_nodes = [16,16], dropout = 0.2, lr = 1e-2, max_epochs = 256. DeepSurv used the complete development risk set in each optimizer step and the complete internal validation set for loss-based early stopping. Configuration selection minimized validation IBS from 12 to 120 months using development-only censoring weights, followed by refitting the selected configuration on the complete outer training set for the selected epoch count.'}, {'Model': 'CoxTime', 'Implementation and role': 'CoxTime (neural non-proportional-hazards Cox extension; continuous predictors standardized within each training split)', 'Prespecified tuning strategy': 'Six prespecified multilayer perceptron configurations were evaluated: (1) num_nodes = [32,16], dropout = 0.1, lr = 1e-2, batch_size = 64, max_epochs = 256; (2) num_nodes = [32,16], dropout = 0.2, lr = 1e-2, batch_size = 64, max_epochs = 256; (3) num_nodes = [32,32], dropout = 0.1, lr = 1e-2, batch_size = 64, max_epochs = 256; (4) num_nodes = [64,32], dropout = 0.2, lr = 1e-3, batch_size = 64, max_epochs = 256; (5) num_nodes = [64,32], dropout = 0.1, lr = 1e-3, batch_size = 128, max_epochs = 256; and (6) num_nodes = [16,16], dropout = 0.2, lr = 1e-2, batch_size = 128, max_epochs = 256. The internal stratified 80/20 validation split used loss-based early stopping; configuration selection minimized validation IBS from 12 to 120 months using development-only censoring weights, followed by refitting the selected configuration on the full outer training set for the selected epoch count.'}, {'Model': 'XGBoost', 'Implementation and role': 'XGBRegressor with survival:cox objective (boosted tree survival model)', 'Prespecified tuning strategy': 'Eight prespecified configurations spanning 100-300 trees, learning rate 0.01-0.10, depth 1-3, subsampling, minimum child weight and L2 penalty; candidate fits are single-threaded'}, {'Model': 'Survival SVM', 'Implementation and role': 'FastSurvivalSVM (linear ranking model; rank_ratio = 1.0)', 'Prespecified tuning strategy': 'Four alpha values {0.01, 0.1, 1.0, 10.0}; AVL-tree optimizer, max_iter = 500 and rank_ratio = 1 fixed as numerical controls; continuous predictors standardized within each training split'}]
    out = pd.DataFrame(rows)
    selection_targets = {'CoxPH': 'Fixed specification; no hyperparameter selection', 'ElasticNetCox': 'Lowest mean validation IBS from 12 to 120 months across three inner folds', 'RSF': 'Lowest mean validation IBS from 12 to 120 months across three inner folds', 'GBSA': 'Lowest mean validation IBS from 12 to 120 months across three inner folds', 'DeepSurv': 'Lowest validation IBS from 12 to 120 months on the internal development-validation split', 'CoxTime': 'Lowest validation IBS from 12 to 120 months on the internal development-validation split', 'XGBoost': 'Lowest mean validation IBS from 12 to 120 months across three inner folds', 'Survival SVM': 'Highest mean Harrell C across three inner folds'}
    out['Selection target'] = out['Model'].map(selection_targets)
    if not primary_tuning_df.empty:

        def configuration_without_refit_epochs(value: Any) -> str:
            parameters = json.loads(str(value))
            if not isinstance(parameters, dict):
                raise ValueError('Selected neural parameters were not a JSON object.')
            parameters.pop('refit_epochs', None)
            return json.dumps(parameters, sort_keys=True)
        deep_primary = primary_tuning_df[primary_tuning_df['model'] == 'DeepSurv']
        if not deep_primary.empty:
            mode_params = deep_primary['best_params_json'].map(configuration_without_refit_epochs).mode().iloc[0]
            out.attrs['primary_deepsurv'] = mode_params
        ct_primary = primary_tuning_df[primary_tuning_df['model'] == 'CoxTime']
        if not ct_primary.empty:
            out.attrs['primary_coxtime'] = ct_primary['best_params_json'].map(configuration_without_refit_epochs).mode().iloc[0]
    return out

def set_repeat_table_header(row: Any) -> None:
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)

def set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def style_docx_table(table: Any, header_fill: str='D9EAF7') -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                if row_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if len(cell.text) <= 18:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if row_idx == 0:
            for cell in row.cells:
                set_cell_shading(cell, header_fill)
            set_repeat_table_header(row)

def add_dataframe_as_table(doc: Document, df: pd.DataFrame, title: str, note: Optional[str]=None, landscape: bool=False, font_size: float=9.0) -> None:
    if landscape:
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = (section.page_height, section.page_width)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
    else:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.bold = True
    table = doc.add_table(rows=1, cols=len(df.columns))
    for j, col in enumerate(df.columns):
        table.cell(0, j).text = str(col)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            val = row[col]
            if isinstance(val, float) and np.isnan(val):
                txt = 'NA'
            else:
                txt = str(val)
            cells[j].text = txt
    style_docx_table(table)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size)
    if note:
        note_p = doc.add_paragraph()
        note_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        note_run = note_p.add_run(note)
        note_run.italic = True
        note_run.font.size = Pt(8.5)

def save_single_table_docx(df: pd.DataFrame, title: str, path: str, note: Optional[str]=None, landscape: bool=False) -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Inches(0.6)
    sec.right_margin = Inches(0.6)
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    title_p = doc.add_paragraph()
    title_p.style = doc.styles['Title']
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.add_run(title)
    add_dataframe_as_table(doc, df, title='', note=note, landscape=landscape)
    doc.save(path)

def write_tables_docx(out_dir: str, table1: pd.DataFrame, table2: pd.DataFrame, table3: pd.DataFrame, table4: pd.DataFrame, supp_s1: pd.DataFrame, supp_s2: pd.DataFrame, supp_s3: pd.DataFrame) -> None:
    docs_dir = ensure_dir(os.path.join(out_dir, 'docx_tables'))
    combined = Document()
    sec = combined.sections[0]
    sec.left_margin = Inches(0.6)
    sec.right_margin = Inches(0.6)
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    title = combined.add_paragraph()
    title.style = combined.styles['Title']
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run('ASCVD Survival Models - Tables')
    subtitle = combined.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run('Generated from 200 repeated 80/20 splits and supplementary fixed-split analyses').italic = True
    main_tables = [(table1, 'Table 1. Baseline characteristics of the overall cohort and the fixed training and test sets.', 'Values for variables with missing data are averaged across the 10 completed fixed-split datasets. Event counts and follow-up summaries are from the raw filtered dataset.', True), (table2, 'Table 2. Multivariable Cox proportional hazards model for incident ASCVD in the fixed training set.', 'Cox coefficients were pooled across 10 imputations using basic large-sample Rubin-style pooling. They describe the prediction benchmark and are not etiologic claims. Predictor-specific proportional-hazards tests are summarized descriptively.', True), (table4, 'Table 3. Performance across 200 stratified random 80/20 splits.', 'The table reports the mean and empirical 2.5th-97.5th percentile range across 200 held-out estimates. O:E and slope use one risk per participant, averaged only across splits in which that participant was held out. Every outer and inner imputer is fit only on the corresponding training predictors. The ranges are descriptive and are not confidence intervals.', True)]
    supplementary_tables = [(supp_s1, 'Supplementary Table S1. Variables with missing baseline data before imputation.', None, False), (supp_s2, 'Supplementary Table S2. Prespecified tuning strategies used in internal validation.', 'CoxPH has one fixed specification. ElasticNetCox, RSF, GBSA and XGBoost select the lowest mean IBS from 12 to 120 months across three stratified inner folds. DeepSurv and CoxTime select the lowest IBS on an internal stratified 80/20 validation split; early stopping remains based on validation loss, followed by refitting on the full outer training fold. Survival SVM selects the highest mean Harrell C because it does not produce survival probabilities.', True), (supp_s3, 'Supplementary Table S3. Paired performance differences relative to CoxPH.', 'Differences are calculated within each identical held-out split and summarized across 200 splits. Empirical 2.5th-97.5th percentile ranges are descriptive. The splits overlap and are not independent cohorts. No P values or formal hypothesis tests were calculated, so no multiplicity adjustment was applied.', True), (table3, 'Supplementary Table S5. Illustrative performance in the prespecified fixed train-test split.', 'This partition is split 1 of the repeated 80/20 stream. It is retained for coefficients, proportional-hazards checks, SHAP and fixed-split calibration. Metrics were averaged across 10 fixed-split imputations and do not lead the comparative interpretation.', True)]
    all_tables = main_tables + supplementary_tables
    for frame, title_text, note_text, landscape in all_tables:
        add_dataframe_as_table(combined, frame, title_text, note=note_text, landscape=landscape)
    combined_path = os.path.join(docs_dir, 'ASCVD_survival_models_tables.docx')
    combined.save(combined_path)
    reopened_combined = Document(combined_path)
    if len(reopened_combined.tables) != len(all_tables):
        raise RuntimeError('Combined DOCX did not reopen with every requested table.')
    for frame, title_text, note_text, landscape in all_tables:
        safe_name = title_text.split('. ', 1)[0].replace(' ', '_').replace('/', '_') + '.docx'
        single_path = os.path.join(docs_dir, safe_name)
        save_single_table_docx(frame, title_text, single_path, note=note_text, landscape=landscape)
        reopened_single = Document(single_path)
        if len(reopened_single.tables) != 1:
            raise RuntimeError(f'DOCX did not reopen with one table: {single_path}')

def write_methods_notes(out_dir: str, args: argparse.Namespace, package_versions: Dict[str, Optional[str]]) -> None:
    notes = f"Methods and reproducibility notes\n=================================\n\nOutcome and predictors\n----------------------\n- Event indicator: {EVENT_COL}\n- Follow-up time: {TIME_COL} in months\n- Prediction horizon: {int(args.horizon_months)} months\n- Integrated Brier-score window: {int(args.ibs_start_month)} to {int(args.horizon_months)} months\n- The 15 prespecified baseline predictors are: {', '.join(PREDICTOR_COLUMNS)}.\n- No follow-up variable is supplied as a predictor or to an imputation model.\n\nFixed 80/20 analysis\n--------------------\n- The fixed partition uses the same joint event/sex stratification, 20% test fraction and seed as repeated split 1. Its participant membership is validated against split 1.\n- It retains {args.fixed_split_imputations} predictor-only stochastic imputations for coefficients, proportional-hazards checks, SHAP and fixed-split calibration.\n- Each imputer is fit only on training predictors and then applied to held-out predictors.\n- CoxPH uses one fixed numerical specification and has no hyperparameter selection. The four configurable classical absolute-risk models and Survival SVM use three independently imputed inner training-validation folds.\n- Neural selection uses a separate predictor-only imputer fitted on the internal development subset.\n\nMain repeated validation\n------------------------\n- The 200 repeated stratified 80/20 splits are the primary internal-validation analysis in this release.\n- Exactly {args.n_splits} unique stratified random 80/20 splits are generated by one StratifiedShuffleSplit stream with random_state={args.seed}.\n- Joint event status and sex define the strata. Every split contains 992 training participants with 92 events and 249 held-out participants with 24 events.\n- The same train-test assignments are used for all eight models.\n- In split s, split_base_seed={args.seed}+100000*s. The Bayesian outer imputer uses split_base_seed; model fitting and tuning use split_base_seed+1; classical inner imputers use split_base_seed+100+inner_fold; and the neural development imputer uses split_base_seed+500.\n- One predictor-only Bayesian iterative imputation is fit on the full training predictors in each split and applied to that split's held-out predictors.\n- Each split uses one seeded posterior-sampling completion as a fixed prediction-pipeline rule, not conventional multiple-imputation pooling. Split-level dispersion therefore includes allocation changes and the seeded imputation draw. Ten-imputation coefficient pooling is confined to the supplementary fixed split.\n- Classical tuning refits a predictor-only imputer within each of three inner training folds. The three completed inner folds are shared by elastic-net Cox, RSF, GBSA, XGBoost and survival SVM. CoxPH has one fixed specification and is fitted directly to each complete outer training split.\n- DeepSurv and CoxTime assess six prespecified configurations on one internal stratified 80/20 development-validation split. Its imputer is fit only on the development predictors. The chosen configuration is refit on the complete outer training split.\n- A single internal development-validation partition can make neural-model selection noisier than inner-fold selection. Selection frequencies across all outer splits are retained in metadata/repeated_200_tuning_summary.csv.\n- DeepSurv uses the complete development or refit risk set in every Cox-loss update. This prevents undefined loss in event-free minibatches and retains the complete risk set.\n- Elastic-net Cox, RSF, GBSA and XGBoost select the configuration with the lowest mean validation integrated Brier score from 12 to 120 months across the three inner folds. DeepSurv and CoxTime select the lowest validation IBS over the same monthly window on their internal development-validation split. For every IBS calculation, the censoring distribution is estimated only from the corresponding inner-training or development outcomes. Neural early stopping remains based on validation loss, and the resulting epoch count is refitted with the selected configuration on the complete outer training split.\n- Survival SVM selects the configuration with the highest mean inner-fold Harrell C because it provides risk scores rather than survival probabilities. CoxPH and the separate spline Cox comparator have fixed specifications and are not tuned.\n- Harrell's C, Uno's C at 120 months, Brier score at 120 months and integrated Brier score from 12 to 120 months are calculated on every held-out split. Survival SVM provides risk scores but not survival probabilities and therefore has no Brier or calibration output.\n- Brier measures use Graf IPCW with the censoring distribution estimated from the corresponding outer training outcomes. Uno's C uses the scikit-survival implementation and includes events at exactly month 120.\n- The main repeated stage contains {args.n_splits * len(MODEL_ORDER)} metric rows, {args.n_splits * len(MODEL_ORDER)} tuning rows and {args.n_splits * 249} held-out prediction rows.\n\nParticipant-level calibration\n-----------------------------\n- Each participant's risk is averaged only across splits in which that participant was held out.\n- All {EXPECTED_FILTERED_ROWS} participants must be represented. With this fixed split stream, held-out coverage ranges from 21 to 57 appearances per participant.\n- The averaged dataset contains one row per participant. The 49,800 repeated prediction rows are never stacked or treated as independent observations for calibration.\n- O:E, calibration slope, five-group Kaplan-Meier calibration and flexible IPCW spline calibration are calculated once from the participant-level averaged predictions.\n- Flexible-calibration pointwise bands use participant bootstrap resampling and are conditional on the averaged predictions.\n- These calibration summaries estimate a cross-fitted repeated-split ensemble, not one final deployable fitted model. The participant-bootstrap bands condition on the saved averaged predictions and do not include model-development or split-stream-selection uncertainty.\n- Grouped and flexible calibration composites are landscape 2x4 figures with common square axes, 600-dpi PNG and LZW-compressed TIFF output, panel labels, and no overall image title or footer.\n- Survival SVM is excluded from absolute calibration because it does not provide survival probabilities.\n- Non-ASCVD death is treated as censoring here; the competing-risk analysis is produced by the secondary-analysis stage.\n\nSummaries and inference\n-----------------------\n- Split-level estimates are summarized by the mean and empirical 2.5th-97.5th percentile range across 200 splits.\n- The ranges describe variability across this prespecified, overlapping split stream. They are not confidence intervals for population performance.\n- Model-minus-CoxPH differences are calculated within each identical held-out split and summarized by their mean and empirical 2.5th-97.5th percentiles.\n- All 5,200 available split-level model-minus-CoxPH differences are retained as an audit table; SVM has no Brier-score differences because it does not produce absolute risk.\n- No P values or formal hypothesis tests are calculated, so no multiplicity adjustment is applied.\n- A fixed-versus-repeated comparison is exported separately. It is descriptive because fixed split 1 is part of the repeated split stream and is not independent.\n- That comparison reports the repeated-method split-1 value and percentile rank separately. Fixed versus repeated-method split 1 still differs in imputation/refitting procedure, so that difference is not an allocation effect; repeated mean versus repeated-method split 1 isolates split position under the repeated-analysis procedure.\n\nCheckpointing and safe pause\n----------------------------\n- The fixed analysis is checkpointed after every fitted model and each completed imputation.\n- The repeated analysis has a fresh split_id checkpoint namespace. It writes atomic per-model checkpoints and an atomic complete-split checkpoint with hash-validated completion markers.\n- The repeated manifest contains the dataset, script, locked package versions, split stream, seed schedule and scientific settings. A mismatch stops resume.\n- A positive --max-session-hours value is checked only at safe model or complete-split boundaries. Rerunning the unchanged script with --resume reuses only validated checkpoints.\n\nModels\n------\n- CoxPH\n- ElasticNetCox\n- RSF\n- GBSA\n- DeepSurv\n- CoxTime\n- XGBoost survival:cox\n- Survival SVM\n\nModel explanation\n-----------------\n- SHAP uses imputation 1 of the fixed split only. It is illustrative and does not lead the repeated validation comparison.\n- SHAP values are descriptive model attributions on each model's output scale. They are not causal and magnitudes must not be compared between models.\n\nPackages detected\n-----------------\n{json.dumps(package_versions, indent=2)}\n"
    with open(os.path.join(out_dir, 'metadata', 'methods_notes.txt'), 'w', encoding='utf-8') as handle:
        handle.write(notes)

def primary_checkpoint_configuration(df: pd.DataFrame, args: argparse.Namespace) -> Dict[str, Any]:
    return {'identity': checkpoint_common_identity(df), 'settings': {'seed': int(args.seed), 'test_size': float(args.test_size), 'inner_folds': int(args.inner_folds), 'n_jobs': int(args.n_jobs), 'm_imputations': int(args.fixed_split_imputations), 'impute_max_iter': int(args.impute_max_iter), 'horizon_months': float(args.horizon_months), 'ibs_start_month': float(args.ibs_start_month), 'selection_grid_months': list(range(int(args.ibs_start_month), int(args.horizon_months) + 1)), 'model_selection_policy': {model_name: model_selection_metadata(model_name) for model_name in MODEL_ORDER}, 'calibration_groups': int(args.calibration_groups), 'calibration_bootstrap_reps': int(args.calibration_bootstrap_reps), 'shap_requested': bool(not args.skip_shap), 'shap_max_background': int(args.shap_max_background), 'shap_max_explain': int(args.shap_max_explain)}}

def _primary_reload_paths(out_dir: str) -> Dict[str, str]:
    return {'table1': os.path.join(out_dir, 'tables', 'table1_fixed_split.csv'), 'table2': os.path.join(out_dir, 'tables', 'table2_cox_hazard_ratios_fixed_split.csv'), 'table3': os.path.join(out_dir, 'tables', 'supplementary_table_s5_fixed_split_metrics.csv'), 'tuning': os.path.join(out_dir, 'metadata', 'tuning_records_fixed_split.csv')}

def _primary_artifact_snapshot(out_dir: str, shap_requested: bool, m_imputations: int) -> Dict[str, Dict[str, Any]]:
    required = set(_primary_reload_paths(out_dir).values())
    required.update({os.path.join(out_dir, 'tables', 'supplementary_table_s1_missingness.csv'), os.path.join(out_dir, 'tables', 'supplementary_table_s5_fixed_split_by_imputation.csv'), os.path.join(out_dir, 'tables', 'proportional_hazards_tests_by_imputation.csv'), os.path.join(out_dir, 'tables', 'absolute_calibration_fixed_120m.csv'), os.path.join(out_dir, 'tables', 'grouped_calibration_fixed_all_models.csv'), os.path.join(out_dir, 'tables', 'flexible_calibration_fixed_all_models.csv'), os.path.join(out_dir, 'figures', 'supplementary_fixed_split_grouped_calibration_all_models.png'), os.path.join(out_dir, 'figures', 'supplementary_fixed_split_grouped_calibration_all_models.tiff'), os.path.join(out_dir, 'figures', 'supplementary_fixed_split_flexible_calibration_all_models.png'), os.path.join(out_dir, 'figures', 'supplementary_fixed_split_flexible_calibration_all_models.tiff'), os.path.join(out_dir, 'metadata', 'metric_completeness_fixed_split.csv'), os.path.join(out_dir, 'metadata', 'flexible_calibration_bootstrap_completeness.csv'), os.path.join(out_dir, 'metadata', 'grouped_calibration_completeness.csv'), os.path.join(out_dir, 'metadata', 'hyperparameter_search_spaces.json')})
    for model_name in CALIBRATION_MODELS:
        required.add(os.path.join(out_dir, 'tables', f'calibration_points_fixed_{model_name}.csv'))
        required.add(os.path.join(out_dir, 'tables', f'flexible_calibration_fixed_{model_name}.csv'))
    for imputation in range(1, m_imputations + 1):
        required.update(_primary_imputation_checkpoint_paths(os.path.join(out_dir, 'checkpoints', 'primary_split_ibs_v1'), imputation).values())
    if shap_requested:
        required.update({os.path.join(out_dir, 'tables', 'shap_importance_fixed_split_all_models.csv'), os.path.join(out_dir, 'figures', 'figure2_shap_summary_models_a_to_d.png'), os.path.join(out_dir, 'figures', 'figure3_shap_summary_models_e_to_h.png'), os.path.join(out_dir, 'figures', 'supplementary_figure_s1_local_shap_force.png'), os.path.join(out_dir, 'figures', 'shap_force_CoxPH.png'), os.path.join(out_dir, 'figures', 'shap_force_DeepSurv.png')})
        for model_name in SHAP_MODELS:
            required.add(os.path.join(out_dir, 'figures', f'shap_importance_{model_name}.csv'))
            required.add(os.path.join(out_dir, 'figures', f'shap_summary_{model_name}.png'))
    missing = sorted((path for path in required if not os.path.isfile(path)))
    if missing:
        raise RuntimeError(f'Fixed-split stage returned without required saved outputs: {missing}')
    return {os.path.relpath(path, out_dir).replace(os.sep, '/'): {'sha256': file_sha256(path), 'size_bytes': int(os.path.getsize(path))} for path in sorted(required)}

def _validate_primary_metrics(table3: pd.DataFrame, metrics_long: pd.DataFrame, completeness: pd.DataFrame, m_imputations: int) -> None:
    aggregate_map = {'harrell_c': 'Harrell_C', 'uno_c_tau': 'Uno_C_tau', 'brier_score_at_horizon': 'Brier_120m', 'integrated_brier_score': 'IBS_12_120m', 'calibration_slope': 'Calibration_slope_120m', 'calibration_oe_ratio': 'O_E_120m'}
    required_table3 = {'Model', *aggregate_map.values()}
    if not required_table3.issubset(table3.columns):
        missing = sorted(required_table3.difference(table3.columns))
        raise RuntimeError(f'Fixed-split Table S5 metric columns were missing: {missing}')
    if len(table3) != len(MODEL_ORDER) or set(table3['Model'].astype(str)) != set(MODEL_ORDER):
        raise RuntimeError('Fixed-split Table S5 did not contain the complete eight-model set.')
    if table3.duplicated(['Model']).any():
        raise RuntimeError('Fixed-split Table S5 contained duplicate model rows.')
    for _, row in table3.iterrows():
        model_name = str(row['Model'])
        for raw_name in _required_metric_columns(model_name):
            value = pd.to_numeric(pd.Series([row[aggregate_map[raw_name]]]), errors='coerce').iloc[0]
            if not np.isfinite(value):
                raise RuntimeError(f'Fixed-split Table S5 contained a non-finite required metric for {model_name}: {aggregate_map[raw_name]}.')
    required_long = {'imputation', 'model', 'metrics_complete', 'metric_errors_json', *FIT_REQUIRED_METRICS}
    if not required_long.issubset(metrics_long.columns):
        missing = sorted(required_long.difference(metrics_long.columns))
        raise RuntimeError(f'Fixed-split imputation-level metric columns were missing: {missing}')
    expected_rows = int(len(MODEL_ORDER) * m_imputations)
    if len(metrics_long) != expected_rows:
        raise RuntimeError(f'Fixed-split imputation-level metrics contained {len(metrics_long)} rows; expected {expected_rows}.')
    expected_keys = {(imputation, model_name) for imputation in range(1, m_imputations + 1) for model_name in MODEL_ORDER}
    observed_keys = set(zip(pd.to_numeric(metrics_long['imputation'], errors='raise').astype(int), metrics_long['model'].astype(str)))
    if observed_keys != expected_keys or metrics_long.duplicated(['imputation', 'model']).any():
        raise RuntimeError('Fixed-split imputation-level metrics did not contain the exact planned key grid.')
    if not _validated_boolean_series(metrics_long['metrics_complete'], 'fixed-split metrics_complete').all():
        raise RuntimeError('Fixed-split imputation-level metrics included an incomplete fit.')
    for value in metrics_long['metric_errors_json'].fillna('{}').astype(str):
        parsed = json.loads(value)
        if not isinstance(parsed, dict) or parsed:
            raise RuntimeError('Fixed-split imputation-level metrics included a logged metric error.')
    for model_name, model_rows in metrics_long.groupby('model'):
        for column in _required_metric_columns(str(model_name)):
            values = pd.to_numeric(model_rows[column], errors='coerce').to_numpy(float)
            if not np.isfinite(values).all():
                raise RuntimeError(f'Fixed-split imputation-level metrics contained a non-finite required value for {model_name}: {column}.')
    required_completeness = {'model', 'n_model_fits', 'n_complete_fits', 'n_fits_with_logged_metric_error'}
    if not required_completeness.issubset(completeness.columns):
        raise RuntimeError('Fixed-split metric-completeness summary columns were missing.')
    if len(completeness) != len(MODEL_ORDER) or set(completeness['model'].astype(str)) != set(MODEL_ORDER):
        raise RuntimeError('Fixed-split metric-completeness summary model set was incomplete.')
    for _, row in completeness.iterrows():
        if int(row['n_model_fits']) != m_imputations or int(row['n_complete_fits']) != m_imputations:
            raise RuntimeError(f"Fixed-split metric completeness failed for {row['model']}.")
        if int(row['n_fits_with_logged_metric_error']) != 0:
            raise RuntimeError(f"Fixed-split metrics logged an error for {row['model']}.")
        for column in _required_metric_columns(str(row['model'])):
            availability = f'{column}_n_available'
            if availability not in completeness.columns or int(row[availability]) != m_imputations:
                raise RuntimeError(f"Fixed-split metric availability was incomplete for {row['model']}: {column}.")

def write_primary_checkpoint_marker(out_dir: str, checkpoint_dir: str, manifest_sha256: str, shap_requested: bool, m_imputations: int) -> None:
    paths = _primary_reload_paths(out_dir)
    tuning = pd.read_csv(paths['tuning'])
    validate_model_imputation_tuning_frame(tuning, repeat=0, expected_fit_rows=int(len(MODEL_ORDER) * m_imputations))
    _validate_primary_metrics(pd.read_csv(paths['table3']), pd.read_csv(os.path.join(out_dir, 'tables', 'supplementary_table_s5_fixed_split_by_imputation.csv'), float_precision='round_trip'), pd.read_csv(os.path.join(out_dir, 'metadata', 'metric_completeness_fixed_split.csv')), m_imputations=m_imputations)
    ph_path = os.path.join(out_dir, 'tables', 'proportional_hazards_tests_by_imputation.csv')
    ph_tests = pd.read_csv(ph_path)
    required_ph_columns = {'imputation', 'Variable', 'variable', 'ph_test_statistic', 'ph_test_p'}
    if not required_ph_columns.issubset(ph_tests.columns):
        raise RuntimeError('The proportional-hazards diagnostic table was incomplete.')
    expected_ph_keys = {(imputation, variable) for imputation in range(1, m_imputations + 1) for variable in PREDICTOR_COLUMNS}
    observed_ph_keys = set(zip(pd.to_numeric(ph_tests['imputation'], errors='raise').astype(int), ph_tests['variable'].astype(str)))
    if observed_ph_keys != expected_ph_keys or ph_tests.duplicated(['imputation', 'variable']).any():
        raise RuntimeError('The proportional-hazards diagnostic key grid was incomplete.')
    ph_numeric = ph_tests[['ph_test_statistic', 'ph_test_p']].apply(pd.to_numeric, errors='coerce')
    if not np.isfinite(ph_numeric.to_numpy(float)).all():
        raise RuntimeError('The proportional-hazards diagnostic table contained a non-finite value.')
    if not ph_numeric['ph_test_p'].between(0.0, 1.0).all():
        raise RuntimeError('The proportional-hazards diagnostic table contained an invalid P value.')
    marker = {'checkpoint_schema_version': CHECKPOINT_SCHEMA_VERSION, 'status': 'complete', 'manifest_sha256': manifest_sha256, 'artifacts': _primary_artifact_snapshot(out_dir, shap_requested=shap_requested, m_imputations=m_imputations)}
    atomic_write_json(marker, os.path.join(checkpoint_dir, 'primary_complete.json'))

def _load_primary_checkpoint_strict(out_dir: str, checkpoint_dir: str, manifest_sha256: str, m_imputations: int) -> Optional[Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]]:
    marker_path = os.path.join(checkpoint_dir, 'primary_complete.json')
    if not os.path.exists(marker_path):
        return None
    with open(marker_path, 'r', encoding='utf-8') as handle:
        marker = json.load(handle)
    if marker.get('checkpoint_schema_version') != CHECKPOINT_SCHEMA_VERSION or marker.get('status') != 'complete' or marker.get('manifest_sha256') != manifest_sha256:
        raise RuntimeError('Fixed-split completion marker did not match the validated run manifest.')
    artifacts = marker.get('artifacts')
    if not isinstance(artifacts, dict) or not artifacts:
        raise RuntimeError('Fixed-split completion marker did not list its saved artifacts.')
    for relative, record in artifacts.items():
        path = os.path.join(out_dir, *relative.split('/'))
        if not os.path.exists(path):
            raise RuntimeError(f'A checkpointed fixed-split artifact is missing: {relative}')
        if int(os.path.getsize(path)) != int(record.get('size_bytes', -1)):
            raise RuntimeError(f'A checkpointed fixed-split artifact changed size: {relative}')
        if file_sha256(path) != record.get('sha256'):
            raise RuntimeError(f'A checkpointed fixed-split artifact changed content: {relative}')
    paths = _primary_reload_paths(out_dir)
    table1 = pd.read_csv(paths['table1'], keep_default_na=False)
    table2 = pd.read_csv(paths['table2'], keep_default_na=False)
    table3 = pd.read_csv(paths['table3'])
    tuning = pd.read_csv(paths['tuning'])
    metrics_long = pd.read_csv(os.path.join(out_dir, 'tables', 'supplementary_table_s5_fixed_split_by_imputation.csv'), float_precision='round_trip')
    completeness = pd.read_csv(os.path.join(out_dir, 'metadata', 'metric_completeness_fixed_split.csv'))
    if len(table1) != 17:
        raise RuntimeError(f'Fixed-split Table 1 contained {len(table1)} rows; expected 17.')
    if len(table2) != len(PREDICTOR_COLUMNS):
        raise RuntimeError(f'Fixed-split Table 2 contained {len(table2)} rows; expected {len(PREDICTOR_COLUMNS)}.')
    _validate_primary_metrics(table3, metrics_long, completeness, m_imputations)
    expected_tuning_rows = int(len(MODEL_ORDER) * m_imputations)
    if len(tuning) != expected_tuning_rows or set(tuning['model']) != set(MODEL_ORDER) or set(tuning['imputation'].astype(int)) != set(range(1, m_imputations + 1)) or (not (tuning['repeat'].astype(int) == 0).all()) or tuning.duplicated(['imputation', 'model']).any():
        raise RuntimeError('Fixed-split tuning checkpoint did not contain every model/imputation key.')
    for value in tuning['best_params_json']:
        json.loads(value)
    LOGGER.info('Resumed the fully validated supplementary fixed split and its saved artifacts.')
    return (table1, table2, table3, tuning)

def load_primary_checkpoint(out_dir: str, checkpoint_dir: str, manifest_sha256: str, m_imputations: int) -> Optional[Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]]:
    try:
        return _load_primary_checkpoint_strict(out_dir=out_dir, checkpoint_dir=checkpoint_dir, manifest_sha256=manifest_sha256, m_imputations=m_imputations)
    except Exception as exc:
        LOGGER.warning('Fixed-split checkpoint was not fully valid and will be recomputed: %s', exc)
        return None

def run_functional_smoke_test(csv_path: str, seed: int, n_jobs: int) -> None:
    print('Starting functional model/metric/SHAP preflight...', flush=True)
    set_global_seed(seed)
    df = load_dataset(csv_path)
    support_splits: Dict[int, Tuple[np.ndarray, np.ndarray]] = {}
    validated_split_ids = {1, 23, 200}
    for split_id, support_train_idx, support_test_idx in iter_repeated_splits(df, seed=seed, n_splits=PRESPECIFIED_N_SPLITS, test_size=PRESPECIFIED_TEST_SIZE):
        if float(df.iloc[support_train_idx][TIME_COL].max()) <= 120.0:
            raise RuntimeError(f'Repeated split {split_id} lacks training censoring support beyond 120 months.')
        if split_id in validated_split_ids:
            support_splits[split_id] = (support_train_idx, support_test_idx)
    if set(support_splits) != validated_split_ids:
        raise RuntimeError('Could not reconstruct every validated repeated split.')
    for support_split_id in sorted(validated_split_ids):
        support_train_idx, support_test_idx = support_splits[support_split_id]
        support_train = df.iloc[support_train_idx].reset_index(drop=True)
        support_test = df.iloc[support_test_idx].reset_index(drop=True)
        if support_split_id == 23 and (not float(support_test[TIME_COL].max()) > float(support_train[TIME_COL].max())):
            raise RuntimeError('The validated upper-support split no longer matches the data.')
        support_times = np.arange(12.0, 121.0, 1.0, dtype=float)
        baseline_risk = np.clip(support_test['age'].fillna(support_train['age'].median()).to_numpy(dtype=float) / 500.0, 0.01, 0.5)
        support_survival = np.exp(-baseline_risk.reshape(-1, 1) * (support_times / 120.0).reshape(1, -1))
        support_metrics = compute_brier_metrics(get_surv_array(support_train), get_surv_array(support_test), support_survival, support_times, 120.0)
        support_uno = uno_c_index(get_surv_array(support_train), get_surv_array(support_test), baseline_risk, tau=np.nextafter(120.0, np.inf))
        if not np.isfinite([support_metrics['brier_score_at_horizon'], support_metrics['integrated_brier_score'], support_uno]).all():
            raise RuntimeError(f'IPCW support preflight was incomplete for repeated split {support_split_id}.')
    print('  repeated-split stream and IPCW support cases passed', flush=True)
    train_idx, test_idx = primary_split(df, seed=seed, test_size=0.2)
    raw_train_df = df.iloc[train_idx].reset_index(drop=True)
    raw_test_df = df.iloc[test_idx].reset_index(drop=True)
    smoke_split_base_seed = seed + 100000
    completed = generate_imputed_split_datasets(raw_train_df, raw_test_df, m=1, seed=smoke_split_base_seed, max_iter=3)[0]
    train_df = completed['train']
    test_df = completed['test']
    smoke_inner_splits = generate_nested_inner_imputed_splits(raw_outer_train_df=raw_train_df, inner_folds=3, split_seed=smoke_split_base_seed + 1, imputation_seed=smoke_split_base_seed + 100, max_iter=3)
    smoke_neural_selection = generate_nested_neural_selection_split(raw_outer_train_df=raw_train_df, split_seed=smoke_split_base_seed + 1, imputation_seed=smoke_split_base_seed + 500, max_iter=3)
    smoke_parameters: Dict[str, Dict[str, Any]] = {'CoxPH': {'model__ties': 'breslow', 'model__tol': 1e-07, 'model__n_iter': 100}, 'ElasticNetCox': {'model__penalizer': 0.01, 'model__l1_ratio': 0.5}, 'RSF': {'model__n_estimators': 12, 'model__min_samples_split': 10, 'model__min_samples_leaf': 5, 'model__max_features': 'sqrt', 'model__max_depth': 3}, 'GBSA': {'model__learning_rate': 0.05, 'model__n_estimators': 12, 'model__min_samples_leaf': 5, 'model__max_depth': 1, 'model__subsample': 0.7}, 'CoxTime': {'num_nodes': [8], 'dropout': 0.1, 'lr': 0.01, 'batch_size': 64, 'epochs': 12, 'refit_epochs': 12}, 'XGBoost': {'model__n_estimators': 12, 'model__learning_rate': 0.1, 'model__max_depth': 2, 'model__subsample': 0.7, 'model__colsample_bytree': 0.7, 'model__min_child_weight': 1.0, 'model__reg_lambda': 1.0}, 'SVM': {'model__alpha': 0.1, 'model__optimizer': 'avltree', 'model__max_iter': 100, 'model__rank_ratio': 1.0}}
    locked_specs = get_classical_model_specs(random_state=seed + 2000, n_jobs=max(1, int(n_jobs)))
    locked_candidate_counts = {model_name: len(list(ParameterGrid(locked_specs[model_name]['param_grid']))) for model_name in ['CoxPH', 'ElasticNetCox', 'RSF', 'GBSA', 'XGBoost', 'SVM']}
    if locked_candidate_counts != {'CoxPH': 1, 'ElasticNetCox': 8, 'RSF': 6, 'GBSA': 6, 'XGBoost': 8, 'SVM': 4}:
        raise RuntimeError(f'Functional preflight found an unlocked classical search: {locked_candidate_counts}.')
    deepsurv_production_configs = get_deepsurv_configs()
    if len(deepsurv_production_configs) != 6 or {config.get('batch_size') for config in deepsurv_production_configs} != {'full_training_risk_set'}:
        raise RuntimeError('Functional preflight did not find exactly six DeepSurv configurations using the complete training risk set.')
    deepsurv_smoke_configs: List[Dict[str, Any]] = []
    for production_config in deepsurv_production_configs:
        smoke_config = dict(production_config)
        smoke_config['epochs'] = 12
        smoke_config.pop('refit_epochs', None)
        deepsurv_smoke_configs.append(smoke_config)
    coxtime_production_configs = get_coxtime_configs()
    if len(coxtime_production_configs) != 6:
        raise RuntimeError('Functional preflight did not find exactly six CoxTime configurations.')
    coxtime_smoke_configs: List[Dict[str, Any]] = []
    for production_config in coxtime_production_configs:
        smoke_config = dict(production_config)
        smoke_config['epochs'] = 12
        smoke_config.pop('refit_epochs', None)
        coxtime_smoke_configs.append(smoke_config)
    fitted_models: Dict[str, Any] = {}
    absolute_risk_predictions: Dict[str, np.ndarray] = {}
    smoke_metrics_by_model: Dict[str, Dict[str, Any]] = {}
    smoke_tuning_by_model: Dict[str, Dict[str, Any]] = {}
    for model_name in MODEL_ORDER:
        model_seed = smoke_split_base_seed + 1
        if model_name == 'DeepSurv':
            fitted, best_params, selection_score = fit_deepsurv(train_df, random_state=model_seed, configs=deepsurv_smoke_configs, selection_split=smoke_neural_selection, require_all_candidates=True, ibs_start_month=12.0, horizon_months=120.0)
        elif model_name == 'CoxTime':
            fitted, best_params, selection_score = fit_coxtime(train_df, random_state=model_seed, configs=coxtime_smoke_configs, selection_split=smoke_neural_selection, ibs_start_month=12.0, horizon_months=120.0, require_all_candidates=True)
        else:
            smoke_specs = get_classical_model_specs(random_state=model_seed, n_jobs=max(1, int(n_jobs)))
            smoke_specs[model_name]['param_grid'] = {parameter: [value] for parameter, value in smoke_parameters[model_name].items()}
            fitted, best_params, selection_score = fit_gridsearched_model(model_name, train_df, inner_folds=3, random_state=model_seed, n_jobs=max(1, int(n_jobs)), specs=smoke_specs, inner_completed_splits=smoke_inner_splits, ibs_start_month=12.0, horizon_months=120.0)
        selection_metadata = tuning_output_metadata(model_name, selection_score)
        if model_name == 'DeepSurv' and best_params.get('batch_size') != 'full_training_risk_set':
            raise RuntimeError('Functional DeepSurv preflight did not retain full-risk-set batching.')
        risk = np.asarray(fitted.predict(test_df[PREDICTOR_COLUMNS]), dtype=float).reshape(-1)
        if risk.shape != (len(test_df),) or not np.isfinite(risk).all():
            raise RuntimeError(f'Functional preflight returned invalid risk predictions for {model_name}.')
        metrics, _, predicted_event_probability = evaluate_model(model_name, fitted, train_df, test_df, horizon_months=120.0, ibs_start_month=12.0, calibration_groups=5)
        require_complete_metrics(model_name, metrics, 'Functional preflight')
        smoke_metrics_by_model[model_name] = metrics
        smoke_tuning_by_model[model_name] = {'split': 'repeated_200_stratified_80_20', 'split_id': 1, 'model': model_name, 'tuning_method': main_tuning_method(model_name), **selection_metadata, 'best_params_json': json.dumps(best_params, sort_keys=True), 'split_base_seed': int(smoke_split_base_seed), 'outer_imputation_seed': int(smoke_split_base_seed), 'model_seed': int(smoke_split_base_seed + 1), 'classical_inner_imputation_seed_base': int(smoke_split_base_seed + 100), 'neural_development_imputation_seed': int(smoke_split_base_seed + 500)}
        if predicted_event_probability is not None:
            absolute_risk_predictions[model_name] = np.asarray(predicted_event_probability, dtype=float)
        if model_name != 'SVM':
            evaluation_times = np.asarray([12.0, 60.0, 120.0])
            survival = predict_survival_matrix(fitted, test_df[PREDICTOR_COLUMNS], evaluation_times)
            if survival.shape != (len(test_df), len(evaluation_times)):
                raise RuntimeError(f'Functional preflight returned the wrong survival shape for {model_name}.')
            if not np.isfinite(survival).all() or np.any(survival < -1e-08) or np.any(survival > 1.0 + 1e-08) or np.any(np.diff(survival, axis=1) > 1e-06):
                raise RuntimeError(f'Functional preflight returned invalid survival probabilities for {model_name}.')
        fitted_models[model_name] = fitted
        print(f'  {model_name}: fit, prediction and metric APIs passed', flush=True)
    from sklearn.ensemble import ExtraTreesRegressor
    flexible_imputer = IterativeImputer(estimator=ExtraTreesRegressor(n_estimators=8, min_samples_leaf=2, max_features=1.0, random_state=seed + 3000, n_jobs=max(1, int(n_jobs))), max_iter=2, sample_posterior=False, random_state=seed + 3000, initial_strategy='median', imputation_order='ascending')
    flexible_train = flexible_imputer.fit_transform(raw_train_df[PREDICTOR_COLUMNS])
    flexible_test = flexible_imputer.transform(raw_test_df[PREDICTOR_COLUMNS])
    if not np.isfinite(flexible_train).all() or not np.isfinite(flexible_test).all():
        raise RuntimeError('Functional ExtraTrees-imputation preflight returned non-finite values.')
    print('  ExtraTrees iterative-imputation API passed', flush=True)
    smoke_split_map = build_repeated_split_map(df, seed=seed, n_splits=PRESPECIFIED_N_SPLITS, test_size=PRESPECIFIED_TEST_SIZE)
    risk_basis = df['age'].fillna(df['age'].median()).rank(method='average', pct=True)
    risk_lookup = pd.DataFrame({STUDY_ID_COL: df[STUDY_ID_COL].to_numpy(), 'risk_basis': 0.02 + 0.18 * risk_basis.to_numpy(dtype=float)})
    smoke_predictions = smoke_split_map[['split_id', STUDY_ID_COL]].merge(risk_lookup, on=STUDY_ID_COL, how='left', validate='many_to_one')
    for model_index, model_name in enumerate(CALIBRATION_MODELS):
        smoke_predictions[model_name] = np.clip(smoke_predictions['risk_basis'] * (0.9 + 0.02 * model_index) + 2e-06 * smoke_predictions['split_id'], 0.0001, 0.95)
    smoke_predictions = smoke_predictions.drop(columns=['risk_basis'])
    validate_repeated_held_out_predictions(smoke_predictions, split_map=smoke_split_map, df=df, n_splits=PRESPECIFIED_N_SPLITS)
    smoke_performance_rows: List[Dict[str, Any]] = []
    for split_id in range(1, PRESPECIFIED_N_SPLITS + 1):
        for model_index, model_name in enumerate(MODEL_ORDER):
            smoke_performance_rows.append({'split_id': split_id, 'n_test': 249, 'model': model_name, 'harrell_c': 0.6 + 0.0001 * split_id + 0.001 * model_index, 'uno_c_tau': 0.59 + 0.0001 * split_id + 0.001 * model_index, 'brier_score_at_horizon': np.nan if model_name == 'SVM' else 0.1 + 1e-05 * split_id, 'integrated_brier_score': np.nan if model_name == 'SVM' else 0.08 + 1e-05 * split_id})
    smoke_performance = pd.DataFrame(smoke_performance_rows)
    smoke_summary, smoke_formatted = summarise_performance(smoke_performance, n_splits=PRESPECIFIED_N_SPLITS)
    smoke_paired_by_split = build_paired_differences_by_split(smoke_performance)
    smoke_paired = compare_models_vs_reference(smoke_performance)
    if len(smoke_summary) != len(MODEL_ORDER) or len(smoke_formatted) != len(MODEL_ORDER) or len(smoke_paired_by_split) != 5200 or (len(smoke_paired) != 26):
        raise RuntimeError('Functional repeated-split summary or paired comparison is incomplete.')
    print('  exact 200-split map, 49,800 predictions and split aggregation passed', flush=True)
    with tempfile.TemporaryDirectory(prefix='ascvd_functional_preflight_') as temporary:
        checkpoint_dir = ensure_dir(os.path.join(temporary, 'repeated_split_checkpoint_smoke'))
        smoke_manifest = 'functional-smoke-manifest'
        split_metric_frames: List[pd.DataFrame] = []
        split_tuning_frames: List[pd.DataFrame] = []
        for model_name in MODEL_ORDER:
            metric_frame = pd.DataFrame([{'split_id': 1, 'n_train': int(len(train_df)), 'n_test': int(len(test_df)), 'train_events': int(raw_train_df[EVENT_COL].sum()), 'test_events': int(raw_test_df[EVENT_COL].sum()), 'model': model_name, **smoke_metrics_by_model[model_name]}])
            tuning_frame = pd.DataFrame([smoke_tuning_by_model[model_name]])
            prediction_frame = pd.DataFrame({STUDY_ID_COL: test_df[STUDY_ID_COL].to_numpy(), 'predicted_event_probability': np.full(len(test_df), np.nan, dtype=float) if model_name == 'SVM' else absolute_risk_predictions[model_name]})
            write_repeated_split_model_checkpoint(checkpoint_dir=checkpoint_dir, manifest_sha256=smoke_manifest, split_id=1, model_name=model_name, metrics_df=metric_frame, tuning_df=tuning_frame, predictions_df=prediction_frame, expected_test_ids=test_df[STUDY_ID_COL])
            loaded_model = load_repeated_split_model_checkpoint(checkpoint_dir=checkpoint_dir, manifest_sha256=smoke_manifest, split_id=1, model_name=model_name, expected_test_ids=test_df[STUDY_ID_COL])
            if loaded_model is None:
                raise RuntimeError(f'Functional partial-checkpoint preflight failed for {model_name}.')
            split_metric_frames.append(metric_frame)
            split_tuning_frames.append(tuning_frame)
        split_predictions = pd.DataFrame({'split_id': 1, STUDY_ID_COL: test_df[STUDY_ID_COL].to_numpy(), **absolute_risk_predictions})
        split_metrics = pd.concat(split_metric_frames, ignore_index=True)
        split_tuning = pd.concat(split_tuning_frames, ignore_index=True)
        write_repeated_split_checkpoint(checkpoint_dir=checkpoint_dir, manifest_sha256=smoke_manifest, split_id=1, metrics_df=split_metrics, tuning_df=split_tuning, predictions_df=split_predictions, expected_test_ids=test_df[STUDY_ID_COL])
        if load_repeated_split_checkpoint(checkpoint_dir=checkpoint_dir, manifest_sha256=smoke_manifest, split_id=1, expected_test_ids=test_df[STUDY_ID_COL]) is None:
            raise RuntimeError('Functional complete-split checkpoint preflight failed.')
        print('  per-model and complete-split checkpoint APIs passed', flush=True)
        smoke_repeated_out = ensure_dir(os.path.join(temporary, 'repeated_calibration_smoke'))
        smoke_calibration_summary = create_repeated_split_calibration_outputs(df=df, held_out_predictions=smoke_predictions, split_map=smoke_split_map, n_splits=PRESPECIFIED_N_SPLITS, horizon_months=120.0, calibration_groups=5, bootstrap_reps=5, out_dir=smoke_repeated_out, seed=seed)
        if len(smoke_calibration_summary) != len(CALIBRATION_MODELS):
            raise RuntimeError('Functional participant-level calibration aggregation was incomplete.')
        print('  participant-averaged calibration and conditional bootstrap passed', flush=True)
        for model_name in MODEL_ORDER:
            require_force = model_name in {'CoxPH', 'DeepSurv'}
            shap_table = make_shap_outputs(model_name=model_name, fitted_model=fitted_models[model_name], train_df=train_df, test_df=test_df, out_dir=temporary, make_force_plot=require_force, max_background=8, max_explain=2)
            summary_path = os.path.join(temporary, f'shap_summary_{model_name}.png')
            if shap_table is None or len(shap_table) != len(PREDICTOR_COLUMNS) or (not np.isfinite(shap_table['mean_abs_shap'].to_numpy(float)).all()) or (not os.path.isfile(summary_path)) or (os.path.getsize(summary_path) == 0):
                raise RuntimeError(f'Functional SHAP preflight failed for {model_name}.')
            if require_force:
                force_path = os.path.join(temporary, f'shap_force_{model_name}.png')
                if not os.path.isfile(force_path) or os.path.getsize(force_path) == 0:
                    raise RuntimeError(f'Functional SHAP force-plot preflight failed for {model_name}.')
            print(f'  {model_name}: SHAP API passed', flush=True)
        flexible_curve = flexible_calibration_curve_with_bootstrap(test_df, absolute_risk_predictions['CoxPH'], time_point=120.0, n_bootstrap=5, random_state=seed + 4000, n_grid=21)
        if int(flexible_curve['n_bootstrap_successful'].iloc[0]) < 4 or not np.isfinite(flexible_curve[['predicted_risk', 'observed_risk_smoothed']].to_numpy(float)).all():
            raise RuntimeError('Functional flexible-calibration preflight was incomplete.')
        calibration_path = os.path.join(temporary, 'flexible_calibration_smoke.csv')
        flexible_curve.to_csv(calibration_path, index=False)
        grouped_tables = {model_name: grouped_calibration_table(test_df, predicted_probability, time_point=120.0, n_groups=5) for model_name, predicted_probability in absolute_risk_predictions.items()}
        grouped_figure_path = os.path.join(temporary, 'combined_grouped_calibration_smoke.png')
        make_combined_calibration_figure(grouped_tables, grouped_figure_path, horizon_months=120.0)
        flexible_figure_path = os.path.join(temporary, 'combined_flexible_calibration_smoke.png')
        make_combined_flexible_calibration_figure({model_name: flexible_curve.copy() for model_name in absolute_risk_predictions}, flexible_figure_path, horizon_months=120.0)
        for figure_path in [grouped_figure_path, flexible_figure_path]:
            if not os.path.isfile(figure_path) or os.path.getsize(figure_path) == 0:
                raise RuntimeError('Functional combined-calibration figure preflight failed.')
            tiff_path = os.path.splitext(figure_path)[0] + '.tiff'
            if not os.path.isfile(tiff_path) or os.path.getsize(tiff_path) == 0:
                raise RuntimeError('Functional LZW TIFF calibration figure was not created.')
            with Image.open(tiff_path) as figure_image:
                if figure_image.width <= figure_image.height:
                    raise RuntimeError('Functional calibration composite is not landscape.')
                compression = str(figure_image.info.get('compression', '')).lower()
                if compression not in {'tiff_lzw', 'lzw'}:
                    raise RuntimeError('Functional calibration TIFF is not LZW compressed.')
        table_frame = flexible_curve.head(5).copy()
        write_tables_docx(temporary, table_frame, table_frame, table_frame, table_frame, table_frame, table_frame, table_frame)
        docx_dir = os.path.join(temporary, 'docx_tables')
        expected_docx_names = ['ASCVD_survival_models_tables.docx', 'Table_1.docx', 'Table_2.docx', 'Table_3.docx', 'Supplementary_Table_S1.docx', 'Supplementary_Table_S2.docx', 'Supplementary_Table_S3.docx', 'Supplementary_Table_S5.docx']
        for docx_name in expected_docx_names:
            docx_path = os.path.join(docx_dir, docx_name)
            if not os.path.isfile(docx_path) or os.path.getsize(docx_path) == 0:
                raise RuntimeError(f'Functional DOCX-generation preflight failed for {docx_name}.')
            reopened = Document(docx_path)
            if not reopened.tables:
                raise RuntimeError(f'Functional DOCX reopen preflight failed for {docx_name}.')
            document_text = '\n'.join((paragraph.text for paragraph in reopened.paragraphs))
            if docx_name == 'Table_3.docx' and '200 stratified random 80/20 splits' not in document_text:
                raise RuntimeError('Functional DOCX preflight did not map repeated 80/20 splits to main Table 3.')
            if docx_name == 'Supplementary_Table_S3.docx' and 'Paired performance differences relative to CoxPH' not in document_text:
                raise RuntimeError('Functional DOCX preflight did not map paired comparisons to Supplementary Table S3.')
            if docx_name == 'Supplementary_Table_S5.docx' and 'fixed train-test split' not in document_text:
                raise RuntimeError('Functional DOCX preflight did not map fixed-split performance to Supplementary Table S5.')
        combined_shap_path = os.path.join(temporary, 'combined_shap_smoke.png')
        combine_shap_summary_images([os.path.join(temporary, 'shap_summary_CoxPH.png'), os.path.join(temporary, 'shap_summary_DeepSurv.png')], ['CoxPH', 'DeepSurv'], combined_shap_path, ncols=2)
        if not os.path.isfile(combined_shap_path) or os.path.getsize(combined_shap_path) == 0:
            raise RuntimeError('Functional combined-image preflight failed.')
        print('  grouped/flexible calibration composites, combined DOCX tables, and combined-image APIs passed', flush=True)
    print('Functional model, metric, flexible-imputation, calibration, SHAP and DOCX preflight passed.', flush=True)

def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description='Fixed-split multiple-imputation analysis plus 200 repeated stratified 80/20 splits for survival-model comparison.')
    parser.add_argument('--csv', default=DEFAULT_CSV_PATH, help='Path to the raw CSV with missing baseline predictors.')
    parser.add_argument('--outdir', default=DEFAULT_OUTDIR, help='Output directory.')
    parser.add_argument('--seed', type=int, default=PRESPECIFIED_SPLITTER_SEED, help='Master splitter seed; analysis setting is exactly 20260320.')
    parser.add_argument('--test-size', type=float, default=PRESPECIFIED_TEST_SIZE, help='Held-out fraction for both the repeated and fixed-split analyses; exactly 0.20.')
    parser.add_argument('--n-splits', type=int, default=PRESPECIFIED_N_SPLITS, help='Repeated 80/20 splits; analysis setting is exactly 200.')
    parser.add_argument('--inner-folds', type=int, default=3, help='Inner folds for classical models; analysis setting is exactly 3.')
    parser.add_argument('--n-jobs', type=int, default=1, help='Parallel grid-search workers. The conservative laptop default is 1.')
    parser.add_argument('--fixed-split-imputations', type=int, default=10, help='Imputations for the separate fixed 80/20 analysis only.')
    parser.add_argument('--impute-max-iter', type=int, default=20, help='Maximum iterations for each predictor-only iterative imputer.')
    parser.add_argument('--horizon-months', type=float, default=120.0, help='Prediction horizon in months.')
    parser.add_argument('--ibs-start-month', type=float, default=12.0, help='Start month for integrated Brier score.')
    parser.add_argument('--calibration-groups', type=int, default=5, help='Risk groups for grouped calibration.')
    parser.add_argument('--calibration-bootstrap-reps', type=int, default=1000, help='Participant-bootstrap replicates for fixed-split and participant-averaged flexible-calibration bands.')
    parser.add_argument('--max-session-hours', type=float, default=0.0, help='Optional wall-time limit. Zero is unlimited. A positive value is checked only after a fitted model or complete repeated split; a planned pause exits with code 75.')
    parser.add_argument('--skip-shap', action='store_true', help='Skip SHAP outputs for the fixed split.')
    parser.add_argument('--shap-max-background', type=int, default=100, help='Maximum SHAP background size.')
    parser.add_argument('--shap-max-explain', type=int, default=150, help='Maximum SHAP explanation rows.')
    parser.add_argument('--resume', action='store_true', help='Resume only hash-validated fixed-imputation and repeated-split checkpoints.')
    parser.add_argument('--validate-only', action='store_true', help='Validate packages and input data, then exit without fitting.')
    parser.add_argument('--functional-smoke-test', action='store_true', help='Exercise all model, metric, split aggregation, checkpoint, figure and DOCX APIs, then exit.')
    return parser.parse_args()

def main() -> None:
    args = parse_args()
    if int(args.seed) != PRESPECIFIED_SPLITTER_SEED:
        raise ValueError(f'--seed must be {PRESPECIFIED_SPLITTER_SEED} for the analysis.')
    if not math.isclose(float(args.test_size), PRESPECIFIED_TEST_SIZE, rel_tol=0.0, abs_tol=1e-12):
        raise ValueError(f'--test-size must be exactly {PRESPECIFIED_TEST_SIZE:.2f}.')
    if args.n_splits != PRESPECIFIED_N_SPLITS:
        raise ValueError(f'--n-splits must be {PRESPECIFIED_N_SPLITS} for the analysis.')
    if args.inner_folds != 3:
        raise ValueError('--inner-folds must be 3 for classical-model tuning.')
    if args.n_jobs == 0:
        raise ValueError('--n-jobs cannot be zero.')
    if args.fixed_split_imputations != 10:
        raise ValueError('--fixed-split-imputations must be 10 for the fixed-split analysis.')
    if args.impute_max_iter < 1:
        raise ValueError('--impute-max-iter must be at least 1.')
    if float(args.horizon_months) != PRESPECIFIED_HORIZON_MONTHS:
        raise ValueError('--horizon-months must be exactly 120 for the analysis.')
    if float(args.ibs_start_month) != PRESPECIFIED_IBS_START_MONTH:
        raise ValueError('--ibs-start-month must be exactly 12 for the analysis.')
    if args.calibration_groups < 2:
        raise ValueError('--calibration-groups must be at least 2.')
    if args.calibration_bootstrap_reps < 1:
        raise ValueError('--calibration-bootstrap-reps must be at least 1.')
    if not np.isfinite(args.max_session_hours) or args.max_session_hours < 0.0:
        raise ValueError('--max-session-hours must be a finite non-negative value.')
    if not os.path.exists(args.csv):
        raise FileNotFoundError('Could not find the analysis CSV.')
    validate_runtime_and_data(args.csv)
    if args.validate_only:
        return
    if args.functional_smoke_test:
        run_functional_smoke_test(args.csv, seed=args.seed, n_jobs=args.n_jobs)
        return
    deadline_monotonic = None if args.max_session_hours == 0.0 else time.monotonic() + 3600.0 * float(args.max_session_hours)
    out_dir = ensure_dir(args.outdir)
    ensure_dir(os.path.join(out_dir, 'tables'))
    ensure_dir(os.path.join(out_dir, 'figures'))
    meta_dir = ensure_dir(os.path.join(out_dir, 'metadata'))
    setup_logging(os.path.join(meta_dir, 'run.log'))
    set_global_seed(args.seed)
    LOGGER.info('Loading analysis data.')
    df = load_dataset(args.csv)
    LOGGER.info('Loaded %d eligible participants.', len(df))
    repeated_checkpoint_dir = ensure_dir(os.path.join(out_dir, 'checkpoints', 'repeated_200_splits_ibs_v1'))
    repeated_configuration = repeated_split_checkpoint_configuration(df=df, seed=args.seed, n_splits=args.n_splits, test_size=args.test_size, inner_folds=args.inner_folds, horizon_months=args.horizon_months, ibs_start_month=args.ibs_start_month, calibration_groups=args.calibration_groups, n_jobs=args.n_jobs, impute_max_iter=args.impute_max_iter, calibration_bootstrap_reps=args.calibration_bootstrap_reps)
    repeated_manifest_sha256 = prepare_checkpoint_manifest(repeated_checkpoint_dir, run_type='repeated_200_splits_ibs_tuning', configuration=repeated_configuration, resume=args.resume)
    primary_checkpoint_dir = ensure_dir(os.path.join(out_dir, 'checkpoints', 'primary_split_ibs_v1'))
    primary_configuration = primary_checkpoint_configuration(df, args)
    primary_manifest_sha256 = prepare_checkpoint_manifest(primary_checkpoint_dir, run_type='primary_split_ibs_tuning', configuration=primary_configuration, resume=args.resume)
    status_path = os.path.join(meta_dir, 'analysis_status.json')
    completion_path = os.path.join(meta_dir, 'analysis_complete.json')
    if args.resume and os.path.isfile(completion_path):
        try:
            with open(completion_path, 'r', encoding='utf-8') as handle:
                prior_completion = json.load(handle)
        except Exception as exc:
            LOGGER.warning('Ignoring a malformed completion marker and rebuilding final outputs: %s', exc)
        else:
            if not (prior_completion.get('status') == 'complete' and prior_completion.get('repeated_manifest_sha256') == repeated_manifest_sha256 and (prior_completion.get('fixed_split_manifest_sha256') == primary_manifest_sha256)):
                raise RuntimeError('An analysis_complete.json file exists but does not match the current manifests. Use a new output directory.')
            LOGGER.info('A matching completion marker exists, but the backend was invoked because saved outputs require revalidation. Rebuilding aggregates from validated checkpoints.')
    atomic_write_json({'status': 'running', 'repeated_manifest_sha256': repeated_manifest_sha256, 'fixed_split_manifest_sha256': primary_manifest_sha256, 'completion_marker': 'analysis_complete.json is written only after all stages validate.'}, status_path)
    missingness_df = create_missingness_table(df)
    atomic_write_dataframe_csv(missingness_df, os.path.join(out_dir, 'tables', 'supplementary_table_s1_missingness.csv'))
    package_versions = get_package_versions()
    save_json(package_versions, os.path.join(meta_dir, 'package_versions.json'))
    config_dump = vars(args).copy()
    config_dump['csv'] = 'analysis_data.csv'
    config_dump['outdir'] = '.'
    config_dump['predictors'] = PREDICTOR_COLUMNS
    config_dump['continuous_columns'] = CONTINUOUS_COLUMNS
    config_dump['binary_columns'] = BINARY_COLUMNS
    config_dump['model_descriptions'] = MODEL_DESCRIPTIONS
    config_dump['model_selection_policy'] = {model_name: model_selection_metadata(model_name) for model_name in MODEL_ORDER}
    save_json(config_dump, os.path.join(meta_dir, 'analysis_config.json'))
    try:
        primary_loaded = None
        if args.resume:
            primary_loaded = load_primary_checkpoint(out_dir=out_dir, checkpoint_dir=primary_checkpoint_dir, manifest_sha256=primary_manifest_sha256, m_imputations=args.fixed_split_imputations)
        if primary_loaded is not None:
            table1, table2, table3, primary_tuning_df = primary_loaded
        else:
            table1, table2, table3, primary_tuning_df = run_primary_split_analysis(df=df, out_dir=out_dir, seed=args.seed, test_size=args.test_size, inner_folds=args.inner_folds, horizon_months=args.horizon_months, ibs_start_month=args.ibs_start_month, calibration_groups=args.calibration_groups, n_jobs=args.n_jobs, m_imputations=args.fixed_split_imputations, impute_max_iter=args.impute_max_iter, make_shap=not args.skip_shap, shap_max_background=args.shap_max_background, shap_max_explain=args.shap_max_explain, calibration_bootstrap_reps=args.calibration_bootstrap_reps, resume=args.resume, checkpoint_dir=primary_checkpoint_dir, manifest_sha256=primary_manifest_sha256, deadline_monotonic=deadline_monotonic)
            write_primary_checkpoint_marker(out_dir=out_dir, checkpoint_dir=primary_checkpoint_dir, manifest_sha256=primary_manifest_sha256, shap_requested=bool(not args.skip_shap), m_imputations=args.fixed_split_imputations)
            LOGGER.info('Checkpointed all fixed-split outputs.')
        performance_df, table4, supplementary_s3 = run_repeated_200_splits(df=df, out_dir=out_dir, seed=args.seed, n_splits=args.n_splits, test_size=args.test_size, inner_folds=args.inner_folds, horizon_months=args.horizon_months, ibs_start_month=args.ibs_start_month, calibration_groups=args.calibration_groups, n_jobs=args.n_jobs, impute_max_iter=args.impute_max_iter, calibration_bootstrap_reps=args.calibration_bootstrap_reps, resume=args.resume, prepared_manifest_sha256=repeated_manifest_sha256, deadline_monotonic=deadline_monotonic)
    except SafeStopRequested as exc:
        atomic_write_json({'status': 'paused', 'exit_code': 75, 'last_complete_boundary': exc.boundary, 'resume_instruction': 'Rerun the unchanged script and data with --resume. Only hash-validated completed units will be reused.', 'repeated_manifest_sha256': repeated_manifest_sha256, 'fixed_split_manifest_sha256': primary_manifest_sha256}, status_path)
        LOGGER.info('Planned safe pause after %s.', exc.boundary)
        raise
    supplementary_s2 = build_supplementary_tuning_table(primary_tuning_df)
    atomic_write_dataframe_csv(supplementary_s2, os.path.join(out_dir, 'tables', 'supplementary_table_s2_tuning_strategies.csv'))
    fixed_vs_repeated = build_fixed_vs_repeated_comparison(fixed_metrics=table3, repeated_performance=performance_df, n_splits=args.n_splits)
    atomic_write_dataframe_csv(fixed_vs_repeated, os.path.join(out_dir, 'tables', 'repeated_200_fixed_vs_repeated_comparison.csv'))
    write_tables_docx(out_dir=out_dir, table1=table1, table2=table2, table3=table3, table4=table4, supp_s1=missingness_df, supp_s2=supplementary_s2, supp_s3=supplementary_s3)
    write_methods_notes(out_dir=out_dir, args=args, package_versions=package_versions)
    completion_payload = {'status': 'complete', 'primary_analysis_status': 'primary_internal_validation_200_repeated_stratified_80_20_splits', 'repeated_manifest_sha256': repeated_manifest_sha256, 'fixed_split_manifest_sha256': primary_manifest_sha256, 'filtered_participants': int(len(df)), 'ascvd_events': int(df[EVENT_COL].sum()), 'n_repeated_splits': int(args.n_splits), 'repeated_test_fraction': float(args.test_size), 'completed_repeated_splits': int(performance_df['split_id'].nunique()), 'fixed_split_imputations': int(args.fixed_split_imputations), 'imputations_per_outer_split': 1, 'classical_inner_folds': int(args.inner_folds), 'models_refitted_per_split': len(MODEL_ORDER), 'models_with_hyperparameter_selection_per_split': 7, 'fixed_specification_models': ['CoxPH'], 'model_selection_policy': {model_name: model_selection_metadata(model_name) for model_name in MODEL_ORDER}, 'neural_tuning': 'internal_stratified_80_20_validation_minimum_IBS_12_120_then_full_outer_refit; early_stopping_by_validation_loss', 'main_repeated_validation_table': 'tables/repeated_200_table3_formatted.csv', 'paired_comparison_table': 'tables/repeated_200_paired_differences_vs_coxph_summary.csv', 'paired_comparison_by_split': 'tables/repeated_200_paired_differences_by_split.csv', 'supplementary_missingness_table': 'tables/supplementary_table_s1_missingness.csv', 'main_grouped_calibration_figure': 'figures/repeated_200_grouped_calibration.png', 'main_grouped_calibration_figure_tiff': 'figures/repeated_200_grouped_calibration.tiff', 'repeated_flexible_calibration_figure': 'figures/repeated_200_flexible_calibration.png', 'repeated_flexible_calibration_figure_tiff': 'figures/repeated_200_flexible_calibration.tiff', 'fixed_split_grouped_calibration_figure': 'figures/supplementary_fixed_split_grouped_calibration_all_models.png', 'fixed_split_grouped_calibration_figure_tiff': 'figures/supplementary_fixed_split_grouped_calibration_all_models.tiff', 'fixed_split_flexible_calibration_figure': 'figures/supplementary_fixed_split_flexible_calibration_all_models.png', 'fixed_split_flexible_calibration_figure_tiff': 'figures/supplementary_fixed_split_flexible_calibration_all_models.tiff', 'held_out_prediction_table': 'tables/repeated_200_held_out_predictions.csv', 'split_level_performance_table': 'tables/repeated_200_performance_by_split.csv', 'split_level_tuning_table': 'tables/repeated_200_tuning_by_split.csv', 'split_map': 'metadata/repeated_200_split_map.csv', 'participant_calibration_summary': 'tables/repeated_200_participant_level_calibration_summary.csv', 'participant_averaged_risk_table': 'tables/repeated_200_mean_held_out_risk_per_participant.csv', 'participant_coverage_table': 'metadata/repeated_200_held_out_coverage.csv', 'inference_note': 'metadata/repeated_200_inference_and_multiplicity_note.txt', 'fixed_vs_repeated_comparison': 'tables/repeated_200_fixed_vs_repeated_comparison.csv', 'calibration_bootstrap_replicates_requested': int(args.calibration_bootstrap_reps), 'all_calibration_figure_layout': 'landscape_panels_no_overall_title_or_footer_common_square_axes', 'calibration_figure_formats': 'PNG_600dpi_and_LZW_TIFF_600dpi', 'shap_requested': bool(not args.skip_shap), 'instruction': 'Review run.log, the metric-completeness file and methods_notes.txt before using results.'}
    atomic_write_json(completion_payload, completion_path)
    atomic_write_json({'status': 'complete', 'completion_marker': os.path.basename(completion_path)}, status_path)
    LOGGER.info('Fixed-split metrics:\n%s', table3.to_string(index=False))
    LOGGER.info('Repeated 200-split summary saved to %s', os.path.join(out_dir, 'tables', 'repeated_200_table3_formatted.csv'))
    LOGGER.info('Done.')
if __name__ == '__main__':
    try:
        main()
    except SafeStopRequested:
        raise SystemExit(75)
''',
    '_analysis_sensitivity.py': r'''from __future__ import annotations
import argparse
import hashlib
import importlib
import importlib.util
import inspect
import json
import logging
import math
import sys
import time
import warnings
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from joblib import Parallel, delayed
from PIL import Image
from sklearn.ensemble import ExtraTreesRegressor
from sklearn.exceptions import ConvergenceWarning
from sklearn.experimental import enable_iterative_imputer
from sklearn.impute import IterativeImputer
from sklearn.model_selection import StratifiedKFold, StratifiedShuffleSplit
LOGGER = logging.getLogger('extra_trees_matched_200_split_sensitivity')
TOTAL_MAIN_SPLITS = 200
SENSITIVITY_SPLITS = TOTAL_MAIN_SPLITS
TEST_SIZE = 0.2
EXPECTED_TEST_EVENTS = 24
EXPECTED_HELDOUT_PREDICTION_ROWS = 49800
EXPECTED_MIN_HELDOUT_PREDICTIONS = 21
EXPECTED_MAX_HELDOUT_PREDICTIONS = 57
INNER_FOLDS = 3
FIXED_IMPUTATION_CYCLES = 10
STABILITY_COMPARISON_CYCLES = 20
STABILITY_SPLIT_IDS = tuple(range(1, SENSITIVITY_SPLITS + 1))
MIN_STABILITY_CELLS_PER_VARIABLE = 10
SENSITIVITY_SEED_OFFSET = 50000000
CHECKPOINT_SCHEMA_VERSION = 5
EXTRA_TREES_N_ESTIMATORS = 100
EXTRA_TREES_MIN_SAMPLES_LEAF = 2
EXTRA_TREES_MAX_FEATURES = 1.0
FLEXIBLE_CALIBRATION_GRID_POINTS = 101
CLASSICAL_IBS_TUNING_METHOD = 'nested_extra_trees_imputation_stratified_3_fold_mean_ibs_12_120'
SVM_TUNING_METHOD = 'nested_extra_trees_imputation_stratified_3_fold_mean_harrell_c'
NEURAL_IBS_TUNING_METHOD = 'development_only_extra_trees_imputation_internal_stratified_80_20_validation_ibs_12_120'
FIXED_TUNING_METHOD = 'fixed_specification_no_hyperparameter_selection'
CLASSICAL_IMPUTATION_POLICY = 'inner_split_specific_extra_trees_shared_across_classical_models'
FIXED_IMPUTATION_POLICY = 'outer_split_extra_trees_completion_fixed_coxph'
NEURAL_IMPUTATION_POLICY = 'development_only_extra_trees_shared_across_neural_models'
FIXED_CYCLE_POLICY = 'prespecified_fixed_10_cycles_tol_0_no_convergence_claim'
EXPECTED_ITERATION_WARNING_TEXT = 'Early stopping criterion not reached'
FOLD_PERFORMANCE_METRICS = ['harrell_c', 'uno_c_tau', 'brier_score_at_horizon', 'integrated_brier_score']
DISCRIMINATION_METRICS = ['harrell_c', 'uno_c_tau']
IMPUTER_DIAGNOSTIC_COLUMNS = ['split_id', 'role', 'inner_fold', 'imputer_seed', 'n_iter', 'max_iter', 'fixed_cycle_policy', 'warning_count', 'warning', 'warning_classification', 'unexpected_warning_count', 'train_missing_cells_before', 'apply_missing_cells_before', 'train_missing_cells_after', 'apply_missing_cells_after', 'observed_cells_preserved', 'all_values_finite', 'plausibility_passed']
STABILITY_COLUMNS = ['split_id', 'imputer_seed', 'variable', 'variable_type', 'n_missing_train', 'n_missing_test', 'n_compared', 'mean_absolute_difference_10_vs_20', 'observed_training_sd', 'normalized_mean_absolute_difference', 'binary_disagreement_rate', 'acceptance_threshold', 'assessment_status', 'threshold_applicable', 'stability_passed', 'cycles_primary', 'cycles_comparator', 'primary_warning', 'comparator_warning']
USER_CSV_PATH = 'analysis_data.csv'
USER_MAIN_SCRIPT_PATH = '_analysis_main.py'
USER_MAIN_OUTPUT_DIR = 'main_outputs'
USER_OUTDIR = 'sensitivity_outputs'
USER_MASTER_SEED = 20260320
USER_HORIZON_MONTHS = 120.0
USER_IBS_START_MONTH = 12.0
USER_CALIBRATION_GROUPS = 5
USER_CALIBRATION_BOOTSTRAP_REPS = 1000
USER_N_JOBS = 2
USER_MAX_SESSION_HOURS = 0.0

class PlannedSensitivityPause(RuntimeError):

    def __init__(self, split_id: int, model_name: str, completed_models: Sequence[str]) -> None:
        super().__init__(f'Safe session limit reached after split {split_id} model {model_name}.')
        self.split_id = int(split_id)
        self.model_name = str(model_name)
        self.completed_models = [str(value) for value in completed_models]
        self.completed_splits: List[int] = []

def setup_logging(log_path: Path) -> None:
    log_path.parent.mkdir(parents=True, exist_ok=True)
    logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s', handlers=[logging.FileHandler(log_path, encoding='utf-8'), logging.StreamHandler()], force=True)

def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description='Matched ExtraTrees sensitivity on all 200 prespecified 80/20 splits from the main 200-split analysis.')
    parser.add_argument('--csv', default=USER_CSV_PATH, help='Raw analysis CSV.')
    parser.add_argument('--main-script', dest='main_script', default=USER_MAIN_SCRIPT_PATH, help='Main 200-split analysis script.')
    parser.add_argument('--main-output-dir', dest='main_output_dir', default=USER_MAIN_OUTPUT_DIR, help='Completed main 200-split output directory.')
    parser.add_argument('--outdir', default=USER_OUTDIR)
    parser.add_argument('--seed', type=int, default=USER_MASTER_SEED)
    parser.add_argument('--impute-cycles', '--impute-max-iter', dest='impute_cycles', type=int, default=FIXED_IMPUTATION_CYCLES, help='Fixed ExtraTrees cycles; the validated value is exactly 10.')
    parser.add_argument('--horizon-months', type=float, default=USER_HORIZON_MONTHS)
    parser.add_argument('--ibs-start-month', type=float, default=USER_IBS_START_MONTH)
    parser.add_argument('--calibration-groups', type=int, default=USER_CALIBRATION_GROUPS)
    parser.add_argument('--calibration-bootstrap-reps', type=int, default=USER_CALIBRATION_BOOTSTRAP_REPS)
    parser.add_argument('--n-jobs', type=int, default=USER_N_JOBS)
    parser.add_argument('--max-session-hours', type=float, default=USER_MAX_SESSION_HOURS, help='Stop safely after the first new model checkpoint at or beyond this time; zero means unlimited.')
    parser.add_argument('--resume', action='store_true', help='Reuse only checkpoints that pass manifest, hash, and value checks.')
    parser.add_argument('--worker-import-smoke-test', action='store_true', help='Exercise the Windows spawned-worker import and ElasticNetCox IBS candidate path.')
    return parser.parse_args()

def load_main_module(script_path: Path) -> Any:
    script_path = script_path.resolve()
    if not script_path.is_file():
        raise FileNotFoundError(f'Main analysis script was not found: {script_path}')
    module_name = script_path.stem
    if not module_name.isidentifier():
        raise ImportError(f'Main script name is not a valid Python module name: {script_path.name}')
    module_directory_path = script_path.parent.resolve()
    retained_sys_path: List[str] = []
    for entry in sys.path:
        try:
            entry_path = Path(entry or '.').resolve()
        except (OSError, RuntimeError):
            retained_sys_path.append(entry)
            continue
        if entry_path != module_directory_path:
            retained_sys_path.append(entry)
    sys.path[:] = [str(module_directory_path), *retained_sys_path]
    existing = sys.modules.get(module_name)
    if existing is not None:
        existing_file = getattr(existing, '__file__', None)
        if existing_file is None or Path(existing_file).resolve() != script_path:
            raise ImportError(f'Module name {module_name!r} is already bound to another file.')
        return existing
    spec = importlib.util.spec_from_file_location(module_name, script_path)
    if spec is None or spec.loader is None:
        raise ImportError(f'Could not import the main script from {script_path}')
    module = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = module
    try:
        spec.loader.exec_module(module)
    except Exception:
        sys.modules.pop(module_name, None)
        raise
    return module

def run_worker_import_smoke_test(main_module: Any, script_path: Path, csv_path: Optional[Path]=None, seed: int=USER_MASTER_SEED) -> None:
    script_path = script_path.resolve()
    expected_sha256 = main_module.file_sha256(str(script_path))
    observed = Parallel(n_jobs=2, backend='loky')((delayed(main_module.file_sha256)(str(script_path)) for _ in range(2)))
    if observed != [expected_sha256, expected_sha256]:
        raise RuntimeError('Spawned workers returned an unexpected main-script hash.')
    if csv_path is not None:
        main_module.set_global_seed(seed)
        df = main_module.load_dataset(str(csv_path.resolve()))
        splitter = StratifiedShuffleSplit(n_splits=1, test_size=TEST_SIZE, random_state=seed)
        train_idx, test_idx = next(splitter.split(df, main_module.make_strata(df)))
        raw_train_df = df.iloc[train_idx].reset_index(drop=True)
        raw_test_df = df.iloc[test_idx].reset_index(drop=True)
        _, _, sensitivity_diagnostic = generate_extra_trees_imputed_split(raw_train_df, raw_test_df, main_module=main_module, imputer_seed=extra_trees_outer_seed(seed, 1), cycles=FIXED_IMPUTATION_CYCLES, n_jobs=2, split_id=1, role='outer_evaluation', inner_fold=None)
        if not (sensitivity_diagnostic['n_iter'] == FIXED_IMPUTATION_CYCLES and sensitivity_diagnostic['all_values_finite'] and sensitivity_diagnostic['plausibility_passed']):
            raise RuntimeError('ExtraTrees fixed-cycle functional smoke did not pass diagnostics.')
        synthetic_stability_rows: List[Dict[str, Any]] = []
        for variable in main_module.PREDICTOR_COLUMNS:
            is_continuous = variable in main_module.CONTINUOUS_COLUMNS
            synthetic_stability_rows.append({'split_id': 1, 'imputer_seed': extra_trees_outer_seed(seed, 1), 'variable': variable, 'variable_type': 'continuous' if is_continuous else 'binary', 'n_missing_train': 0, 'n_missing_test': 0, 'n_compared': 0, 'mean_absolute_difference_10_vs_20': 0.0, 'observed_training_sd': 1.0 if is_continuous else np.nan, 'normalized_mean_absolute_difference': 0.0 if is_continuous else np.nan, 'binary_disagreement_rate': np.nan if is_continuous else 0.0, 'acceptance_threshold': 0.1 if is_continuous else 0.05, 'assessment_status': 'not_applicable_no_missing_cells', 'threshold_applicable': False, 'stability_passed': np.nan, 'cycles_primary': FIXED_IMPUTATION_CYCLES, 'cycles_comparator': STABILITY_COMPARISON_CYCLES, 'primary_warning': EXPECTED_ITERATION_WARNING_TEXT, 'comparator_warning': EXPECTED_ITERATION_WARNING_TEXT})
        for variable, variable_type, normalized, disagreement, threshold in [('__global_continuous__', 'aggregate_continuous', 0.0, np.nan, 0.1), ('__global_binary__', 'aggregate_binary', np.nan, 0.0, 0.05)]:
            synthetic_stability_rows.append({'split_id': 1, 'imputer_seed': extra_trees_outer_seed(seed, 1), 'variable': variable, 'variable_type': variable_type, 'n_missing_train': 1, 'n_missing_test': 0, 'n_compared': 1, 'mean_absolute_difference_10_vs_20': np.nan, 'observed_training_sd': np.nan, 'normalized_mean_absolute_difference': normalized, 'binary_disagreement_rate': disagreement, 'acceptance_threshold': threshold, 'assessment_status': 'aggregate_assessed', 'threshold_applicable': True, 'stability_passed': True, 'cycles_primary': FIXED_IMPUTATION_CYCLES, 'cycles_comparator': STABILITY_COMPARISON_CYCLES, 'primary_warning': EXPECTED_ITERATION_WARNING_TEXT, 'comparator_warning': EXPECTED_ITERATION_WARNING_TEXT})
        validate_stability_frame(pd.DataFrame(synthetic_stability_rows, columns=STABILITY_COLUMNS), 1, main_module.PREDICTOR_COLUMNS, main_module)
        completed = main_module.generate_imputed_split_datasets(raw_train_df, raw_test_df, m=1, seed=seed + 1000, max_iter=1)[0]
        inner_completed_splits = main_module.generate_nested_inner_imputed_splits(raw_outer_train_df=raw_train_df, inner_folds=INNER_FOLDS, split_seed=seed + 2000, imputation_seed=seed + 3000, max_iter=1)
        specs = main_module.get_classical_model_specs(random_state=seed + 2000, n_jobs=2)
        specs['ElasticNetCox']['param_grid'] = {'model__penalizer': [0.01], 'model__l1_ratio': [0.0]}
        _, _, selection_score = main_module.fit_gridsearched_model('ElasticNetCox', completed['train'], inner_folds=INNER_FOLDS, random_state=seed + 2000, n_jobs=2, specs=specs, inner_completed_splits=inner_completed_splits, ibs_start_month=12.0, horizon_months=120.0)
        if not np.isfinite(float(selection_score)):
            raise RuntimeError('Spawned ElasticNetCox IBS smoke returned a non-finite score.')
    print('ExtraTrees fixed-cycle/stability-validator, spawned-worker import, and ElasticNetCox IBS smoke passed.', flush=True)

def validate_dependencies(main_module: Any) -> None:
    required_parameters = {'fit_gridsearched_model': {'inner_completed_splits', 'ibs_start_month', 'horizon_months'}, 'fit_deepsurv': {'selection_split', 'ibs_start_month', 'horizon_months', 'require_all_candidates'}, 'fit_coxtime': {'selection_split', 'ibs_start_month', 'horizon_months', 'require_all_candidates'}}
    missing_interfaces: List[str] = []
    for function_name, parameter_names in required_parameters.items():
        function = getattr(main_module, function_name, None)
        observed = set(inspect.signature(function).parameters) if callable(function) else set()
        for parameter_name in sorted(parameter_names.difference(observed)):
            missing_interfaces.append(f'{function_name}(..., {parameter_name}=...)')
    for function_name in ['model_selection_metadata', 'tuning_output_metadata', 'validate_tuning_selection_record']:
        if not callable(getattr(main_module, function_name, None)):
            missing_interfaces.append(f'{function_name}(...)')
    if missing_interfaces:
        raise RuntimeError('The main pipeline lacks required nested-imputation interfaces: ' + ', '.join(missing_interfaces))

def dataframe_sha256(df: pd.DataFrame) -> str:
    payload = df.to_csv(index=False, na_rep='<NA>', float_format='%.17g', lineterminator='\n')
    return hashlib.sha256(payload.encode('utf-8')).hexdigest()

def canonical_study_id(value: Any) -> str:
    if pd.isna(value):
        raise ValueError('A study ID was missing.')
    if isinstance(value, (int, np.integer)):
        return str(int(value))
    if isinstance(value, (float, np.floating)) and float(value).is_integer():
        return str(int(value))
    return str(value).strip()

def study_id_keys(series: pd.Series) -> pd.Series:
    keys = series.map(canonical_study_id)
    if (keys == '').any():
        raise ValueError('A study ID was empty after normalization.')
    return keys

def split_base_seed(seed: int, split_id: int) -> int:
    return int(seed + 100000 * split_id)

def matched_model_seed(seed: int, split_id: int) -> int:
    return int(split_base_seed(seed, split_id) + 1)

def extra_trees_outer_seed(seed: int, split_id: int) -> int:
    return int(SENSITIVITY_SEED_OFFSET + split_base_seed(seed, split_id))

def extra_trees_inner_seed(seed: int, split_id: int, inner_fold: int) -> int:
    return int(SENSITIVITY_SEED_OFFSET + split_base_seed(seed, split_id) + 100 + inner_fold)

def extra_trees_neural_seed(seed: int, split_id: int) -> int:
    return int(SENSITIVITY_SEED_OFFSET + split_base_seed(seed, split_id) + 500)

def required_split_metrics(model_name: str) -> List[str]:
    return list(DISCRIMINATION_METRICS) if model_name == 'SVM' else list(FOLD_PERFORMANCE_METRICS)

def validate_performance_metric(metric: str, value: float, context: str) -> None:
    if not np.isfinite(value):
        raise ValueError(f'{context}: {metric} is non-finite.')
    if metric in FOLD_PERFORMANCE_METRICS and (not -1e-12 <= value <= 1.0 + 1e-12):
        raise ValueError(f'{context}: {metric}={value} lies outside [0,1].')

def expected_tuning_method(model_name: str) -> str:
    if model_name == 'CoxPH':
        return FIXED_TUNING_METHOD
    if model_name in {'DeepSurv', 'CoxTime'}:
        return NEURAL_IBS_TUNING_METHOD
    if model_name == 'SVM':
        return SVM_TUNING_METHOD
    return CLASSICAL_IBS_TUNING_METHOD

def expected_tuning_imputation_policy(model_name: str) -> str:
    if model_name == 'CoxPH':
        return FIXED_IMPUTATION_POLICY
    if model_name in {'DeepSurv', 'CoxTime'}:
        return NEURAL_IMPUTATION_POLICY
    return CLASSICAL_IMPUTATION_POLICY

def _normalize_saved_map(saved: pd.DataFrame, main_module: Any) -> pd.DataFrame:
    expected_columns = ['split_id', main_module.STUDY_ID_COL, 'role', 'split_seed', 'test_order']
    if list(saved.columns) != expected_columns:
        raise RuntimeError(f'The main split map columns must be exactly {expected_columns}.')
    normalized = saved[expected_columns].copy()
    normalized['split_id'] = pd.to_numeric(normalized['split_id'], errors='raise').astype('int64')
    normalized['study_id_key'] = study_id_keys(normalized[main_module.STUDY_ID_COL])
    normalized['role'] = normalized['role'].astype(str).str.strip().str.lower()
    normalized['split_seed'] = pd.to_numeric(normalized['split_seed'], errors='raise').astype('int64')
    normalized['test_order'] = pd.to_numeric(normalized['test_order'], errors='raise').astype('int64')
    return normalized

def validate_main_reference(main_module: Any, df: pd.DataFrame, main_output_dir: Path, seed: int) -> Tuple[List[Tuple[int, np.ndarray, np.ndarray]], pd.DataFrame, pd.DataFrame, Dict[str, Any]]:
    completion_path = main_output_dir / 'metadata' / 'analysis_complete.json'
    map_path = main_output_dir / 'metadata' / 'repeated_200_split_map.csv'
    performance_path = main_output_dir / 'tables' / 'repeated_200_performance_by_split.csv'
    tuning_path = main_output_dir / 'tables' / 'repeated_200_tuning_by_split.csv'
    for path in [completion_path, map_path, performance_path, tuning_path]:
        if not path.is_file():
            raise FileNotFoundError(f'Required completed main-analysis file is missing: {path}')
    with open(completion_path, 'r', encoding='utf-8') as handle:
        completion = json.load(handle)
    if completion.get('status') != 'complete':
        raise RuntimeError('The main 200-split completion marker is not complete.')
    saved = _normalize_saved_map(pd.read_csv(map_path), main_module)
    if set(saved['role']) != {'test'}:
        raise RuntimeError('The main split map must contain test rows only.')
    if set(saved['split_id']) != set(range(1, TOTAL_MAIN_SPLITS + 1)):
        raise RuntimeError('The main split map does not contain split IDs 1 through 200.')
    if set(saved['split_seed']) != {int(seed)}:
        raise RuntimeError('The saved split_seed does not match the validated master seed.')
    if saved.duplicated(['split_id', 'study_id_key']).any():
        raise RuntimeError('The main split map duplicates a held-out participant within a split.')
    source_ids = study_id_keys(df[main_module.STUDY_ID_COL]).reset_index(drop=True)
    if source_ids.duplicated().any():
        raise RuntimeError('The eligible dataset contains duplicate study IDs.')
    source_index_by_id = {value: idx for idx, value in enumerate(source_ids.tolist())}
    test_n = int(math.ceil(TEST_SIZE * len(df)))
    expected_total_rows = TOTAL_MAIN_SPLITS * test_n
    if len(saved) != expected_total_rows:
        raise RuntimeError(f'The main split map has {len(saved)} rows; expected {expected_total_rows}.')
    split_helper = getattr(main_module, 'iter_repeated_splits', None)
    if not callable(split_helper):
        raise RuntimeError('The main pipeline does not expose the validated 200-split generator.')
    definitions: List[Tuple[int, np.ndarray, np.ndarray]] = []
    matched_map_frames: List[pd.DataFrame] = []
    all_source_indices = np.arange(len(df), dtype=int)
    generated = split_helper(df, seed=seed, n_splits=TOTAL_MAIN_SPLITS, test_size=TEST_SIZE)
    for split_id, train_idx, generated_test_idx in generated:
        split_id = int(split_id)
        rows = saved.loc[saved['split_id'].eq(split_id)].sort_values('test_order')
        if len(rows) != test_n:
            raise RuntimeError(f'Split {split_id} has the wrong held-out row count.')
        if rows['test_order'].tolist() != list(range(1, test_n + 1)):
            raise RuntimeError(f'Split {split_id} test_order is not exactly 1..{test_n}.')
        unexpected_ids = sorted(set(rows['study_id_key']).difference(source_index_by_id))
        if unexpected_ids:
            raise RuntimeError(f'Split {split_id} contains unexpected study IDs: {unexpected_ids[:10]}')
        saved_test_idx = np.asarray([source_index_by_id[value] for value in rows['study_id_key']], dtype=int)
        generated_test_idx = np.asarray(generated_test_idx, dtype=int)
        if not np.array_equal(saved_test_idx, generated_test_idx):
            raise RuntimeError(f'Split {split_id} assignments or test order do not match the prespecified StratifiedShuffleSplit generator.')
        expected_train_idx = all_source_indices[~np.isin(all_source_indices, saved_test_idx, assume_unique=True)]
        generated_train_idx = np.asarray(train_idx, dtype=int)
        if set(generated_train_idx.tolist()) != set(expected_train_idx.tolist()):
            raise RuntimeError(f'Split {split_id} training complement is invalid.')
        if len(np.unique(generated_train_idx)) != len(generated_train_idx):
            raise RuntimeError(f'Split {split_id} repeats a training index.')
        definitions.append((split_id, generated_train_idx, saved_test_idx))
        matched_map_frames.append(rows.copy())
    if len(definitions) != SENSITIVITY_SPLITS:
        raise RuntimeError('Could not recover all 200 matched split assignments.')
    matched_map = pd.concat(matched_map_frames, ignore_index=True)
    if len(matched_map) != EXPECTED_HELDOUT_PREDICTION_ROWS:
        raise RuntimeError(f'The matched 200-split map has {len(matched_map)} rows; expected {EXPECTED_HELDOUT_PREDICTION_ROWS}.')
    coverage = matched_map.groupby('study_id_key').size()
    missing_ids = sorted(set(source_ids).difference(coverage.index))
    if missing_ids:
        raise RuntimeError(f'The 200 prespecified splits do not hold out every participant; {len(missing_ids)} IDs have no held-out prediction.')
    if int(coverage.min()) != EXPECTED_MIN_HELDOUT_PREDICTIONS or int(coverage.max()) != EXPECTED_MAX_HELDOUT_PREDICTIONS:
        raise RuntimeError(f'The 200-split held-out coverage range changed: observed {int(coverage.min())}-{int(coverage.max())}, expected {EXPECTED_MIN_HELDOUT_PREDICTIONS}-{EXPECTED_MAX_HELDOUT_PREDICTIONS}.')
    main_performance = pd.read_csv(performance_path, float_precision='round_trip')
    required_perf = {'split_id', 'model', *FOLD_PERFORMANCE_METRICS}
    missing_perf = sorted(required_perf.difference(main_performance.columns))
    if missing_perf:
        raise RuntimeError(f'The main performance table is missing: {missing_perf}')
    main_performance['split_id'] = pd.to_numeric(main_performance['split_id'], errors='raise').astype(int)
    expected_keys = {(split_id, model) for split_id in range(1, TOTAL_MAIN_SPLITS + 1) for model in main_module.MODEL_ORDER}
    observed_keys = set(zip(main_performance['split_id'], main_performance['model'].astype(str)))
    if len(main_performance) != len(expected_keys) or observed_keys != expected_keys:
        raise RuntimeError('The main performance table has an incomplete or duplicate grid.')
    for model_name, rows in main_performance.groupby('model'):
        for metric in required_split_metrics(str(model_name)):
            values = pd.to_numeric(rows[metric], errors='coerce').to_numpy(float)
            if not np.isfinite(values).all():
                raise RuntimeError(f'The main performance table has non-finite {metric} for {model_name}.')
            for value in values:
                validate_performance_metric(metric, float(value), f'Main performance {model_name}')
    main_tuning = pd.read_csv(tuning_path, float_precision='round_trip')
    tuning_keys = set(zip(pd.to_numeric(main_tuning['split_id'], errors='raise').astype(int), main_tuning['model'].astype(str)))
    if len(main_tuning) != len(expected_keys) or tuning_keys != expected_keys or main_tuning.duplicated(['split_id', 'model']).any():
        raise RuntimeError('The main tuning table has an incomplete or duplicate grid.')
    for _, row in main_tuning.iterrows():
        model_name = str(row['model'])
        if str(row['tuning_method']) != main_module.main_tuning_method(model_name):
            raise RuntimeError(f'The main tuning method is wrong for {model_name}.')
        main_module.validate_tuning_selection_record(row, model_name, f'Main reference tuning {model_name}')
    selected_performance = main_performance.copy()
    provenance = {'main_reference_verified': True, 'main_completion_sha256': main_module.file_sha256(str(completion_path)), 'main_split_map_sha256': main_module.file_sha256(str(map_path)), 'main_performance_sha256': main_module.file_sha256(str(performance_path)), 'main_tuning_sha256': main_module.file_sha256(str(tuning_path)), 'verified_total_main_splits': TOTAL_MAIN_SPLITS, 'matched_sensitivity_split_ids': list(range(1, SENSITIVITY_SPLITS + 1)), 'matched_map_rows': int(len(matched_map)), 'held_out_coverage_minimum': int(coverage.min()), 'held_out_coverage_maximum': int(coverage.max())}
    return (definitions, matched_map, selected_performance, provenance)

def prepare_checkpoint_manifest(main_module: Any, df: pd.DataFrame, out_dir: Path, args: argparse.Namespace, reference_provenance: Dict[str, Any]) -> Tuple[Path, str]:
    checkpoint_dir = out_dir / 'checkpoints' / 'extra_trees_matched_200_splits_v2_ibs_tuning'
    checkpoint_dir.mkdir(parents=True, exist_ok=True)
    configuration = {'identity': main_module.checkpoint_common_identity(df), 'sensitivity_script_sha256': main_module.file_sha256(str(Path(__file__).resolve())), 'main_reference_provenance': reference_provenance, 'settings': {'outer_design': 'all_200_prespecified_stratified_80_20_splits', 'test_size': TEST_SIZE, 'imputer': 'ExtraTreesRegressor_within_IterativeImputer', 'extra_trees_n_estimators': EXTRA_TREES_N_ESTIMATORS, 'extra_trees_min_samples_leaf': EXTRA_TREES_MIN_SAMPLES_LEAF, 'extra_trees_max_features': EXTRA_TREES_MAX_FEATURES, 'fixed_cycle_policy': FIXED_CYCLE_POLICY, 'impute_cycles': FIXED_IMPUTATION_CYCLES, 'imputer_tol': 0.0, 'seed': int(args.seed), 'inner_folds_classical': INNER_FOLDS, 'classical_absolute_risk_tuning': CLASSICAL_IBS_TUNING_METHOD, 'svm_tuning': SVM_TUNING_METHOD, 'coxph_tuning': FIXED_TUNING_METHOD, 'classical_tuning_imputation': CLASSICAL_IMPUTATION_POLICY, 'coxph_imputation_policy': FIXED_IMPUTATION_POLICY, 'neural_tuning': NEURAL_IBS_TUNING_METHOD, 'neural_tuning_imputation': NEURAL_IMPUTATION_POLICY, 'model_selection_policy': {model_name: main_module.model_selection_metadata(model_name) for model_name in main_module.MODEL_ORDER}, 'horizon_months': float(args.horizon_months), 'ibs_start_month': float(args.ibs_start_month), 'calibration_groups': int(args.calibration_groups), 'calibration_bootstrap_reps': int(args.calibration_bootstrap_reps), 'n_jobs': int(args.n_jobs), 'stability_split_ids': list(STABILITY_SPLIT_IDS), 'stability_cycles': [FIXED_IMPUTATION_CYCLES, STABILITY_COMPARISON_CYCLES], 'stability_thresholds': {'minimum_cells_for_individual_variable_assessment': MIN_STABILITY_CELLS_PER_VARIABLE, 'continuous_normalized_mad': 0.1, 'binary_disagreement': 0.05, 'count_weighted_global_continuous_normalized_mad': 0.1, 'global_binary_disagreement': 0.05}, 'stability_locking_policy': 'prespecified_before_inspecting_any_outcome_model_performance_or_calibration; predictor_missing_cells_only; no_event_status_follow_up_performance_or_calibration_used', 'seed_scheme': {'split_base_seed': 'seed + 100000 * split_id', 'model_tuning_seed': 'split_base_seed + 1', 'outer_extra_trees_seed': f'{SENSITIVITY_SEED_OFFSET} + split_base_seed', 'classical_inner_extra_trees_seed': f'{SENSITIVITY_SEED_OFFSET} + split_base_seed + 100 + inner_fold', 'neural_extra_trees_seed': f'{SENSITIVITY_SEED_OFFSET} + split_base_seed + 500'}, 'participant_prediction_policy': 'checkpoint_only; final calibration uses one mean held-out risk per participant', 'plausibility_policy': 'finite_complete_binary_valid_observed_cells_preserved_and_each_imputed_continuous_value_within_its_imputer_training_observed_min_max', 'warning_policy': 'expected_iteration_cap_ConvergenceWarning_classified_as_fixed_cycle_notice_not_convergence_assessment; any_unexpected_warning_halts'}}
    manifest_sha256 = main_module.prepare_checkpoint_manifest(str(checkpoint_dir), run_type='extra_trees_predictor_only_matched_all_200_splits', configuration=configuration, resume=args.resume)
    return (checkpoint_dir, manifest_sha256)

def build_imputation_matrix(df: pd.DataFrame, main_module: Any) -> pd.DataFrame:
    return df.loc[:, main_module.PREDICTOR_COLUMNS].copy()

def _bool_value(value: Any, context: str) -> bool:
    if isinstance(value, (bool, np.bool_)):
        return bool(value)
    normalized = str(value).strip().lower()
    if normalized in {'true', '1'}:
        return True
    if normalized in {'false', '0'}:
        return False
    raise ValueError(f'{context} is not a valid Boolean value: {value!r}')

def _validate_completed_predictors(completed_train: pd.DataFrame, completed_apply: pd.DataFrame, original_train: pd.DataFrame, original_apply: pd.DataFrame, main_module: Any) -> Tuple[bool, bool, bool]:
    combined = pd.concat([completed_train, completed_apply], ignore_index=True)
    finite = bool(np.isfinite(combined.to_numpy(dtype=float)).all())
    no_missing = not bool(combined.isna().any().any())
    binary_ok = all((set(pd.to_numeric(combined[column], errors='raise').unique()).issubset({0, 1}) for column in main_module.BINARY_COLUMNS))
    lower_bounds = {'age': 18.0, 'sbp': 40.0, 'dbp': 20.0, 'bmi': 10.0, 'chol': 0.1, 'hdl': 0.1, 'trig': 0.05, 'hba1c': 2.0, 'egfr': 1.0}
    continuous_ok = all((bool((pd.to_numeric(combined[column], errors='raise') >= lower_bounds[column]).all()) for column in main_module.CONTINUOUS_COLUMNS))
    imputed_within_training_support = True
    for column in main_module.CONTINUOUS_COLUMNS:
        observed_training = pd.to_numeric(original_train.loc[original_train[column].notna(), column], errors='raise').to_numpy(float)
        if observed_training.size == 0:
            imputed_within_training_support = False
            continue
        imputed_values = np.concatenate([completed_train.loc[original_train[column].isna().to_numpy(), column].to_numpy(dtype=float), completed_apply.loc[original_apply[column].isna().to_numpy(), column].to_numpy(dtype=float)])
        if imputed_values.size and (np.any(imputed_values < float(np.min(observed_training)) - 1e-10) or np.any(imputed_values > float(np.max(observed_training)) + 1e-10)):
            imputed_within_training_support = False
    observed_preserved = True
    for original, completed in [(original_train, completed_train), (original_apply, completed_apply)]:
        for column in main_module.CONTINUOUS_COLUMNS:
            observed = original[column].notna().to_numpy()
            if observed.any() and (not np.allclose(original.loc[observed, column].to_numpy(dtype=float), completed.loc[observed, column].to_numpy(dtype=float), rtol=0.0, atol=1e-10)):
                observed_preserved = False
        for column in main_module.BINARY_COLUMNS:
            observed = original[column].notna().to_numpy()
            if observed.any() and (not np.array_equal(original.loc[observed, column].to_numpy(dtype=int), completed.loc[observed, column].to_numpy(dtype=int))):
                observed_preserved = False
    plausibility = bool(finite and no_missing and binary_ok and continuous_ok and imputed_within_training_support)
    return (finite, plausibility, observed_preserved)

def generate_extra_trees_imputed_split(raw_train_df: pd.DataFrame, raw_apply_df: pd.DataFrame, main_module: Any, imputer_seed: int, cycles: int, n_jobs: int, split_id: int, role: str, inner_fold: Optional[int]) -> Tuple[pd.DataFrame, pd.DataFrame, Dict[str, Any]]:
    train_matrix = build_imputation_matrix(raw_train_df, main_module)
    apply_matrix = build_imputation_matrix(raw_apply_df, main_module)
    columns = list(train_matrix.columns)
    estimator = ExtraTreesRegressor(n_estimators=EXTRA_TREES_N_ESTIMATORS, min_samples_leaf=EXTRA_TREES_MIN_SAMPLES_LEAF, max_features=EXTRA_TREES_MAX_FEATURES, random_state=imputer_seed, n_jobs=n_jobs)
    imputer = IterativeImputer(estimator=estimator, max_iter=int(cycles), tol=0.0, sample_posterior=False, random_state=imputer_seed, initial_strategy='median', imputation_order='ascending', skip_complete=False)
    with warnings.catch_warnings(record=True) as caught:
        warnings.simplefilter('always')
        completed_train_raw = pd.DataFrame(imputer.fit_transform(train_matrix[columns]), columns=columns, index=raw_train_df.index)
        completed_apply_raw = pd.DataFrame(imputer.transform(apply_matrix[columns]), columns=columns, index=raw_apply_df.index)
    unexpected = [item for item in caught if not (issubclass(item.category, ConvergenceWarning) and EXPECTED_ITERATION_WARNING_TEXT in str(item.message))]
    if unexpected:
        details = [f'{item.category.__name__}: {item.message}' for item in unexpected]
        raise RuntimeError(f'Split {split_id} {role} imputation produced unexpected warnings: {details}')
    expected_warnings = [item for item in caught if issubclass(item.category, ConvergenceWarning) and EXPECTED_ITERATION_WARNING_TEXT in str(item.message)]
    if int(imputer.n_iter_) != int(cycles):
        raise RuntimeError(f'Split {split_id} {role} ran {imputer.n_iter_} cycles; expected {cycles}.')
    if not np.isfinite(completed_train_raw.to_numpy(dtype=float)).all() or not np.isfinite(completed_apply_raw.to_numpy(dtype=float)).all():
        raise RuntimeError(f'Split {split_id} {role} imputation returned non-finite values.')
    train_predictors = main_module.postprocess_imputed_predictors(completed_train_raw)
    apply_predictors = main_module.postprocess_imputed_predictors(completed_apply_raw)
    finite, plausible, observed_preserved = _validate_completed_predictors(train_predictors, apply_predictors, train_matrix, apply_matrix, main_module)
    if not finite or not plausible or (not observed_preserved):
        raise RuntimeError(f'Split {split_id} {role} failed finite, plausibility, or observed-value checks.')
    core_columns = [main_module.STUDY_ID_COL, main_module.TIME_COL, main_module.EVENT_COL]
    imputed_train = pd.concat([raw_train_df[core_columns].reset_index(drop=True), train_predictors.reset_index(drop=True)], axis=1)
    imputed_apply = pd.concat([raw_apply_df[core_columns].reset_index(drop=True), apply_predictors.reset_index(drop=True)], axis=1)
    warning_text = ' | '.join((str(item.message) for item in expected_warnings))
    diagnostic = {'split_id': int(split_id), 'role': str(role), 'inner_fold': int(inner_fold) if inner_fold is not None else np.nan, 'imputer_seed': int(imputer_seed), 'n_iter': int(imputer.n_iter_), 'max_iter': int(cycles), 'fixed_cycle_policy': FIXED_CYCLE_POLICY if cycles == FIXED_IMPUTATION_CYCLES else f'prespecified_{cycles}_cycle_stability_comparator_tol_0', 'warning_count': int(len(expected_warnings)), 'warning': warning_text, 'warning_classification': 'expected_fixed_cycle_notice_not_convergence_assessment' if expected_warnings else 'no_notice_fixed_cycle_completed', 'unexpected_warning_count': 0, 'train_missing_cells_before': int(train_matrix.isna().sum().sum()), 'apply_missing_cells_before': int(apply_matrix.isna().sum().sum()), 'train_missing_cells_after': int(train_predictors.isna().sum().sum()), 'apply_missing_cells_after': int(apply_predictors.isna().sum().sum()), 'observed_cells_preserved': bool(observed_preserved), 'all_values_finite': bool(finite), 'plausibility_passed': bool(plausible)}
    return (imputed_train, imputed_apply, diagnostic)

def build_stability_diagnostic(raw_train_df: pd.DataFrame, raw_test_df: pd.DataFrame, primary_train_df: pd.DataFrame, primary_test_df: pd.DataFrame, primary_diagnostic: Dict[str, Any], main_module: Any, split_id: int, imputer_seed: int, n_jobs: int) -> pd.DataFrame:
    comparator_train, comparator_test, comparator_diagnostic = generate_extra_trees_imputed_split(raw_train_df, raw_test_df, main_module=main_module, imputer_seed=imputer_seed, cycles=STABILITY_COMPARISON_CYCLES, n_jobs=n_jobs, split_id=split_id, role='outer_evaluation_stability_comparator', inner_fold=None)
    rows: List[Dict[str, Any]] = []
    standardized_continuous_cell_differences: List[np.ndarray] = []
    binary_cell_disagreements: List[np.ndarray] = []
    for variable in main_module.PREDICTOR_COLUMNS:
        train_missing = raw_train_df[variable].isna().to_numpy()
        test_missing = raw_test_df[variable].isna().to_numpy()
        primary_values = np.concatenate([primary_train_df.loc[train_missing, variable].to_numpy(dtype=float), primary_test_df.loc[test_missing, variable].to_numpy(dtype=float)])
        comparator_values = np.concatenate([comparator_train.loc[train_missing, variable].to_numpy(dtype=float), comparator_test.loc[test_missing, variable].to_numpy(dtype=float)])
        n_compared = int(len(primary_values))
        if n_compared != int(train_missing.sum() + test_missing.sum()):
            raise RuntimeError(f'Split {split_id} stability mask failed for {variable}.')
        if not (np.isfinite(primary_values).all() and np.isfinite(comparator_values).all()):
            raise RuntimeError(f'Split {split_id} stability values are non-finite for {variable}.')
        mean_absolute_difference = float(np.mean(np.abs(primary_values - comparator_values))) if n_compared else 0.0
        observed_training_sd = np.nan
        normalized_difference = np.nan
        binary_disagreement = np.nan
        if variable in main_module.CONTINUOUS_COLUMNS:
            observed = pd.to_numeric(raw_train_df.loc[raw_train_df[variable].notna(), variable], errors='raise').to_numpy(float)
            observed_training_sd = float(np.std(observed, ddof=1))
            if n_compared and (not np.isfinite(observed_training_sd) or observed_training_sd <= 0.0):
                raise RuntimeError(f'Split {split_id} cannot standardize {variable}: the observed training standard deviation is not positive and finite.')
            else:
                normalized_difference = mean_absolute_difference / observed_training_sd if n_compared else 0.0
            if n_compared:
                standardized_continuous_cell_differences.append(np.abs(primary_values - comparator_values) / observed_training_sd)
            variable_type = 'continuous'
            threshold = 0.1
        elif variable in main_module.BINARY_COLUMNS:
            binary_disagreement = float(np.mean(primary_values != comparator_values)) if n_compared else 0.0
            if n_compared:
                binary_cell_disagreements.append((primary_values != comparator_values).astype(float))
            variable_type = 'binary'
            threshold = 0.05
        else:
            raise RuntimeError(f'Predictor type is not defined for {variable}.')
        if n_compared == 0:
            assessment_status = 'not_applicable_no_missing_cells'
            threshold_applicable = False
            passed: Any = np.nan
        elif n_compared < MIN_STABILITY_CELLS_PER_VARIABLE:
            assessment_status = 'sparse_not_individually_assessable'
            threshold_applicable = False
            passed = np.nan
        else:
            assessment_status = 'individually_assessed'
            threshold_applicable = True
            if variable_type == 'continuous':
                passed = bool(np.isfinite(normalized_difference) and normalized_difference <= threshold + 1e-12)
            else:
                passed = bool(binary_disagreement <= threshold + 1e-12)
        rows.append({'split_id': int(split_id), 'imputer_seed': int(imputer_seed), 'variable': variable, 'variable_type': variable_type, 'n_missing_train': int(train_missing.sum()), 'n_missing_test': int(test_missing.sum()), 'n_compared': n_compared, 'mean_absolute_difference_10_vs_20': mean_absolute_difference, 'observed_training_sd': observed_training_sd, 'normalized_mean_absolute_difference': normalized_difference, 'binary_disagreement_rate': binary_disagreement, 'acceptance_threshold': threshold, 'assessment_status': assessment_status, 'threshold_applicable': threshold_applicable, 'stability_passed': passed, 'cycles_primary': FIXED_IMPUTATION_CYCLES, 'cycles_comparator': STABILITY_COMPARISON_CYCLES, 'primary_warning': str(primary_diagnostic['warning']), 'comparator_warning': str(comparator_diagnostic['warning'])})
    continuous_cells = np.concatenate(standardized_continuous_cell_differences) if standardized_continuous_cell_differences else np.asarray([], dtype=float)
    binary_cells = np.concatenate(binary_cell_disagreements) if binary_cell_disagreements else np.asarray([], dtype=float)
    if continuous_cells.size == 0 or binary_cells.size == 0:
        raise RuntimeError(f'Split {split_id} lacks evaluable continuous or binary missing cells.')
    global_continuous = float(np.mean(continuous_cells))
    global_binary = float(np.mean(binary_cells))
    aggregate_specs = [('__global_continuous__', 'aggregate_continuous', int(continuous_cells.size), global_continuous, np.nan, 0.1), ('__global_binary__', 'aggregate_binary', int(binary_cells.size), np.nan, global_binary, 0.05)]
    predictor_rows = pd.DataFrame(rows)
    for variable, variable_type, n_compared, normalized, disagreement, threshold in aggregate_specs:
        source_type = main_module.CONTINUOUS_COLUMNS if variable_type == 'aggregate_continuous' else main_module.BINARY_COLUMNS
        type_rows = predictor_rows.loc[predictor_rows['variable'].isin(source_type)]
        rows.append({'split_id': int(split_id), 'imputer_seed': int(imputer_seed), 'variable': variable, 'variable_type': variable_type, 'n_missing_train': int(type_rows['n_missing_train'].sum()), 'n_missing_test': int(type_rows['n_missing_test'].sum()), 'n_compared': int(n_compared), 'mean_absolute_difference_10_vs_20': np.nan, 'observed_training_sd': np.nan, 'normalized_mean_absolute_difference': normalized, 'binary_disagreement_rate': disagreement, 'acceptance_threshold': threshold, 'assessment_status': 'aggregate_assessed', 'threshold_applicable': True, 'stability_passed': bool((normalized if variable_type == 'aggregate_continuous' else disagreement) <= threshold + 1e-12), 'cycles_primary': FIXED_IMPUTATION_CYCLES, 'cycles_comparator': STABILITY_COMPARISON_CYCLES, 'primary_warning': str(primary_diagnostic['warning']), 'comparator_warning': str(comparator_diagnostic['warning'])})
    output = pd.DataFrame(rows, columns=STABILITY_COLUMNS)
    if len(output) != len(main_module.PREDICTOR_COLUMNS) + 2:
        raise RuntimeError(f'Split {split_id} stability table is incomplete.')
    applicable = output.loc[output['threshold_applicable'].astype(bool)]
    if not applicable['stability_passed'].astype(bool).all():
        failed = applicable.loc[~applicable['stability_passed'].astype(bool), ['variable', 'normalized_mean_absolute_difference', 'binary_disagreement_rate']]
        raise RuntimeError(f"Split {split_id} failed the prespecified 10-versus-20-cycle stability thresholds: {failed.to_dict('records')}")
    return output

def require_complete_selection_predictors(main_module: Any, train_df: pd.DataFrame, validation_df: pd.DataFrame, context: str) -> None:
    if train_df.empty or validation_df.empty:
        raise RuntimeError(f'{context} produced an empty completed subset.')
    if train_df[main_module.PREDICTOR_COLUMNS].isna().any().any() or validation_df[main_module.PREDICTOR_COLUMNS].isna().any().any():
        raise RuntimeError(f'{context} retained a missing predictor.')
    train_ids = set(study_id_keys(train_df[main_module.STUDY_ID_COL]))
    validation_ids = set(study_id_keys(validation_df[main_module.STUDY_ID_COL]))
    if train_ids.intersection(validation_ids):
        raise RuntimeError(f'{context} has overlapping development and validation IDs.')

def generate_extra_trees_inner_completed_splits(raw_outer_train_df: pd.DataFrame, main_module: Any, split_id: int, model_seed: int, master_seed: int, n_jobs: int) -> Tuple[List[Tuple[pd.DataFrame, pd.DataFrame]], List[Dict[str, Any]]]:
    splitter = StratifiedKFold(n_splits=INNER_FOLDS, shuffle=True, random_state=model_seed)
    strata = main_module.make_strata(raw_outer_train_df)
    completed_splits: List[Tuple[pd.DataFrame, pd.DataFrame]] = []
    diagnostics: List[Dict[str, Any]] = []
    held_out_ids: List[str] = []
    for inner_fold, (inner_train_idx, inner_validation_idx) in enumerate(splitter.split(raw_outer_train_df, strata), start=1):
        inner_train, inner_validation, diagnostic = generate_extra_trees_imputed_split(raw_outer_train_df.iloc[inner_train_idx].reset_index(drop=True), raw_outer_train_df.iloc[inner_validation_idx].reset_index(drop=True), main_module=main_module, imputer_seed=extra_trees_inner_seed(master_seed, split_id, inner_fold), cycles=FIXED_IMPUTATION_CYCLES, n_jobs=n_jobs, split_id=split_id, role='classical_inner', inner_fold=inner_fold)
        require_complete_selection_predictors(main_module, inner_train, inner_validation, context=f'Split {split_id} classical inner fold {inner_fold}')
        completed_splits.append((inner_train, inner_validation))
        diagnostics.append(diagnostic)
        held_out_ids.extend(study_id_keys(inner_validation[main_module.STUDY_ID_COL]).tolist())
    expected_ids = sorted(study_id_keys(raw_outer_train_df[main_module.STUDY_ID_COL]).tolist())
    if sorted(held_out_ids) != expected_ids:
        raise RuntimeError(f'Split {split_id} inner validation did not hold out each outer-training participant exactly once.')
    return (completed_splits, diagnostics)

def generate_extra_trees_neural_selection_split(raw_outer_train_df: pd.DataFrame, main_module: Any, split_id: int, model_seed: int, master_seed: int, n_jobs: int) -> Tuple[Tuple[pd.DataFrame, pd.DataFrame], Dict[str, Any]]:
    splitter = StratifiedShuffleSplit(n_splits=1, test_size=0.2, random_state=model_seed)
    development_idx, validation_idx = next(splitter.split(raw_outer_train_df, main_module.make_strata(raw_outer_train_df)))
    development, validation, diagnostic = generate_extra_trees_imputed_split(raw_outer_train_df.iloc[development_idx].reset_index(drop=True), raw_outer_train_df.iloc[validation_idx].reset_index(drop=True), main_module=main_module, imputer_seed=extra_trees_neural_seed(master_seed, split_id), cycles=FIXED_IMPUTATION_CYCLES, n_jobs=n_jobs, split_id=split_id, role='neural_development', inner_fold=None)
    require_complete_selection_predictors(main_module, development, validation, context=f'Split {split_id} neural development/validation')
    completed_ids = sorted(study_id_keys(development[main_module.STUDY_ID_COL]).tolist() + study_id_keys(validation[main_module.STUDY_ID_COL]).tolist())
    expected_ids = sorted(study_id_keys(raw_outer_train_df[main_module.STUDY_ID_COL]).tolist())
    if completed_ids != expected_ids:
        raise RuntimeError(f'Split {split_id} neural selection did not partition outer training.')
    return ((development, validation), diagnostic)

def tune_and_fit_model(main_module: Any, model_name: str, train_df: pd.DataFrame, model_seed: int, n_jobs: int, inner_completed_splits: Sequence[Tuple[pd.DataFrame, pd.DataFrame]], neural_selection_split: Tuple[pd.DataFrame, pd.DataFrame], ibs_start_month: float, horizon_months: float) -> Tuple[Any, Dict[str, Any], Optional[float], str]:
    if model_name == 'DeepSurv':
        fitted, params, score = main_module.fit_deepsurv(train_df, random_state=model_seed, selection_split=neural_selection_split, ibs_start_month=ibs_start_month, horizon_months=horizon_months, require_all_candidates=True)
        return (fitted, params, float(score), expected_tuning_method(model_name))
    if model_name == 'CoxTime':
        fitted, params, score = main_module.fit_coxtime(train_df, random_state=model_seed, selection_split=neural_selection_split, ibs_start_month=ibs_start_month, horizon_months=horizon_months, require_all_candidates=True)
        return (fitted, params, float(score), expected_tuning_method(model_name))
    specs = main_module.get_classical_model_specs(random_state=model_seed, n_jobs=1)
    fitted, params, score = main_module.fit_gridsearched_model(model_name, train_df, inner_folds=INNER_FOLDS, random_state=model_seed, n_jobs=n_jobs, specs=specs, inner_completed_splits=inner_completed_splits, ibs_start_month=ibs_start_month, horizon_months=horizon_months)
    return (fitted, params, None if score is None else float(score), expected_tuning_method(model_name))

def validate_one_model_split_result(model_name: str, metrics: Dict[str, Any], predicted_event_probability: Optional[np.ndarray], expected_test_n: int, requested_horizon: float, context: str) -> Dict[str, str]:
    errors: Dict[str, str] = {}
    for metric in required_split_metrics(model_name):
        try:
            value = float(metrics.get(metric, np.nan))
        except Exception:
            value = np.nan
        if not np.isfinite(value):
            errors[metric] = 'Required held-out metric was non-finite.'
        else:
            try:
                validate_performance_metric(metric, value, context)
            except ValueError as exc:
                errors[metric] = str(exc)
    if model_name != 'SVM':
        probabilities = np.asarray(predicted_event_probability, dtype=float)
        if probabilities.shape != (expected_test_n,):
            errors['predicted_event_probability'] = f'Expected {expected_test_n} values; got {probabilities.shape}.'
        elif not np.isfinite(probabilities).all():
            errors['predicted_event_probability'] = 'Probabilities were non-finite.'
        elif np.any((probabilities < -1e-08) | (probabilities > 1.0 + 1e-08)):
            errors['predicted_event_probability'] = 'Probabilities were outside [0,1].'
        horizon_used = float(metrics.get('brier_horizon_months', np.nan))
        if not np.isfinite(horizon_used) or not np.isclose(horizon_used, requested_horizon, rtol=0.0, atol=1e-12):
            errors['brier_horizon_months'] = f'Used {horizon_used}; expected {requested_horizon}.'
    if errors:
        raise RuntimeError(f'{context}: {model_name} result is incomplete: {errors}')
    return errors

def session_deadline(max_session_hours: float) -> Optional[float]:
    if max_session_hours <= 0.0:
        return None
    return time.monotonic() + 3600.0 * float(max_session_hours)

def deadline_reached(deadline: Optional[float]) -> bool:
    return deadline is not None and time.monotonic() >= deadline

def split_checkpoint_paths(checkpoint_dir: Path, split_id: int) -> Dict[str, Path]:
    directory = checkpoint_dir / 'splits'
    directory.mkdir(parents=True, exist_ok=True)
    prefix = f'split_{split_id:03d}'
    return {'metrics': directory / f'{prefix}_metrics.csv', 'tuning': directory / f'{prefix}_tuning.csv', 'predictions': directory / f'{prefix}_predictions.csv', 'diagnostics': directory / f'{prefix}_imputation_diagnostics.csv', 'stability': directory / f'{prefix}_cycle_stability.csv', 'marker': directory / f'{prefix}_complete.json'}

def model_checkpoint_paths(checkpoint_dir: Path, split_id: int, model_name: str) -> Dict[str, Path]:
    directory = checkpoint_dir / 'models' / f'split_{split_id:03d}'
    directory.mkdir(parents=True, exist_ok=True)
    safe_model = ''.join((character.lower() if character.isalnum() else '_' for character in model_name))
    prefix = f'split_{split_id:03d}_{safe_model}'
    return {'metrics': directory / f'{prefix}_metrics.csv', 'tuning': directory / f'{prefix}_tuning.csv', 'predictions': directory / f'{prefix}_predictions.csv', 'marker': directory / f'{prefix}_complete.json'}

def validate_imputation_diagnostics(diagnostics_df: pd.DataFrame, split_id: int) -> None:
    if list(diagnostics_df.columns) != IMPUTER_DIAGNOSTIC_COLUMNS:
        raise ValueError('Imputation diagnostic columns are incomplete or reordered.')
    if len(diagnostics_df) != 5:
        raise ValueError('Each split must contain five 10-cycle imputer diagnostics.')
    if set(pd.to_numeric(diagnostics_df['split_id'], errors='raise').astype(int)) != {split_id}:
        raise ValueError('Imputation diagnostics have the wrong split ID.')
    role_counts = diagnostics_df['role'].astype(str).value_counts().to_dict()
    if role_counts != {'classical_inner': 3, 'outer_evaluation': 1, 'neural_development': 1}:
        raise ValueError(f'Imputation roles are incomplete: {role_counts}')
    inner_rows = diagnostics_df.loc[diagnostics_df['role'].astype(str).eq('classical_inner')]
    if set(pd.to_numeric(inner_rows['inner_fold'], errors='raise').astype(int)) != {1, 2, 3}:
        raise ValueError('Classical inner-imputer diagnostics lack folds 1-3.')
    non_inner = diagnostics_df.loc[~diagnostics_df['role'].astype(str).eq('classical_inner')]
    if non_inner['inner_fold'].notna().any():
        raise ValueError('A non-inner imputer was assigned an inner-fold number.')
    outer_row = diagnostics_df.loc[diagnostics_df['role'].astype(str).eq('outer_evaluation')].iloc[0]
    neural_row = diagnostics_df.loc[diagnostics_df['role'].astype(str).eq('neural_development')].iloc[0]
    if int(outer_row['imputer_seed']) != extra_trees_outer_seed(USER_MASTER_SEED, split_id):
        raise ValueError('Outer ExtraTrees imputer seed changed.')
    if int(neural_row['imputer_seed']) != extra_trees_neural_seed(USER_MASTER_SEED, split_id):
        raise ValueError('Neural-development ExtraTrees imputer seed changed.')
    for _, row in inner_rows.iterrows():
        inner_fold = int(row['inner_fold'])
        if int(row['imputer_seed']) != extra_trees_inner_seed(USER_MASTER_SEED, split_id, inner_fold):
            raise ValueError('A classical inner ExtraTrees imputer seed changed.')
    if not (pd.to_numeric(diagnostics_df['n_iter'], errors='raise').astype(int).eq(FIXED_IMPUTATION_CYCLES).all() and pd.to_numeric(diagnostics_df['max_iter'], errors='raise').astype(int).eq(FIXED_IMPUTATION_CYCLES).all() and diagnostics_df['fixed_cycle_policy'].astype(str).eq(FIXED_CYCLE_POLICY).all()):
        raise ValueError('A sensitivity imputer did not use the fixed 10-cycle policy.')
    for column in ['observed_cells_preserved', 'all_values_finite', 'plausibility_passed']:
        if not all((_bool_value(value, column) for value in diagnostics_df[column])):
            raise ValueError(f'Imputation diagnostic {column} did not pass.')
    for column in ['train_missing_cells_after', 'apply_missing_cells_after']:
        if not pd.to_numeric(diagnostics_df[column], errors='raise').astype(int).eq(0).all():
            raise ValueError('Completed imputation data retain missing cells.')
    warnings_count = pd.to_numeric(diagnostics_df['warning_count'], errors='raise').astype(int)
    if not warnings_count.eq(1).all():
        raise ValueError('Every fixed-cycle imputer must record exactly one expected cycle-limit notice.')
    for _, row in diagnostics_df.iterrows():
        message = '' if pd.isna(row['warning']) else str(row['warning'])
        if EXPECTED_ITERATION_WARNING_TEXT not in message:
            raise ValueError('Recorded imputation warning was not the expected cap warning.')
        if str(row['warning_classification']) != 'expected_fixed_cycle_notice_not_convergence_assessment':
            raise ValueError('Imputation warning classification is inconsistent.')
        if int(row['unexpected_warning_count']) != 0:
            raise ValueError('Imputation diagnostics contain an unexpected warning.')

def validate_stability_frame(stability_df: pd.DataFrame, split_id: int, predictor_columns: Sequence[str], main_module: Any) -> None:
    if list(stability_df.columns) != STABILITY_COLUMNS:
        raise ValueError('Cycle-stability columns are incomplete or reordered.')
    if split_id not in STABILITY_SPLIT_IDS:
        if len(stability_df) != 0:
            raise ValueError('A split outside the fixed stability-audit set has cycle-stability rows.')
        return
    if len(stability_df) != len(predictor_columns) + 2:
        raise ValueError('A fixed stability-audit split has incomplete cycle-stability rows.')
    if set(pd.to_numeric(stability_df['split_id'], errors='raise').astype(int)) != {split_id}:
        raise ValueError('Cycle-stability rows have the wrong split ID.')
    if not pd.to_numeric(stability_df['imputer_seed'], errors='raise').astype(int).eq(extra_trees_outer_seed(USER_MASTER_SEED, split_id)).all():
        raise ValueError('Cycle-stability rows have the wrong imputer seed.')
    expected_variables = {*predictor_columns, '__global_continuous__', '__global_binary__'}
    if set(stability_df['variable'].astype(str)) != expected_variables:
        raise ValueError('Cycle-stability rows do not cover every predictor.')
    predictor_rows = stability_df.loc[stability_df['variable'].astype(str).isin(predictor_columns)].copy()
    n_compared = pd.to_numeric(predictor_rows['n_compared'], errors='raise').astype(int)
    for index, row in predictor_rows.iterrows():
        count = int(n_compared.loc[index])
        variable = str(row['variable'])
        if variable in main_module.CONTINUOUS_COLUMNS:
            expected_type = 'continuous'
            expected_threshold = 0.1
        elif variable in main_module.BINARY_COLUMNS:
            expected_type = 'binary'
            expected_threshold = 0.05
        else:
            raise ValueError('A cycle-stability predictor has no declared type.')
        if str(row['variable_type']) != expected_type or not math.isclose(float(row['acceptance_threshold']), expected_threshold, rel_tol=0.0, abs_tol=1e-12):
            raise ValueError('A predictor-level cycle-stability type or threshold changed.')
        status = str(row['assessment_status'])
        applicable = _bool_value(row['threshold_applicable'], 'threshold_applicable')
        passed_missing = pd.isna(row['stability_passed'])
        if count == 0:
            if status != 'not_applicable_no_missing_cells' or applicable or (not passed_missing):
                raise ValueError('A no-missing stability row has an invalid assessment label.')
        elif count < MIN_STABILITY_CELLS_PER_VARIABLE:
            if status != 'sparse_not_individually_assessable' or applicable or (not passed_missing):
                raise ValueError('A sparse stability row was incorrectly labelled stable.')
        else:
            if status != 'individually_assessed' or not applicable:
                raise ValueError('An evaluable variable lacks an individual assessment.')
            if not _bool_value(row['stability_passed'], 'stability_passed'):
                raise ValueError('An individually assessed stability threshold failed.')
            observed_value = float(row['normalized_mean_absolute_difference' if expected_type == 'continuous' else 'binary_disagreement_rate'])
            if not np.isfinite(observed_value) or observed_value > expected_threshold + 1e-12:
                raise ValueError('An individually assessed numeric stability value exceeds its threshold.')
    aggregate_rows = stability_df.loc[stability_df['variable'].astype(str).isin({'__global_continuous__', '__global_binary__'})]
    if len(aggregate_rows) != 2:
        raise ValueError('Cycle stability lacks both aggregate assessments.')
    for _, row in aggregate_rows.iterrows():
        if str(row['assessment_status']) != 'aggregate_assessed' or not _bool_value(row['threshold_applicable'], 'threshold_applicable') or (not _bool_value(row['stability_passed'], 'stability_passed')):
            raise ValueError('An aggregate cycle-stability threshold failed.')
    continuous_aggregate = aggregate_rows.loc[aggregate_rows['variable'].astype(str).eq('__global_continuous__')].iloc[0]
    binary_aggregate = aggregate_rows.loc[aggregate_rows['variable'].astype(str).eq('__global_binary__')].iloc[0]
    if str(continuous_aggregate['variable_type']) != 'aggregate_continuous' or str(binary_aggregate['variable_type']) != 'aggregate_binary' or (not math.isclose(float(continuous_aggregate['acceptance_threshold']), 0.1, rel_tol=0.0, abs_tol=1e-12)) or (not math.isclose(float(binary_aggregate['acceptance_threshold']), 0.05, rel_tol=0.0, abs_tol=1e-12)):
        raise ValueError('Aggregate cycle-stability types or thresholds changed.')
    if not (float(continuous_aggregate['normalized_mean_absolute_difference']) <= 0.1 + 1e-12 and float(binary_aggregate['binary_disagreement_rate']) <= 0.05 + 1e-12):
        raise ValueError('Aggregate cycle-stability values exceed thresholds.')
    if not (pd.to_numeric(stability_df['cycles_primary'], errors='raise').astype(int).eq(FIXED_IMPUTATION_CYCLES).all() and pd.to_numeric(stability_df['cycles_comparator'], errors='raise').astype(int).eq(STABILITY_COMPARISON_CYCLES).all()):
        raise ValueError('Cycle-stability iteration counts changed.')
    for column in ['primary_warning', 'comparator_warning']:
        for value in stability_df[column]:
            message = '' if pd.isna(value) else str(value)
            if EXPECTED_ITERATION_WARNING_TEXT not in message:
                raise ValueError('Cycle-stability output lacks its expected cycle-limit notice.')

def validate_model_checkpoint_frames(main_module: Any, split_id: int, model_name: str, metrics_df: pd.DataFrame, tuning_df: pd.DataFrame, predictions_df: pd.DataFrame, expected_test_ids: pd.Series, expected_model_seed: int) -> None:
    required_metric_columns = {'split_id', 'split_seed', 'model', 'n_train', 'n_test', 'n_test_events', 'split_metrics_complete', 'split_metric_errors_json', 'tau_months', 'brier_horizon_months', *FOLD_PERFORMANCE_METRICS}
    if not required_metric_columns.issubset(metrics_df.columns) or len(metrics_df) != 1:
        raise ValueError('Model metric checkpoint columns or row count are incomplete.')
    metric_row = metrics_df.iloc[0]
    if int(metric_row['split_id']) != split_id or int(metric_row['split_seed']) != USER_MASTER_SEED or str(metric_row['model']) != model_name:
        raise ValueError('Model metric checkpoint has the wrong key.')
    if not _bool_value(metric_row['split_metrics_complete'], 'split_metrics_complete'):
        raise ValueError('Model metric checkpoint is marked incomplete.')
    n_train = int(metric_row['n_train'])
    n_test = int(metric_row['n_test'])
    n_test_events = int(metric_row['n_test_events'])
    if n_train <= 0 or n_test != len(expected_test_ids) or n_train + n_test != int(main_module.EXPECTED_FILTERED_ROWS) or (n_test_events != EXPECTED_TEST_EVENTS):
        raise ValueError('Model metric checkpoint has invalid sample/event counts.')
    parsed_errors = json.loads(str(metric_row['split_metric_errors_json']))
    if not isinstance(parsed_errors, dict) or parsed_errors:
        raise ValueError('Model metric checkpoint contains a required-metric error.')
    for metric in required_split_metrics(model_name):
        validate_performance_metric(metric, float(metric_row[metric]), f'Split {split_id} {model_name} checkpoint')
    tau_months = float(metric_row['tau_months'])
    if not np.isfinite(tau_months) or tau_months <= 0.0:
        raise ValueError('Model metric checkpoint has an invalid evaluation tau.')
    brier_horizon = float(metric_row['brier_horizon_months'])
    if model_name == 'SVM':
        if not np.isnan(brier_horizon):
            raise ValueError('SVM checkpoint unexpectedly has a Brier horizon.')
    elif not np.isfinite(brier_horizon) or brier_horizon <= 0.0:
        raise ValueError('Absolute-risk checkpoint has an invalid Brier horizon.')
    required_tuning_columns = {'split_id', 'split_seed', 'model', 'tuning_method', 'best_params_json', 'model_seed', 'tuning_imputation_policy', 'inner_folds_classical', *main_module.TUNING_SELECTION_COLUMNS}
    if not required_tuning_columns.issubset(tuning_df.columns) or len(tuning_df) != 1:
        raise ValueError('Model tuning checkpoint columns or row count are incomplete.')
    tuning_row = tuning_df.iloc[0]
    if int(tuning_row['split_id']) != split_id or int(tuning_row['split_seed']) != USER_MASTER_SEED or str(tuning_row['model']) != model_name or (int(tuning_row['model_seed']) != expected_model_seed) or (str(tuning_row['tuning_method']) != expected_tuning_method(model_name)) or (str(tuning_row['tuning_imputation_policy']) != expected_tuning_imputation_policy(model_name)):
        raise ValueError('Model tuning checkpoint has an invalid key, seed, or score.')
    main_module.validate_tuning_selection_record(tuning_row, model_name, f'ExtraTrees split {split_id} {model_name}')
    inner_value = pd.to_numeric(pd.Series([tuning_row['inner_folds_classical']]), errors='coerce').iloc[0]
    if model_name in {'CoxPH', 'DeepSurv', 'CoxTime'}:
        if not pd.isna(inner_value):
            raise ValueError('A fixed or neural checkpoint contains a classical inner-fold count.')
    elif not np.isfinite(inner_value) or int(inner_value) != INNER_FOLDS:
        raise ValueError('A classical checkpoint has the wrong inner-fold count.')
    if not isinstance(json.loads(str(tuning_row['best_params_json'])), dict):
        raise ValueError('Best-parameter JSON is not an object.')
    expected_columns = [main_module.STUDY_ID_COL, 'split_id', 'test_order']
    if model_name in main_module.CALIBRATION_MODELS:
        expected_columns.append('predicted_event_probability')
    if list(predictions_df.columns) != expected_columns:
        raise ValueError('Model prediction checkpoint columns changed.')
    if len(predictions_df) != len(expected_test_ids):
        raise ValueError('Model prediction checkpoint has the wrong row count.')
    if study_id_keys(predictions_df[main_module.STUDY_ID_COL]).tolist() != study_id_keys(expected_test_ids).tolist():
        raise ValueError('Model prediction checkpoint study IDs or order changed.')
    if not pd.to_numeric(predictions_df['split_id'], errors='raise').astype(int).eq(split_id).all():
        raise ValueError('Model predictions have the wrong split ID.')
    if pd.to_numeric(predictions_df['test_order'], errors='raise').astype(int).tolist() != list(range(1, len(expected_test_ids) + 1)):
        raise ValueError('Model predictions have the wrong test order.')
    if model_name in main_module.CALIBRATION_MODELS:
        probabilities = pd.to_numeric(predictions_df['predicted_event_probability'], errors='raise').to_numpy(float)
        if not np.isfinite(probabilities).all() or np.any((probabilities < -1e-08) | (probabilities > 1.0 + 1e-08)):
            raise ValueError('Model prediction checkpoint contains invalid probabilities.')

def validate_split_checkpoint_frames(main_module: Any, split_id: int, metrics_df: pd.DataFrame, tuning_df: pd.DataFrame, predictions_df: pd.DataFrame, diagnostics_df: pd.DataFrame, stability_df: pd.DataFrame, expected_test_ids: pd.Series, expected_model_seed: int) -> None:
    model_order = list(main_module.MODEL_ORDER)
    if len(metrics_df) != len(model_order) or metrics_df.duplicated(['split_id', 'model']).any() or set(metrics_df['model'].astype(str)) != set(model_order):
        raise ValueError('Split metric checkpoint does not contain all eight models once.')
    if len(tuning_df) != len(model_order) or tuning_df.duplicated(['split_id', 'model']).any() or set(tuning_df['model'].astype(str)) != set(model_order):
        raise ValueError('Split tuning checkpoint does not contain all eight models once.')
    prediction_base = predictions_df[[main_module.STUDY_ID_COL, 'split_id', 'test_order']]
    for model_name in model_order:
        model_metrics = metrics_df.loc[metrics_df['model'].astype(str).eq(model_name)]
        model_tuning = tuning_df.loc[tuning_df['model'].astype(str).eq(model_name)]
        model_predictions = prediction_base.copy()
        if model_name in main_module.CALIBRATION_MODELS:
            if model_name not in predictions_df.columns:
                raise ValueError(f'Split predictions lack {model_name}.')
            model_predictions['predicted_event_probability'] = predictions_df[model_name]
        validate_model_checkpoint_frames(main_module, split_id, model_name, model_metrics, model_tuning, model_predictions, expected_test_ids, expected_model_seed)
    expected_prediction_columns = [main_module.STUDY_ID_COL, 'split_id', 'test_order', *main_module.CALIBRATION_MODELS]
    if list(predictions_df.columns) != expected_prediction_columns:
        raise ValueError('Split prediction columns are incomplete or reordered.')
    validate_imputation_diagnostics(diagnostics_df, split_id)
    validate_stability_frame(stability_df, split_id, main_module.PREDICTOR_COLUMNS, main_module)

def write_model_checkpoint(main_module: Any, checkpoint_dir: Path, manifest_sha256: str, split_id: int, model_name: str, metrics_df: pd.DataFrame, tuning_df: pd.DataFrame, predictions_df: pd.DataFrame, expected_test_ids: pd.Series, model_seed: int, imputed_train_sha256: str, imputed_test_sha256: str) -> None:
    validate_model_checkpoint_frames(main_module, split_id, model_name, metrics_df, tuning_df, predictions_df, expected_test_ids, model_seed)
    paths = model_checkpoint_paths(checkpoint_dir, split_id, model_name)
    frames = {'metrics': metrics_df, 'tuning': tuning_df, 'predictions': predictions_df}
    for name, frame in frames.items():
        main_module.atomic_write_dataframe_csv(frame, str(paths[name]))
    marker = {'checkpoint_schema_version': CHECKPOINT_SCHEMA_VERSION, 'complete': True, 'manifest_sha256': manifest_sha256, 'split_id': int(split_id), 'model': model_name, 'model_seed': int(model_seed), 'imputed_train_sha256': imputed_train_sha256, 'imputed_test_sha256': imputed_test_sha256, 'tuning_method': expected_tuning_method(model_name), 'tuning_imputation_policy': expected_tuning_imputation_policy(model_name), 'selection_policy': main_module.model_selection_metadata(model_name), 'files': {name: {'name': paths[name].name, 'sha256': main_module.file_sha256(str(paths[name])), 'n_rows': int(len(frame))} for name, frame in frames.items()}}
    main_module.atomic_write_json(marker, str(paths['marker']))

def load_model_checkpoint(main_module: Any, checkpoint_dir: Path, manifest_sha256: str, split_id: int, model_name: str, expected_test_ids: pd.Series, model_seed: int, imputed_train_sha256: str, imputed_test_sha256: str) -> Optional[Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]]:
    paths = model_checkpoint_paths(checkpoint_dir, split_id, model_name)
    if not paths['marker'].is_file():
        return None
    try:
        with open(paths['marker'], 'r', encoding='utf-8') as handle:
            marker = json.load(handle)
        if marker.get('checkpoint_schema_version') != CHECKPOINT_SCHEMA_VERSION or marker.get('complete') is not True or marker.get('manifest_sha256') != manifest_sha256 or (int(marker.get('split_id', -1)) != split_id) or (marker.get('model') != model_name) or (int(marker.get('model_seed', -1)) != model_seed) or (marker.get('imputed_train_sha256') != imputed_train_sha256) or (marker.get('imputed_test_sha256') != imputed_test_sha256) or (marker.get('tuning_method') != expected_tuning_method(model_name)) or (marker.get('tuning_imputation_policy') != expected_tuning_imputation_policy(model_name)) or (marker.get('selection_policy') != main_module.model_selection_metadata(model_name)):
            raise ValueError('Model completion marker metadata did not validate.')
        frames: Dict[str, pd.DataFrame] = {}
        for name in ['metrics', 'tuning', 'predictions']:
            file_meta = marker['files'][name]
            if file_meta.get('name') != paths[name].name or not paths[name].is_file():
                raise ValueError(f'Model {name} checkpoint is missing or renamed.')
            if main_module.file_sha256(str(paths[name])) != file_meta.get('sha256'):
                raise ValueError(f'Model {name} checkpoint hash changed.')
            frame = pd.read_csv(paths[name], float_precision='round_trip')
            if len(frame) != int(file_meta.get('n_rows', -1)):
                raise ValueError(f'Model {name} checkpoint row count changed.')
            frames[name] = frame
        validate_model_checkpoint_frames(main_module, split_id, model_name, frames['metrics'], frames['tuning'], frames['predictions'], expected_test_ids, model_seed)
        return (frames['metrics'], frames['tuning'], frames['predictions'])
    except Exception as exc:
        LOGGER.warning('Split %d %s model checkpoint is invalid and will be recomputed: %s', split_id, model_name, exc)
        return None

def write_split_checkpoint(main_module: Any, checkpoint_dir: Path, manifest_sha256: str, split_id: int, metrics_df: pd.DataFrame, tuning_df: pd.DataFrame, predictions_df: pd.DataFrame, diagnostics_df: pd.DataFrame, stability_df: pd.DataFrame, expected_test_ids: pd.Series, model_seed: int) -> None:
    validate_split_checkpoint_frames(main_module, split_id, metrics_df, tuning_df, predictions_df, diagnostics_df, stability_df, expected_test_ids, model_seed)
    paths = split_checkpoint_paths(checkpoint_dir, split_id)
    frames = {'metrics': metrics_df, 'tuning': tuning_df, 'predictions': predictions_df, 'diagnostics': diagnostics_df, 'stability': stability_df}
    for name, frame in frames.items():
        main_module.atomic_write_dataframe_csv(frame, str(paths[name]))
    marker = {'checkpoint_schema_version': CHECKPOINT_SCHEMA_VERSION, 'complete': True, 'manifest_sha256': manifest_sha256, 'split_id': int(split_id), 'model_seed': int(model_seed), 'files': {name: {'name': paths[name].name, 'sha256': main_module.file_sha256(str(paths[name])), 'n_rows': int(len(frame))} for name, frame in frames.items()}}
    main_module.atomic_write_json(marker, str(paths['marker']))

def load_split_checkpoint(main_module: Any, checkpoint_dir: Path, manifest_sha256: str, split_id: int, expected_test_ids: pd.Series, model_seed: int) -> Optional[Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]]:
    paths = split_checkpoint_paths(checkpoint_dir, split_id)
    if not paths['marker'].is_file():
        return None
    try:
        with open(paths['marker'], 'r', encoding='utf-8') as handle:
            marker = json.load(handle)
        if marker.get('checkpoint_schema_version') != CHECKPOINT_SCHEMA_VERSION or marker.get('complete') is not True or marker.get('manifest_sha256') != manifest_sha256 or (int(marker.get('split_id', -1)) != split_id) or (int(marker.get('model_seed', -1)) != model_seed):
            raise ValueError('Split completion marker metadata did not validate.')
        frames: Dict[str, pd.DataFrame] = {}
        for name in ['metrics', 'tuning', 'predictions', 'diagnostics', 'stability']:
            file_meta = marker['files'][name]
            if file_meta.get('name') != paths[name].name or not paths[name].is_file():
                raise ValueError(f'Split {name} checkpoint is missing or renamed.')
            if main_module.file_sha256(str(paths[name])) != file_meta.get('sha256'):
                raise ValueError(f'Split {name} checkpoint hash changed.')
            frame = pd.read_csv(paths[name], float_precision='round_trip')
            if len(frame) != int(file_meta.get('n_rows', -1)):
                raise ValueError(f'Split {name} checkpoint row count changed.')
            frames[name] = frame
        validate_split_checkpoint_frames(main_module, split_id, frames['metrics'], frames['tuning'], frames['predictions'], frames['diagnostics'], frames['stability'], expected_test_ids, model_seed)
        return (frames['metrics'], frames['tuning'], frames['predictions'], frames['diagnostics'], frames['stability'])
    except Exception as exc:
        LOGGER.warning('Split %d checkpoint is invalid and will be recomputed: %s', split_id, exc)
        return None

def fit_one_outer_split(main_module: Any, df: pd.DataFrame, train_idx: np.ndarray, test_idx: np.ndarray, split_id: int, args: argparse.Namespace, checkpoint_dir: Path, manifest_sha256: str, deadline: Optional[float]) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    raw_train_df = df.iloc[train_idx].copy().reset_index(drop=True)
    raw_test_df = df.iloc[test_idx].copy().reset_index(drop=True)
    model_seed = matched_model_seed(args.seed, split_id)
    outer_seed = extra_trees_outer_seed(args.seed, split_id)
    train_df, test_df, outer_diagnostic = generate_extra_trees_imputed_split(raw_train_df, raw_test_df, main_module=main_module, imputer_seed=outer_seed, cycles=FIXED_IMPUTATION_CYCLES, n_jobs=args.n_jobs, split_id=split_id, role='outer_evaluation', inner_fold=None)
    inner_completed_splits, inner_diagnostics = generate_extra_trees_inner_completed_splits(raw_outer_train_df=raw_train_df, main_module=main_module, split_id=split_id, model_seed=model_seed, master_seed=args.seed, n_jobs=args.n_jobs)
    neural_selection_split, neural_diagnostic = generate_extra_trees_neural_selection_split(raw_outer_train_df=raw_train_df, main_module=main_module, split_id=split_id, model_seed=model_seed, master_seed=args.seed, n_jobs=args.n_jobs)
    diagnostics_df = pd.DataFrame([outer_diagnostic, *inner_diagnostics, neural_diagnostic], columns=IMPUTER_DIAGNOSTIC_COLUMNS)
    validate_imputation_diagnostics(diagnostics_df, split_id)
    if split_id in STABILITY_SPLIT_IDS:
        stability_df = build_stability_diagnostic(raw_train_df, raw_test_df, train_df, test_df, outer_diagnostic, main_module, split_id, outer_seed, args.n_jobs)
    else:
        stability_df = pd.DataFrame(columns=STABILITY_COLUMNS)
    validate_stability_frame(stability_df, split_id, main_module.PREDICTOR_COLUMNS, main_module)
    imputed_train_sha256 = dataframe_sha256(train_df)
    imputed_test_sha256 = dataframe_sha256(test_df)
    expected_test_ids = raw_test_df[main_module.STUDY_ID_COL].reset_index(drop=True)
    metric_rows: List[Dict[str, Any]] = []
    tuning_rows: List[Dict[str, Any]] = []
    predictions = pd.DataFrame({main_module.STUDY_ID_COL: raw_test_df[main_module.STUDY_ID_COL].to_numpy(), 'split_id': int(split_id), 'test_order': np.arange(1, len(raw_test_df) + 1, dtype=int)})
    fit_order = ['CoxPH', 'ElasticNetCox', 'RSF', 'GBSA', 'XGBoost', 'SVM', 'DeepSurv', 'CoxTime']
    if set(fit_order) != set(main_module.MODEL_ORDER):
        raise RuntimeError('The main eight-model set changed unexpectedly.')
    for model_name in fit_order:
        loaded_model = None
        if args.resume:
            loaded_model = load_model_checkpoint(main_module, checkpoint_dir, manifest_sha256, split_id, model_name, expected_test_ids, model_seed, imputed_train_sha256, imputed_test_sha256)
        if loaded_model is not None:
            metric_rows.extend(loaded_model[0].to_dict('records'))
            tuning_rows.extend(loaded_model[1].to_dict('records'))
            if model_name in main_module.CALIBRATION_MODELS:
                predictions[model_name] = pd.to_numeric(loaded_model[2]['predicted_event_probability'], errors='raise').to_numpy(float)
            LOGGER.info('Split %d: resumed completed %s model checkpoint', split_id, model_name)
            continue
        LOGGER.info('Split %d: fitting %s', split_id, model_name)
        fitted, params, selection_score, tuning_method = tune_and_fit_model(main_module, model_name, train_df, model_seed=model_seed, n_jobs=args.n_jobs, inner_completed_splits=inner_completed_splits, neural_selection_split=neural_selection_split, ibs_start_month=args.ibs_start_month, horizon_months=args.horizon_months)
        metrics, _, predicted_event_probability = main_module.evaluate_model(model_name, fitted, train_df, test_df, args.horizon_months, args.ibs_start_month, args.calibration_groups, compute_auxiliary_calibration=False)
        required_errors = validate_one_model_split_result(model_name, metrics, predicted_event_probability, expected_test_n=len(test_df), requested_horizon=args.horizon_months, context=f'ExtraTrees sensitivity split {split_id}')
        metric_row = {'split_id': int(split_id), 'split_seed': int(args.seed), 'model': model_name, 'n_train': int(len(train_df)), 'n_test': int(len(test_df)), 'n_test_events': int(test_df[main_module.EVENT_COL].sum()), 'split_metrics_complete': True, 'split_metric_errors_json': json.dumps(required_errors, sort_keys=True), **metrics}
        tuning_row = {'split_id': int(split_id), 'split_seed': int(args.seed), 'model': model_name, 'tuning_method': tuning_method, 'inner_folds_classical': INNER_FOLDS if model_name not in {'CoxPH', 'DeepSurv', 'CoxTime'} else np.nan, **main_module.tuning_output_metadata(model_name, selection_score), 'best_params_json': json.dumps(params, sort_keys=True), 'model_seed': int(model_seed), 'tuning_imputation_policy': expected_tuning_imputation_policy(model_name)}
        metric_rows.append(metric_row)
        tuning_rows.append(tuning_row)
        model_predictions = predictions[[main_module.STUDY_ID_COL, 'split_id', 'test_order']].copy()
        if model_name in main_module.CALIBRATION_MODELS:
            model_probability = np.clip(np.asarray(predicted_event_probability, dtype=float), 0.0, 1.0)
            predictions[model_name] = model_probability
            model_predictions['predicted_event_probability'] = model_probability
        write_model_checkpoint(main_module, checkpoint_dir, manifest_sha256, split_id, model_name, pd.DataFrame([metric_row]), pd.DataFrame([tuning_row]), model_predictions, expected_test_ids, model_seed, imputed_train_sha256, imputed_test_sha256)
        LOGGER.info('Split %d: checkpointed %s', split_id, model_name)
        if deadline_reached(deadline):
            raise PlannedSensitivityPause(split_id, model_name, [str(row['model']) for row in tuning_rows])
    predictions = predictions[[main_module.STUDY_ID_COL, 'split_id', 'test_order', *main_module.CALIBRATION_MODELS]]
    return (pd.DataFrame(metric_rows), pd.DataFrame(tuning_rows), predictions, diagnostics_df, stability_df)

def write_planned_pause_payload(main_module: Any, metadata_dir: Path, manifest_sha256: str, completed_splits: Sequence[int], reason: str, next_split: int, partial_split: Optional[int]=None, last_checkpointed_model: Optional[str]=None, partial_split_completed_models: Optional[Sequence[str]]=None) -> Dict[str, Any]:
    normalized = sorted({int(value) for value in completed_splits})
    payload: Dict[str, Any] = {'status': 'paused', 'reason': str(reason), 'completed_splits_this_resume_scan': normalized, 'completed_split_count': int(len(normalized)), 'next_split': int(next_split), 'manifest_sha256': manifest_sha256, 'model_checkpoint_resume_supported': True, 'instruction': 'Rerun the unchanged one-click launcher; validated model and split checkpoints will resume.'}
    if partial_split is not None:
        payload['partial_split'] = int(partial_split)
    if last_checkpointed_model is not None:
        payload['last_checkpointed_model'] = str(last_checkpointed_model)
    if partial_split_completed_models is not None:
        payload['partial_split_completed_models'] = [str(value) for value in partial_split_completed_models]
        payload['partial_split_completed_model_count'] = int(len(payload['partial_split_completed_models']))
    main_module.atomic_write_json(payload, str(metadata_dir / 'analysis_complete.json'))
    return payload

def run_matched_200_split_sensitivity(main_module: Any, df: pd.DataFrame, definitions: Sequence[Tuple[int, np.ndarray, np.ndarray]], checkpoint_dir: Path, manifest_sha256: str, args: argparse.Namespace) -> Tuple[bool, List[pd.DataFrame], List[pd.DataFrame], List[pd.DataFrame], List[pd.DataFrame], List[pd.DataFrame]]:
    all_metrics: List[pd.DataFrame] = []
    all_tuning: List[pd.DataFrame] = []
    all_predictions: List[pd.DataFrame] = []
    all_diagnostics: List[pd.DataFrame] = []
    all_stability: List[pd.DataFrame] = []
    deadline = session_deadline(args.max_session_hours)
    for split_id, train_idx, test_idx in definitions:
        expected_test_ids = df.iloc[test_idx][main_module.STUDY_ID_COL].reset_index(drop=True)
        model_seed = matched_model_seed(args.seed, split_id)
        loaded = None
        if args.resume:
            loaded = load_split_checkpoint(main_module, checkpoint_dir, manifest_sha256, split_id, expected_test_ids, model_seed)
        if loaded is None:
            LOGGER.info('Running matched ExtraTrees split %d / %d', split_id, SENSITIVITY_SPLITS)
            try:
                loaded = fit_one_outer_split(main_module, df, train_idx, test_idx, split_id, args, checkpoint_dir, manifest_sha256, deadline)
            except PlannedSensitivityPause as exc:
                exc.completed_splits = sorted({int(frame['split_id'].iloc[0]) for frame in all_metrics if not frame.empty})
                raise
            write_split_checkpoint(main_module, checkpoint_dir, manifest_sha256, split_id, loaded[0], loaded[1], loaded[2], loaded[3], loaded[4], expected_test_ids, model_seed)
            LOGGER.info('Split %d checkpoint was written and verified.', split_id)
        else:
            LOGGER.info('Resumed split %d from a validated checkpoint.', split_id)
        all_metrics.append(loaded[0])
        all_tuning.append(loaded[1])
        all_predictions.append(loaded[2])
        all_diagnostics.append(loaded[3])
        all_stability.append(loaded[4])
        if split_id < SENSITIVITY_SPLITS and deadline_reached(deadline):
            return (False, all_metrics, all_tuning, all_predictions, all_diagnostics, all_stability)
    return (True, all_metrics, all_tuning, all_predictions, all_diagnostics, all_stability)

def verify_prediction_coverage_and_average(main_module: Any, df: pd.DataFrame, prediction_frames: Sequence[pd.DataFrame], matched_map: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    if len(prediction_frames) != SENSITIVITY_SPLITS:
        raise RuntimeError('Prediction verification requires all 200 split checkpoints.')
    stacked = pd.concat(prediction_frames, ignore_index=True)
    stacked['study_id_key'] = study_id_keys(stacked[main_module.STUDY_ID_COL])
    expected = matched_map[['split_id', 'study_id_key', 'test_order']].copy()
    expected['split_id'] = pd.to_numeric(expected['split_id'], errors='raise').astype(int)
    expected['test_order'] = pd.to_numeric(expected['test_order'], errors='raise').astype(int)
    observed = stacked[['split_id', 'study_id_key', 'test_order']].copy()
    observed['split_id'] = pd.to_numeric(observed['split_id'], errors='raise').astype(int)
    observed['test_order'] = pd.to_numeric(observed['test_order'], errors='raise').astype(int)
    if len(observed) != len(expected):
        raise RuntimeError(f'Prediction checkpoints contain {len(observed)} rows; expected {len(expected)}.')
    if len(observed) != EXPECTED_HELDOUT_PREDICTION_ROWS:
        raise RuntimeError(f'Held-out prediction rows are {len(observed)}; expected {EXPECTED_HELDOUT_PREDICTION_ROWS}.')
    if observed.duplicated(['split_id', 'study_id_key']).any():
        raise RuntimeError('Predictions duplicate a participant within a split.')
    expected_sorted = expected.sort_values(['split_id', 'test_order']).reset_index(drop=True)
    observed_sorted = observed.sort_values(['split_id', 'test_order']).reset_index(drop=True)
    if not expected_sorted.equals(observed_sorted):
        mismatch = expected_sorted.merge(observed_sorted, on=['split_id', 'test_order'], how='outer', suffixes=('_expected', '_observed'), indicator=True)
        mismatch = mismatch.loc[mismatch['_merge'].ne('both') | mismatch['study_id_key_expected'].ne(mismatch['study_id_key_observed'])]
        raise RuntimeError(f"Prediction participant assignments changed from the main map. First differences: {mismatch.head(10).to_dict('records')}")
    probabilities = stacked[main_module.CALIBRATION_MODELS].apply(pd.to_numeric, errors='raise')
    probability_values = probabilities.to_numpy(float)
    if not np.isfinite(probability_values).all() or np.any((probability_values < -1e-08) | (probability_values > 1.0 + 1e-08)):
        raise RuntimeError('Assembled held-out probabilities are invalid.')
    source_ids = study_id_keys(df[main_module.STUDY_ID_COL])
    counts = stacked.groupby('study_id_key').size()
    missing = sorted(set(source_ids).difference(counts.index))
    if missing:
        raise RuntimeError(f'Participant-mean calibration is impossible: {len(missing)} participants were never held out in the 200 splits.')
    if int(counts.min()) != EXPECTED_MIN_HELDOUT_PREDICTIONS or int(counts.max()) != EXPECTED_MAX_HELDOUT_PREDICTIONS:
        raise RuntimeError(f'Held-out prediction counts per participant changed: observed {int(counts.min())}-{int(counts.max())}, expected {EXPECTED_MIN_HELDOUT_PREDICTIONS}-{EXPECTED_MAX_HELDOUT_PREDICTIONS}.')
    mean_predictions = stacked.groupby('study_id_key', as_index=False)[main_module.CALIBRATION_MODELS].mean()
    if len(mean_predictions) != len(df):
        raise RuntimeError('Mean held-out predictions do not cover the eligible cohort.')
    mean_values = mean_predictions[main_module.CALIBRATION_MODELS].to_numpy(float)
    if not np.isfinite(mean_values).all() or np.any((mean_values < -1e-08) | (mean_values > 1.0 + 1e-08)):
        raise RuntimeError('Participant-mean held-out probabilities are invalid.')
    coverage_rows: List[Dict[str, Any]] = []
    for model_name in main_module.CALIBRATION_MODELS:
        coverage_rows.append({'model': model_name, 'eligible_participants': int(len(df)), 'held_out_prediction_rows': int(len(stacked)), 'unique_participants': int(len(counts)), 'minimum_predictions_per_participant': int(counts.min()), 'maximum_predictions_per_participant': int(counts.max()), 'mean_predictions_per_participant': float(counts.mean()), 'participants_without_prediction': int(len(missing)), 'coverage_complete': True, 'calibration_prediction_basis': 'one_mean_held_out_prediction_per_participant'})
    return (stacked, mean_predictions, pd.DataFrame(coverage_rows))

def final_split_metrics(main_module: Any, checkpoint_metrics: pd.DataFrame) -> pd.DataFrame:
    columns = ['split_id', 'split_seed', 'model', 'n_train', 'n_test', 'n_test_events', 'split_metrics_complete', 'split_metric_errors_json', 'tau_months', *FOLD_PERFORMANCE_METRICS, 'brier_horizon_months']
    output = checkpoint_metrics.loc[:, columns].copy()
    order = {model: index for index, model in enumerate(main_module.MODEL_ORDER)}
    output['_model_order'] = output['model'].map(order)
    output = output.sort_values(['split_id', '_model_order']).drop(columns='_model_order')
    return output.reset_index(drop=True)

def descriptive_summary(values: np.ndarray) -> Dict[str, Any]:
    values = np.asarray(values, dtype=float)
    if values.size == 0 or not np.isfinite(values).all():
        raise ValueError('A descriptive summary received incomplete values.')
    lower, upper = np.quantile(values, [0.025, 0.975], method='linear')
    return {'n_splits': int(values.size), 'mean': float(np.mean(values)), 'empirical_2_5_percentile': float(lower), 'empirical_97_5_percentile': float(upper), 'range_interpretation': 'descriptive_empirical_percentile_range_not_confidence_interval'}

def summarise_split_metrics(main_module: Any, split_metrics: pd.DataFrame) -> pd.DataFrame:
    rows: List[Dict[str, Any]] = []
    for model_name in main_module.MODEL_ORDER:
        model_rows = split_metrics.loc[split_metrics['model'].eq(model_name)]
        if len(model_rows) != SENSITIVITY_SPLITS:
            raise RuntimeError(f'{model_name} lacks {SENSITIVITY_SPLITS} sensitivity split estimates.')
        row: Dict[str, Any] = {'model': model_name, 'summary_type': 'descriptive_across_200_matched_held_out_splits'}
        for metric in required_split_metrics(model_name):
            summary = descriptive_summary(pd.to_numeric(model_rows[metric], errors='raise').to_numpy(float))
            for label, value in summary.items():
                row[f'{metric}_{label}'] = value
        rows.append(row)
    return pd.DataFrame(rows)

def paired_differences_vs_coxph(main_module: Any, split_metrics: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame]:
    detail_rows: List[Dict[str, Any]] = []
    reference = split_metrics.loc[split_metrics['model'].eq('CoxPH')].set_index('split_id')
    for model_name in main_module.MODEL_ORDER:
        if model_name == 'CoxPH':
            continue
        model_rows = split_metrics.loc[split_metrics['model'].eq(model_name)].set_index('split_id')
        for metric in required_split_metrics(model_name):
            aligned = model_rows[[metric]].join(reference[[metric]], how='inner', lsuffix='_model', rsuffix='_coxph', validate='one_to_one')
            if len(aligned) != SENSITIVITY_SPLITS:
                raise RuntimeError(f'Paired {metric} comparison for {model_name} lacks {SENSITIVITY_SPLITS} splits.')
            for split_id, row in aligned.iterrows():
                model_value = float(row[f'{metric}_model'])
                reference_value = float(row[f'{metric}_coxph'])
                detail_rows.append({'split_id': int(split_id), 'model': model_name, 'reference_model': 'CoxPH', 'metric': metric, 'model_value': model_value, 'coxph_value': reference_value, 'difference_model_minus_coxph': model_value - reference_value})
    detailed = pd.DataFrame(detail_rows)
    summary_rows: List[Dict[str, Any]] = []
    for (model_name, metric), rows in detailed.groupby(['model', 'metric'], sort=False):
        summary_rows.append({'model': model_name, 'reference_model': 'CoxPH', 'metric': metric, **{f'difference_{key}': value for key, value in descriptive_summary(rows['difference_model_minus_coxph'].to_numpy(float)).items()}, 'inference': 'descriptive_no_confidence_interval_or_p_value'})
    return (detailed, pd.DataFrame(summary_rows))

def compare_with_main_bayesian(main_module: Any, main_performance: pd.DataFrame, sensitivity_performance: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame]:
    key_columns = ['split_id', 'model']
    merged = sensitivity_performance[key_columns + FOLD_PERFORMANCE_METRICS].merge(main_performance[key_columns + FOLD_PERFORMANCE_METRICS], on=key_columns, how='inner', validate='one_to_one', suffixes=('_extra_trees', '_bayesian'))
    expected_rows = SENSITIVITY_SPLITS * len(main_module.MODEL_ORDER)
    if len(merged) != expected_rows:
        raise RuntimeError('Matched ExtraTrees-versus-Bayesian comparison is incomplete.')
    detail_rows: List[Dict[str, Any]] = []
    for _, row in merged.iterrows():
        model_name = str(row['model'])
        for metric in required_split_metrics(model_name):
            extra_value = float(row[f'{metric}_extra_trees'])
            bayesian_value = float(row[f'{metric}_bayesian'])
            detail_rows.append({'split_id': int(row['split_id']), 'model': model_name, 'metric': metric, 'extra_trees_value': extra_value, 'bayesian_value': bayesian_value, 'difference_extra_trees_minus_bayesian': extra_value - bayesian_value})
    detailed = pd.DataFrame(detail_rows)
    summary_rows: List[Dict[str, Any]] = []
    for (model_name, metric), rows in detailed.groupby(['model', 'metric'], sort=False):
        summary_rows.append({'model': model_name, 'metric': metric, **{f'difference_{key}': value for key, value in descriptive_summary(rows['difference_extra_trees_minus_bayesian'].to_numpy(float)).items()}, 'inference': 'descriptive_no_confidence_interval_or_p_value'})
    return (detailed, pd.DataFrame(summary_rows))

def participant_mean_calibration(main_module: Any, df: pd.DataFrame, mean_predictions: pd.DataFrame, horizon_months: float, calibration_groups: int, bootstrap_reps: int, seed: int) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    source = df[[main_module.STUDY_ID_COL, main_module.TIME_COL, main_module.EVENT_COL]].copy()
    source['study_id_key'] = study_id_keys(source[main_module.STUDY_ID_COL])
    merged = source.merge(mean_predictions, on='study_id_key', how='left', validate='one_to_one')
    if len(merged) != len(df) or merged[main_module.CALIBRATION_MODELS].isna().any().any():
        raise RuntimeError('Participant-mean calibration data are incomplete.')
    outcome_df = merged[[main_module.STUDY_ID_COL, main_module.TIME_COL, main_module.EVENT_COL]].copy()
    metric_rows: List[Dict[str, Any]] = []
    grouped_rows: List[pd.DataFrame] = []
    flexible_rows: List[pd.DataFrame] = []
    for model_name in main_module.CALIBRATION_MODELS:
        probabilities = merged[model_name].to_numpy(dtype=float)
        oe = main_module.calibration_oe_from_probs(outcome_df, probabilities, time_point=horizon_months)
        slope = main_module.calibration_slope_from_probs(outcome_df, probabilities, time_point=horizon_months)
        required = {'calibration_observed_km_risk': float(oe['calibration_observed_km_risk']), 'calibration_mean_predicted_risk': float(oe['calibration_mean_predicted_risk']), 'calibration_oe_ratio': float(oe['calibration_oe_ratio']), 'calibration_slope': float(slope['calibration_slope'])}
        if not np.isfinite(list(required.values())).all():
            raise RuntimeError(f'Participant-mean calibration is incomplete for {model_name}.')
        metric_rows.append({'model': model_name, 'n_participants': int(len(outcome_df)), 'n_ascvd_events': int(outcome_df[main_module.EVENT_COL].sum()), 'horizon_months': float(horizon_months), **required, 'prediction_basis': 'one_mean_held_out_prediction_per_participant_across_all_200_splits', 'calibration_estimand': 'cross_fitted_repeated_split_ensemble_not_one_final_deployable_model', 'uncertainty_scope': 'conditional_participant_bootstrap_does_not_include_model_development_or_split_stream_selection'})
        grouped = main_module.grouped_calibration_table(outcome_df, probabilities, time_point=horizon_months, n_groups=calibration_groups).copy()
        if len(grouped) != calibration_groups:
            raise RuntimeError(f'Grouped participant-mean calibration is incomplete for {model_name}.')
        grouped.insert(0, 'model', model_name)
        grouped['prediction_basis'] = 'one_mean_held_out_prediction_per_participant'
        grouped_rows.append(grouped)
        flexible = main_module.flexible_calibration_curve_with_bootstrap(outcome_df, probabilities, time_point=horizon_months, n_bootstrap=bootstrap_reps, random_state=seed + 700000 + main_module.MODEL_ORDER.index(model_name), n_grid=FLEXIBLE_CALIBRATION_GRID_POINTS).copy()
        minimum_successful = int(math.ceil(0.9 * bootstrap_reps))
        successful = int(flexible['n_bootstrap_successful'].iloc[0])
        if successful < minimum_successful:
            raise RuntimeError(f'Flexible participant-mean calibration for {model_name} completed {successful}/{bootstrap_reps} bootstraps; at least {minimum_successful} are required.')
        flexible.insert(0, 'model', model_name)
        flexible['prediction_basis'] = 'one_mean_held_out_prediction_per_participant'
        flexible['band_interpretation'] = 'pointwise_participant_bootstrap_conditional_on_averaged_predictions'
        flexible_rows.append(flexible)
    metrics_df = pd.DataFrame(metric_rows)
    grouped_df = pd.concat(grouped_rows, ignore_index=True)
    flexible_df = pd.concat(flexible_rows, ignore_index=True)
    if len(metrics_df) != len(main_module.CALIBRATION_MODELS):
        raise RuntimeError('Calibration summary does not include all absolute-risk models.')
    if len(grouped_df) != len(main_module.CALIBRATION_MODELS) * calibration_groups:
        raise RuntimeError('Grouped calibration has an incomplete model-group grid.')
    return (metrics_df, grouped_df, flexible_df)

def _calibration_axis_max(tables: Iterable[pd.DataFrame], columns: Sequence[str]) -> float:
    maxima: List[float] = []
    for table in tables:
        for column in columns:
            if column in table.columns:
                values = pd.to_numeric(table[column], errors='coerce').to_numpy(float)
                finite = values[np.isfinite(values)]
                if finite.size:
                    maxima.append(float(np.max(finite)))
    raw_max = max(maxima, default=0.05) * 1.1
    return min(1.0, max(0.05, math.ceil(raw_max / 0.05) * 0.05))

def _save_landscape_figure(fig: Any, png_path: Path, tiff_path: Path) -> None:
    if getattr(fig, '_suptitle', None) is not None or list(getattr(fig, 'texts', [])):
        raise RuntimeError('Calibration composites must not contain an overall title or figure-level footer.')
    png_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(png_path, dpi=600, facecolor='white')
    fig.savefig(tiff_path, format='tiff', dpi=600, facecolor='white', pil_kwargs={'compression': 'tiff_lzw'})
    plt.close(fig)
    for path in [png_path, tiff_path]:
        if not path.is_file() or path.stat().st_size == 0:
            raise RuntimeError(f'Calibration figure was not created: {path}')
        with Image.open(path) as image:
            if image.width <= image.height:
                raise RuntimeError(f'Calibration figure is not landscape: {path}')
            dpi = image.info.get('dpi', (0, 0))
            if min(float(dpi[0]), float(dpi[1])) < 599.0:
                raise RuntimeError(f'Calibration figure is not 600 dpi: {path} ({dpi})')
            if path.suffix.lower() in {'.tif', '.tiff'}:
                compression = str(image.info.get('compression', '')).lower()
                if compression not in {'tiff_lzw', 'lzw', '5'}:
                    raise RuntimeError(f'Calibration TIFF is not LZW-compressed: {path} ({compression})')

def make_grouped_calibration_figure(main_module: Any, grouped_df: pd.DataFrame, figure_dir: Path, horizon_months: float) -> None:
    tables = {model: grouped_df.loc[grouped_df['model'].eq(model)].copy() for model in main_module.CALIBRATION_MODELS}
    axis_max = _calibration_axis_max(tables.values(), ['mean_predicted_risk', 'observed_km_risk', 'observed_km_risk_95ci_high'])
    fig, axes = plt.subplots(2, 4, figsize=(14.5, 7.2))
    axes = axes.ravel()
    for panel_index, (ax, model_name) in enumerate(zip(axes, main_module.MODEL_ORDER)):
        panel_label = f"{chr(ord('A') + panel_index)}. {model_name}"
        table = tables.get(model_name)
        if table is None or table.empty:
            ax.axis('off')
            ax.text(0.5, 0.5, f'{panel_label}\nNot available (no survival probabilities)', ha='center', va='center', fontsize=10)
            continue
        ax.plot([0, axis_max], [0, axis_max], linestyle='--', linewidth=1, color='0.4')
        x = table['mean_predicted_risk'].to_numpy(dtype=float)
        y = table['observed_km_risk'].to_numpy(dtype=float)
        low = table['observed_km_risk_95ci_low'].to_numpy(dtype=float)
        high = table['observed_km_risk_95ci_high'].to_numpy(dtype=float)
        error = np.vstack([y - low, high - y])
        finite_interval = np.isfinite(error).all(axis=0)
        if finite_interval.any():
            ax.errorbar(x[finite_interval], y[finite_interval], yerr=np.maximum(error[:, finite_interval], 0.0), marker='o', linestyle='none', linewidth=1, capsize=2, color='#1f77b4')
        if (~finite_interval).any():
            ax.scatter(x[~finite_interval], y[~finite_interval], color='#1f77b4')
        for _, row in table.iterrows():
            ax.annotate(str(int(row['group'])), (row['mean_predicted_risk'], row['observed_km_risk']), textcoords='offset points', xytext=(4, 4), fontsize=7)
        count_columns = {'n', 'events_by_horizon', 'n_at_risk_at_horizon'}
        if not count_columns.issubset(table.columns):
            raise RuntimeError(f'Grouped calibration counts are missing for {model_name}.')
        count_lines = ['Group  N/events/at risk']
        for _, row in table.sort_values('group').iterrows():
            count_lines.append(f"{int(row['group']):>2}   {int(row['n'])}/{int(row['events_by_horizon'])}/{int(row['n_at_risk_at_horizon'])}")
        ax.text(0.98, 0.02, '\n'.join(count_lines), transform=ax.transAxes, ha='right', va='bottom', fontsize=5.8, family='monospace', linespacing=1.0, bbox={'facecolor': 'white', 'edgecolor': '0.8', 'alpha': 0.88, 'pad': 1.2})
        ax.set_title(panel_label, fontsize=10)
        ax.set_xlabel(f'Predicted {int(horizon_months)}-month risk', fontsize=8)
        ax.set_ylabel('Observed KM risk', fontsize=8)
        ax.set_xlim(0, axis_max)
        ax.set_ylim(0, axis_max)
        ax.set_aspect('equal', adjustable='box')
        ax.tick_params(axis='both', labelsize=7)
    fig.tight_layout(w_pad=1.0, h_pad=1.0)
    _save_landscape_figure(fig, figure_dir / 'sensitivity_extra_trees_grouped_calibration_mean_heldout.png', figure_dir / 'sensitivity_extra_trees_grouped_calibration_mean_heldout.tiff')

def make_flexible_calibration_figure(main_module: Any, flexible_df: pd.DataFrame, figure_dir: Path, horizon_months: float) -> None:
    tables = {model: flexible_df.loc[flexible_df['model'].eq(model)].copy() for model in main_module.CALIBRATION_MODELS}
    axis_max = _calibration_axis_max(tables.values(), ['predicted_risk', 'observed_risk_smoothed', 'bootstrap_95ci_high'])
    fig, axes = plt.subplots(2, 4, figsize=(14.5, 7.2))
    axes = axes.ravel()
    for panel_index, (ax, model_name) in enumerate(zip(axes, main_module.MODEL_ORDER)):
        panel_label = f"{chr(ord('A') + panel_index)}. {model_name}"
        table = tables.get(model_name)
        if table is None or table.empty:
            ax.axis('off')
            ax.text(0.5, 0.5, f'{panel_label}\nNot available (no survival probabilities)', ha='center', va='center', fontsize=10)
            continue
        x = table['predicted_risk'].to_numpy(dtype=float)
        y = table['observed_risk_smoothed'].to_numpy(dtype=float)
        low = table['bootstrap_95ci_low'].to_numpy(dtype=float)
        high = table['bootstrap_95ci_high'].to_numpy(dtype=float)
        ax.plot([0, axis_max], [0, axis_max], linestyle='--', linewidth=1, color='0.4')
        ax.plot(x, y, linewidth=1.5, color='#1f77b4')
        finite_band = np.isfinite(low) & np.isfinite(high)
        if finite_band.any():
            ax.fill_between(x[finite_band], low[finite_band], high[finite_band], color='#1f77b4', alpha=0.2, linewidth=0)
        ax.set_title(panel_label, fontsize=10)
        ax.set_xlabel(f'Predicted {int(horizon_months)}-month risk', fontsize=8)
        ax.set_ylabel('Smoothed observed risk', fontsize=8)
        ax.set_xlim(0, axis_max)
        ax.set_ylim(0, axis_max)
        ax.set_aspect('equal', adjustable='box')
        ax.tick_params(axis='both', labelsize=7)
    fig.tight_layout(w_pad=1.0, h_pad=1.0)
    _save_landscape_figure(fig, figure_dir / 'sensitivity_extra_trees_flexible_calibration_mean_heldout.png', figure_dir / 'sensitivity_extra_trees_flexible_calibration_mean_heldout.tiff')

def write_interpretation_report(out_path: Path) -> None:
    lines = ['ExtraTrees predictor-only imputation sensitivity', '=' * 55, '', 'Design', '------', 'This sensitivity uses the exact participant assignments and test-row order from all 200 splits of the main stratified 80/20 analysis. All eight models are refitted within each outer training sample. Six configurable absolute-risk models are selected by the lowest training-only validation IBS from 12 to 120 months. Every IBS censoring distribution is estimated only from the corresponding inner-training or neural-development outcomes. Neural early stopping remains based on validation loss. Survival SVM is selected by the highest validation Harrell C, and CoxPH has one fixed specification.', 'The model and tuning seed for each split is identical to the corresponding main Bayesian-imputation analysis. Only the imputer family and its separate deterministic seed differ.', '', 'Fixed-cycle ExtraTrees imputation', '----------------------------------', 'For this analysis, every outer, inner, and neural-development ExtraTrees IterativeImputer uses a prespecified 10-cycle procedure with tol=0. Reaching the iteration cap is expected and is recorded. It is not presented as a failed stopping criterion or as evidence of convergence.', 'Only baseline predictors enter the imputation matrix. The outer imputer is fitted on outer-training predictors and applied to held-out predictors. Each classical inner-fold imputer is fitted only on its inner-training predictors. The neural-development imputer is fitted only on the development predictors.', "The diagnostics report all five imputer roles in every split: outer evaluation, three classical inner folds, and neural development. The expected scikit-learn iteration-cap message is classified as a fixed-cycle notice, not as an unresolved warning or a convergence assessment. The unexpected-warning count must remain zero. Any unexpected warning, non-finite value, retained missing cell, changed observed value, or failed plausibility check stops the analysis. Each imputed continuous value must lie within that predictor's observed training minimum and maximum. This bound is applied only to originally missing cells, not to observed apply-set values.", '', 'Cycle-stability diagnostic', '--------------------------', 'For all 200 matched sensitivity splits, the outer imputation is repeated for 20 fixed cycles. Every predictor is reported. A variable-level threshold is applied only when at least 10 cells are imputed, avoiding a pass/fail decision based on only 1-4 cells. Sparse rows are labelled as not individually assessable and remain descriptive, but their cells still enter the global assessment. Among assessable variables, continuous normalized mean absolute difference must be no greater than 0.10 and binary disagreement no greater than 0.05. A count-weighted global standardized difference across all missing continuous cells must be no greater than 0.10, and global binary disagreement must be no greater than 0.05.', 'This stability policy was prespecified before inspecting any outcome-model performance or calibration result. It uses predictor values and originally missing cells only. Event status, follow-up time, model performance, and calibration do not enter the stability calculation or decision. The 0.10 and 0.05 cutoffs are operational quality thresholds for this analysis, not general convergence standards.', '', 'Reporting', '---------', 'Performance and paired differences are summarized by their mean and descriptive empirical 2.5th and 97.5th percentiles across all 200 matched splits. These percentiles are not confidence intervals. No P values are produced.', 'The ExtraTrees results are paired directly with all 200 matching Bayesian-imputation results from the main analysis.', 'Calibration uses one risk per participant, calculated by averaging only the predictions from splits in which that participant was held out. Every eligible participant must have at least one held-out prediction. Repeated prediction rows are not treated as independent observations.', 'Grouped and flexible calibration figures are landscape 2-by-4 panels. They have no overall title or footer and are saved as 600-dpi PNG and LZW-compressed TIFF files.', 'Participant-level prediction files contain individual records and should remain in the protected analysis environment.']
    out_path.write_text('\n'.join(lines) + '\n', encoding='utf-8')

def save_analysis_config(main_module: Any, out_dir: Path, args: argparse.Namespace, reference_provenance: Dict[str, Any], manifest_sha256: str) -> None:
    payload = {'analysis': 'matched_extra_trees_predictor_only_all_200_splits', 'main_outer_design': '200_stratified_random_80_20_splits', 'sensitivity_split_ids': list(range(1, SENSITIVITY_SPLITS + 1)), 'test_size': TEST_SIZE, 'seed': int(args.seed), 'inner_folds_classical': INNER_FOLDS, 'classical_absolute_risk_tuning': CLASSICAL_IBS_TUNING_METHOD, 'svm_tuning': SVM_TUNING_METHOD, 'coxph_tuning': FIXED_TUNING_METHOD, 'classical_tuning_imputation': CLASSICAL_IMPUTATION_POLICY, 'coxph_imputation_policy': FIXED_IMPUTATION_POLICY, 'neural_tuning': NEURAL_IBS_TUNING_METHOD, 'neural_tuning_imputation': NEURAL_IMPUTATION_POLICY, 'model_selection_policy': {model_name: main_module.model_selection_metadata(model_name) for model_name in main_module.MODEL_ORDER}, 'imputer': 'ExtraTreesRegressor within IterativeImputer', 'imputer_n_estimators': EXTRA_TREES_N_ESTIMATORS, 'imputer_min_samples_leaf': EXTRA_TREES_MIN_SAMPLES_LEAF, 'imputer_max_features': EXTRA_TREES_MAX_FEATURES, 'imputer_cycles': FIXED_IMPUTATION_CYCLES, 'imputer_tol': 0.0, 'fixed_cycle_policy': FIXED_CYCLE_POLICY, 'stopping_criterion_interpretation': 'not_applicable_prespecified_fixed_cycle_rule', 'expected_iteration_cap_warning': EXPECTED_ITERATION_WARNING_TEXT, 'horizon_months': float(args.horizon_months), 'ibs_start_month': float(args.ibs_start_month), 'selection_grid_months': list(range(int(args.ibs_start_month), int(args.horizon_months) + 1)), 'calibration_groups': int(args.calibration_groups), 'calibration_bootstrap_reps': int(args.calibration_bootstrap_reps), 'flexible_calibration_grid_points': FLEXIBLE_CALIBRATION_GRID_POINTS, 'n_jobs': int(args.n_jobs), 'classical_candidate_internal_threads': 1, 'seed_scheme': {'split_base_seed': 'seed + 100000 * split_id', 'model_tuning_seed': 'split_base_seed + 1; identical to main analysis', 'outer_extra_trees_seed': f'{SENSITIVITY_SEED_OFFSET} + split_base_seed', 'classical_inner_extra_trees_seed': f'{SENSITIVITY_SEED_OFFSET} + split_base_seed + 100 + inner_fold', 'neural_extra_trees_seed': f'{SENSITIVITY_SEED_OFFSET} + split_base_seed + 500'}, 'imputer_diagnostic_roles_per_split': {'outer_evaluation': 1, 'classical_inner': 3, 'neural_development': 1}, 'stability': {'split_ids': list(STABILITY_SPLIT_IDS), 'acceptance_scope': 'all_200_outer_evaluation_imputers_dense_applicable_and_aggregate_thresholds_must_pass', 'cycles': [FIXED_IMPUTATION_CYCLES, STABILITY_COMPARISON_CYCLES], 'minimum_imputed_cells_for_individual_variable_assessment': MIN_STABILITY_CELLS_PER_VARIABLE, 'continuous_normalized_mean_absolute_difference_maximum': 0.1, 'binary_disagreement_rate_maximum': 0.05, 'count_weighted_global_continuous_normalized_difference_maximum': 0.1, 'global_binary_disagreement_rate_maximum': 0.05, 'sparse_variable_policy': 'reported_not_individually_assessable_included_in_global_assessment', 'locking_policy': 'prespecified_before_inspecting_any_outcome_model_performance_or_calibration; predictor_values_and_missing_cells_only; excludes_event_status_follow_up_performance_and_calibration', 'minimum_cell_rationale': 'avoids_pass_fail_calls_based_on_only_1_to_4_imputed_cells'}, 'summary_policy': 'mean_and_descriptive_empirical_2_5th_97_5th_percentiles_no_CI_or_P_value', 'calibration_policy': 'one_mean_held_out_prediction_per_participant_all_participants_required', 'figure_policy': 'landscape_2_by_4_no_overall_title_no_footer_600dpi_PNG_and_LZW_TIFF', 'manifest_sha256': manifest_sha256, 'main_reference_provenance': reference_provenance, 'model_order': list(main_module.MODEL_ORDER), 'calibration_models': list(main_module.CALIBRATION_MODELS), 'imputation_matrix': list(main_module.PREDICTOR_COLUMNS), 'outcome_fields_excluded_from_imputation': [main_module.EVENT_COL, main_module.TIME_COL, 'Nelson-Aalen term', 'competing-event status'], 'participant_level_output_policy': 'predictions retained only in protected checkpoints', 'plausibility_policy': 'finite_complete_binary_valid_observed_cells_preserved; originally_missing_continuous_values_within_predictor_specific_imputer_training_min_max; observed_apply_values_not_bounded', 'warning_policy': 'exactly_one_expected_fixed_cycle_notice_required_per_imputer; warning_classification_marks_notice_not_a_convergence_assessment; unexpected_warning_count_must_equal_zero'}
    main_module.atomic_write_json(payload, str(out_dir / 'metadata' / 'sensitivity_extra_trees_analysis_config.json'))

def _read_required_csv(path: Path) -> pd.DataFrame:
    if not path.is_file():
        raise RuntimeError(f'Required sensitivity output is missing: {path}')
    return pd.read_csv(path, float_precision='round_trip')

def validate_final_outputs(main_module: Any, out_dir: Path, expected_participants: int) -> None:
    table_dir = out_dir / 'tables'
    metadata_dir = out_dir / 'metadata'
    figure_dir = out_dir / 'figures'
    paths = {'performance': table_dir / 'sensitivity_extra_trees_200split_performance_by_split.csv', 'summary': table_dir / 'sensitivity_extra_trees_200split_performance_summary.csv', 'cox_detail': table_dir / 'sensitivity_extra_trees_paired_differences_vs_coxph_by_split.csv', 'cox_summary': table_dir / 'sensitivity_extra_trees_paired_differences_vs_coxph_summary.csv', 'main_detail': table_dir / 'sensitivity_extra_trees_comparison_vs_main_bayesian_by_split.csv', 'main_summary': table_dir / 'sensitivity_extra_trees_comparison_vs_main_bayesian_summary.csv', 'calibration': table_dir / 'sensitivity_extra_trees_participant_mean_calibration.csv', 'grouped': table_dir / 'sensitivity_extra_trees_participant_mean_grouped_calibration.csv', 'flexible': table_dir / 'sensitivity_extra_trees_participant_mean_flexible_calibration.csv', 'tuning': metadata_dir / 'sensitivity_extra_trees_tuning_records.csv', 'diagnostics': metadata_dir / 'sensitivity_extra_trees_imputation_diagnostics_by_role.csv', 'stability': metadata_dir / 'sensitivity_extra_trees_cycle_stability_10_vs_20.csv', 'coverage': metadata_dir / 'sensitivity_extra_trees_heldout_coverage_summary.csv'}
    loaded = {name: _read_required_csv(path) for name, path in paths.items()}
    expected_model_rows = SENSITIVITY_SPLITS * len(main_module.MODEL_ORDER)
    if len(loaded['performance']) != expected_model_rows:
        raise RuntimeError('Final performance output has an incomplete split-model grid.')
    if len(loaded['tuning']) != expected_model_rows:
        raise RuntimeError('Final tuning output has an incomplete split-model grid.')
    expected_keys = {(split_id, model) for split_id in range(1, SENSITIVITY_SPLITS + 1) for model in main_module.MODEL_ORDER}
    for name in ['performance', 'tuning']:
        frame = loaded[name]
        keys = set(zip(pd.to_numeric(frame['split_id'], errors='raise').astype(int), frame['model'].astype(str)))
        if keys != expected_keys or frame.duplicated(['split_id', 'model']).any():
            raise RuntimeError(f'Final {name} output keys are incomplete or duplicated.')
        if not pd.to_numeric(frame['split_seed'], errors='raise').astype(int).eq(USER_MASTER_SEED).all():
            raise RuntimeError(f'Final {name} output has the wrong split seed.')
    for model_name, rows in loaded['performance'].groupby('model'):
        expected_test_n = int(math.ceil(TEST_SIZE * expected_participants))
        if not (pd.to_numeric(rows['n_train'], errors='raise').astype(int).eq(expected_participants - expected_test_n).all() and pd.to_numeric(rows['n_test'], errors='raise').astype(int).eq(expected_test_n).all() and pd.to_numeric(rows['n_test_events'], errors='raise').astype(int).eq(EXPECTED_TEST_EVENTS).all()):
            raise RuntimeError(f'Final performance has wrong sample counts for {model_name}.')
        if not all((_bool_value(value, 'split_metrics_complete') for value in rows['split_metrics_complete'])):
            raise RuntimeError(f'Final performance marks {model_name} incomplete.')
        for value in rows['split_metric_errors_json'].astype(str):
            parsed = json.loads(value)
            if not isinstance(parsed, dict) or parsed:
                raise RuntimeError(f'Final performance contains a required-metric error for {model_name}.')
        tau = pd.to_numeric(rows['tau_months'], errors='raise').to_numpy(float)
        if not np.isfinite(tau).all() or np.any(tau <= 0.0) or np.any(tau > USER_HORIZON_MONTHS + 1e-12):
            raise RuntimeError(f'Final performance has an invalid tau for {model_name}.')
        brier_horizon = pd.to_numeric(rows['brier_horizon_months'], errors='coerce').to_numpy(float)
        if str(model_name) == 'SVM':
            if not np.isnan(brier_horizon).all():
                raise RuntimeError('Final SVM performance unexpectedly has a Brier horizon.')
        elif not (np.isfinite(brier_horizon).all() and np.allclose(brier_horizon, USER_HORIZON_MONTHS, rtol=0.0, atol=1e-12)):
            raise RuntimeError(f'Final performance has an invalid Brier horizon for {model_name}.')
        for metric in required_split_metrics(str(model_name)):
            for value in pd.to_numeric(rows[metric], errors='raise').to_numpy(float):
                validate_performance_metric(metric, float(value), f'Final sensitivity performance {model_name}')
    for _, row in loaded['tuning'].iterrows():
        split_id = int(row['split_id'])
        model_name = str(row['model'])
        if int(row['model_seed']) != matched_model_seed(USER_MASTER_SEED, split_id):
            raise RuntimeError('Final tuning output has a mismatched model seed.')
        if str(row['tuning_method']) != expected_tuning_method(model_name):
            raise RuntimeError(f'Final tuning method is wrong for {model_name}.')
        if str(row['tuning_imputation_policy']) != expected_tuning_imputation_policy(model_name):
            raise RuntimeError(f'Final imputation policy is wrong for {model_name}.')
        main_module.validate_tuning_selection_record(row, model_name, f'Final ExtraTrees tuning split {split_id} {model_name}')
        inner_value = pd.to_numeric(pd.Series([row['inner_folds_classical']]), errors='coerce').iloc[0]
        if model_name in {'CoxPH', 'DeepSurv', 'CoxTime'}:
            if not pd.isna(inner_value):
                raise RuntimeError(f'Final tuning output gives {model_name} a classical fold count.')
        elif not np.isfinite(inner_value) or int(inner_value) != INNER_FOLDS:
            raise RuntimeError(f'Final tuning output has the wrong inner-fold count for {model_name}.')
    if len(loaded['summary']) != len(main_module.MODEL_ORDER):
        raise RuntimeError('Performance summary does not include all eight models.')
    expected_cox_comparisons = sum((len(required_split_metrics(model_name)) for model_name in main_module.MODEL_ORDER if model_name != 'CoxPH'))
    expected_main_comparisons = sum((len(required_split_metrics(model_name)) for model_name in main_module.MODEL_ORDER))
    if not (len(loaded['cox_detail']) == SENSITIVITY_SPLITS * expected_cox_comparisons and len(loaded['cox_summary']) == expected_cox_comparisons):
        raise RuntimeError('CoxPH paired-comparison outputs have incomplete grids.')
    if not (len(loaded['main_detail']) == SENSITIVITY_SPLITS * expected_main_comparisons and len(loaded['main_summary']) == expected_main_comparisons):
        raise RuntimeError('Bayesian-imputer paired-comparison outputs have incomplete grids.')
    for name in ['cox_detail', 'main_detail']:
        if loaded[name].duplicated(['split_id', 'model', 'metric']).any():
            raise RuntimeError(f'{name} contains duplicate split-model-metric rows.')
    for name in ['cox_summary', 'main_summary']:
        if loaded[name].duplicated(['model', 'metric']).any():
            raise RuntimeError(f'{name} contains duplicate model-metric rows.')
    percentile_columns = [column for column in loaded['summary'].columns if column.endswith('empirical_2_5_percentile') or column.endswith('empirical_97_5_percentile')]
    if not percentile_columns:
        raise RuntimeError('Performance summary lacks empirical percentile ranges.')
    if len(loaded['diagnostics']) != SENSITIVITY_SPLITS * 5:
        raise RuntimeError('Imputation diagnostics do not contain five roles per split.')
    for split_id, rows in loaded['diagnostics'].groupby('split_id', sort=True):
        validate_imputation_diagnostics(rows.reset_index(drop=True), int(split_id))
    expected_stability_rows = len(STABILITY_SPLIT_IDS) * (len(main_module.PREDICTOR_COLUMNS) + 2)
    if len(loaded['stability']) != expected_stability_rows:
        raise RuntimeError('Cycle-stability output has an incomplete split-variable grid.')
    for split_id in STABILITY_SPLIT_IDS:
        rows = loaded['stability'].loc[pd.to_numeric(loaded['stability']['split_id'], errors='raise').astype(int).eq(split_id)].reset_index(drop=True)
        validate_stability_frame(rows, split_id, main_module.PREDICTOR_COLUMNS, main_module)
    if len(loaded['coverage']) != len(main_module.CALIBRATION_MODELS):
        raise RuntimeError('Coverage summary does not include all risk models.')
    if not pd.to_numeric(loaded['coverage']['unique_participants'], errors='raise').astype(int).eq(expected_participants).all():
        raise RuntimeError('Coverage summary does not include every participant.')
    if not all((_bool_value(value, 'coverage_complete') for value in loaded['coverage']['coverage_complete'])):
        raise RuntimeError('Coverage summary reports incomplete participant coverage.')
    if not (pd.to_numeric(loaded['coverage']['held_out_prediction_rows'], errors='raise').astype(int).eq(EXPECTED_HELDOUT_PREDICTION_ROWS).all() and pd.to_numeric(loaded['coverage']['minimum_predictions_per_participant'], errors='raise').astype(int).eq(EXPECTED_MIN_HELDOUT_PREDICTIONS).all() and pd.to_numeric(loaded['coverage']['maximum_predictions_per_participant'], errors='raise').astype(int).eq(EXPECTED_MAX_HELDOUT_PREDICTIONS).all()):
        raise RuntimeError('Coverage summary does not match the validated 49,800-row, 21-to-57-predictions-per-participant contract.')
    if len(loaded['calibration']) != len(main_module.CALIBRATION_MODELS):
        raise RuntimeError('Calibration summary does not include all risk models.')
    if len(loaded['grouped']) != len(main_module.CALIBRATION_MODELS) * USER_CALIBRATION_GROUPS:
        raise RuntimeError('Grouped calibration output has an incomplete grid.')
    expected_flexible_models = set(main_module.CALIBRATION_MODELS)
    if set(loaded['flexible']['model'].astype(str)) != expected_flexible_models:
        raise RuntimeError('Flexible calibration output lacks a risk model.')
    if len(loaded['flexible']) != len(main_module.CALIBRATION_MODELS) * FLEXIBLE_CALIBRATION_GRID_POINTS:
        raise RuntimeError('Flexible calibration output has an incomplete model-grid table.')
    forbidden_columns = {main_module.STUDY_ID_COL, 'study_id_key'}
    for name, frame in loaded.items():
        if forbidden_columns.intersection(frame.columns):
            raise RuntimeError(f'Participant identifiers leaked into final output {name}.')
    figure_paths = [figure_dir / 'sensitivity_extra_trees_grouped_calibration_mean_heldout.png', figure_dir / 'sensitivity_extra_trees_grouped_calibration_mean_heldout.tiff', figure_dir / 'sensitivity_extra_trees_flexible_calibration_mean_heldout.png', figure_dir / 'sensitivity_extra_trees_flexible_calibration_mean_heldout.tiff']
    for path in figure_paths:
        if not path.is_file() or path.stat().st_size == 0:
            raise RuntimeError(f'Required calibration figure is missing: {path}')
        with Image.open(path) as image:
            if image.width <= image.height:
                raise RuntimeError(f'Calibration figure is not landscape: {path}')
            dpi = image.info.get('dpi', (0, 0))
            if min(float(dpi[0]), float(dpi[1])) < 599.0:
                raise RuntimeError(f'Calibration figure is below 600 dpi: {path}')
            if path.suffix.lower() in {'.tif', '.tiff'}:
                compression = str(image.info.get('compression', '')).lower()
                if compression not in {'tiff_lzw', 'lzw', '5'}:
                    raise RuntimeError(f'Calibration TIFF is not LZW: {path}')

def print_startup_summary(args: argparse.Namespace) -> None:
    lines = ['=' * 88, 'Matched ExtraTrees predictor-only imputation sensitivity', '=' * 88, 'Raw CSV:                    supplied by command line', 'Main script:                internal analysis module', 'Main output directory:      supplied by command line', 'Sensitivity output:         supplied by command line', f'Master seed:                {args.seed}', 'Outer design:               exact split IDs 1-200 of main stratified 80/20 splits', f'Classical IBS tuning:       {INNER_FOLDS} stratified folds, monthly 12-120 months', 'Neural IBS tuning:          development-only internal stratified 80/20 validation', 'IBS censoring weights:      corresponding inner-training/development outcomes only', 'Neural early stopping:      validation loss (configuration selection uses IBS)', 'Survival SVM tuning:        mean inner-fold Harrell C', 'CoxPH:                      one fixed specification; no selection', 'ExtraTrees policy:          prespecified fixed 10 cycles, tol=0', 'Stability split IDs:        all matched sensitivity splits 1-200', 'Stability comparison:       10 versus 20 fixed cycles', f'Calibration bootstraps:     {args.calibration_bootstrap_reps}', f'n_jobs:                     {args.n_jobs}', 'Candidate model threads:    1', f'Max session hours:          {args.max_session_hours}', f'Resume:                     {args.resume}', '', 'All imputers use baseline predictors only.', 'Every split refits all eight models; seven have configuration selection and CoxPH is fixed.', 'Calibration uses one mean held-out risk per participant; all must be covered.', 'Participant-level predictions remain only in local checkpoints.', '=' * 88]
    summary = '\n'.join(lines)
    print(summary, flush=True)
    LOGGER.info('\n%s', summary)

def main() -> None:
    args = parse_args()
    if args.seed != USER_MASTER_SEED:
        raise ValueError(f'--seed is locked at the main splitter seed {USER_MASTER_SEED}.')
    if args.n_jobs == 0:
        raise ValueError('--n-jobs cannot be zero.')
    if args.impute_cycles != FIXED_IMPUTATION_CYCLES:
        raise ValueError(f'--impute-cycles is locked at {FIXED_IMPUTATION_CYCLES}; received {args.impute_cycles}.')
    if args.calibration_groups != USER_CALIBRATION_GROUPS:
        raise ValueError(f'--calibration-groups is locked at {USER_CALIBRATION_GROUPS}.')
    if args.calibration_bootstrap_reps < 100:
        raise ValueError('--calibration-bootstrap-reps must be at least 100.')
    if not np.isclose(args.horizon_months, 120.0, rtol=0.0, atol=1e-12):
        raise ValueError('--horizon-months is locked at 120 for the analysis.')
    if not np.isclose(args.ibs_start_month, 12.0, rtol=0.0, atol=1e-12):
        raise ValueError('--ibs-start-month is locked at 12 for the analysis.')
    if not np.isfinite(args.max_session_hours) or args.max_session_hours < 0.0:
        raise ValueError('--max-session-hours must be finite and non-negative.')
    if args.worker_import_smoke_test:
        main_script = Path(args.main_script)
        main_module = load_main_module(main_script)
        run_worker_import_smoke_test(main_module, main_script, csv_path=Path(args.csv), seed=args.seed)
        return
    out_dir = Path(args.outdir)
    table_dir = out_dir / 'tables'
    metadata_dir = out_dir / 'metadata'
    figure_dir = out_dir / 'figures'
    for directory in [table_dir, metadata_dir, figure_dir]:
        directory.mkdir(parents=True, exist_ok=True)
    setup_logging(metadata_dir / 'run_extra_trees_matched_200_splits.log')
    print_startup_summary(args)
    main_script = Path(args.main_script)
    csv_path = Path(args.csv)
    main_output_dir = Path(args.main_output_dir)
    LOGGER.info('Loading the main analysis module.')
    main_module = load_main_module(main_script)
    main_module.set_global_seed(args.seed)
    validate_dependencies(main_module)
    LOGGER.info('Loading and validating the analysis data.')
    df = main_module.load_dataset(str(csv_path))
    if study_id_keys(df[main_module.STUDY_ID_COL]).duplicated().any():
        raise RuntimeError('The eligible dataset contains duplicate study IDs.')
    definitions, matched_map, main_performance, reference_provenance = validate_main_reference(main_module, df, main_output_dir, args.seed)
    checkpoint_dir, manifest_sha256 = prepare_checkpoint_manifest(main_module, df, out_dir, args, reference_provenance)
    try:
        complete, metric_frames, tuning_frames, prediction_frames, diagnostic_frames, stability_frames = run_matched_200_split_sensitivity(main_module, df, definitions, checkpoint_dir, manifest_sha256, args)
    except PlannedSensitivityPause as exc:
        write_planned_pause_payload(main_module, metadata_dir, manifest_sha256, exc.completed_splits, 'max_session_hours_reached_after_checkpointed_model', next_split=exc.split_id, partial_split=exc.split_id, last_checkpointed_model=exc.model_name, partial_split_completed_models=exc.completed_models)
        LOGGER.info('Sensitivity paused safely after split %d model %s; %d complete splits retained.', exc.split_id, exc.model_name, len(exc.completed_splits))
        raise SystemExit(75)
    if not complete:
        completed_splits = sorted({int(frame['split_id'].iloc[0]) for frame in metric_frames if not frame.empty})
        next_split = next((split_id for split_id in range(1, SENSITIVITY_SPLITS + 1) if split_id not in completed_splits), SENSITIVITY_SPLITS)
        write_planned_pause_payload(main_module, metadata_dir, manifest_sha256, completed_splits, 'max_session_hours_reached_after_completed_split', next_split=next_split)
        raise SystemExit(75)
    checkpoint_metrics = pd.concat(metric_frames, ignore_index=True)
    tuning = pd.concat(tuning_frames, ignore_index=True)
    diagnostics = pd.concat(diagnostic_frames, ignore_index=True)
    stability = pd.concat(stability_frames, ignore_index=True)
    _, mean_predictions, coverage = verify_prediction_coverage_and_average(main_module, df, prediction_frames, matched_map)
    split_metrics = final_split_metrics(main_module, checkpoint_metrics)
    performance_summary = summarise_split_metrics(main_module, split_metrics)
    cox_detail, cox_summary = paired_differences_vs_coxph(main_module, split_metrics)
    main_detail, main_summary = compare_with_main_bayesian(main_module, main_performance, split_metrics)
    calibration, grouped_calibration, flexible_calibration = participant_mean_calibration(main_module, df, mean_predictions, args.horizon_months, args.calibration_groups, args.calibration_bootstrap_reps, args.seed)
    make_grouped_calibration_figure(main_module, grouped_calibration, figure_dir, args.horizon_months)
    make_flexible_calibration_figure(main_module, flexible_calibration, figure_dir, args.horizon_months)
    outputs = {table_dir / 'sensitivity_extra_trees_200split_performance_by_split.csv': split_metrics, table_dir / 'sensitivity_extra_trees_200split_performance_summary.csv': performance_summary, table_dir / 'sensitivity_extra_trees_paired_differences_vs_coxph_by_split.csv': cox_detail, table_dir / 'sensitivity_extra_trees_paired_differences_vs_coxph_summary.csv': cox_summary, table_dir / 'sensitivity_extra_trees_comparison_vs_main_bayesian_by_split.csv': main_detail, table_dir / 'sensitivity_extra_trees_comparison_vs_main_bayesian_summary.csv': main_summary, table_dir / 'sensitivity_extra_trees_participant_mean_calibration.csv': calibration, table_dir / 'sensitivity_extra_trees_participant_mean_grouped_calibration.csv': grouped_calibration, table_dir / 'sensitivity_extra_trees_participant_mean_flexible_calibration.csv': flexible_calibration, metadata_dir / 'sensitivity_extra_trees_tuning_records.csv': tuning, metadata_dir / 'sensitivity_extra_trees_imputation_diagnostics_by_role.csv': diagnostics, metadata_dir / 'sensitivity_extra_trees_cycle_stability_10_vs_20.csv': stability, metadata_dir / 'sensitivity_extra_trees_heldout_coverage_summary.csv': coverage}
    for path, frame in outputs.items():
        main_module.atomic_write_dataframe_csv(frame, str(path))
    write_interpretation_report(out_dir / 'sensitivity_extra_trees_interpretation_report.txt')
    save_analysis_config(main_module, out_dir, args, reference_provenance, manifest_sha256)
    validate_final_outputs(main_module, out_dir, expected_participants=len(df))
    held_out_counts = matched_map.groupby('study_id_key').size()
    completion_payload = {'status': 'complete', 'analysis': 'matched_extra_trees_predictor_only_all_200_splits', 'imputer': 'extra_trees', 'fixed_imputation_cycles': FIXED_IMPUTATION_CYCLES, 'imputer_tol': 0.0, 'stopping_criterion_claimed': False, 'seed': int(args.seed), 'matched_split_ids': list(range(1, SENSITIVITY_SPLITS + 1)), 'outer_splits_completed': SENSITIVITY_SPLITS, 'models_refitted_per_split': int(len(main_module.MODEL_ORDER)), 'models_with_hyperparameter_selection_per_split': 7, 'fixed_specification_models': ['CoxPH'], 'model_selection_policy': {model_name: main_module.model_selection_metadata(model_name) for model_name in main_module.MODEL_ORDER}, 'model_level_safe_pause_enabled': True, 'split_level_safe_pause_enabled': True, 'imputer_diagnostics_expected': SENSITIVITY_SPLITS * 5, 'stability_split_ids': list(STABILITY_SPLIT_IDS), 'cycle_stability_scope': 'all_200_outer_evaluation_imputers_only_inner_and_neural_roles_receive_strict_completion_diagnostics', 'aggregate_cycle_stability_all_200_splits_passed': True, 'dense_applicable_variable_cycle_thresholds_passed': True, 'stability_sparse_variables_labelled_not_individually_assessable': True, 'eligible_participants_with_mean_heldout_prediction': int(len(df)), 'minimum_heldout_predictions_per_participant': int(held_out_counts.min()), 'maximum_heldout_predictions_per_participant': int(held_out_counts.max()), 'participant_level_outputs_in_final_tables': False, 'participant_level_predictions_retained_in_protected_checkpoints': True, 'calibration_figure_layout': 'landscape_2_by_4_no_overall_title_or_footer', 'calibration_figure_formats': 'PNG_600dpi_and_LZW_TIFF_600dpi', 'calibration_bootstrap_replicates_requested': int(args.calibration_bootstrap_reps), 'manifest_sha256': manifest_sha256, 'instruction': 'Use the aggregate tables, figures, and metadata. Checkpoint files are not required to interpret the reported results.'}
    main_module.atomic_write_json(completion_payload, str(metadata_dir / 'analysis_complete.json'))
    LOGGER.info('Matched ExtraTrees sensitivity complete: 200 splits and full participant coverage.')
if __name__ == '__main__':
    main()
''',
    '_analysis_secondary_core.py': r'''from __future__ import annotations
import argparse
import hashlib
import importlib.metadata
import json
import math
import os
import platform
import re
import tempfile
import time
import warnings
from dataclasses import dataclass
from pathlib import Path
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from scipy.optimize import minimize
from sklearn.ensemble import ExtraTreesRegressor
from sklearn.experimental import enable_iterative_imputer
from sklearn.impute import IterativeImputer
from sklearn.linear_model import BayesianRidge, LinearRegression
from sklearn.model_selection import StratifiedShuffleSplit
from sklearn.exceptions import ConvergenceWarning
ID = 'studyid'
TIME = 'time_to_event'
EVENT = 'event_status'
SEX = 'sex'
MASTER_ID = 'Study ID'
MASTER_DEATH_FLAG = 'Death other cause'
MASTER_DEATH_TIME = 'F/U Time'
MASTER_DM_FIELD = 'Hx of DM 0=no'
MASTER_SHEET = os.environ.get('ANALYSIS_MASTER_SHEET', '')
CSV_TIME = 'csv_time_to_event'
COMPETING_EVENT = 'competing_death'
EXPECTED_SOURCE_ROWS = 1247
EXPECTED_ELIGIBLE_ROWS = 1241
EXPECTED_ASCVD_EVENTS = 116
EXPECTED_COMPETING_DEATHS = 46
EXPECTED_OTHER_CENSORED = 1079
EXPECTED_DEATH_FLAGS = 56
EXPECTED_POTENTIAL_COMPETING_DEATHS = 47
EXPECTED_NONPOSITIVE_COMPETING_EXCLUSIONS = 1
EXPECTED_ASCVD_FIRST_LATER_DEATHS = 9
COMPETING_DEATH_RIDGE = 1.0
OOF_CALIBRATION_GROUPS = 5
EXTRA_TREES_SENSITIVITY_SEED_OFFSET = 50000000
MAIN_SPLIT_MAP_RELATIVE = Path('metadata/repeated_200_split_map.csv')
MAIN_PERFORMANCE_RELATIVE = Path('tables/repeated_200_performance_by_split.csv')
MAIN_TUNING_RELATIVE = Path('tables/repeated_200_tuning_by_split.csv')
MAIN_MISSINGNESS_RELATIVE = Path('tables/supplementary_table_s1_missingness.csv')
MAIN_POOLED_COX_RELATIVE = Path('tables/table2_cox_hazard_ratios_fixed_split.csv')
MAIN_PH_RELATIVE = Path('tables/proportional_hazards_tests_by_imputation.csv')
EXPECTED_PRIMARY_SPLITS = 200
EXPECTED_TREE_SPLITS = 200
EXPECTED_TEST_ROWS_PER_SPLIT = 249
EXPECTED_MAIN_SPLIT_MAP_ROWS = EXPECTED_PRIMARY_SPLITS * EXPECTED_TEST_ROWS_PER_SPLIT
EXPECTED_MAIN_MODELS = {'CoxPH', 'ElasticNetCox', 'RSF', 'GBSA', 'DeepSurv', 'CoxTime', 'XGBoost', 'SVM'}
IBS_SELECTED_MAIN_MODELS = {'ElasticNetCox', 'RSF', 'GBSA', 'DeepSurv', 'CoxTime', 'XGBoost'}
MAIN_CANDIDATE_COUNTS = {'CoxPH': 1, 'ElasticNetCox': 8, 'RSF': 6, 'GBSA': 6, 'DeepSurv': 6, 'CoxTime': 6, 'XGBoost': 8, 'SVM': 4}
MAIN_IBS_SELECTION_METRIC = 'integrated_brier_score_12_120_months'
EXTRA_TREES_FIXED_CYCLES = 10

def expected_main_tuning_method(model_name: str) -> str:
    if model_name == 'CoxPH':
        return 'fixed_specification_no_hyperparameter_selection'
    if model_name in {'DeepSurv', 'CoxTime'}:
        return 'nested_imputation_internal_stratified_80_20_validation_ibs_12_120'
    if model_name in IBS_SELECTED_MAIN_MODELS:
        return 'nested_imputation_stratified_3_fold_mean_ibs_12_120'
    if model_name == 'SVM':
        return 'nested_imputation_stratified_3_fold_mean_harrell_c'
    raise RuntimeError(f'Unknown main-model tuning method for {model_name!r}.')

def expected_main_selection_metadata(model_name: str) -> dict[str, object]:
    if model_name in IBS_SELECTED_MAIN_MODELS:
        return {'selection_metric': MAIN_IBS_SELECTION_METRIC, 'selection_direction': 'minimize', 'model_selection_performed': True, 'validation_design': 'internal_stratified_80_20_validation' if model_name in {'DeepSurv', 'CoxTime'} else 'stratified_3_fold_inner_cross_validation', 'selection_window_start_months': 12.0, 'selection_window_end_months': 120.0, 'selection_grid_step_months': 1.0, 'selection_grid_points': 109, 'n_prespecified_configurations': MAIN_CANDIDATE_COUNTS[model_name]}
    if model_name == 'SVM':
        return {'selection_metric': 'harrell_c', 'selection_direction': 'maximize', 'model_selection_performed': True, 'validation_design': 'stratified_3_fold_inner_cross_validation', 'selection_window_start_months': None, 'selection_window_end_months': None, 'selection_grid_step_months': None, 'selection_grid_points': None, 'n_prespecified_configurations': MAIN_CANDIDATE_COUNTS[model_name]}
    if model_name == 'CoxPH':
        return {'selection_metric': 'not_applicable_fixed_specification', 'selection_direction': 'not_applicable', 'model_selection_performed': False, 'validation_design': 'not_applicable_fixed_specification', 'selection_window_start_months': None, 'selection_window_end_months': None, 'selection_grid_step_months': None, 'selection_grid_points': None, 'n_prespecified_configurations': MAIN_CANDIDATE_COUNTS[model_name]}
    raise RuntimeError(f'Unknown main-model selection metadata for {model_name!r}.')

def _saved_boolean(value: object, label: str) -> bool:
    if isinstance(value, (bool, np.bool_)):
        return bool(value)
    normalized = str(value).strip().lower()
    if normalized not in {'true', 'false', '1', '0'}:
        raise RuntimeError(f'{label} is not a saved boolean.')
    return normalized in {'true', '1'}

def validate_main_tuning_selection_record(row: pd.Series, model_name: str) -> None:
    expected = expected_main_selection_metadata(model_name)
    required = {'tuning_method', 'best_params_json', 'selection_score', *expected.keys()}
    missing = sorted(required.difference(row.index))
    if missing:
        raise RuntimeError(f'Main {model_name} tuning lacks fields: {missing}.')
    if str(row['tuning_method']) != expected_main_tuning_method(model_name):
        raise RuntimeError(f'Main {model_name} tuning method is wrong.')
    for field in ['selection_metric', 'selection_direction', 'validation_design']:
        if str(row[field]) != str(expected[field]):
            raise RuntimeError(f'Main {model_name} tuning has the wrong {field}.')
    selected = _saved_boolean(row['model_selection_performed'], f'Main {model_name} model_selection_performed')
    if selected != bool(expected['model_selection_performed']):
        raise RuntimeError(f'Main {model_name} tuning has the wrong selection flag.')
    candidate_count = pd.to_numeric(pd.Series([row['n_prespecified_configurations']]), errors='raise').iloc[0]
    if not np.isfinite(candidate_count) or float(candidate_count) != float(expected['n_prespecified_configurations']):
        raise RuntimeError(f'Main {model_name} tuning has the wrong candidate count.')
    score = pd.to_numeric(pd.Series([row['selection_score']]), errors='coerce').iloc[0]
    if selected:
        if not np.isfinite(score) or not 0.0 <= float(score) <= 1.0:
            raise RuntimeError(f'Main {model_name} selection score is invalid.')
    elif not pd.isna(score):
        raise RuntimeError(f'Fixed main {model_name} unexpectedly reports a score.')
    window_fields = ['selection_window_start_months', 'selection_window_end_months', 'selection_grid_step_months', 'selection_grid_points']
    if expected['selection_metric'] == MAIN_IBS_SELECTION_METRIC:
        for field in window_fields:
            observed = pd.to_numeric(pd.Series([row[field]]), errors='coerce').iloc[0]
            if not np.isfinite(observed) or not math.isclose(float(observed), float(expected[field]), rel_tol=0.0, abs_tol=1e-12):
                raise RuntimeError(f'Main {model_name} tuning has the wrong {field}.')
    elif not all((pd.isna(row[field]) for field in window_fields)):
        raise RuntimeError(f'Main {model_name} unexpectedly records an IBS window.')
    try:
        parameters = json.loads(str(row['best_params_json']))
    except json.JSONDecodeError as exc:
        raise RuntimeError(f'Main {model_name} parameters are not valid JSON.') from exc
    if not isinstance(parameters, dict):
        raise RuntimeError(f'Main {model_name} parameters are not a JSON object.')

class PlannedSessionPause(RuntimeError):
    pass
PREDICTORS = ['age', 'sex', 'smoke', 'fhxcvd', 'dmmed', 'htnmed', 'dldmed', 'sbp', 'dbp', 'bmi', 'chol', 'hdl', 'trig', 'hba1c', 'egfr']
CONTINUOUS = ['age', 'sbp', 'dbp', 'bmi', 'chol', 'hdl', 'trig', 'hba1c', 'egfr']
BINARY = ['sex', 'smoke', 'fhxcvd', 'dmmed', 'htnmed', 'dldmed']
RCS_VARS = ['age', 'sbp', 'bmi', 'hba1c', 'egfr']
DISPLAY = {'age': 'Age', 'sex': 'Sex', 'smoke': 'Smoking history', 'fhxcvd': 'Family history of CVD', 'dmmed': 'Glucose-lowering medication', 'htnmed': 'Blood pressure-lowering medication', 'dldmed': 'Lipid-lowering medication', 'sbp': 'SBP', 'dbp': 'DBP', 'bmi': 'BMI', 'chol': 'TC', 'hdl': 'HDL-C', 'trig': 'TG', 'hba1c': 'HbA1c', 'egfr': 'eGFR'}
COX_METRIC_COLUMNS = ['harrell_c', 'uno_c', 'brier_120', 'ibs_12_120', 'oe_ratio', 'observed_km_120', 'mean_predicted_120', 'calibration_slope']

def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open('rb') as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b''):
            digest.update(chunk)
    return digest.hexdigest()

def canonical_json(value: object) -> str:
    return json.dumps(value, sort_keys=True, separators=(',', ':'), default=str)

def lpa_candidate_columns(columns: list[object]) -> list[str]:
    candidates: list[str] = []
    for column in columns:
        label = str(column)
        normalized = re.sub('[^a-z0-9]+', '', label.lower())
        if normalized == 'lpa' or 'lipoproteina' in normalized:
            candidates.append(label)
    return sorted(set(candidates))

def secondary_package_versions() -> dict[str, str | None]:
    versions: dict[str, str | None] = {}
    for distribution in ['numpy', 'pandas', 'scipy', 'matplotlib', 'scikit-learn', 'joblib', 'scikit-survival', 'lifelines', 'openpyxl']:
        try:
            versions[distribution] = importlib.metadata.version(distribution)
        except importlib.metadata.PackageNotFoundError:
            versions[distribution] = None
    return versions

def validate_secondary_dependencies() -> None:
    try:
        from lifelines import CoxPHFitter
    except Exception as exc:
        raise RuntimeError('Secondary OOF calibration requires lifelines. Install the exact bundled requirements before starting the analysis.') from exc
    if secondary_package_versions().get('lifelines') is None:
        raise RuntimeError('The installed lifelines distribution could not be identified for the run manifest.')
    try:
        from sksurv.metrics import concordance_index_censored
        from sksurv.nonparametric import CensoringDistributionEstimator
    except Exception as exc:
        raise RuntimeError('Secondary performance metrics require scikit-survival from the exact bundled environment.') from exc

def atomic_write_json(value: object, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fd, temporary_name = tempfile.mkstemp(prefix=f'.{path.name}.', suffix='.tmp', dir=path.parent)
    try:
        with os.fdopen(fd, 'w', encoding='utf-8') as handle:
            json.dump(value, handle, indent=2, sort_keys=True, default=str)
            handle.write('\n')
            handle.flush()
            os.fsync(handle.fileno())
        os.replace(temporary_name, path)
    except Exception:
        try:
            os.unlink(temporary_name)
        except FileNotFoundError:
            pass
        raise

def atomic_write_csv(frame: pd.DataFrame, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fd, temporary_name = tempfile.mkstemp(prefix=f'.{path.name}.', suffix='.tmp', dir=path.parent)
    try:
        with os.fdopen(fd, 'w', encoding='utf-8', newline='') as handle:
            frame.to_csv(handle, index=False, float_format='%.17g')
            handle.flush()
            os.fsync(handle.fileno())
        os.replace(temporary_name, path)
    except Exception:
        try:
            os.unlink(temporary_name)
        except FileNotFoundError:
            pass
        raise

def prepare_competing_data(raw_all: pd.DataFrame, master: pd.DataFrame) -> pd.DataFrame:
    raw_required = [ID, TIME, EVENT, *PREDICTORS, 'ldl']
    master_required = [MASTER_ID, MASTER_DEATH_FLAG, MASTER_DEATH_TIME, MASTER_DM_FIELD, 'Fatal MI - HARD CHD', 'Fatal CVA - HARD Stroke', 'Seen in clinic after Jan 2019(1=Yes)', 'CRP']
    missing_raw = [column for column in raw_required if column not in raw_all.columns]
    missing_master = [column for column in master_required if column not in master.columns]
    if missing_raw:
        raise ValueError(f'CSV is missing required columns: {missing_raw}')
    if missing_master:
        raise ValueError(f'Master worksheet {MASTER_SHEET!r} is missing required columns: {missing_master}')
    raw = raw_all.copy()
    master_subset = master.loc[:, master_required].copy()
    if raw[ID].isna().any() or master_subset[MASTER_ID].isna().any():
        raise ValueError('Study IDs must not be missing in either input file.')
    if raw[ID].duplicated().any():
        raise ValueError('The CSV contains duplicated study IDs.')
    if master_subset[MASTER_ID].duplicated().any():
        raise ValueError(f'The master worksheet {MASTER_SHEET!r} contains duplicated study IDs.')
    raw[ID] = pd.to_numeric(raw[ID], errors='raise')
    master_subset[MASTER_ID] = pd.to_numeric(master_subset[MASTER_ID], errors='raise')
    raw_ids = set(raw[ID].tolist())
    master_ids = set(master_subset[MASTER_ID].tolist())
    missing_from_master = sorted(raw_ids.difference(master_ids))
    extra_in_master = sorted(master_ids.difference(raw_ids))
    if missing_from_master or extra_in_master:
        raise ValueError(f'CSV and master worksheet Study IDs do not have exact one-to-one coverage. Missing from master (first 10): {missing_from_master[:10]}; extra in master (first 10): {extra_in_master[:10]}.')
    raw[TIME] = pd.to_numeric(raw[TIME], errors='raise')
    raw[EVENT] = pd.to_numeric(raw[EVENT], errors='raise')
    if raw[EVENT].isna().any() or not raw[EVENT].isin([0, 1]).all():
        raise ValueError(f'{EVENT!r} must contain only non-missing 0 and 1 values.')
    raw[EVENT] = raw[EVENT].astype(int)
    for column in PREDICTORS:
        raw[column] = pd.to_numeric(raw[column], errors='raise')
    raw['ldl'] = pd.to_numeric(raw['ldl'], errors='coerce')
    master_subset[MASTER_DEATH_FLAG] = pd.to_numeric(master_subset[MASTER_DEATH_FLAG], errors='raise')
    if master_subset[MASTER_DEATH_FLAG].isna().any() or not master_subset[MASTER_DEATH_FLAG].isin([0, 1]).all():
        raise ValueError(f'{MASTER_DEATH_FLAG!r} must contain only non-missing 0 and 1 values.')
    master_subset[MASTER_DEATH_FLAG] = master_subset[MASTER_DEATH_FLAG].astype(int)
    master_subset[MASTER_DEATH_TIME] = pd.to_numeric(master_subset[MASTER_DEATH_TIME], errors='coerce')
    master_subset[MASTER_DM_FIELD] = pd.to_numeric(master_subset[MASTER_DM_FIELD], errors='raise')
    if master_subset[MASTER_DM_FIELD].isna().any() or not master_subset[MASTER_DM_FIELD].isin([0, 1]).all():
        raise ValueError(f'{MASTER_DM_FIELD!r} must contain only non-missing 0 and 1 values.')
    master_subset[MASTER_DM_FIELD] = master_subset[MASTER_DM_FIELD].astype(int)
    merged = raw.merge(master_subset, left_on=ID, right_on=MASTER_ID, how='left', validate='one_to_one', indicator=True)
    unmatched = merged.loc[merged['_merge'] != 'both', ID].tolist()
    if unmatched:
        preview = unmatched[:10]
        raise ValueError(f'The master worksheet did not match {len(unmatched)} CSV study IDs. First unmatched IDs: {preview}')
    merged = merged.drop(columns=['_merge'])
    if merged[MASTER_DEATH_FLAG].isna().any():
        raise ValueError('Death other cause was missing after the CSV/master merge.')
    dm_mismatch = merged['dmmed'].astype(int).ne(merged[MASTER_DM_FIELD].astype(int))
    if dm_mismatch.any():
        raise ValueError(f'CSV dmmed disagreed with master {MASTER_DM_FIELD!r} for {int(dm_mismatch.sum())} participants.')
    merged[CSV_TIME] = merged[TIME].astype(float)
    potential_competing = merged[EVENT].eq(0) & merged[MASTER_DEATH_FLAG].eq(1)
    missing_verified = potential_competing & merged[MASTER_DEATH_TIME].isna()
    if missing_verified.any():
        raise ValueError(f'{int(missing_verified.sum())} potential non-ASCVD deaths lack a verified {MASTER_DEATH_TIME!r} value.')
    different_verified_time = potential_competing & ~np.isclose(merged[MASTER_DEATH_TIME].to_numpy(dtype=float), merged[CSV_TIME].to_numpy(dtype=float), rtol=0.0, atol=1e-12)
    if different_verified_time.any():
        raise ValueError(f'{int(different_verified_time.sum())} potential competing-death rows have different master F/U Time and CSV time_to_event values. The main eight-model pipeline reads CSV time_to_event, so the run was stopped rather than using inconsistent censoring times.')
    later_death_after_ascvd = merged[EVENT].eq(1) & merged[MASTER_DEATH_FLAG].eq(1)
    missing_later_death_time = later_death_after_ascvd & merged[MASTER_DEATH_TIME].isna()
    if missing_later_death_time.any():
        raise ValueError(f'{int(missing_later_death_time.sum())} ASCVD-event rows with a later death flag lack {MASTER_DEATH_TIME!r}.')
    event_order_conflict = later_death_after_ascvd & (merged[MASTER_DEATH_TIME] <= merged[CSV_TIME])
    if event_order_conflict.any():
        raise ValueError(f'{int(event_order_conflict.sum())} rows have a non-ASCVD death time that is not later than the recorded ASCVD event time. First-event order is unclear.')
    merged.loc[potential_competing, TIME] = merged.loc[potential_competing, MASTER_DEATH_TIME].astype(float)
    merged[COMPETING_EVENT] = potential_competing.astype(int)
    return merged

def validate_expected_study_counts(prepared_all: pd.DataFrame) -> pd.DataFrame:
    eligible = prepared_all.loc[prepared_all[TIME] > 0].copy().reset_index(drop=True)
    competing = eligible[COMPETING_EVENT].eq(1)
    other_censored = eligible[EVENT].eq(0) & ~competing
    potential_all = prepared_all[EVENT].eq(0) & prepared_all[MASTER_DEATH_FLAG].eq(1)
    later_death_all = prepared_all[EVENT].eq(1) & prepared_all[MASTER_DEATH_FLAG].eq(1)
    checks = {'source rows': (len(prepared_all), EXPECTED_SOURCE_ROWS), 'eligible rows': (len(eligible), EXPECTED_ELIGIBLE_ROWS), 'ASCVD events': (int(eligible[EVENT].sum()), EXPECTED_ASCVD_EVENTS), 'unique eligible IDs': (int(eligible[ID].nunique()), EXPECTED_ELIGIBLE_ROWS), 'eligible non-ASCVD deaths': (int(competing.sum()), EXPECTED_COMPETING_DEATHS), 'other right-censored observations': (int(other_censored.sum()), EXPECTED_OTHER_CENSORED), 'all other-cause death flags': (int(prepared_all[MASTER_DEATH_FLAG].sum()), EXPECTED_DEATH_FLAGS), 'potential first non-ASCVD deaths': (int(potential_all.sum()), EXPECTED_POTENTIAL_COMPETING_DEATHS), 'non-positive competing-death exclusions': (int((potential_all & (prepared_all[TIME] <= 0)).sum()), EXPECTED_NONPOSITIVE_COMPETING_EXCLUSIONS), 'ASCVD-first later-death rows': (int(later_death_all.sum()), EXPECTED_ASCVD_FIRST_LATER_DEATHS)}
    mismatches = [f'{label}: found {found}, expected {expected}' for label, (found, expected) in checks.items() if found != expected]
    if mismatches:
        raise ValueError('Input files do not match the validated study dataset: ' + '; '.join(mismatches))
    if eligible.loc[competing, MASTER_DEATH_TIME].isna().any():
        raise ValueError('An eligible non-ASCVD death lacks verified F/U Time.')
    if (eligible.loc[competing, MASTER_DEATH_TIME] <= 0).any():
        raise ValueError('An eligible non-ASCVD death has non-positive verified F/U Time.')
    return eligible

def postprocess_imputation(x: pd.DataFrame) -> pd.DataFrame:
    out = x.copy()
    for col in BINARY:
        out[col] = (out[col].to_numpy(float) >= 0.5).astype(int)
    lower = {'age': 18.0, 'sbp': 40.0, 'dbp': 20.0, 'bmi': 10.0, 'chol': 0.1, 'hdl': 0.1, 'trig': 0.05, 'hba1c': 2.0, 'egfr': 1.0}
    for col, value in lower.items():
        out[col] = np.clip(out[col].to_numpy(float), value, None)
    return out

def imputation_matrix(df: pd.DataFrame) -> pd.DataFrame:
    return df[PREDICTORS].copy()

def impute_split(train: pd.DataFrame, test: pd.DataFrame, m: int, seed: int, method: str, competing_col: str | None=None, n_jobs: int=1, diagnostics: list[dict[str, object]] | None=None, diagnostic_context: dict[str, object] | None=None) -> list[tuple[pd.DataFrame, pd.DataFrame]]:
    train_matrix = imputation_matrix(train)
    test_matrix = imputation_matrix(test)
    train_missing_before = int(train_matrix.isna().sum().sum())
    test_missing_before = int(test_matrix.isna().sum().sum())
    cols = list(train_matrix.columns)
    completed = []
    max_iter = EXTRA_TREES_FIXED_CYCLES if method == 'extra_trees' else 20
    for i in range(m):
        if method == 'bayesian':
            estimator = BayesianRidge()
            posterior = True
        elif method == 'extra_trees':
            estimator = ExtraTreesRegressor(n_estimators=100, min_samples_leaf=2, max_features=1.0, random_state=seed + i, n_jobs=n_jobs)
            posterior = False
        else:
            raise ValueError(method)
        imp = IterativeImputer(estimator=estimator, max_iter=max_iter, sample_posterior=posterior, random_state=seed + i, initial_strategy='median', imputation_order='ascending', tol=0.0 if method == 'extra_trees' else 0.001)
        with warnings.catch_warnings(record=True) as caught_warnings:
            if method == 'extra_trees':
                warnings.simplefilter('always')
            else:
                warnings.simplefilter('always', ConvergenceWarning)
            tr = pd.DataFrame(imp.fit_transform(train_matrix[cols]), columns=cols)
            te = pd.DataFrame(imp.transform(test_matrix[cols]), columns=cols)
        convergence_warnings = [warning for warning in caught_warnings if issubclass(warning.category, ConvergenceWarning)]
        expected_fixed_cycle_warnings = [warning for warning in convergence_warnings if method == 'extra_trees' and 'early stopping criterion not reached' in str(warning.message).lower()]
        unexpected_convergence_warnings = [warning for warning in convergence_warnings if warning not in expected_fixed_cycle_warnings]
        unexpected_fixed_cycle_warnings = [warning for warning in caught_warnings if warning not in expected_fixed_cycle_warnings]
        if method == 'extra_trees' and (int(imp.n_iter_) != EXTRA_TREES_FIXED_CYCLES or len(expected_fixed_cycle_warnings) != 1 or unexpected_fixed_cycle_warnings):
            warning_details = ' | '.join((f'{warning.category.__name__}: {warning.message}' for warning in unexpected_fixed_cycle_warnings))
            raise RuntimeError(f"ExtraTrees did not complete the fixed 10-cycle policy with exactly the expected IterativeImputer cycle-limit notice and no other warnings. Unexpected warnings: {warning_details or 'none'}.")
        train_observed = train_matrix[cols].notna().to_numpy()
        test_observed = test_matrix[cols].notna().to_numpy()
        observed_cells_preserved = bool(np.allclose(tr.to_numpy(dtype=float)[train_observed], train_matrix[cols].to_numpy(dtype=float)[train_observed], rtol=0.0, atol=0.0) and np.allclose(te.to_numpy(dtype=float)[test_observed], test_matrix[cols].to_numpy(dtype=float)[test_observed], rtol=0.0, atol=0.0))
        tr_pred = postprocess_imputation(tr[PREDICTORS])
        te_pred = postprocess_imputation(te[PREDICTORS])
        all_values_finite = bool(np.isfinite(tr_pred.to_numpy(dtype=float)).all() and np.isfinite(te_pred.to_numpy(dtype=float)).all())
        plausibility_passed = bool(all_values_finite and all((tr_pred[column].isin([0, 1]).all() for column in BINARY)) and all((te_pred[column].isin([0, 1]).all() for column in BINARY)))
        if not observed_cells_preserved or not all_values_finite or (not plausibility_passed):
            raise RuntimeError(f'{method} imputation failed observed-cell, finite-value or plausibility checks.')
        if diagnostics is not None:
            diagnostic_row = dict(diagnostic_context or {})
            diagnostic_row.update({'imputer': method, 'imputation': int(i + 1), 'imputer_seed': int(seed + i), 'imputer_n_iter': int(imp.n_iter_), 'imputer_max_iter': int(max_iter), 'imputer_iteration_limit_reached': bool(imp.n_iter_ >= max_iter), 'imputer_convergence_warning_count': int(len(convergence_warnings)), 'imputer_convergence_warning_messages': ' | '.join((str(warning.message) for warning in convergence_warnings)), 'imputer_policy': 'prespecified_fixed_10_cycles_tol_0_no_convergence_claim' if method == 'extra_trees' else 'prespecified_20_cycles_posterior_sampling_no_early_stopping', 'imputer_fixed_cycle_policy': True, 'imputer_fixed_cycle_completed': bool(imp.n_iter_ == (EXTRA_TREES_FIXED_CYCLES if method == 'extra_trees' else 20)), 'imputer_expected_fixed_cycle_warning_count': int(len(expected_fixed_cycle_warnings)), 'imputer_unexpected_convergence_warning_count': int(len(unexpected_convergence_warnings)), 'train_missing_cells_before': train_missing_before, 'test_missing_cells_before': test_missing_before, 'train_missing_cells_after': int(tr_pred.isna().sum().sum()), 'test_missing_cells_after': int(te_pred.isna().sum().sum()), 'observed_cells_preserved': observed_cells_preserved, 'all_values_finite': all_values_finite, 'plausibility_passed': plausibility_passed})
            diagnostics.append(diagnostic_row)
        tr_out = pd.concat([train[[ID, TIME, EVENT] + ([competing_col] if competing_col else [])].reset_index(drop=True), tr_pred], axis=1)
        te_out = pd.concat([test[[ID, TIME, EVENT] + ([competing_col] if competing_col else [])].reset_index(drop=True), te_pred], axis=1)
        completed.append((tr_out, te_out))
    return completed

def make_strata(df: pd.DataFrame) -> pd.Series:
    strata = df[EVENT].astype(str) + '_' + df[SEX].astype(str)
    return strata if strata.value_counts().min() >= 2 else df[EVENT].astype(str)

def rcs_basis(x: np.ndarray, knots: np.ndarray) -> np.ndarray:
    x = np.asarray(x, dtype=float)
    k = np.asarray(knots, dtype=float)
    if len(k) < 3 or not np.all(np.diff(k) > 0):
        raise ValueError('RCS knots must contain at least three distinct values')
    scale = (k[-1] - k[0]) ** 2
    cols = [x]
    for j in range(len(k) - 2):
        term = np.maximum(x - k[j], 0.0) ** 3
        term -= np.maximum(x - k[-2], 0.0) ** 3 * (k[-1] - k[j]) / (k[-1] - k[-2])
        term += np.maximum(x - k[-1], 0.0) ** 3 * (k[-2] - k[j]) / (k[-1] - k[-2])
        cols.append(term / scale)
    return np.column_stack(cols)

@dataclass
class DesignBuilder:
    spline: bool
    knots: dict[str, np.ndarray] | None = None
    names: list[str] | None = None

    def fit_transform(self, df: pd.DataFrame) -> np.ndarray:
        self.knots = {}
        arrays = []
        names = []
        for col in PREDICTORS:
            x = df[col].to_numpy(float)
            if self.spline and col in RCS_VARS:
                knots = np.unique(np.quantile(x, [0.1, 0.5, 0.9]))
                if len(knots) == 3:
                    self.knots[col] = knots
                    b = rcs_basis(x, knots)
                    arrays.append(b)
                    names.extend([col, f'{col}_rcs1'])
                    continue
            arrays.append(x[:, None])
            names.append(col)
        self.names = names
        return np.column_stack(arrays)

    def transform(self, df: pd.DataFrame) -> np.ndarray:
        arrays = []
        for col in PREDICTORS:
            x = df[col].to_numpy(float)
            if self.spline and self.knots is not None and (col in self.knots):
                arrays.append(rcs_basis(x, self.knots[col]))
            else:
                arrays.append(x[:, None])
        return np.column_stack(arrays)

class CoxPH:

    def __init__(self, ridge: float=1e-07):
        self.ridge = ridge

    def fit(self, x: np.ndarray, time: np.ndarray, event: np.ndarray) -> 'CoxPH':
        x = np.asarray(x, float)
        time = np.asarray(time, float)
        event = np.asarray(event, int)
        self.mean_ = x.mean(axis=0)
        self.scale_ = x.std(axis=0, ddof=0)
        self.scale_[self.scale_ < 1e-10] = 1.0
        z = (x - self.mean_) / self.scale_
        unique_event_times = np.sort(np.unique(time[event == 1]))

        def value_grad(beta: np.ndarray) -> tuple[float, np.ndarray]:
            eta = np.clip(z @ beta, -40, 40)
            exp_eta = np.exp(eta)
            loglik = 0.0
            grad = np.zeros(z.shape[1])
            for t in unique_event_times:
                ev = (time == t) & (event == 1)
                risk = time >= t
                d = int(ev.sum())
                s0 = exp_eta[risk].sum()
                s1 = (exp_eta[risk, None] * z[risk]).sum(axis=0)
                loglik += (z[ev] @ beta).sum() - d * math.log(s0)
                grad += z[ev].sum(axis=0) - d * s1 / s0
            loglik -= 0.5 * self.ridge * float(beta @ beta)
            grad -= self.ridge * beta
            return (-loglik, -grad)
        opt = minimize(lambda b: value_grad(b), np.zeros(z.shape[1]), method='L-BFGS-B', jac=True, options={'maxiter': 500, 'ftol': 1e-11, 'gtol': 1e-07})
        if not opt.success and np.linalg.norm(opt.jac) > 0.001:
            raise RuntimeError(f'Cox fit failed: {opt.message}')
        self.beta_ = opt.x
        eta = np.clip(z @ self.beta_, -40, 40)
        exp_eta = np.exp(eta)
        increments = []
        for t in unique_event_times:
            d = np.sum((time == t) & (event == 1))
            increments.append(d / exp_eta[time >= t].sum())
        self.event_times_ = unique_event_times
        self.baseline_increments_ = np.asarray(increments, float)
        return self

    def linear_predictor(self, x: np.ndarray) -> np.ndarray:
        z = (np.asarray(x, float) - self.mean_) / self.scale_
        return z @ self.beta_

    def baseline_cumulative_hazard(self, times: np.ndarray) -> np.ndarray:
        idx = np.searchsorted(self.event_times_, np.asarray(times, float), side='right')
        cs = np.r_[0.0, np.cumsum(self.baseline_increments_)]
        return cs[idx]

    def survival(self, x: np.ndarray, times: np.ndarray) -> np.ndarray:
        h0 = self.baseline_cumulative_hazard(np.asarray(times, float))
        rr = np.exp(np.clip(self.linear_predictor(x), -40, 40))
        return np.exp(-rr[:, None] * h0[None, :])

def km_survival(time: np.ndarray, event: np.ndarray, query: np.ndarray, left: bool=False) -> np.ndarray:
    time = np.asarray(time, float)
    event = np.asarray(event, int)
    event_times = np.sort(np.unique(time[event == 1]))
    surv = 1.0
    survival_after = []
    for t in event_times:
        n = np.sum(time >= t)
        d = np.sum((time == t) & (event == 1))
        surv *= 1.0 - d / n
        survival_after.append(surv)
    if not len(event_times):
        return np.ones(len(np.atleast_1d(query)))
    side = 'left' if left else 'right'
    idx = np.searchsorted(event_times, np.asarray(query, float), side=side) - 1
    out = np.ones(len(np.atleast_1d(query)))
    mask = idx >= 0
    out[mask] = np.asarray(survival_after)[idx[mask]]
    return out

def km_risk(time: np.ndarray, event: np.ndarray, horizon: float) -> float:
    return 1.0 - float(km_survival(time, event, np.asarray([horizon]))[0])

def aalen_johansen_risks(time: np.ndarray, status: np.ndarray, horizon: float) -> tuple[float, float]:
    time = np.asarray(time, float)
    status = np.asarray(status, int)
    event_times = np.sort(np.unique(time[(status > 0) & (time <= horizon)]))
    s = 1.0
    f1 = 0.0
    f2 = 0.0
    for t in event_times:
        n = np.sum(time >= t)
        d1 = np.sum((time == t) & (status == 1))
        d2 = np.sum((time == t) & (status == 2))
        f1 += s * d1 / n
        f2 += s * d2 / n
        s *= 1.0 - (d1 + d2) / n
    return (float(f1), float(f2))

def bootstrap_competing_risk_estimates(time: np.ndarray, status: np.ndarray, horizon: float, seed: int=20260320, n_bootstrap: int=1000) -> dict[str, float | int | str]:
    time = np.asarray(time, dtype=float)
    status = np.asarray(status, dtype=int)
    rng = np.random.default_rng(seed)
    estimates: list[tuple[float, float, float]] = []
    for _ in range(n_bootstrap):
        index = rng.integers(0, len(time), len(time))
        sampled_time = time[index]
        sampled_status = status[index]
        aj_ascvd, aj_death = aalen_johansen_risks(sampled_time, sampled_status, horizon)
        km_ascvd = km_risk(sampled_time, (sampled_status == 1).astype(int), horizon)
        values = (aj_ascvd, aj_death, km_ascvd - aj_ascvd)
        if np.isfinite(values).all():
            estimates.append(values)
    minimum_successful = math.ceil(0.95 * n_bootstrap)
    if len(estimates) < minimum_successful:
        raise RuntimeError(f'Only {len(estimates)}/{n_bootstrap} competing-risk bootstrap estimates were successful; at least {minimum_successful} were required.')
    array = np.asarray(estimates, dtype=float)
    low, high = np.percentile(array, [2.5, 97.5], axis=0)
    return {'competing_risk_bootstrap_requested': int(n_bootstrap), 'competing_risk_bootstrap_successful': int(len(estimates)), 'competing_risk_bootstrap_seed': int(seed), 'aalen_johansen_ascvd_cif_120_bootstrap_95_low': float(low[0]), 'aalen_johansen_ascvd_cif_120_bootstrap_95_high': float(high[0]), 'aalen_johansen_non_ascvd_death_cif_120_bootstrap_95_low': float(low[1]), 'aalen_johansen_non_ascvd_death_cif_120_bootstrap_95_high': float(high[1]), 'absolute_km_minus_aj_bootstrap_95_low': float(low[2]), 'absolute_km_minus_aj_bootstrap_95_high': float(high[2]), 'bootstrap_interval_interpretation': 'pointwise participant-bootstrap percentile 95% intervals'}

def cause_specific_cif(event_model: CoxPH, death_model: CoxPH, x: np.ndarray, horizon: float) -> np.ndarray:
    times = np.unique(np.r_[event_model.event_times_[event_model.event_times_ <= horizon], death_model.event_times_[death_model.event_times_ <= horizon]])
    event_inc = np.zeros(len(times))
    death_inc = np.zeros(len(times))
    event_map = dict(zip(event_model.event_times_, event_model.baseline_increments_))
    death_map = dict(zip(death_model.event_times_, death_model.baseline_increments_))
    for j, t in enumerate(times):
        event_inc[j] = event_map.get(t, 0.0)
        death_inc[j] = death_map.get(t, 0.0)
    rr1 = np.exp(np.clip(event_model.linear_predictor(x), -40, 40))
    rr2 = np.exp(np.clip(death_model.linear_predictor(x), -40, 40))
    s = np.ones(len(x))
    cif = np.zeros(len(x))
    for dh10, dh20 in zip(event_inc, death_inc):
        dh1 = rr1 * dh10
        dh2 = rr2 * dh20
        total_hazard = dh1 + dh2
        event_share = np.divide(dh1, total_hazard, out=np.zeros_like(dh1), where=total_hazard > 0.0)
        event_probability = -np.expm1(-total_hazard) * event_share
        cif += s * event_probability
        s *= np.exp(-total_hazard)
    return np.clip(cif, 0.0, 1.0)

def evaluate_cox(train: pd.DataFrame, test: pd.DataFrame, spline: bool, horizon: float=120.0) -> tuple[dict[str, float], np.ndarray, CoxPH, DesignBuilder]:
    from _analysis_main import compute_brier_metrics, get_surv_array, harrell_c_index, uno_c_index
    design = DesignBuilder(spline=spline)
    x_train = design.fit_transform(train)
    x_test = design.transform(test)
    model = CoxPH(ridge=1e-05).fit(x_train, train[TIME], train[EVENT])
    eval_times = np.arange(12.0, horizon + 1.0, 1.0)
    surv = model.survival(x_test, eval_times)
    risk = 1.0 - surv[:, -1]
    y_train = get_surv_array(train)
    y_test = get_surv_array(test)
    risk_score = model.linear_predictor(x_test)
    brier = compute_brier_metrics(y_train, y_test, surv, eval_times, horizon_months=horizon)
    observed = km_risk(test[TIME], test[EVENT], horizon)
    metrics = {'harrell_c': harrell_c_index(y_test, risk_score), 'uno_c': uno_c_index(y_train, y_test, risk_score, tau=np.nextafter(float(horizon), np.inf)), 'brier_120': float(brier['brier_score_at_horizon']), 'ibs_12_120': float(brier['integrated_brier_score']), 'oe_ratio': float(observed / np.mean(risk)), 'observed_km_120': float(observed), 'mean_predicted_120': float(np.mean(risk)), 'calibration_slope': calibration_slope_from_probs(test, risk, horizon=horizon)}
    return (metrics, risk, model, design)

def vif_summary(imputed_trains: list[pd.DataFrame]) -> pd.DataFrame:
    rows = []
    for imp_no, df in enumerate(imputed_trains, 1):
        x = df[PREDICTORS].to_numpy(float)
        for j, col in enumerate(PREDICTORS):
            y = x[:, j]
            others = np.delete(x, j, axis=1)
            r2 = LinearRegression().fit(others, y).score(others, y)
            rows.append({'imputation': imp_no, 'predictor': DISPLAY[col], 'vif': 1.0 / max(1.0 - r2, 1e-12)})
    long = pd.DataFrame(rows)
    return long.groupby('predictor', as_index=False).agg(mean_vif=('vif', 'mean'), min_vif=('vif', 'min'), max_vif=('vif', 'max')).sort_values('mean_vif', ascending=False).reset_index(drop=True)

def ldl_component_multicollinearity_vif(eligible: pd.DataFrame) -> pd.DataFrame:
    columns = ['chol', 'hdl', 'trig', 'ldl']
    labels = {'chol': 'TC', 'hdl': 'HDL-C', 'trig': 'Triglycerides', 'ldl': 'LDL-C'}
    complete = eligible.loc[:, columns].dropna().copy()
    if len(complete) != 1206:
        raise RuntimeError(f'The LDL-C complete-case diagnostic contained {len(complete)} rows; expected 1206.')
    matrix = complete.to_numpy(dtype=float)
    if not np.isfinite(matrix).all():
        raise RuntimeError('The LDL-C complete-case diagnostic contained a non-finite value.')
    rows: list[dict[str, object]] = []
    for index, column in enumerate(columns):
        outcome = matrix[:, index]
        other_predictors = np.delete(matrix, index, axis=1)
        r_squared = float(LinearRegression().fit(other_predictors, outcome).score(other_predictors, outcome))
        vif = float(1.0 / max(1.0 - r_squared, np.finfo(float).eps))
        if not np.isfinite(r_squared) or not np.isfinite(vif) or vif <= 0.0:
            raise RuntimeError(f'The LDL-C component VIF was invalid for {column}.')
        rows.append({'predictor': labels[column], 'complete_case_n': int(len(complete)), 'r_squared': r_squared, 'vif': vif})
    return pd.DataFrame(rows)

def calibration_curve_with_bootstrap(test: pd.DataFrame, risk: np.ndarray, seed: int, n_boot: int=300) -> tuple[pd.DataFrame, dict[str, object]]:
    p = np.clip(np.asarray(risk, float), 1e-05, 1 - 1e-05)
    lp = np.log(-np.log(1.0 - p))
    knots = np.unique(np.quantile(lp, [0.1, 0.5, 0.9]))
    grid_p = np.linspace(max(0.001, p.min()), min(1.0 - 1e-05, p.max()), 101)
    grid_lp = np.log(-np.log(1.0 - grid_p))
    original_time = test[TIME].to_numpy(dtype=float)
    original_event = test[EVENT].to_numpy(dtype=int)
    calibration_time = np.minimum(original_time, 120.0)
    calibration_event = ((original_event == 1) & (original_time <= 120.0)).astype(int)

    def fit_predict(index: np.ndarray) -> np.ndarray:
        basis = rcs_basis(lp[index], knots)
        model = CoxPH(ridge=0.0001).fit(basis, calibration_time[index], calibration_event[index])
        g = rcs_basis(grid_lp, knots)
        return 1.0 - model.survival(g, np.asarray([120.0]))[:, 0]
    main = fit_predict(np.arange(len(test)))
    rng = np.random.default_rng(seed)
    boot = []
    for _ in range(n_boot):
        idx = rng.integers(0, len(test), len(test))
        if calibration_event[idx].sum() < 3:
            continue
        try:
            boot.append(fit_predict(idx))
        except Exception:
            continue
    minimum_successful = max(1, int(math.ceil(0.9 * n_boot)))
    if len(boot) < minimum_successful:
        raise RuntimeError(f'Flexible calibration completed only {len(boot)}/{n_boot} bootstrap fits; at least {minimum_successful} were required.')
    boot_arr = np.asarray(boot)
    low, high = np.nanpercentile(boot_arr, [2.5, 97.5], axis=0)
    curve = pd.DataFrame({'predicted_risk': grid_p, 'observed_risk': main, 'bootstrap_95_low': low, 'bootstrap_95_high': high})
    if not np.isfinite(curve.to_numpy(dtype=float)).all():
        raise RuntimeError('Flexible calibration produced a non-finite curve or band.')
    metadata: dict[str, object] = {'n_bootstrap_requested': int(n_boot), 'n_bootstrap_successful': int(len(boot)), 'minimum_successful_required': int(minimum_successful), 'complete': bool(len(boot) >= minimum_successful), 'interval_interpretation': 'pointwise participant-bootstrap 95% bands conditional on fixed test predictions'}
    return (curve, metadata)

def plot_combined_calibration_curves(cox_curve: pd.DataFrame, spline_curve: pd.DataFrame, png_path: Path, tiff_path: Path, axis_max: float) -> None:
    axis_max = float(np.clip(axis_max, 0.05, 1.0))
    fig, axes = plt.subplots(1, 2, figsize=(11.4, 5.3), sharex=True, sharey=True)
    for ax, curve, panel_title in zip(axes, [cox_curve, spline_curve], ['A. CoxPH', 'B. Spline Cox']):
        ax.plot([0, axis_max], [0, axis_max], '--', color='#4C78A8', lw=1.2, label='Ideal')
        ax.fill_between(curve['predicted_risk'], curve['bootstrap_95_low'], curve['bootstrap_95_high'], color='#F28E2B', alpha=0.22, linewidth=0, label='Bootstrap 95% band')
        ax.plot(curve['predicted_risk'], curve['observed_risk'], color='#F28E2B', lw=2.0, label='Flexible calibration')
        ax.set(xlim=(0, axis_max), ylim=(0, axis_max), xlabel='Predicted 120-month risk', ylabel='Observed 120-month risk')
        ax.set_title(panel_title, loc='left', fontweight='bold')
        ax.set_aspect('equal', adjustable='box')
        ax.legend(frameon=False, loc='upper left')
    fig.tight_layout()
    if getattr(fig, '_suptitle', None) is not None or list(getattr(fig, 'texts', [])):
        raise RuntimeError('Secondary calibration composite must not contain an overall title or footer.')
    fig.savefig(png_path, dpi=600, bbox_inches='tight')
    fig.savefig(tiff_path, dpi=600, bbox_inches='tight', pil_kwargs={'compression': 'tiff_lzw'})
    plt.close(fig)

def data_audit(prepared_all: pd.DataFrame, source_column_names: list[object] | None=None) -> dict[str, object]:
    eligible = prepared_all.loc[prepared_all[TIME] > 0].copy()
    fatal = (eligible['Fatal MI - HARD CHD'] == 1) | (eligible['Fatal CVA - HARD Stroke'] == 1)
    competing = eligible[COMPETING_EVENT].eq(1)
    other_censored = eligible[EVENT].eq(0) & ~competing
    recent = 'Seen in clinic after Jan 2019(1=Yes)'
    km_ascvd = km_risk(eligible[TIME].to_numpy(), eligible[EVENT].to_numpy(), 120.0)
    status = np.where(eligible[EVENT] == 1, 1, np.where(competing, 2, 0))
    aj_ascvd, aj_death = aalen_johansen_risks(eligible[TIME].to_numpy(), status, 120.0)
    deaths_by_120 = int((competing & (eligible[TIME] <= 120)).sum())
    km_minus_aj = km_ascvd - aj_ascvd
    competing_uncertainty = bootstrap_competing_risk_estimates(eligible[TIME].to_numpy(dtype=float), status, 120.0)
    ldl = eligible['ldl'].astype(float)
    ldl_numeric = ldl.dropna()
    friedewald = eligible['chol'].astype(float) - eligible['hdl'].astype(float) - eligible['trig'].astype(float) / 2.2
    complete = ldl.notna() & friedewald.notna()
    crp_recorded = int(eligible['CRP'].notna().sum())
    crp_numeric_series = pd.to_numeric(eligible['CRP'], errors='coerce')
    crp_numeric_values = crp_numeric_series.dropna()
    crp_numeric = int(len(crp_numeric_values))
    crp_unrecorded = int(len(eligible) - crp_recorded)
    crp_recorded_nonnumeric = eligible.loc[eligible['CRP'].notna() & crp_numeric_series.isna(), 'CRP']
    crp_recorded_nonnumeric_normalized = crp_recorded_nonnumeric.astype(str).str.replace('\\s+', '', regex=True)
    if not crp_recorded_nonnumeric_normalized.isin({'<1', '<5'}).all():
        raise RuntimeError('The recorded nonnumeric CRP values are not exclusively the validated below-reporting-limit values.')
    crp_recorded_below_reporting_limit = int(len(crp_recorded_nonnumeric))
    potential_all = prepared_all[EVENT].eq(0) & prepared_all[MASTER_DEATH_FLAG].eq(1)
    competing_time_matches = np.isclose(prepared_all.loc[potential_all, MASTER_DEATH_TIME].to_numpy(float), prepared_all.loc[potential_all, CSV_TIME].to_numpy(float), rtol=0.0, atol=1e-12)
    ascvd_then_later_death = prepared_all[EVENT].eq(1) & prepared_all[MASTER_DEATH_FLAG].eq(1)
    lpa_candidates = lpa_candidate_columns(source_column_names or [])
    if lpa_candidates:
        raise RuntimeError('Unexpected possible Lp(a) source fields were found: ' + ', '.join(lpa_candidates))
    return {'source_rows': int(len(prepared_all)), 'eligible_n': int(len(eligible)), 'ascvd_events': int(eligible[EVENT].sum()), 'fatal_ascvd_events': int((fatal & (eligible[EVENT] == 1)).sum()), 'coronary_heart_disease_death_events': int(((eligible['Fatal MI - HARD CHD'] == 1) & (eligible[EVENT] == 1)).sum()), 'fatal_stroke_events': int(((eligible['Fatal CVA - HARD Stroke'] == 1) & (eligible[EVENT] == 1)).sum()), 'nonfatal_ascvd_events': int((~fatal & (eligible[EVENT] == 1)).sum()), 'death_other_cause_flags_all_rows': int(prepared_all[MASTER_DEATH_FLAG].sum()), 'potential_non_ascvd_deaths_before_followup_exclusion': int(potential_all.sum()), 'potential_non_ascvd_deaths_excluded_for_nonpositive_time': int((potential_all & (prepared_all[TIME] <= 0)).sum()), 'documented_non_ascvd_deaths_before_ascvd': int(competing.sum()), 'documented_non_ascvd_deaths_by_120_months': deaths_by_120, 'documented_non_ascvd_deaths_after_120_months': int((competing & (eligible[TIME] > 120)).sum()), 'other_right_censored': int(other_censored.sum()), 'other_right_censored_seen_after_jan_2019': int((other_censored & (eligible[recent] == 1)).sum()), 'other_right_censored_without_record_after_jan_2019': int((other_censored & (eligible[recent] == 0)).sum()), 'ascvd_first_then_later_non_ascvd_death_flags': int(ascvd_then_later_death.sum()), 'event_order_conflicts': 0, 'verified_death_time_source_workbook_sheet': MASTER_SHEET, 'verified_death_time_source_column': MASTER_DEATH_TIME, 'verified_death_time_unit': 'months', 'verified_death_time_source_definition': 'Master-sheet F/U Time was used as the time to non-ASCVD death for rows with Death other cause=1.', 'verified_death_time_used_for_competing_events': True, 'potential_competing_times_matching_csv_time_to_event': int(competing_time_matches.sum()), 'potential_competing_times_different_from_csv_time_to_event': int((~competing_time_matches).sum()), 'eligible_competing_time_min_months': float(eligible.loc[competing, TIME].min()), 'eligible_competing_time_max_months': float(eligible.loc[competing, TIME].max()), 'km_ascvd_risk_120_treating_death_as_censoring': km_ascvd, 'aalen_johansen_ascvd_cif_120': aj_ascvd, 'aalen_johansen_non_ascvd_death_cif_120': aj_death, 'absolute_km_minus_aj': km_minus_aj, **competing_uncertainty, 'competing_event_time_source_used': True, 'ldl_nonmissing_n': int(ldl.notna().sum()), 'ldl_missing_n': int(ldl.isna().sum()), 'ldl_mean_mmol_l': float(ldl_numeric.mean()), 'ldl_sd_mmol_l': float(ldl_numeric.std(ddof=1)), 'ldl_median_mmol_l': float(ldl_numeric.median()), 'ldl_q1_mmol_l': float(ldl_numeric.quantile(0.25)), 'ldl_q3_mmol_l': float(ldl_numeric.quantile(0.75)), 'ldl_min_mmol_l': float(ldl_numeric.min()), 'ldl_max_mmol_l': float(ldl_numeric.max()), 'ldl_correlation_with_friedewald': float(np.corrcoef(ldl[complete], friedewald[complete])[0, 1]), 'ldl_within_0_11_mmol_l_of_friedewald_pct': float(100 * np.mean(np.abs(ldl[complete] - friedewald[complete]) <= 0.11)), 'generic_crp_recorded_n': crp_recorded, 'generic_crp_numeric_n': crp_numeric, 'generic_crp_unrecorded_n': crp_unrecorded, 'generic_crp_recorded_below_reporting_limit_n': crp_recorded_below_reporting_limit, 'generic_crp_unrecorded_pct': float(100 * crp_unrecorded / len(eligible)), 'generic_crp_numeric_mean': float(crp_numeric_values.mean()), 'generic_crp_numeric_sd': float(crp_numeric_values.std(ddof=1)), 'generic_crp_numeric_median': float(crp_numeric_values.median()), 'generic_crp_numeric_q1': float(crp_numeric_values.quantile(0.25)), 'generic_crp_numeric_q3': float(crp_numeric_values.quantile(0.75)), 'generic_crp_numeric_min': float(crp_numeric_values.min()), 'generic_crp_numeric_max': float(crp_numeric_values.max()), 'generic_crp_is_hscrp': False, 'generic_crp_units_verified': False, 'dmmed_master_source_column': MASTER_DM_FIELD, 'dmmed_master_mismatch_n': int(prepared_all['dmmed'].astype(int).ne(prepared_all[MASTER_DM_FIELD].astype(int)).sum()), 'dmmed_source_definition_interpretation': "Master field 'Hx of DM 0=no' was interpreted as glucose-lowering medication use for this analysis.", 'lpa_schema_search_performed': True, 'lpa_schema_sources': ['validated analysis CSV', f'selected master worksheet: {MASTER_SHEET}'], 'lpa_candidate_columns': lpa_candidates, 'lpa_available': bool(lpa_candidates)}

def prepare_run_manifest(outdir: Path, csv_path: Path, master_path: Path, main_dependencies: dict[str, Path], args: argparse.Namespace) -> str:
    configuration = {'schema_version': 8, 'analysis': 'secondary_analyses_repeated_200_splits', 'analysis_script_sha256': {'_analysis_secondary.py': file_sha256(Path(__file__).with_name('_analysis_secondary.py')), '_analysis_secondary_core.py': file_sha256(Path(__file__).resolve()), '_analysis_main.py': file_sha256(Path(__file__).with_name('_analysis_main.py'))}, 'python_version': platform.python_version(), 'package_versions': secondary_package_versions(), 'csv_sha256': file_sha256(csv_path), 'master_workbook_sha256': file_sha256(master_path), 'master_sheet': MASTER_SHEET, 'verified_non_ascvd_death_time_column': MASTER_DEATH_TIME, 'verified_non_ascvd_death_time_unit': 'months', 'verified_non_ascvd_death_time_confirmed_by_source_definition': True, 'main_analysis_dependencies': {name: {'relative_name': path.relative_to(Path(args.main_output_dir).expanduser().resolve()).as_posix(), 'sha256': file_sha256(path)} for name, path in sorted(main_dependencies.items())}, 'settings': {'seed': int(args.seed), 'bayesian_split_ids': [1, int(args.n_splits)], 'bayesian_n_stratified_80_20_splits': int(args.n_splits), 'outer_imputations_per_split': 1, 'fixed_80_20_imputations': int(args.m_imputations), 'extra_trees_split_ids': [1, int(args.tree_splits)], 'extra_trees_n_stratified_80_20_splits': int(args.tree_splits), 'extra_trees_imputations_per_split': 1, 'extra_trees_fixed_cycles': EXTRA_TREES_FIXED_CYCLES, 'extra_trees_imputer_seed_offset': EXTRA_TREES_SENSITIVITY_SEED_OFFSET, 'n_jobs': int(args.n_jobs), 'calibration_bootstrap_reps': int(args.calibration_bootstrap_reps)}}
    configuration_sha256 = hashlib.sha256(canonical_json(configuration).encode('utf-8')).hexdigest()
    expected = {'schema_version': 8, 'configuration_sha256': configuration_sha256, 'configuration': configuration}
    checkpoint_dir = outdir / 'checkpoints'
    checkpoint_dir.mkdir(parents=True, exist_ok=True)
    manifest_path = checkpoint_dir / 'run_manifest.json'
    if manifest_path.exists():
        existing = json.loads(manifest_path.read_text(encoding='utf-8'))
        if canonical_json(existing) != canonical_json(expected):
            raise RuntimeError('Secondary-analysis checkpoint manifest mismatch. Do not combine different data, scripts, environments, or settings in one output folder.')
        if not args.resume:
            raise RuntimeError(f'A compatible secondary-analysis run already exists at {outdir}. Use --resume.')
    else:
        existing_items = [path for path in outdir.iterdir() if path.name != 'checkpoints']
        checkpoint_items = [path for path in checkpoint_dir.iterdir() if not path.name.endswith('.tmp')]
        if existing_items or checkpoint_items:
            raise RuntimeError(f'Output directory {outdir} contains files but no valid run manifest. Use a new output directory.')
        atomic_write_json(expected, manifest_path)
    return configuration_sha256

def write_artifact_marker(marker_path: Path, outdir: Path, files: list[Path], metadata: dict[str, object]) -> None:
    records: dict[str, dict[str, object]] = {}
    for path in files:
        if not path.exists() or not path.is_file():
            raise RuntimeError(f'Cannot mark a stage complete because this output is missing: {path}')
        relative = path.relative_to(outdir).as_posix()
        records[relative] = {'sha256': file_sha256(path), 'size_bytes': int(path.stat().st_size)}
    atomic_write_json({'schema_version': 1, 'status': 'complete', 'metadata': metadata, 'files': records}, marker_path)

def artifact_marker_valid(marker_path: Path, outdir: Path, expected_metadata: dict[str, object]) -> bool:
    if not marker_path.exists():
        return False
    try:
        marker = json.loads(marker_path.read_text(encoding='utf-8'))
        if marker.get('status') != 'complete':
            return False
        if canonical_json(marker.get('metadata')) != canonical_json(expected_metadata):
            return False
        files = marker.get('files')
        if not isinstance(files, dict) or not files:
            return False
        for relative, record in files.items():
            path = outdir / relative
            if not path.exists() or path.stat().st_size != int(record['size_bytes']):
                return False
            if file_sha256(path) != record['sha256']:
                return False
        return True
    except Exception:
        return False

def validate_cox_metric_record(metrics: dict[str, float], context: str) -> None:
    missing = [column for column in COX_METRIC_COLUMNS if column not in metrics]
    if missing:
        raise RuntimeError(f'{context} is missing Cox metrics: {missing}.')
    values = np.asarray([float(metrics[column]) for column in COX_METRIC_COLUMNS])
    if not np.isfinite(values).all():
        failed = [column for column, value in zip(COX_METRIC_COLUMNS, values) if not np.isfinite(value)]
        raise RuntimeError(f'{context} produced non-finite Cox metrics: {failed}.')

def _validate_serialized_boolean(series: pd.Series, label: str) -> None:
    if pd.api.types.is_bool_dtype(series):
        if series.isna().any():
            raise RuntimeError(f'{label} contains a missing value.')
        return
    normalized = series.fillna('').astype(str).str.strip().str.lower()
    if not normalized.isin({'true', 'false', '1', '0'}).all():
        raise RuntimeError(f'{label} contains a value that is not a serialized boolean.')

def _canonical_id_set(series: pd.Series, label: str) -> set[str]:
    if series.isna().any():
        raise RuntimeError(f'{label} contains a missing participant ID.')
    numeric = pd.to_numeric(series, errors='coerce')
    if numeric.notna().all():
        values = numeric.to_numpy(dtype=float)
        if not np.isfinite(values).all():
            raise RuntimeError(f'{label} contains a non-finite participant ID.')
        return {format(value, '.17g') for value in values}
    normalized = series.astype(str).str.strip()
    if normalized.eq('').any():
        raise RuntimeError(f'{label} contains an empty participant ID.')
    return set(normalized)

def _canonical_id(value: object) -> str:
    numeric = pd.to_numeric(pd.Series([value]), errors='coerce').iloc[0]
    if pd.notna(numeric):
        numeric_float = float(numeric)
        if not np.isfinite(numeric_float):
            raise RuntimeError('Participant IDs must be finite.')
        return format(numeric_float, '.17g')
    normalized = str(value).strip()
    if not normalized:
        raise RuntimeError('Participant IDs must not be empty.')
    return normalized

def _strict_integer_series(series: pd.Series, label: str) -> pd.Series:
    numeric = pd.to_numeric(series, errors='raise')
    values = numeric.to_numpy(dtype=float)
    if not np.isfinite(values).all() or not np.equal(values, np.floor(values)).all():
        raise RuntimeError(f'{label} must contain only finite integers.')
    return numeric.astype(int)

def validate_main_split_map(split_map: pd.DataFrame, df: pd.DataFrame, expected_n_splits: int, expected_seed: int) -> None:
    required_columns = {'split_id', ID, 'role', 'split_seed', 'test_order'}
    if set(split_map.columns) != required_columns:
        raise RuntimeError('The main split map must contain exactly these columns: ' + ', '.join(sorted(required_columns)) + '.')
    expected_rows = expected_n_splits * EXPECTED_TEST_ROWS_PER_SPLIT
    if len(split_map) != expected_rows:
        raise RuntimeError(f'The main split map has {len(split_map)} rows; expected {expected_rows}.')
    split_ids = _strict_integer_series(split_map['split_id'], 'split_id')
    split_seeds = _strict_integer_series(split_map['split_seed'], 'split_seed')
    test_orders = _strict_integer_series(split_map['test_order'], 'test_order')
    if set(split_ids) != set(range(1, expected_n_splits + 1)):
        raise RuntimeError('The main split map does not contain the exact split_id grid.')
    if set(split_seeds) != {int(expected_seed)}:
        raise RuntimeError('The main split map split_seed does not match the master StratifiedShuffleSplit seed.')
    if set(split_map['role'].astype(str)) != {'test'}:
        raise RuntimeError("The main split map contains a role other than 'test'.")
    if split_map.duplicated(['split_id', ID]).any():
        raise RuntimeError('A participant appears twice in one main test split.')
    eligible_ids = _canonical_id_set(df[ID], 'eligible secondary data')
    mapped_ids = _canonical_id_set(split_map[ID], 'main split map')
    if not mapped_ids.issubset(eligible_ids):
        raise RuntimeError('The main split map contains an ID outside the eligible cohort.')
    validated = split_map.assign(_split_id=split_ids, _test_order=test_orders)
    for split_id, rows in validated.groupby('_split_id', sort=True):
        if len(rows) != EXPECTED_TEST_ROWS_PER_SPLIT:
            raise RuntimeError(f'Main split {split_id} has {len(rows)} test rows; expected {EXPECTED_TEST_ROWS_PER_SPLIT}.')
        if set(rows['_test_order']) != set(range(1, EXPECTED_TEST_ROWS_PER_SPLIT + 1)):
            raise RuntimeError(f'Main split {split_id} has an invalid test_order grid.')
        test_keys = {_canonical_id(value) for value in rows[ID]}
        event_count = int(df.loc[df[ID].map(_canonical_id).isin(test_keys), EVENT].sum())
        if event_count != 24:
            raise RuntimeError(f'Main split {split_id} has {event_count} ASCVD events; expected 24.')

def load_and_validate_main_analysis_dependencies(main_output_dir: Path, df: pd.DataFrame, expected_n_splits: int, expected_seed: int, require_fixed_split_evidence: bool=True) -> tuple[pd.DataFrame, dict[str, Path]]:
    split_map_path = main_output_dir / MAIN_SPLIT_MAP_RELATIVE
    performance_path = main_output_dir / MAIN_PERFORMANCE_RELATIVE
    tuning_path = main_output_dir / MAIN_TUNING_RELATIVE
    missingness_path = main_output_dir / MAIN_MISSINGNESS_RELATIVE
    if not split_map_path.is_file():
        raise FileNotFoundError(f'Main split map is missing: {split_map_path}')
    if not performance_path.is_file():
        raise FileNotFoundError(f'Main repeated-split performance is missing: {performance_path}')
    if not tuning_path.is_file():
        raise FileNotFoundError(f'Main repeated-split tuning is missing: {tuning_path}')
    if not missingness_path.is_file():
        raise FileNotFoundError(f'Main supplementary missingness table is missing: {missingness_path}')
    split_map = pd.read_csv(split_map_path, float_precision='round_trip')
    validate_main_split_map(split_map, df, expected_n_splits, expected_seed)
    build_verified_split_indices(df, split_map, list(range(1, expected_n_splits + 1)), seed=expected_seed)
    performance = pd.read_csv(performance_path, float_precision='round_trip')
    required_performance = {'split_id', 'n_train', 'n_test', 'train_events', 'test_events', 'model', 'harrell_c', 'uno_c_tau', 'tau_months', 'brier_score_at_horizon', 'brier_horizon_months', 'integrated_brier_score'}
    missing_performance = sorted(required_performance.difference(performance.columns))
    if missing_performance:
        raise RuntimeError(f'Main repeated-split performance lacks columns: {missing_performance}.')
    perf_split = _strict_integer_series(performance['split_id'], 'performance split_id')
    perf_n_train = _strict_integer_series(performance['n_train'], 'performance n_train')
    perf_n_test = _strict_integer_series(performance['n_test'], 'performance n_test')
    perf_train_events = _strict_integer_series(performance['train_events'], 'performance train_events')
    perf_test_events = _strict_integer_series(performance['test_events'], 'performance test_events')
    observed_keys = set(zip(perf_split, performance['model'].astype(str)))
    expected_keys = {(split_id, model) for split_id in range(1, expected_n_splits + 1) for model in EXPECTED_MAIN_MODELS}
    if len(performance) != len(expected_keys) or len(observed_keys) != len(performance) or observed_keys != expected_keys or (not perf_n_train.eq(EXPECTED_ELIGIBLE_ROWS - EXPECTED_TEST_ROWS_PER_SPLIT).all()) or (not perf_n_test.eq(EXPECTED_TEST_ROWS_PER_SPLIT).all()) or (not perf_train_events.eq(EXPECTED_ASCVD_EVENTS - 24).all()) or (not perf_test_events.eq(24).all()):
        raise RuntimeError('Main repeated-split performance has an invalid key grid.')
    common_metrics = performance[['harrell_c', 'uno_c_tau', 'tau_months']].apply(pd.to_numeric, errors='raise')
    if not np.isfinite(common_metrics.to_numpy(dtype=float)).all() or not common_metrics['harrell_c'].between(0.0, 1.0).all() or (not common_metrics['uno_c_tau'].between(0.0, 1.0).all()) or (not common_metrics['tau_months'].eq(120.0).all()):
        raise RuntimeError('Main repeated-split discrimination metrics are invalid.')
    absolute_columns = ['brier_score_at_horizon', 'brier_horizon_months', 'integrated_brier_score']
    absolute_metrics = performance[absolute_columns].apply(pd.to_numeric, errors='coerce')
    svm = performance['model'].astype(str).eq('SVM')
    if not absolute_metrics.loc[~svm].notna().all().all() or not np.isfinite(absolute_metrics.loc[~svm].to_numpy(dtype=float)).all() or (not absolute_metrics.loc[~svm, 'brier_score_at_horizon'].between(0.0, 1.0).all()) or (not absolute_metrics.loc[~svm, 'integrated_brier_score'].between(0.0, 1.0).all()) or (not absolute_metrics.loc[~svm, 'brier_horizon_months'].eq(120.0).all()) or (not absolute_metrics.loc[svm].isna().all().all()):
        raise RuntimeError('Main repeated-split absolute-risk metrics must be finite for seven models and missing only for Survival SVM.')
    tuning = pd.read_csv(tuning_path, float_precision='round_trip')
    required_tuning = {'split_id', 'model', 'split_base_seed', 'outer_imputation_seed', 'model_seed', 'tuning_method', 'best_params_json', 'selection_score', 'selection_metric', 'selection_direction', 'model_selection_performed', 'validation_design', 'selection_window_start_months', 'selection_window_end_months', 'selection_grid_step_months', 'selection_grid_points', 'n_prespecified_configurations'}
    missing_tuning = sorted(required_tuning.difference(tuning.columns))
    if missing_tuning:
        raise RuntimeError(f'Main repeated-split tuning lacks columns: {missing_tuning}.')
    tuning_split = _strict_integer_series(tuning['split_id'], 'tuning split_id')
    tuning_keys = set(zip(tuning_split, tuning['model'].astype(str)))
    tuning_base = _strict_integer_series(tuning['split_base_seed'], 'tuning split_base_seed')
    tuning_outer_imputer = _strict_integer_series(tuning['outer_imputation_seed'], 'tuning outer_imputation_seed')
    tuning_model_seed = _strict_integer_series(tuning['model_seed'], 'tuning model_seed')
    expected_base = int(expected_seed) + 100000 * tuning_split
    if len(tuning) != len(expected_keys) or len(tuning_keys) != len(tuning) or tuning_keys != expected_keys or (not tuning_base.eq(expected_base).all()) or (not tuning_outer_imputer.eq(expected_base).all()) or (not tuning_model_seed.eq(expected_base + 1).all()):
        raise RuntimeError('Main repeated-split tuning seeds or key grid are invalid.')
    for _, tuning_row in tuning.iterrows():
        validate_main_tuning_selection_record(tuning_row, str(tuning_row['model']))
    missingness = pd.read_csv(missingness_path, float_precision='round_trip')
    expected_missingness = {'fhxcvd': 259, 'sbp': 1, 'dbp': 1, 'bmi': 4, 'trig': 35, 'hba1c': 56, 'egfr': 11}
    if set(missingness.columns) != {'Variable', 'n_missing', 'pct_missing', 'dtype'} or len(missingness) != len(expected_missingness) or missingness['Variable'].astype(str).duplicated().any() or (set(missingness['Variable'].astype(str)) != set(expected_missingness)):
        raise RuntimeError('Supplementary Table S1 has an invalid row or column grid.')
    for _, row in missingness.iterrows():
        variable = str(row['Variable'])
        n_missing = int(pd.to_numeric(row['n_missing'], errors='raise'))
        pct_missing = float(pd.to_numeric(row['pct_missing'], errors='raise'))
        if n_missing != expected_missingness[variable] or not np.isclose(pct_missing, 100.0 * n_missing / EXPECTED_ELIGIBLE_ROWS, rtol=1e-10, atol=1e-12) or (not str(row['dtype']).strip()):
            raise RuntimeError(f'Supplementary Table S1 is invalid for {variable}.')
    dependencies = {'split_map': split_map_path, 'performance': performance_path, 'tuning': tuning_path, 'missingness': missingness_path}
    if require_fixed_split_evidence:
        pooled_path = main_output_dir / MAIN_POOLED_COX_RELATIVE
        ph_path = main_output_dir / MAIN_PH_RELATIVE
        if not pooled_path.is_file() or not ph_path.is_file():
            raise FileNotFoundError('The completed main analysis must include the fixed-split pooled Cox and proportional-hazards diagnostic tables.')
        pooled = pd.read_csv(pooled_path)
        ph = pd.read_csv(ph_path)
        if len(pooled) != len(PREDICTORS) or 'variable' not in pooled.columns or set(pooled['variable'].astype(str)) != set(PREDICTORS):
            raise RuntimeError('The main pooled Cox table is incomplete.')
        required_ph = {'imputation', 'variable', 'ph_test_statistic', 'ph_test_p'}
        if not required_ph.issubset(ph.columns):
            raise RuntimeError('The main proportional-hazards table is incomplete.')
        ph_imp = pd.to_numeric(ph['imputation'], errors='raise').astype(int)
        ph_keys = set(zip(ph_imp, ph['variable'].astype(str)))
        expected_ph_keys = {(imputation, variable) for imputation in range(1, 11) for variable in PREDICTORS}
        ph_values = ph[['ph_test_statistic', 'ph_test_p']].apply(pd.to_numeric, errors='raise')
        if len(ph) != len(expected_ph_keys) or ph_keys != expected_ph_keys or (not np.isfinite(ph_values.to_numpy(dtype=float)).all()) or (not ph_values['ph_test_p'].between(0.0, 1.0).all()):
            raise RuntimeError('The main proportional-hazards key grid is invalid.')
        dependencies['pooled_cox'] = pooled_path
        dependencies['ph_diagnostics'] = ph_path
    return (split_map, dependencies)

def build_verified_split_indices(df: pd.DataFrame, split_map: pd.DataFrame, split_ids: list[int], seed: int) -> dict[int, tuple[np.ndarray, np.ndarray]]:
    if not split_ids or sorted(set(split_ids)) != split_ids or split_ids[0] < 1:
        raise RuntimeError('Requested split IDs must be sorted, unique positive integers.')
    requested = set(split_ids)
    splitter = StratifiedShuffleSplit(n_splits=split_ids[-1], test_size=0.2, random_state=seed)
    validated_map = split_map.assign(_split_id=_strict_integer_series(split_map['split_id'], 'split map split_id'), _test_order=_strict_integer_series(split_map['test_order'], 'split map test_order'))
    verified: dict[int, tuple[np.ndarray, np.ndarray]] = {}
    for split_id, (train_index, test_index) in enumerate(splitter.split(df, make_strata(df)), start=1):
        if split_id not in requested:
            continue
        rows = validated_map.loc[validated_map['_split_id'].eq(split_id)].copy()
        rows = rows.sort_values('_test_order')
        saved_test_ids = [_canonical_id(value) for value in rows[ID]]
        generated_test_ids = [_canonical_id(value) for value in df.iloc[test_index][ID]]
        if saved_test_ids != generated_test_ids:
            raise RuntimeError(f'Saved split {split_id} does not match the regenerated StratifiedShuffleSplit test order.')
        if len(test_index) != EXPECTED_TEST_ROWS_PER_SPLIT or len(train_index) != len(df) - EXPECTED_TEST_ROWS_PER_SPLIT or np.intersect1d(train_index, test_index).size:
            raise RuntimeError(f'Generated split {split_id} does not partition the cohort.')
        verified[split_id] = (np.asarray(train_index, dtype=int), np.asarray(test_index, dtype=int))
    if set(verified) != requested:
        raise RuntimeError('Could not regenerate every requested main split.')
    return verified

def split_indices_from_map(df: pd.DataFrame, split_map: pd.DataFrame, split_id: int, seed: int) -> tuple[np.ndarray, np.ndarray]:
    return build_verified_split_indices(df, split_map, [split_id], seed)[split_id]

def validate_split_checkpoint_frames(split_results: pd.DataFrame, split_diagnostics: pd.DataFrame, split_predictions: pd.DataFrame, split_id: int, method: str, expected_ids: pd.Series, expected_imputer_seed: int, expected_split_base_seed: int, expected_model_seed: int) -> None:
    context = f'{method} split {split_id}'
    required_result_columns = {'split_id', 'n_test', 'model', 'imputer', 'split_base_seed', 'model_seed', *COX_METRIC_COLUMNS}
    missing_result_columns = sorted(required_result_columns.difference(split_results.columns))
    if missing_result_columns:
        raise RuntimeError(f'Validated {context} checkpoint is missing result columns: {missing_result_columns}.')
    result_split = pd.to_numeric(split_results['split_id'], errors='raise')
    if len(split_results) != 2 or set(split_results['model'].astype(str)) != {'CoxPH', 'Spline Cox'} or split_results['model'].astype(str).duplicated().any() or (set(result_split.astype(int)) != {split_id}) or (set(split_results['imputer'].astype(str)) != {method}) or (not (pd.to_numeric(split_results['n_test'], errors='raise').astype(int) == len(expected_ids)).all()) or (not pd.to_numeric(split_results['split_base_seed'], errors='raise').astype(int).eq(expected_split_base_seed).all()) or (not pd.to_numeric(split_results['model_seed'], errors='raise').astype(int).eq(expected_model_seed).all()):
        raise RuntimeError(f'Validated {context} checkpoint has an unexpected result key grid.')
    result_values = split_results[COX_METRIC_COLUMNS].apply(pd.to_numeric, errors='raise')
    if not np.isfinite(result_values.to_numpy(dtype=float)).all():
        raise RuntimeError(f'Validated {context} checkpoint contains non-finite results.')
    required_diagnostic_columns = {'analysis', 'split_id', 'split_base_seed', 'model_seed', 'imputer', 'imputation', 'imputer_seed', 'imputer_n_iter', 'imputer_max_iter', 'imputer_iteration_limit_reached', 'imputer_convergence_warning_count', 'imputer_convergence_warning_messages', 'imputer_policy', 'imputer_fixed_cycle_policy', 'imputer_fixed_cycle_completed', 'imputer_expected_fixed_cycle_warning_count', 'imputer_unexpected_convergence_warning_count', 'train_missing_cells_before', 'test_missing_cells_before', 'train_missing_cells_after', 'test_missing_cells_after', 'observed_cells_preserved', 'all_values_finite', 'plausibility_passed'}
    missing_diagnostic_columns = sorted(required_diagnostic_columns.difference(split_diagnostics.columns))
    if missing_diagnostic_columns:
        raise RuntimeError(f'Validated {context} checkpoint is missing diagnostic columns: {missing_diagnostic_columns}.')
    diagnostic_split = pd.to_numeric(split_diagnostics['split_id'], errors='raise').astype(int)
    diagnostic_imputation = pd.to_numeric(split_diagnostics['imputation'], errors='raise').astype(int)
    if len(split_diagnostics) != 1 or set(diagnostic_split) != {split_id} or set(diagnostic_imputation) != {1} or (set(split_diagnostics['analysis'].astype(str)) != {'repeated_stratified_80_20_split_validation'}) or (set(split_diagnostics['imputer'].astype(str)) != {method}):
        raise RuntimeError(f'Validated {context} checkpoint has an unexpected diagnostic key grid.')
    diagnostic_numeric_columns = ['split_id', 'split_base_seed', 'model_seed', 'imputation', 'imputer_seed', 'imputer_n_iter', 'imputer_max_iter', 'imputer_convergence_warning_count', 'imputer_expected_fixed_cycle_warning_count', 'imputer_unexpected_convergence_warning_count', 'train_missing_cells_before', 'test_missing_cells_before', 'train_missing_cells_after', 'test_missing_cells_after']
    diagnostic_values = split_diagnostics[diagnostic_numeric_columns].apply(pd.to_numeric, errors='raise')
    diagnostic_array = diagnostic_values.to_numpy(dtype=float)
    if not np.isfinite(diagnostic_array).all():
        raise RuntimeError(f'Validated {context} checkpoint contains non-finite diagnostics.')
    if not np.equal(diagnostic_array, np.floor(diagnostic_array)).all():
        raise RuntimeError(f'Validated {context} checkpoint contains non-integer diagnostics.')
    if (diagnostic_values['imputer_n_iter'] < 0).any() or (diagnostic_values['imputer_max_iter'] < 1).any() or (diagnostic_values['imputer_n_iter'] > diagnostic_values['imputer_max_iter']).any() or (diagnostic_values['imputer_convergence_warning_count'] < 0).any():
        raise RuntimeError(f'Validated {context} checkpoint contains out-of-range diagnostics.')
    if not diagnostic_values['imputer_seed'].eq(expected_imputer_seed).all():
        raise RuntimeError(f'Validated {context} checkpoint has the wrong imputer seed; expected {expected_imputer_seed}.')
    _validate_serialized_boolean(split_diagnostics['imputer_iteration_limit_reached'], f'Validated {context} checkpoint iteration-limit flag')
    _validate_serialized_boolean(split_diagnostics['imputer_fixed_cycle_policy'], f'Validated {context} checkpoint fixed-cycle flag')
    _validate_serialized_boolean(split_diagnostics['imputer_fixed_cycle_completed'], f'Validated {context} checkpoint fixed-cycle completion flag')
    for boolean_column in ['observed_cells_preserved', 'all_values_finite', 'plausibility_passed']:
        _validate_serialized_boolean(split_diagnostics[boolean_column], f'Validated {context} checkpoint {boolean_column}')
        normalized = split_diagnostics[boolean_column].astype(str).str.lower().isin({'true', '1'})
        if not normalized.all():
            raise RuntimeError(f'Validated {context} checkpoint failed {boolean_column}.')
    if not diagnostic_values['train_missing_cells_after'].eq(0).all() or not diagnostic_values['test_missing_cells_after'].eq(0).all():
        raise RuntimeError(f'Validated {context} checkpoint retained missing values.')
    if not diagnostic_values['split_base_seed'].eq(expected_split_base_seed).all():
        raise RuntimeError(f'Validated {context} checkpoint has the wrong split seed.')
    if not diagnostic_values['model_seed'].eq(expected_model_seed).all():
        raise RuntimeError(f'Validated {context} checkpoint has the wrong model seed.')
    fixed_cycle = split_diagnostics['imputer_fixed_cycle_policy'].astype(str).str.lower().isin({'true', '1'})
    fixed_completed = split_diagnostics['imputer_fixed_cycle_completed'].astype(str).str.lower().isin({'true', '1'})
    if method == 'extra_trees':
        if not fixed_cycle.all() or not fixed_completed.all() or (not diagnostic_values['imputer_max_iter'].eq(EXTRA_TREES_FIXED_CYCLES).all()) or (not diagnostic_values['imputer_n_iter'].eq(EXTRA_TREES_FIXED_CYCLES).all()) or (not diagnostic_values['imputer_unexpected_convergence_warning_count'].eq(0).all()) or (not diagnostic_values['imputer_expected_fixed_cycle_warning_count'].eq(1).all()) or (not diagnostic_values['imputer_convergence_warning_count'].eq(1).all()) or (not split_diagnostics['imputer_policy'].astype(str).eq('prespecified_fixed_10_cycles_tol_0_no_convergence_claim').all()):
            raise RuntimeError(f'Validated {context} checkpoint did not complete the fixed {EXTRA_TREES_FIXED_CYCLES}-cycle policy.')
    elif not fixed_cycle.all() or not fixed_completed.all() or (not diagnostic_values['imputer_max_iter'].eq(20).all()) or (not diagnostic_values['imputer_n_iter'].eq(20).all()) or (not diagnostic_values['imputer_convergence_warning_count'].eq(0).all()) or (not split_diagnostics['imputer_policy'].astype(str).eq('prespecified_20_cycles_posterior_sampling_no_early_stopping').all()):
        raise RuntimeError(f'Validated {context} did not complete its posterior-sampling 20-cycle policy.')
    required_prediction_columns = {'split_id', 'imputer', ID, TIME, EVENT, 'coxph_risk_120', 'spline_cox_risk_120'}
    missing_prediction_columns = sorted(required_prediction_columns.difference(split_predictions.columns))
    if missing_prediction_columns:
        raise RuntimeError(f'Validated {context} checkpoint is missing prediction columns: {missing_prediction_columns}.')
    if set(pd.to_numeric(split_predictions['split_id'], errors='raise').astype(int)) != {split_id} or set(split_predictions['imputer'].astype(str)) != {method} or split_predictions[ID].duplicated().any() or (_canonical_id_set(split_predictions[ID], context) != _canonical_id_set(expected_ids, f'expected {context}')):
        raise RuntimeError(f'Validated {context} checkpoint does not contain the exact held-out IDs.')
    prediction_values = split_predictions[[TIME, EVENT, 'coxph_risk_120', 'spline_cox_risk_120']].apply(pd.to_numeric, errors='raise')
    if not np.isfinite(prediction_values.to_numpy(dtype=float)).all():
        raise RuntimeError(f'Validated {context} checkpoint contains non-finite prediction data.')
    for risk_column in ['coxph_risk_120', 'spline_cox_risk_120']:
        if not prediction_values[risk_column].between(0.0, 1.0).all():
            raise RuntimeError(f'Validated {context} checkpoint has an out-of-range {risk_column}.')

def calibration_slope_from_probs(data: pd.DataFrame, predicted_event_prob: np.ndarray, horizon: float) -> float:
    from lifelines import CoxPHFitter
    probabilities = np.clip(np.asarray(predicted_event_prob, dtype=float), 1e-06, 1.0 - 1e-06)
    if len(probabilities) != len(data) or not np.isfinite(probabilities).all() or float(np.std(probabilities)) <= 0.0:
        raise RuntimeError('Calibration-slope probabilities are invalid or constant.')
    linear_predictor = np.log(-np.log1p(-probabilities))
    original_time = data[TIME].to_numpy(dtype=float)
    original_event = data[EVENT].to_numpy(dtype=int)
    calibration_frame = pd.DataFrame({TIME: np.minimum(original_time, float(horizon)), EVENT: ((original_event == 1) & (original_time <= float(horizon))).astype(int), 'lp': linear_predictor})
    if int(calibration_frame[EVENT].sum()) < 3:
        raise RuntimeError('Calibration slope has fewer than three events by the horizon.')
    model = CoxPHFitter()
    model.fit(calibration_frame, duration_col=TIME, event_col=EVENT)
    slope = float(model.params_['lp'])
    if not np.isfinite(slope):
        raise RuntimeError('Calibration slope is non-finite.')
    return slope

def participant_averaged_split_calibration(predictions: pd.DataFrame, df: pd.DataFrame, method: str, n_groups: int=OOF_CALIBRATION_GROUPS, require_full_coverage: bool=True) -> tuple[list[dict[str, object]], list[dict[str, object]], dict[str, object]]:
    working = predictions.copy()
    working['_id_key'] = working[ID].map(_canonical_id)
    invariant = working.groupby('_id_key')[[TIME, EVENT]].nunique(dropna=False)
    if (invariant > 1).any().any():
        raise RuntimeError(f'{method} predictions changed outcomes across splits.')
    averaged = working.groupby('_id_key', as_index=False).agg(coxph_risk_120=('coxph_risk_120', 'mean'), spline_cox_risk_120=('spline_cox_risk_120', 'mean'), held_out_count=('split_id', 'size'))
    cohort = df[[ID, TIME, EVENT]].copy()
    cohort['_id_key'] = cohort[ID].map(_canonical_id)
    averaged = cohort.merge(averaged, on='_id_key', how='left', validate='one_to_one')
    missing = int(averaged['held_out_count'].isna().sum())
    if missing and require_full_coverage:
        raise RuntimeError(f'{method} repeated splits left {missing} participants without a held-out prediction.')
    if missing:
        averaged = averaged.loc[averaged['held_out_count'].notna()].copy()
    overall_rows: list[dict[str, object]] = []
    grouped_rows: list[dict[str, object]] = []
    observed_all = km_risk(averaged[TIME].to_numpy(dtype=float), averaged[EVENT].to_numpy(dtype=int), 120.0)
    for model, risk_column in [('CoxPH', 'coxph_risk_120'), ('Spline Cox', 'spline_cox_risk_120')]:
        risk = averaged[risk_column].to_numpy(dtype=float)
        mean_risk = float(np.mean(risk))
        if not np.isfinite(risk).all() or mean_risk <= 0.0:
            raise RuntimeError(f'{method} {model} has invalid participant-averaged held-out risks.')
        calibration_slope = calibration_slope_from_probs(averaged, risk, horizon=120.0)
        overall_rows.append({'imputer': method, 'model': model, 'n_participants': int(len(averaged)), 'calibration_slope': float(calibration_slope), 'calibration_oe_ratio': float(observed_all / mean_risk), 'observed_km_risk': float(observed_all), 'mean_predicted_risk': mean_risk, 'calibration_horizon_months': 120.0, 'calibration_unit': "one participant-level risk averaged across that participant's held-out split predictions", 'calibration_estimand': 'cross-fitted repeated-split ensemble', 'represents_one_final_deployable_model': False, 'uncertainty_scope': 'participant-bootstrap bands, when reported, are conditional on the fixed participant-averaged held-out predictions and exclude model-development and split-stream-selection uncertainty', 'model_development_uncertainty_included': False, 'split_stream_selection_uncertainty_included': False})
        ranked = pd.Series(risk).rank(method='first')
        group = pd.qcut(ranked, q=n_groups, labels=False) + 1
        if int(pd.Series(group).nunique()) != n_groups:
            raise RuntimeError(f'{method} {model} did not produce {n_groups} calibration groups.')
        for group_number in range(1, n_groups + 1):
            in_group = np.asarray(group == group_number)
            group_time = averaged.loc[in_group, TIME].to_numpy(dtype=float)
            group_event = averaged.loc[in_group, EVENT].to_numpy(dtype=int)
            group_risk = risk[in_group]
            grouped_rows.append({'imputer': method, 'model': model, 'risk_group': int(group_number), 'risk_group_order': '1=lowest predicted risk; 5=highest predicted risk', 'n': int(in_group.sum()), 'ascvd_events_total_follow_up': int(group_event.sum()), 'mean_predicted_120': float(np.mean(group_risk)), 'observed_km_120': float(km_risk(group_time, group_event, 120.0)), 'calibration_estimand': 'cross-fitted repeated-split ensemble', 'represents_one_final_deployable_model': False, 'uncertainty_scope': 'participant-bootstrap bands, when reported, are conditional on the fixed participant-averaged held-out predictions and exclude model-development and split-stream-selection uncertainty', 'model_development_uncertainty_included': False, 'split_stream_selection_uncertainty_included': False})
    coverage = averaged['held_out_count'].to_numpy(dtype=int)
    coverage_record = {'method': method, 'eligible_n': int(len(df)), 'prediction_rows': int(len(predictions)), 'unique_participants': int(len(averaged)), 'participants_without_held_out_prediction': int(missing), 'held_out_count_min': int(coverage.min()), 'held_out_count_median': float(np.median(coverage)), 'held_out_count_max': int(coverage.max()), 'one_averaged_prediction_per_participant_for_calibration': True, 'repeated_prediction_rows_treated_as_independent': False, 'calibration_estimand': 'cross-fitted repeated-split ensemble', 'represents_one_final_deployable_model': False, 'uncertainty_scope': 'participant-bootstrap bands, when reported, are conditional on the fixed participant-averaged held-out predictions and exclude model-development and split-stream-selection uncertainty', 'model_development_uncertainty_included': False, 'split_stream_selection_uncertainty_included': False}
    return (overall_rows, grouped_rows, coverage_record)

def repeated_cox_comparison(df: pd.DataFrame, outdir: Path, split_map: pd.DataFrame, split_ids: list[int], seed: int, method: str, n_jobs: int, resume: bool, configuration_sha256: str, deadline_monotonic: float | None=None, functional_smoke: bool=False) -> pd.DataFrame:
    if df[ID].isna().any() or df[ID].duplicated().any():
        raise RuntimeError('Repeated-split validation requires one non-missing row per participant ID.')
    records: list[dict[str, object]] = []
    diagnostics: list[dict[str, object]] = []
    prediction_records: list[dict[str, object]] = []
    expected_split_ids = list(range(1, EXPECTED_PRIMARY_SPLITS + 1))
    if not functional_smoke and method == 'bayesian' and (split_ids != expected_split_ids):
        raise RuntimeError('Bayesian spline validation requires exact split IDs 1-200.')
    if not functional_smoke and method == 'extra_trees' and (split_ids != list(range(1, EXPECTED_TREE_SPLITS + 1))):
        raise RuntimeError('ExtraTrees spline sensitivity requires exact split IDs 1-200.')
    if method not in {'bayesian', 'extra_trees'}:
        raise ValueError(f'Unknown repeated-split imputer: {method}')
    namespace = 'repeated_200_splits_bayesian_v1' if method == 'bayesian' else 'repeated_200_splits_extra_trees_all_200_v1'
    checkpoint_dir = outdir / 'checkpoints' / namespace
    checkpoint_dir.mkdir(parents=True, exist_ok=True)
    verified_indices = build_verified_split_indices(df, split_map, split_ids, seed)
    for split_id in split_ids:
        tr_idx, te_idx = verified_indices[split_id]
        split_dir = checkpoint_dir / f'split_{split_id:03d}'
        result_path = split_dir / 'results.csv'
        diagnostic_path = split_dir / 'imputation_diagnostics.csv'
        prediction_path = split_dir / 'held_out_predictions.csv'
        marker_path = split_dir / 'complete.json'
        split_base_seed = seed + 100000 * split_id
        model_seed = split_base_seed + 1
        imputer_seed = split_base_seed + EXTRA_TREES_SENSITIVITY_SEED_OFFSET if method == 'extra_trees' else split_base_seed
        expected_ids = df.iloc[te_idx][ID]
        test_id_sha256 = hashlib.sha256(canonical_json([_canonical_id(value) for value in expected_ids]).encode('utf-8')).hexdigest()
        marker_metadata = {'configuration_sha256': configuration_sha256, 'method': method, 'checkpoint_namespace': namespace, 'split_id': int(split_id), 'imputations': 1, 'split_base_seed': int(split_base_seed), 'model_seed': int(model_seed), 'imputer_seed': int(imputer_seed), 'test_id_sha256': test_id_sha256}
        resumed = False
        if resume and artifact_marker_valid(marker_path, outdir, marker_metadata):
            split_results = pd.read_csv(result_path)
            split_diagnostics = pd.read_csv(diagnostic_path)
            split_predictions = pd.read_csv(prediction_path)
            try:
                validate_split_checkpoint_frames(split_results, split_diagnostics, split_predictions, split_id=split_id, method=method, expected_ids=expected_ids, expected_imputer_seed=imputer_seed, expected_split_base_seed=split_base_seed, expected_model_seed=model_seed)
            except Exception as exc:
                print(f'{method}: split {split_id}/{split_ids[-1]} checkpoint failed validation and will be recomputed: {exc}', flush=True)
            else:
                resumed = True
        if not resumed:
            train = df.iloc[tr_idx].reset_index(drop=True)
            test = df.iloc[te_idx].reset_index(drop=True)
            split_diagnostic_records: list[dict[str, object]] = []
            completed = impute_split(train, test, m=1, seed=imputer_seed, method=method, n_jobs=n_jobs, diagnostics=split_diagnostic_records, diagnostic_context={'analysis': 'repeated_stratified_80_20_split_validation', 'split_id': split_id, 'split_base_seed': split_base_seed, 'model_seed': model_seed})
            imputed_train, imputed_test = completed[0]
            split_result_records: list[dict[str, object]] = []
            risks: dict[str, np.ndarray] = {}
            for model_name, spline, risk_column in [('CoxPH', False, 'coxph_risk_120'), ('Spline Cox', True, 'spline_cox_risk_120')]:
                metrics, risk, _, _ = evaluate_cox(imputed_train, imputed_test, spline=spline)
                validate_cox_metric_record(metrics, f'{method} split {split_id}, {model_name}')
                split_result_records.append({'split_id': int(split_id), 'n_test': int(len(imputed_test)), 'model': model_name, 'imputer': method, 'split_base_seed': int(split_base_seed), 'model_seed': int(model_seed), **metrics})
                risks[risk_column] = np.asarray(risk, dtype=float)
            split_results = pd.DataFrame(split_result_records)
            split_diagnostics = pd.DataFrame(split_diagnostic_records)
            split_predictions = test[[ID, TIME, EVENT]].copy().reset_index(drop=True)
            split_predictions.insert(0, 'imputer', method)
            split_predictions.insert(0, 'split_id', int(split_id))
            split_predictions['coxph_risk_120'] = risks['coxph_risk_120']
            split_predictions['spline_cox_risk_120'] = risks['spline_cox_risk_120']
            validate_split_checkpoint_frames(split_results, split_diagnostics, split_predictions, split_id=split_id, method=method, expected_ids=expected_ids, expected_imputer_seed=imputer_seed, expected_split_base_seed=split_base_seed, expected_model_seed=model_seed)
            atomic_write_csv(split_results, result_path)
            atomic_write_csv(split_diagnostics, diagnostic_path)
            atomic_write_csv(split_predictions, prediction_path)
            write_artifact_marker(marker_path, outdir, [result_path, diagnostic_path, prediction_path], marker_metadata)
            print(f'{method}: completed split {split_id}/{split_ids[-1]}', flush=True)
        else:
            print(f'{method}: resumed split {split_id}/{split_ids[-1]}', flush=True)
        records.extend(split_results.to_dict('records'))
        diagnostics.extend(split_diagnostics.to_dict('records'))
        prediction_records.extend(split_predictions.to_dict('records'))
        if not resumed and deadline_monotonic is not None and (time.monotonic() >= deadline_monotonic):
            raise PlannedSessionPause(f'Safe session limit reached after {method} split {split_id}.')
    long = pd.DataFrame(records)
    expected_result_rows = len(split_ids) * 2
    if len(long) != expected_result_rows:
        raise RuntimeError(f'{method} produced {len(long)} split/model rows; expected {expected_result_rows}.')
    observed_keys = set(zip(pd.to_numeric(long['split_id'], errors='raise').astype(int), long['model'].astype(str)))
    expected_keys = {(split_id, model) for split_id in split_ids for model in ['CoxPH', 'Spline Cox']}
    if observed_keys != expected_keys or len(observed_keys) != len(long):
        raise RuntimeError(f'{method} split/model result keys are incomplete or duplicated.')
    atomic_write_csv(long, outdir / f'cox_spline_by_split_{method}.csv')
    summary_rows: list[dict[str, object]] = []
    for (imputer, model_name), frame in long.groupby(['imputer', 'model'], sort=True):
        if len(frame) != len(split_ids) or frame['split_id'].nunique() != len(split_ids):
            raise RuntimeError(f'{imputer} {model_name} lacks a complete split set.')
        for metric in COX_METRIC_COLUMNS:
            values = frame[metric].to_numpy(dtype=float)
            low, high = np.percentile(values, [2.5, 97.5])
            summary_rows.append({'imputer': imputer, 'model': model_name, 'metric': metric, 'n_splits': int(len(values)), 'mean': float(np.mean(values)), 'empirical_p2_5': float(low), 'empirical_p97_5': float(high), 'range_interpretation': 'descriptive empirical split distribution; not a confidence interval'})
    summary = pd.DataFrame(summary_rows)
    if len(summary) != 2 * len(COX_METRIC_COLUMNS) or not summary['n_splits'].eq(len(split_ids)).all():
        raise RuntimeError(f'{method} split summary is incomplete.')
    atomic_write_csv(summary, outdir / f'cox_spline_summary_{method}.csv')
    diagnostic_frame = pd.DataFrame(diagnostics)
    if len(diagnostic_frame) != len(split_ids):
        raise RuntimeError(f'{method} imputation diagnostics are incomplete.')
    atomic_write_csv(diagnostic_frame, outdir / f'imputation_diagnostics_by_split_{method}.csv')
    predictions = pd.DataFrame(prediction_records)
    expected_prediction_rows = len(split_ids) * EXPECTED_TEST_ROWS_PER_SPLIT
    if len(predictions) != expected_prediction_rows:
        raise RuntimeError(f'{method} produced {len(predictions)} prediction rows; expected {expected_prediction_rows}.')
    calibration_records, calibration_group_records, coverage = participant_averaged_split_calibration(predictions, df, method, require_full_coverage=not functional_smoke)
    atomic_write_csv(pd.DataFrame(calibration_records), outdir / f'cox_spline_participant_averaged_calibration_{method}.csv')
    atomic_write_csv(pd.DataFrame(calibration_group_records), outdir / f'cox_spline_participant_averaged_calibration_groups_{method}.csv')
    atomic_write_json({'method': method, 'split_ids': split_ids, 'n_splits': int(len(split_ids)), 'test_rows_per_split': EXPECTED_TEST_ROWS_PER_SPLIT, 'participant_level_predictions_location': 'validated local checkpoints only; not a final result table', **coverage}, outdir / f'cox_spline_heldout_coverage_{method}.json')
    return long

def run_primary_secondary_stage(eligible: pd.DataFrame, outdir: Path, args: argparse.Namespace, split_map: pd.DataFrame) -> list[Path]:
    tr_idx, te_idx = split_indices_from_map(eligible, split_map, split_id=1, seed=args.seed)
    train = eligible.iloc[tr_idx].reset_index(drop=True)
    test = eligible.iloc[te_idx].reset_index(drop=True)
    split_counts = {'fixed_split_id': 1, 'fixed_split_seed': int(args.seed), 'fixed_split_is_member_of_repeated_stream': True, 'train_n': int(len(train)), 'test_n': int(len(test)), 'train_ascvd_events': int(train[EVENT].sum()), 'test_ascvd_events': int(test[EVENT].sum()), 'train_competing_deaths': int(train[COMPETING_EVENT].sum()), 'test_competing_deaths': int(test[COMPETING_EVENT].sum())}
    expected_split_counts = {'fixed_split_id': 1, 'fixed_split_seed': int(args.seed), 'fixed_split_is_member_of_repeated_stream': True, 'train_n': 992, 'test_n': 249, 'train_ascvd_events': 92, 'test_ascvd_events': 24, 'train_competing_deaths': 37, 'test_competing_deaths': 9}
    if split_counts != expected_split_counts:
        raise RuntimeError(f'The fixed 80/20 split did not match the validated split: {split_counts}; expected {expected_split_counts}.')
    split_counts_path = outdir / 'fixed_split_event_counts.json'
    atomic_write_json(split_counts, split_counts_path)
    spline_methods_path = outdir / 'spline_cox_method_metadata.json'
    atomic_write_json({'analysis_role': 'separate secondary-requested linear-versus-spline Cox sensitivity; not a ninth model in the main eight-model nested comparison', 'linear_comparator_label_in_output': 'CoxPH', 'linear_comparator_interpretation': 'dedicated linear Cox comparator fitted by the same sensitivity code as the spline arm; numerical values need not be identical to the fixed-specification scikit-survival CoxPH in the main pipeline', 'repeated_evaluation': {'bayesian_imputer': {'split_ids': '1-200', 'n_stratified_80_20_splits': int(args.n_splits), 'imputations_per_split': 1, 'iteration_policy': '20 cycles with posterior sampling; sklearn early stopping is not used when sample_posterior=True'}, 'extra_trees_imputer': {'split_ids': '1-200', 'n_stratified_80_20_splits': int(args.tree_splits), 'imputations_per_split': 1, 'split_membership': 'exactly matched to Bayesian and main-analysis split IDs 1-200', 'iteration_policy': 'fixed 10 cycles with tol=0'}, 'outer_split_method': 'the exact 200 saved test assignments from one stratified 80/20 StratifiedShuffleSplit stream in the main analysis', 'saved_split_map': MAIN_SPLIT_MAP_RELATIVE.as_posix(), 'imputation_seed_scheme': {'split_base_seed': 'seed + 100000 * split_id', 'bayesian_imputer_seed': 'split_base_seed', 'model_seed': 'split_base_seed + 1 (same under both imputers)', 'extra_trees_sensitivity_offset': EXTRA_TREES_SENSITIVITY_SEED_OFFSET, 'extra_trees_imputer_seed': '50000000 + split_base_seed; identical to the matched all-model ExtraTrees-imputation sensitivity'}, 'coverage_check': 'each split contains 249 unique held-out participants and 24 ASCVD events', 'performance_unit': 'one held-out 80/20 split estimate', 'uncertainty_interpretation': 'mean and empirical 2.5th-97.5th percentiles across overlapping split estimates are descriptive and are not confidence intervals', 'calibration': "each participant's risks are averaged only across splits in which that participant was held out; the resulting one risk per participant is used for O:E, slope and five risk groups", 'calibration_slope': 'Cox coefficient for the complementary-log-log transform of predicted 120-month risk after administratively censoring follow-up at 120 months; calculated from participant-averaged held-out predictions'}, 'outer_split_imputation': "one predictor-only imputation fitted on each split's training set and applied to held-out predictors; no outcome or follow-up variable enters the imputation matrix", 'fixed_split_illustration': {'role': 'supplementary illustration for coefficients and calibration', 'allocation': 'seeded stratified 80/20 split used for this illustration', 'imputations': int(args.m_imputations), 'relationship_to_repeated_split_validation': 'the same seed-20260320 assignment as primary repeated-analysis split 1; retained only for coefficients, diagnostics and detailed illustration'}, 'spline_variables': list(RCS_VARS), 'restricted_cubic_spline_knots': 'training-set 10th, 50th and 90th percentiles, recomputed in each completed outer-training dataset', 'design_columns': '15 linear main-effect columns plus one additional restricted-cubic spline basis column for each of the five spline variables (20 total)', 'fixed_split_training_ascvd_events': int(split_counts['train_ascvd_events']), 'small_sample_caution': 'The 20-column spline design is evaluated as an exploratory sensitivity and is not treated as a definitive replacement benchmark.', 'design_scaling': 'each fitted Cox design column standardized using training data', 'ties': 'Breslow partial likelihood', 'ridge_stabilization': {'linear_comparator': 1e-05, 'spline_cox': 1e-05, 'purpose': 'the same very small numerical-stabilization penalty is used for both comparators; it was not selected by outcome-guided tuning'}, 'hyperparameter_tuning': 'none in this dedicated sensitivity', 'interpretation': 'Use paired results only to assess whether adding the candidate spline terms changes performance within this dedicated benchmark.'}, spline_methods_path)
    competing_status_path = outdir / 'competing_risk_status.json'
    competing_status: dict[str, object] = {'status': 'running', 'competing_analysis_completed': False, 'time_source': f'master workbook / {MASTER_SHEET} / {MASTER_DEATH_TIME}', 'time_unit': 'months', 'treated_as_verified': True, 'source_definition': 'Master-sheet F/U Time was used as the time to non-ASCVD death for rows with Death other cause=1.', 'eligible_non_ascvd_deaths': EXPECTED_COMPETING_DEATHS, 'cause_specific_death_model_predictors': len(PREDICTORS), 'cause_specific_death_model_role': 'exploratory sensitivity only; 37 competing deaths in the fixed-split training set for 15 predictors can yield unstable coefficients', 'model_based_sensitivity': 'cause-specific Cox hazards combined into a 120-month cumulative-incidence prediction; observed risk estimated by Aalen-Johansen', 'competing_death_model_ridge_penalty': COMPETING_DEATH_RIDGE, 'interpretation': 'The Aalen-Johansen estimates carry the competing-risk conclusion. The individual model-based CIF and correlation are exploratory.', **split_counts}
    atomic_write_json(competing_status, competing_status_path)
    primary_imputation_diagnostics: list[dict[str, object]] = []
    primary_imputed = impute_split(train, test, m=args.m_imputations, seed=args.seed + 1000, method='bayesian', competing_col=COMPETING_EVENT, n_jobs=args.n_jobs, diagnostics=primary_imputation_diagnostics, diagnostic_context={'analysis': 'supplementary_fixed_80_20', 'split_id': 1})
    primary_diagnostics_path = outdir / 'imputation_diagnostics_fixed_split_bayesian.csv'
    atomic_write_csv(pd.DataFrame(primary_imputation_diagnostics), primary_diagnostics_path)
    vif_path = outdir / 'supplementary_vif.csv'
    atomic_write_csv(vif_summary([x[0] for x in primary_imputed]), vif_path)
    ldl_vif_path = outdir / 'ldl_component_multicollinearity_vif.csv'
    atomic_write_csv(ldl_component_multicollinearity_vif(eligible), ldl_vif_path)
    primary_rows: list[dict[str, object]] = []
    standard_risks: list[np.ndarray] = []
    spline_risks: list[np.ndarray] = []
    competing_cif_risks: list[np.ndarray] = []
    competing_rows: list[dict[str, object]] = []
    for imp_no, (tr, te) in enumerate(primary_imputed, 1):
        for model_name, spline in [('CoxPH', False), ('Spline Cox', True)]:
            metrics, risk, model, design = evaluate_cox(tr, te, spline=spline)
            if not all((np.isfinite(float(value)) for value in metrics.values())):
                raise RuntimeError(f'Fixed-split {model_name} imputation {imp_no} produced a non-finite metric.')
            primary_rows.append({'imputation': imp_no, 'model': model_name, **metrics})
            if spline:
                spline_risks.append(risk)
            else:
                standard_risks.append(risk)
                xtr = design.transform(tr)
                xte = design.transform(te)
                death_model = CoxPH(ridge=COMPETING_DEATH_RIDGE).fit(xtr, tr[TIME], tr[COMPETING_EVENT])
                cif = cause_specific_cif(model, death_model, xte, 120.0)
                competing_cif_risks.append(np.asarray(cif, dtype=float))
                status = np.where(te[EVENT] == 1, 1, np.where(te[COMPETING_EVENT] == 1, 2, 0))
                observed_aj, _ = aalen_johansen_risks(te[TIME].to_numpy(), status, 120.0)
                if not np.isfinite(cif).all() or np.any(cif < -1e-12) or np.any(cif > 1.0 + 1e-12) or (not np.isfinite(observed_aj)) or (not 0.0 <= observed_aj <= 1.0) or (float(np.mean(cif)) <= 0.0) or (float(np.mean(risk)) <= 0.0):
                    raise RuntimeError(f'Competing-risk validation failed in fixed-split imputation {imp_no}.')
                competing_rows.append({'imputation': imp_no, 'competing_death_model_ridge_penalty': COMPETING_DEATH_RIDGE, 'standard_cox_mean_predicted_risk': float(np.mean(risk)), 'cause_specific_cox_mean_predicted_cif': float(np.mean(cif)), 'prediction_difference_standard_minus_competing': float(np.mean(risk) - np.mean(cif)), 'observed_test_aj_cif': observed_aj, 'oe_standard': float(observed_aj / np.mean(risk)), 'oe_competing_risk': float(observed_aj / np.mean(cif)), 'prediction_correlation': float(np.corrcoef(risk, cif)[0, 1])})
    primary_df = pd.DataFrame(primary_rows)
    if len(primary_df) != 2 * args.m_imputations:
        raise RuntimeError('Fixed-split Cox/spline analysis did not produce the expected rows.')
    primary_by_imp_path = outdir / 'fixed_split_cox_spline_by_imputation.csv'
    primary_summary_path = outdir / 'fixed_split_cox_spline_summary.csv'
    atomic_write_csv(primary_df, primary_by_imp_path)
    mean_standard = np.mean(np.vstack(standard_risks), axis=0)
    mean_spline = np.mean(np.vstack(spline_risks), axis=0)
    observed_primary = km_risk(test[TIME].to_numpy(dtype=float), test[EVENT].to_numpy(dtype=int), 120.0)
    primary_summary = primary_df.drop(columns=['imputation']).groupby('model', as_index=False).mean(numeric_only=True)
    for model_name, pooled_risk in [('CoxPH', mean_standard), ('Spline Cox', mean_spline)]:
        primary_summary.loc[primary_summary['model'] == model_name, ['observed_km_120', 'mean_predicted_120', 'oe_ratio']] = [float(observed_primary), float(np.mean(pooled_risk)), float(observed_primary / np.mean(pooled_risk))]
    atomic_write_csv(primary_summary, primary_summary_path)
    competing_df = pd.DataFrame(competing_rows)
    if len(competing_df) != args.m_imputations:
        raise RuntimeError('Competing-risk analysis did not complete every fixed-split imputation.')
    competing_by_imp_path = outdir / 'competing_risk_cox_by_imputation.csv'
    competing_summary_path = outdir / 'competing_risk_cox_summary.csv'
    atomic_write_csv(competing_df, competing_by_imp_path)
    competing_summary_values = competing_df.drop(columns=['imputation']).mean(numeric_only=True)
    mean_competing_cif = np.mean(np.vstack(competing_cif_risks), axis=0)
    observed_aj_values = competing_df['observed_test_aj_cif'].to_numpy(dtype=float)
    if not np.allclose(observed_aj_values, observed_aj_values[0], rtol=0.0, atol=1e-12):
        raise RuntimeError('Observed Aalen-Johansen risk changed across imputations.')
    observed_aj = float(observed_aj_values[0])
    competing_summary_values.loc['standard_cox_mean_predicted_risk'] = float(np.mean(mean_standard))
    competing_summary_values.loc['cause_specific_cox_mean_predicted_cif'] = float(np.mean(mean_competing_cif))
    competing_summary_values.loc['prediction_difference_standard_minus_competing'] = float(np.mean(mean_standard) - np.mean(mean_competing_cif))
    competing_summary_values.loc['observed_test_aj_cif'] = observed_aj
    competing_summary_values.loc['oe_standard'] = float(observed_aj / np.mean(mean_standard))
    competing_summary_values.loc['oe_competing_risk'] = float(observed_aj / np.mean(mean_competing_cif))
    competing_summary_values.loc['prediction_correlation'] = float(np.corrcoef(mean_standard, mean_competing_cif)[0, 1])
    competing_summary = competing_summary_values.rename_axis('metric').reset_index(name='mean')
    atomic_write_csv(competing_summary, competing_summary_path)
    competing_status['status'] = 'complete'
    competing_status['competing_analysis_completed'] = True
    competing_status['imputations_completed'] = int(len(competing_df))
    atomic_write_json(competing_status, competing_status_path)
    combined_calibration_png_path = outdir / 'flexible_calibration_cox_spline_combined.png'
    combined_calibration_tiff_path = outdir / 'flexible_calibration_cox_spline_combined.tiff'
    cox_curve_path = outdir / 'flexible_calibration_coxph_curve.csv'
    spline_curve_path = outdir / 'flexible_calibration_spline_cox_curve.csv'
    calibration_completeness_path = outdir / 'flexible_calibration_bootstrap_completeness.csv'
    cox_curve, cox_bootstrap = calibration_curve_with_bootstrap(test, mean_standard, args.seed, n_boot=args.calibration_bootstrap_reps)
    spline_curve, spline_bootstrap = calibration_curve_with_bootstrap(test, mean_spline, args.seed + 1, n_boot=args.calibration_bootstrap_reps)
    shared_axis_max = min(1.0, max(0.05, float(np.max(mean_standard)), float(np.max(mean_spline)), float(cox_curve['observed_risk'].max()), float(cox_curve['bootstrap_95_high'].max()), float(spline_curve['observed_risk'].max()), float(spline_curve['bootstrap_95_high'].max())) * 1.05)
    plot_combined_calibration_curves(cox_curve, spline_curve, combined_calibration_png_path, combined_calibration_tiff_path, axis_max=shared_axis_max)
    atomic_write_csv(cox_curve, cox_curve_path)
    atomic_write_csv(spline_curve, spline_curve_path)
    bootstrap_rows = []
    for model_name, metadata in [('CoxPH', cox_bootstrap), ('Spline Cox', spline_bootstrap)]:
        bootstrap_rows.append({'model': model_name, **metadata, 'shared_axis_max': shared_axis_max})
    atomic_write_csv(pd.DataFrame(bootstrap_rows), calibration_completeness_path)
    return [split_counts_path, spline_methods_path, competing_status_path, primary_diagnostics_path, vif_path, ldl_vif_path, primary_by_imp_path, primary_summary_path, competing_by_imp_path, competing_summary_path, combined_calibration_png_path, combined_calibration_tiff_path, cox_curve_path, spline_curve_path, calibration_completeness_path]
''',
    '_analysis_secondary.py': r'''from __future__ import annotations
import argparse
import tempfile
import time
from pathlib import Path
import numpy as np
import pandas as pd
from sklearn.model_selection import StratifiedShuffleSplit
from _analysis_secondary_core import COX_METRIC_COLUMNS, EXPECTED_ASCVD_EVENTS, EXPECTED_COMPETING_DEATHS, EXPECTED_ELIGIBLE_ROWS, EXPECTED_MAIN_MODELS, EXPECTED_OTHER_CENSORED, EXPECTED_PRIMARY_SPLITS, EXPECTED_SOURCE_ROWS, EXPECTED_TEST_ROWS_PER_SPLIT, EXPECTED_TREE_SPLITS, EXTRA_TREES_FIXED_CYCLES, EXTRA_TREES_SENSITIVITY_SEED_OFFSET, ID, MASTER_DEATH_TIME, MASTER_SHEET, PlannedSessionPause, artifact_marker_valid, atomic_write_csv, atomic_write_json, data_audit, expected_main_selection_metadata, expected_main_tuning_method, file_sha256, load_and_validate_main_analysis_dependencies, make_strata, prepare_competing_data, prepare_run_manifest, repeated_cox_comparison, run_primary_secondary_stage, validate_expected_study_counts, validate_secondary_dependencies, write_artifact_marker

def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description='Secondary analyses matched to the exact saved primary repeated-analysis 200 stratified 80/20 split assignments.')
    parser.add_argument('--csv', default='analysis_data.csv', help='Raw analysis CSV.')
    parser.add_argument('--master', default='source_data.xlsx', help='Master Excel workbook.')
    parser.add_argument('--main-output-dir', default='main_outputs', help='Completed main-analysis directory containing the locked split map.')
    parser.add_argument('--outdir', default='secondary_outputs')
    parser.add_argument('--seed', type=int, default=20260320)
    parser.add_argument('--n-splits', type=int, default=EXPECTED_PRIMARY_SPLITS, help='Primary Bayesian spline-Cox 80/20 repeated-analysis splits; locked at 200.')
    parser.add_argument('--tree-splits', type=int, default=EXPECTED_TREE_SPLITS, help='Matched ExtraTrees spline-Cox sensitivity split IDs; locked at 1-200.')
    parser.add_argument('--m-imputations', type=int, default=10, help='Imputations for fixed-split coefficients and diagnostics; locked at 10.')
    parser.add_argument('--n-jobs', type=int, default=1)
    parser.add_argument('--calibration-bootstrap-reps', type=int, default=1000)
    parser.add_argument('--max-session-hours', type=float, default=0.0, help='Optional safe session limit. A positive value pauses with exit code 75 only after a completed split; zero runs continuously.')
    parser.add_argument('--resume', action='store_true', help='Resume only checkpoints matching the exact data, scripts, split map and settings.')
    parser.add_argument('--validate-only', action='store_true', help='Validate input data and verified death times, then exit.')
    parser.add_argument('--functional-smoke-test', action='store_true', help='Exercise split loading, spline, imputation, calibration and resume APIs, then exit.')
    return parser.parse_args()

def validate_args(args: argparse.Namespace) -> None:
    if args.n_jobs == 0:
        raise ValueError('--n-jobs cannot be zero.')
    if args.n_splits != EXPECTED_PRIMARY_SPLITS:
        raise ValueError('--n-splits must equal 200 for the locked secondary analysis.')
    if args.tree_splits != EXPECTED_TREE_SPLITS:
        raise ValueError('--tree-splits must equal 200 for the matched sensitivity.')
    if args.m_imputations != 10:
        raise ValueError('--m-imputations must equal 10 for fixed-split diagnostics.')
    if args.calibration_bootstrap_reps < 1:
        raise ValueError('--calibration-bootstrap-reps must be at least 1.')
    if not np.isfinite(args.max_session_hours) or args.max_session_hours < 0.0:
        raise ValueError('--max-session-hours must be a finite non-negative value.')

def _write_smoke_main_dependencies(eligible: pd.DataFrame, main_out: Path, seed: int, n_splits: int) -> None:
    splitter = StratifiedShuffleSplit(n_splits=n_splits, test_size=0.2, random_state=seed)
    rows: list[dict[str, object]] = []
    performance: list[dict[str, object]] = []
    tuning: list[dict[str, object]] = []
    for split_id, (_, test_index) in enumerate(splitter.split(eligible, make_strata(eligible)), start=1):
        for test_order, study_id in enumerate(eligible.iloc[test_index][ID].tolist(), start=1):
            rows.append({'split_id': split_id, ID: study_id, 'role': 'test', 'split_seed': seed, 'test_order': test_order})
        for model in sorted(EXPECTED_MAIN_MODELS):
            split_base_seed = seed + 100000 * split_id
            performance.append({'split_id': split_id, 'n_train': 992, 'n_test': EXPECTED_TEST_ROWS_PER_SPLIT, 'train_events': 92, 'test_events': 24, 'model': model, 'harrell_c': 0.5, 'uno_c_tau': 0.5, 'tau_months': 120.0, 'brier_score_at_horizon': np.nan if model == 'SVM' else 0.1, 'brier_horizon_months': np.nan if model == 'SVM' else 120.0, 'integrated_brier_score': np.nan if model == 'SVM' else 0.1})
            tuning.append({'split_id': split_id, 'model': model, 'split_base_seed': split_base_seed, 'outer_imputation_seed': split_base_seed, 'model_seed': split_base_seed + 1, 'tuning_method': expected_main_tuning_method(model), 'best_params_json': '{}', 'selection_score': np.nan if model == 'CoxPH' else 0.1, **expected_main_selection_metadata(model)})
    (main_out / 'metadata').mkdir(parents=True, exist_ok=True)
    (main_out / 'tables').mkdir(parents=True, exist_ok=True)
    atomic_write_csv(pd.DataFrame(rows), main_out / 'metadata/repeated_200_split_map.csv')
    atomic_write_csv(pd.DataFrame(performance), main_out / 'tables/repeated_200_performance_by_split.csv')
    atomic_write_csv(pd.DataFrame(tuning), main_out / 'tables/repeated_200_tuning_by_split.csv')
    missingness = pd.DataFrame([{'Variable': variable, 'n_missing': n_missing, 'pct_missing': 100.0 * n_missing / EXPECTED_ELIGIBLE_ROWS, 'dtype': str(eligible[variable].dtype)} for variable, n_missing in {'fhxcvd': 259, 'hba1c': 56, 'trig': 35, 'egfr': 11, 'bmi': 4, 'dbp': 1, 'sbp': 1}.items()])
    atomic_write_csv(missingness, main_out / 'tables/supplementary_table_s1_missingness.csv')

def _paired_spline_summary(long_frames: list[pd.DataFrame]) -> pd.DataFrame:
    rows: list[dict[str, object]] = []
    metric_specs = [('harrell_c', 'raw_difference', 'positive_difference_favors_spline'), ('uno_c', 'raw_difference', 'positive_difference_favors_spline'), ('brier_120', 'raw_difference', 'negative_difference_favors_spline'), ('ibs_12_120', 'raw_difference', 'negative_difference_favors_spline'), ('calibration_slope', 'absolute_deviation_from_1', 'negative_difference_favors_spline'), ('oe_ratio', 'absolute_log_deviation_from_0', 'negative_difference_favors_spline')]
    for long in long_frames:
        method = str(long['imputer'].iloc[0])
        for metric, transform, interpretation in metric_specs:
            wide = long.pivot(index='split_id', columns='model', values=metric)
            if set(wide.columns) != {'CoxPH', 'Spline Cox'}:
                raise RuntimeError(f'{method} {metric} lacks both spline comparators.')
            if transform == 'absolute_deviation_from_1':
                compared = (wide - 1.0).abs()
            elif transform == 'absolute_log_deviation_from_0':
                if (wide <= 0.0).any().any():
                    raise RuntimeError(f'{method} O:E contains a non-positive value.')
                compared = np.log(wide).abs()
            else:
                compared = wide
            differences = (compared['Spline Cox'] - compared['CoxPH']).to_numpy(dtype=float)
            if len(differences) != EXPECTED_PRIMARY_SPLITS or not np.isfinite(differences).all():
                raise RuntimeError(f'{method} {metric} paired differences are invalid.')
            low, high = np.percentile(differences, [2.5, 97.5])
            rows.append({'imputer': method, 'metric': metric, 'comparison_scale': transform, 'paired_unit': 'identical held-out 80/20 split', 'n_paired_splits': int(len(differences)), 'mean_difference_spline_minus_coxph': float(np.mean(differences)), 'empirical_p2_5': float(low), 'empirical_p97_5': float(high), 'range_interpretation': 'descriptive empirical paired-split distribution; not a confidence interval', 'direction': interpretation, 'p_value_calculated': False, 'multiplicity_claim_made': False})
    return pd.DataFrame(rows)

def _imputer_comparison_summary(bayesian: pd.DataFrame, extra_trees: pd.DataFrame) -> pd.DataFrame:
    bayes_matched = bayesian.loc[bayesian['split_id'].isin(range(1, EXPECTED_TREE_SPLITS + 1))].copy()
    keys = ['split_id', 'model']
    merged = bayes_matched.merge(extra_trees, on=keys, how='inner', suffixes=('_bayesian', '_extra_trees'), validate='one_to_one')
    expected_keys = {(split_id, model) for split_id in range(1, EXPECTED_TREE_SPLITS + 1) for model in {'CoxPH', 'Spline Cox'}}
    observed_keys = set(zip(pd.to_numeric(merged['split_id'], errors='raise').astype(int), merged['model'].astype(str)))
    if len(merged) != EXPECTED_TREE_SPLITS * 2 or observed_keys != expected_keys or len(observed_keys) != len(merged):
        raise RuntimeError('The all-200 imputer comparison has an incomplete key grid.')
    for column in ['n_test', 'split_base_seed', 'model_seed']:
        if not merged[f'{column}_bayesian'].eq(merged[f'{column}_extra_trees']).all():
            raise RuntimeError(f'ExtraTrees and Bayesian spline analyses do not match on {column}.')
    rows: list[dict[str, object]] = []
    for model, model_frame in merged.groupby('model', sort=True):
        if len(model_frame) != EXPECTED_TREE_SPLITS:
            raise RuntimeError(f'The matched {model} imputer comparison does not contain 200 splits.')
        for metric in COX_METRIC_COLUMNS:
            differences = model_frame[f'{metric}_extra_trees'].to_numpy(dtype=float) - model_frame[f'{metric}_bayesian'].to_numpy(dtype=float)
            if not np.isfinite(differences).all():
                raise RuntimeError(f'The matched {model} {metric} imputer differences are non-finite.')
            low, high = np.percentile(differences, [2.5, 97.5])
            rows.append({'model': model, 'metric': metric, 'n_matched_splits': EXPECTED_TREE_SPLITS, 'mean_difference_extra_trees_minus_bayesian': float(np.mean(differences)), 'empirical_p2_5': float(low), 'empirical_p97_5': float(high), 'range_interpretation': 'descriptive empirical matched-split distribution; not a confidence interval', 'same_test_membership': True, 'same_model_seed': True, 'only_imputer_changed': True, 'p_value_calculated': False})
    summary = pd.DataFrame(rows)
    if len(summary) != 2 * len(COX_METRIC_COLUMNS) or not summary['n_matched_splits'].eq(EXPECTED_TREE_SPLITS).all():
        raise RuntimeError('The all-200 matched imputer summary is incomplete.')
    return summary

def _main_models_vs_spline_comparison(main_performance: pd.DataFrame, main_tuning: pd.DataFrame, bayesian_spline_long: pd.DataFrame, seed: int) -> tuple[pd.DataFrame, pd.DataFrame]:
    model_order = ['CoxPH', 'ElasticNetCox', 'RSF', 'GBSA', 'DeepSurv', 'CoxTime', 'XGBoost', 'SVM']
    keys = ['split_id', 'model']
    tuning_columns = ['split_id', 'model', 'split_base_seed', 'model_seed']
    candidate = main_performance.merge(main_tuning[tuning_columns], on=keys, how='inner', validate='one_to_one')
    expected_candidate_keys = {(split_id, model) for split_id in range(1, EXPECTED_PRIMARY_SPLITS + 1) for model in model_order}
    observed_candidate_keys = set(zip(pd.to_numeric(candidate['split_id'], errors='raise').astype(int), candidate['model'].astype(str)))
    if len(candidate) != len(expected_candidate_keys) or observed_candidate_keys != expected_candidate_keys or candidate.duplicated(keys).any():
        raise RuntimeError('Main-model versus spline comparison has an invalid main grid.')
    spline = bayesian_spline_long.loc[bayesian_spline_long['model'].astype(str).eq('Spline Cox')].copy()
    spline_split = pd.to_numeric(spline['split_id'], errors='raise').astype(int)
    expected_spline_ids = set(range(1, EXPECTED_PRIMARY_SPLITS + 1))
    expected_base = int(seed) + 100000 * spline_split
    if len(spline) != EXPECTED_PRIMARY_SPLITS or set(spline_split) != expected_spline_ids or spline['split_id'].duplicated().any() or (not pd.to_numeric(spline['n_test'], errors='raise').astype(int).eq(EXPECTED_TEST_ROWS_PER_SPLIT).all()) or (not pd.to_numeric(spline['split_base_seed'], errors='raise').astype(int).eq(expected_base).all()) or (not pd.to_numeric(spline['model_seed'], errors='raise').astype(int).eq(expected_base + 1).all()):
        raise RuntimeError('Spline Cox does not contain the exact 200 split/seed grid.')
    spline_columns = {'n_test': 'spline_n_test', 'split_base_seed': 'spline_split_base_seed', 'model_seed': 'spline_model_seed', 'harrell_c': 'spline_harrell_c', 'uno_c': 'spline_uno_c', 'brier_120': 'spline_brier_120', 'ibs_12_120': 'spline_ibs_12_120'}
    compared = candidate.merge(spline[['split_id', *spline_columns]].rename(columns=spline_columns), on='split_id', how='inner', validate='many_to_one')
    expected_candidate_base = int(seed) + 100000 * pd.to_numeric(compared['split_id'], errors='raise').astype(int)
    if len(compared) != len(expected_candidate_keys) or not pd.to_numeric(compared['n_test'], errors='raise').astype(int).eq(EXPECTED_TEST_ROWS_PER_SPLIT).all() or (not pd.to_numeric(compared['spline_n_test'], errors='raise').astype(int).eq(EXPECTED_TEST_ROWS_PER_SPLIT).all()) or (not pd.to_numeric(compared['split_base_seed'], errors='raise').astype(int).eq(expected_candidate_base).all()) or (not pd.to_numeric(compared['model_seed'], errors='raise').astype(int).eq(expected_candidate_base + 1).all()) or (not pd.to_numeric(compared['spline_split_base_seed'], errors='raise').astype(int).eq(expected_candidate_base).all()) or (not pd.to_numeric(compared['spline_model_seed'], errors='raise').astype(int).eq(expected_candidate_base + 1).all()):
        raise RuntimeError('Main models and Spline Cox are not matched on split size or seed identity.')
    metric_specs = [('harrell_c', 'harrell_c', 'spline_harrell_c', 'positive_candidate_minus_spline_favors_candidate'), ('uno_c', 'uno_c_tau', 'spline_uno_c', 'positive_candidate_minus_spline_favors_candidate'), ('brier_120', 'brier_score_at_horizon', 'spline_brier_120', 'negative_candidate_minus_spline_favors_candidate'), ('ibs_12_120', 'integrated_brier_score', 'spline_ibs_12_120', 'negative_candidate_minus_spline_favors_candidate')]
    audit_rows: list[dict[str, object]] = []
    for row in compared.itertuples(index=False):
        candidate_model = str(row.model)
        for metric, candidate_column, spline_column, direction in metric_specs:
            if candidate_model == 'SVM' and metric in {'brier_120', 'ibs_12_120'}:
                continue
            candidate_value = float(getattr(row, candidate_column))
            spline_value = float(getattr(row, spline_column))
            difference = candidate_value - spline_value
            if not np.isfinite([candidate_value, spline_value, difference]).all():
                raise RuntimeError(f'{candidate_model} versus Spline Cox produced invalid {metric}.')
            audit_rows.append({'split_id': int(row.split_id), 'candidate_model': candidate_model, 'reference_model': 'Spline Cox', 'metric': metric, 'candidate_value': candidate_value, 'spline_cox_value': spline_value, 'candidate_minus_spline': float(difference), 'direction': direction, 'n_test': int(row.n_test), 'split_base_seed': int(row.split_base_seed), 'model_seed': int(row.model_seed), 'same_held_out_split': True, 'same_seed_contract': True})
    audit = pd.DataFrame(audit_rows)
    expected_audit_rows = EXPECTED_PRIMARY_SPLITS * (7 * 4 + 2)
    expected_metric_counts = {model: (2 if model == 'SVM' else 4) * EXPECTED_PRIMARY_SPLITS for model in model_order}
    if len(audit) != expected_audit_rows or audit.duplicated(['split_id', 'candidate_model', 'metric']).any() or audit.groupby('candidate_model').size().to_dict() != expected_metric_counts:
        raise RuntimeError('Main-model versus spline audit grid is incomplete.')
    audit['_model_order'] = audit['candidate_model'].map({model: index for index, model in enumerate(model_order)})
    metric_order = {'harrell_c': 0, 'uno_c': 1, 'brier_120': 2, 'ibs_12_120': 3}
    audit['_metric_order'] = audit['metric'].map(metric_order)
    audit = audit.sort_values(['split_id', '_model_order', '_metric_order']).drop(columns=['_model_order', '_metric_order']).reset_index(drop=True)
    summary_rows: list[dict[str, object]] = []
    for (model, metric), frame in audit.groupby(['candidate_model', 'metric'], sort=False):
        values = frame['candidate_minus_spline'].to_numpy(dtype=float)
        low, high = np.percentile(values, [2.5, 97.5])
        summary_rows.append({'candidate_model': model, 'reference_model': 'Spline Cox', 'metric': metric, 'n_paired_splits': int(len(values)), 'mean_candidate_minus_spline': float(np.mean(values)), 'empirical_p2_5': float(low), 'empirical_p97_5': float(high), 'direction': str(frame['direction'].iloc[0]), 'paired_unit': 'identical held-out 80/20 split', 'range_interpretation': 'descriptive empirical paired-split distribution; not a confidence interval', 'p_value_calculated': False, 'multiplicity_claim_made': False})
    summary = pd.DataFrame(summary_rows)
    if len(summary) != 30 or not summary['n_paired_splits'].eq(EXPECTED_PRIMARY_SPLITS).all() or summary.duplicated(['candidate_model', 'metric']).any():
        raise RuntimeError('Main-model versus spline summary grid is incomplete.')
    summary['_model_order'] = summary['candidate_model'].map({model: index for index, model in enumerate(model_order)})
    summary['_metric_order'] = summary['metric'].map(metric_order)
    summary = summary.sort_values(['_model_order', '_metric_order']).drop(columns=['_model_order', '_metric_order']).reset_index(drop=True)
    return (audit, summary)

def _validate_participant_calibration(outdir: Path, method: str, expected_splits: int) -> None:
    calibration_path = outdir / f'cox_spline_participant_averaged_calibration_{method}.csv'
    coverage_path = outdir / f'cox_spline_heldout_coverage_{method}.json'
    calibration = pd.read_csv(calibration_path)
    required = {'imputer', 'model', 'n_participants', 'calibration_slope', 'calibration_oe_ratio', 'observed_km_risk', 'mean_predicted_risk', 'calibration_horizon_months', 'calibration_unit', 'calibration_estimand', 'represents_one_final_deployable_model', 'uncertainty_scope', 'model_development_uncertainty_included', 'split_stream_selection_uncertainty_included'}
    if not required.issubset(calibration.columns):
        raise RuntimeError(f'{method} participant calibration is incomplete.')
    numeric = calibration[['n_participants', 'calibration_slope', 'calibration_oe_ratio', 'observed_km_risk', 'mean_predicted_risk', 'calibration_horizon_months']].apply(pd.to_numeric, errors='raise')
    if len(calibration) != 2 or set(calibration['model'].astype(str)) != {'CoxPH', 'Spline Cox'} or set(calibration['imputer'].astype(str)) != {method} or (not numeric['n_participants'].eq(EXPECTED_ELIGIBLE_ROWS).all()) or (not numeric['calibration_horizon_months'].eq(120.0).all()) or (not numeric['calibration_oe_ratio'].gt(0.0).all()) or (not np.isfinite(numeric.to_numpy(dtype=float)).all()) or (set(calibration['calibration_estimand'].astype(str)) != {'cross-fitted repeated-split ensemble'}) or calibration['represents_one_final_deployable_model'].astype(str).str.lower().isin({'true', '1'}).any() or calibration['model_development_uncertainty_included'].astype(str).str.lower().isin({'true', '1'}).any() or calibration['split_stream_selection_uncertainty_included'].astype(str).str.lower().isin({'true', '1'}).any() or (not coverage_path.is_file()):
        raise RuntimeError(f'{method} participant calibration failed validation.')
    coverage = pd.read_json(coverage_path, typ='series')
    if int(coverage['n_splits']) != expected_splits or int(coverage['prediction_rows']) != expected_splits * EXPECTED_TEST_ROWS_PER_SPLIT or int(coverage['unique_participants']) != EXPECTED_ELIGIBLE_ROWS or (int(coverage['participants_without_held_out_prediction']) != 0) or (int(coverage['held_out_count_min']) != 21) or (int(coverage['held_out_count_max']) != 57):
        raise RuntimeError(f'{method} held-out coverage failed validation.')

def _biomarker_availability_summary(audit: dict[str, object]) -> pd.DataFrame:
    rows = [{'marker': 'LDL-C', 'source_field': 'ldl', 'availability': 'available', 'eligible_n': int(audit['eligible_n']), 'recorded_n': int(audit['ldl_nonmissing_n']), 'numeric_n': int(audit['ldl_nonmissing_n']), 'missing_n': int(audit['ldl_missing_n']), 'recorded_below_reporting_limit_n': 0, 'mean': float(audit['ldl_mean_mmol_l']), 'sd': float(audit['ldl_sd_mmol_l']), 'median': float(audit['ldl_median_mmol_l']), 'q1': float(audit['ldl_q1_mmol_l']), 'q3': float(audit['ldl_q3_mmol_l']), 'minimum': float(audit['ldl_min_mmol_l']), 'maximum': float(audit['ldl_max_mmol_l']), 'units': 'mmol/L', 'interpretation': 'LDL-C was available but was not one of the 15 prespecified model predictors.'}, {'marker': 'Generic CRP (not hsCRP)', 'source_field': 'CRP', 'availability': 'partly available; assay type and units not verified', 'eligible_n': int(audit['eligible_n']), 'recorded_n': int(audit['generic_crp_recorded_n']), 'numeric_n': int(audit['generic_crp_numeric_n']), 'missing_n': int(audit['generic_crp_unrecorded_n']), 'recorded_below_reporting_limit_n': int(audit['generic_crp_recorded_below_reporting_limit_n']), 'mean': float(audit['generic_crp_numeric_mean']), 'sd': float(audit['generic_crp_numeric_sd']), 'median': float(audit['generic_crp_numeric_median']), 'q1': float(audit['generic_crp_numeric_q1']), 'q3': float(audit['generic_crp_numeric_q3']), 'minimum': float(audit['generic_crp_numeric_min']), 'maximum': float(audit['generic_crp_numeric_max']), 'units': 'not verified', 'interpretation': 'The source field is generic CRP and must not be reported as hsCRP. Numerical summaries use the 232 directly numeric results; the 56 recorded below-reporting-limit values were not converted because the assay type and units were not verified.'}, {'marker': 'Lp(a)', 'source_field': f'none found in validated analysis CSV or selected master worksheet {MASTER_SHEET}', 'availability': 'unavailable', 'eligible_n': int(audit['eligible_n']), 'recorded_n': 0, 'numeric_n': 0, 'missing_n': int(audit['eligible_n']), 'recorded_below_reporting_limit_n': 0, 'mean': np.nan, 'sd': np.nan, 'median': np.nan, 'q1': np.nan, 'q3': np.nan, 'minimum': np.nan, 'maximum': np.nan, 'units': 'not available', 'interpretation': f'No Lp(a) field was found in the validated analysis CSV or the selected master worksheet ({MASTER_SHEET}).'}]
    return pd.DataFrame(rows)

def main() -> None:
    args = parse_args()
    validate_args(args)
    validate_secondary_dependencies()
    csv_path = Path(args.csv).expanduser().resolve()
    master_path = Path(args.master).expanduser().resolve()
    if not csv_path.is_file():
        raise FileNotFoundError('Could not find the analysis CSV.')
    if not master_path.is_file():
        raise FileNotFoundError('Could not find the source workbook.')
    raw_all = pd.read_csv(csv_path)
    master = pd.read_excel(master_path, sheet_name=MASTER_SHEET)
    prepared_all = prepare_competing_data(raw_all, master)
    eligible = validate_expected_study_counts(prepared_all)
    audit = data_audit(prepared_all, source_column_names=[*raw_all.columns.tolist(), *master.columns.tolist()])
    print('Secondary-analysis data validation passed.')
    print(f'Source rows: {EXPECTED_SOURCE_ROWS}')
    print(f'Eligible participants: {EXPECTED_ELIGIBLE_ROWS}')
    print(f'ASCVD events: {EXPECTED_ASCVD_EVENTS}')
    print(f'Verified non-ASCVD deaths: {EXPECTED_COMPETING_DEATHS}')
    print(f'Other right-censored observations: {EXPECTED_OTHER_CENSORED}')
    print(f'Competing-event time source: workbook sheet {MASTER_SHEET!r}, column {MASTER_DEATH_TIME!r}')
    if args.validate_only:
        return
    if args.functional_smoke_test:
        smoke_args = argparse.Namespace(**vars(args))
        smoke_args.m_imputations = 1
        smoke_args.calibration_bootstrap_reps = 10
        with tempfile.TemporaryDirectory(prefix='secondary_200_split_preflight_') as temporary:
            root = Path(temporary)
            smoke_main = root / 'main'
            smoke_out = root / 'secondary'
            _write_smoke_main_dependencies(eligible, smoke_main, args.seed, n_splits=2)
            split_map, _ = load_and_validate_main_analysis_dependencies(smoke_main, eligible, expected_n_splits=2, expected_seed=args.seed, require_fixed_split_evidence=False)
            primary_files = run_primary_secondary_stage(eligible, smoke_out, smoke_args, split_map)
            if not primary_files or not all((path.is_file() for path in primary_files)):
                raise RuntimeError('Secondary fixed-split smoke artifacts are incomplete.')
            bayes = repeated_cox_comparison(eligible, smoke_out, split_map=split_map, split_ids=[1, 2], seed=args.seed, method='bayesian', n_jobs=args.n_jobs, resume=False, configuration_sha256='secondary-200-split-functional-preflight', functional_smoke=True)
            resumed = repeated_cox_comparison(eligible, smoke_out, split_map=split_map, split_ids=[1, 2], seed=args.seed, method='bayesian', n_jobs=args.n_jobs, resume=True, configuration_sha256='secondary-200-split-functional-preflight', functional_smoke=True)
            keys = ['split_id', 'model', 'imputer']
            bayes_sorted = bayes.sort_values(keys).reset_index(drop=True)
            resumed_sorted = resumed.sort_values(keys).reset_index(drop=True)
            if bayes_sorted[keys].astype(str).to_dict('records') != resumed_sorted[keys].astype(str).to_dict('records') or not np.allclose(bayes_sorted[COX_METRIC_COLUMNS].to_numpy(dtype=float), resumed_sorted[COX_METRIC_COLUMNS].to_numpy(dtype=float), rtol=1e-12, atol=1e-12):
                raise RuntimeError('Secondary Bayesian split checkpoints failed exact resume.')
            tree = repeated_cox_comparison(eligible, smoke_out, split_map=split_map, split_ids=[1], seed=args.seed, method='extra_trees', n_jobs=args.n_jobs, resume=False, configuration_sha256='secondary-200-split-functional-preflight', functional_smoke=True)
            if len(tree) != 2 or not np.isfinite(tree[COX_METRIC_COLUMNS].to_numpy(dtype=float)).all():
                raise RuntimeError('Secondary ExtraTrees split smoke metrics are incomplete.')
        print('Secondary 200-split/calibration/checkpoint functional preflight passed.', flush=True)
        return
    main_output_dir = Path(args.main_output_dir).expanduser().resolve()
    split_map, main_dependencies = load_and_validate_main_analysis_dependencies(main_output_dir, eligible, expected_n_splits=args.n_splits, expected_seed=args.seed, require_fixed_split_evidence=True)
    outdir = Path(args.outdir).expanduser().resolve()
    outdir.mkdir(parents=True, exist_ok=True)
    configuration_sha256 = prepare_run_manifest(outdir, csv_path, master_path, main_dependencies, args)
    atomic_write_json(audit, outdir / 'data_audit_aggregates.json')
    atomic_write_csv(_biomarker_availability_summary(audit), outdir / 'biomarker_availability_summary.csv')
    deadline_monotonic = None if args.max_session_hours == 0.0 else time.monotonic() + args.max_session_hours * 3600.0
    primary_marker = outdir / 'checkpoints' / 'fixed_split_secondary_stage_complete.json'
    primary_metadata = {'configuration_sha256': configuration_sha256, 'stage': 'fixed_split_10_imputation_secondary_analyses', 'split_id': 1, 'imputations': int(args.m_imputations), 'calibration_bootstrap_reps': int(args.calibration_bootstrap_reps), 'calibration_figure_layout': 'landscape_panels_no_overall_title_or_footer_common_square_axes', 'calibration_figure_formats': 'PNG_600dpi_and_LZW_TIFF_600dpi', 'verified_death_time_column': MASTER_DEATH_TIME}
    if args.resume and artifact_marker_valid(primary_marker, outdir, primary_metadata):
        print('Resumed the completed fixed-split secondary stage.', flush=True)
    else:
        primary_files = run_primary_secondary_stage(eligible, outdir, args, split_map)
        write_artifact_marker(primary_marker, outdir, primary_files, primary_metadata)
        print('Completed and checkpointed the fixed-split secondary stage.', flush=True)
    try:
        bayes_long = repeated_cox_comparison(eligible, outdir, split_map=split_map, split_ids=list(range(1, args.n_splits + 1)), seed=args.seed, method='bayesian', n_jobs=args.n_jobs, resume=args.resume, configuration_sha256=configuration_sha256, deadline_monotonic=deadline_monotonic)
        tree_long = repeated_cox_comparison(eligible, outdir, split_map=split_map, split_ids=list(range(1, args.tree_splits + 1)), seed=args.seed, method='extra_trees', n_jobs=args.n_jobs, resume=args.resume, configuration_sha256=configuration_sha256, deadline_monotonic=deadline_monotonic)
    except PlannedSessionPause as exc:
        pause_status = {'status': 'paused_safely', 'configuration_sha256': configuration_sha256, 'reason': str(exc), 'resume_instruction': 'Run the unchanged launcher again.', 'exit_code': 75}
        atomic_write_json(pause_status, outdir / 'ANALYSIS_PAUSED.json')
        print(str(exc), flush=True)
        raise SystemExit(75)
    paired_path = outdir / 'paired_cox_spline_differences.csv'
    atomic_write_csv(_paired_spline_summary([bayes_long, tree_long]), paired_path)
    imputer_comparison_path = outdir / 'matched_extra_trees_vs_bayesian_all_200.csv'
    atomic_write_csv(_imputer_comparison_summary(bayes_long, tree_long), imputer_comparison_path)
    main_spline_audit, main_spline_summary = _main_models_vs_spline_comparison(pd.read_csv(main_dependencies['performance'], float_precision='round_trip'), pd.read_csv(main_dependencies['tuning'], float_precision='round_trip'), bayes_long, seed=args.seed)
    main_spline_audit_path = outdir / 'main_models_vs_spline_cox_by_split.csv'
    main_spline_summary_path = outdir / 'main_models_vs_spline_cox_summary.csv'
    atomic_write_csv(main_spline_audit, main_spline_audit_path)
    atomic_write_csv(main_spline_summary, main_spline_summary_path)
    _validate_participant_calibration(outdir, 'bayesian', args.n_splits)
    _validate_participant_calibration(outdir, 'extra_trees', args.tree_splits)
    diagnostic_files = [outdir / 'imputation_diagnostics_fixed_split_bayesian.csv', outdir / 'imputation_diagnostics_by_split_bayesian.csv', outdir / 'imputation_diagnostics_by_split_extra_trees.csv']
    diagnostics = pd.concat([pd.read_csv(path) for path in diagnostic_files], ignore_index=True)
    expected_diagnostics = args.m_imputations + args.n_splits + args.tree_splits
    if len(diagnostics) != expected_diagnostics:
        raise RuntimeError(f'Imputation diagnostics contained {len(diagnostics)} rows; expected {expected_diagnostics}.')
    tree_diagnostics = pd.read_csv(diagnostic_files[-1])
    fixed_completed = tree_diagnostics['imputer_fixed_cycle_completed'].astype(str).str.lower().isin({'true', '1'})
    if len(tree_diagnostics) != EXPECTED_TREE_SPLITS or not pd.to_numeric(tree_diagnostics['imputer_n_iter'], errors='raise').eq(EXTRA_TREES_FIXED_CYCLES).all() or (not fixed_completed.all()) or (not pd.to_numeric(tree_diagnostics['imputer_unexpected_convergence_warning_count'], errors='raise').eq(0).all()) or (not pd.to_numeric(tree_diagnostics['imputer_expected_fixed_cycle_warning_count'], errors='raise').eq(1).all()) or (not tree_diagnostics['imputer_policy'].astype(str).eq('prespecified_fixed_10_cycles_tol_0_no_convergence_claim').all()):
        raise RuntimeError('ExtraTrees did not complete its fixed 10-cycle policy.')
    diagnostic_summary = {'completed_datasets': int(len(diagnostics)), 'bayesian_split_datasets': int(args.n_splits), 'extra_trees_split_datasets': int(args.tree_splits), 'extra_trees_policy': 'prespecified_fixed_10_cycles_tol_0_no_convergence_claim', 'extra_trees_fixed_cycles_required': EXTRA_TREES_FIXED_CYCLES, 'extra_trees_fixed_cycle_completed_count': int(fixed_completed.sum()), 'extra_trees_expected_fixed_cycle_notice_count': int(pd.to_numeric(tree_diagnostics['imputer_expected_fixed_cycle_warning_count'], errors='raise').sum()), 'extra_trees_unexpected_convergence_warning_count': int(pd.to_numeric(tree_diagnostics['imputer_unexpected_convergence_warning_count'], errors='raise').sum()), 'interpretation': 'ExtraTrees uses a prespecified 10-cycle policy for this analysis. Reaching cycle 10 is expected and is not interpreted as failure to converge. Unexpected numerical or convergence warnings are counted separately and must be zero.'}
    atomic_write_json(diagnostic_summary, outdir / 'imputation_diagnostics_summary.json')
    completion = {'schema_version': 8, 'status': 'complete', 'configuration_sha256': configuration_sha256, 'seed': int(args.seed), 'primary_repeated_validation': '200 repeated stratified random 80/20 splits', 'bayesian_split_ids': [1, args.n_splits], 'bayesian_n_splits': int(args.n_splits), 'extra_trees_split_ids': [1, args.tree_splits], 'extra_trees_n_splits': int(args.tree_splits), 'same_saved_split_map_used': True, 'same_model_seed_between_imputers': True, 'extra_trees_imputer_seed_offset': EXTRA_TREES_SENSITIVITY_SEED_OFFSET, 'extra_trees_fixed_cycles': EXTRA_TREES_FIXED_CYCLES, 'fixed_split_imputations': int(args.m_imputations), 'fixed_split_role': 'supplementary coefficients and diagnostics', 'dedicated_linear_and_spline_cox_shared_ridge': 1e-05, 'main_model_selection_policy_verified': {model_name: expected_main_selection_metadata(model_name) for model_name in sorted(EXPECTED_MAIN_MODELS)}, 'spline_cox_hyperparameter_selection_performed': False, 'split_summary': 'mean and empirical 2.5th-97.5th percentiles', 'split_ranges_are_confidence_intervals': False, 'paired_p_values_calculated': False, 'multiplicity_inference_claimed': False, 'participant_calibration_unit': 'one risk per participant averaged across held-out appearances only', 'participant_calibration_estimand': 'cross-fitted repeated-split ensemble', 'participant_calibration_represents_one_final_deployable_model': False, 'participant_calibration_uncertainty_scope': 'participant-bootstrap bands are conditional on fixed participant-averaged held-out predictions and exclude model-development and split-stream-selection uncertainty', 'calibration_bootstrap_reps': int(args.calibration_bootstrap_reps), 'calibration_figure_layout': 'landscape_panels_no_overall_title_or_footer_common_square_axes', 'calibration_figure_formats': 'PNG_600dpi_and_LZW_TIFF_600dpi', 'main_dependency_sha256': {name: file_sha256(path) for name, path in sorted(main_dependencies.items())}, 'main_models_vs_spline_cox': {'by_split_file': main_spline_audit_path.name, 'by_split_rows': int(len(main_spline_audit)), 'summary_file': main_spline_summary_path.name, 'summary_rows': int(len(main_spline_summary)), 'difference': 'candidate minus Spline Cox', 'confidence_intervals_calculated': False, 'p_values_calculated': False}, 'competing_event_time_source': f'{MASTER_SHEET} / {MASTER_DEATH_TIME}', 'eligible_non_ascvd_deaths': EXPECTED_COMPETING_DEATHS}
    atomic_write_json(completion, outdir / 'analysis_complete.json')
    pause_marker = outdir / 'ANALYSIS_PAUSED.json'
    if pause_marker.exists():
        pause_marker.unlink()
    print('Secondary analyses complete.', flush=True)
if __name__ == '__main__':
    main()
''',
    '_analysis_final.py': r'''from __future__ import annotations
import argparse
import hashlib
import importlib.metadata
import importlib.util
import json
import math
import os
import sys
import tempfile
import warnings
from pathlib import Path
from typing import Any, Sequence
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
SCHEMA_VERSION = 1
EXPECTED_PARTICIPANTS = 1241
EXPECTED_EVENTS = 116
EXPECTED_FIXED_TRAINING_PARTICIPANTS = 992
EXPECTED_FIXED_TRAINING_EVENTS = 92
EXPECTED_FIXED_TEST_PARTICIPANTS = 249
EXPECTED_FIXED_TEST_EVENTS = 24
EXPECTED_MODELS = ['CoxPH', 'ElasticNetCox', 'RSF', 'GBSA', 'DeepSurv', 'CoxTime', 'XGBoost']
DEFAULT_SEED = 20260320
DEFAULT_M_IMPUTATIONS = 50
DEFAULT_IMPUTE_MAX_ITER = 20
DEFAULT_BOOTSTRAP_REPS = 1000
DEFAULT_HORIZON_MONTHS = 120.0
INFERENCE_IMPUTATION_SEED_OFFSET = 4000000
CALIBRATION_BOOTSTRAP_SEED_OFFSET = 710000
PMM_DONORS = 10
REQUIRED_BOOTSTRAP_SUCCESS_FRACTION = 1.0
LOCKED_ADDITIONAL_PACKAGES = {'statsmodels': '0.14.4', 'patsy': '1.0.1'}
NELSON_AALEN_COLUMN = 'nelson_aalen_cumulative_hazard_auxiliary'
EVENT_AUXILIARY_COLUMN = 'ascvd_event_auxiliary'

def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument('--csv', default='analysis_data.csv')
    parser.add_argument('--main-script', default='_analysis_main.py')
    parser.add_argument('--main-output-dir', default='main_outputs')
    parser.add_argument('--secondary-output-dir', default='secondary_outputs')
    parser.add_argument('--outdir', default='final_outputs')
    parser.add_argument('--seed', type=int, default=DEFAULT_SEED)
    parser.add_argument('--m-imputations', type=int, default=DEFAULT_M_IMPUTATIONS)
    parser.add_argument('--impute-max-iter', type=int, default=DEFAULT_IMPUTE_MAX_ITER)
    parser.add_argument('--calibration-bootstrap-reps', type=int, default=DEFAULT_BOOTSTRAP_REPS)
    parser.add_argument('--horizon-months', type=float, default=DEFAULT_HORIZON_MONTHS)
    parser.add_argument('--max-session-hours', type=float, default=0.0, help='Accepted for launcher consistency; this short stage is atomic.')
    parser.add_argument('--resume', action='store_true')
    parser.add_argument('--validate-only', action='store_true')
    parser.add_argument('--functional-smoke-test', action='store_true')
    return parser.parse_args()

def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open('rb') as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b''):
            digest.update(chunk)
    return digest.hexdigest()

def canonical_json(value: Any) -> str:
    return json.dumps(value, sort_keys=True, separators=(',', ':'), allow_nan=False, default=str)

def atomic_write_text(text: str, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    handle, temporary_name = tempfile.mkstemp(prefix=f'.{path.name}.', suffix='.tmp', dir=path.parent)
    temporary = Path(temporary_name)
    try:
        with os.fdopen(handle, 'w', encoding='utf-8', newline='') as stream:
            stream.write(text)
            stream.flush()
            os.fsync(stream.fileno())
        os.replace(temporary, path)
    except Exception:
        temporary.unlink(missing_ok=True)
        raise

def atomic_write_json(value: Any, path: Path) -> None:
    atomic_write_text(json.dumps(value, indent=2, sort_keys=True, default=str) + '\n', path)

def atomic_write_csv(frame: pd.DataFrame, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    handle, temporary_name = tempfile.mkstemp(prefix=f'.{path.name}.', suffix='.tmp', dir=path.parent)
    temporary = Path(temporary_name)
    try:
        with os.fdopen(handle, 'w', encoding='utf-8', newline='') as stream:
            frame.to_csv(stream, index=False, float_format='%.17g')
            stream.flush()
            os.fsync(stream.fileno())
        os.replace(temporary, path)
    except Exception:
        temporary.unlink(missing_ok=True)
        raise

def require(condition: bool, message: str) -> None:
    if not condition:
        raise RuntimeError(message)

def import_main_module(path: Path) -> Any:
    require(path.is_file(), f'Main analysis script is missing: {path}')
    module_name = 'analysis_main_pipeline_for_final_outputs'
    spec = importlib.util.spec_from_file_location(module_name, path)
    require(spec is not None and spec.loader is not None, 'Could not load the main pipeline.')
    module = importlib.util.module_from_spec(spec)
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    required = {'load_dataset', 'primary_split', 'fit_lifelines_cox_and_ph_test', 'calibration_oe_from_probs', 'PREDICTOR_COLUMNS', 'BINARY_COLUMNS', 'CONTINUOUS_COLUMNS', 'STUDY_ID_COL', 'TIME_COL', 'EVENT_COL', 'DISPLAY_NAMES'}
    missing = sorted((name for name in required if not hasattr(module, name)))
    require(not missing, f'The main pipeline lacks required interfaces: {missing}')
    return module

def validate_locked_args(args: argparse.Namespace) -> None:
    require(args.seed == DEFAULT_SEED, f'--seed must remain {DEFAULT_SEED}.')
    require(args.m_imputations == DEFAULT_M_IMPUTATIONS, f'--m-imputations must remain {DEFAULT_M_IMPUTATIONS}.')
    require(args.impute_max_iter == DEFAULT_IMPUTE_MAX_ITER, f'--impute-max-iter must remain {DEFAULT_IMPUTE_MAX_ITER}.')
    require(args.calibration_bootstrap_reps == DEFAULT_BOOTSTRAP_REPS, f'--calibration-bootstrap-reps must remain {DEFAULT_BOOTSTRAP_REPS}.')
    require(math.isclose(args.horizon_months, DEFAULT_HORIZON_MONTHS, abs_tol=1e-12), f'--horizon-months must remain {DEFAULT_HORIZON_MONTHS:g}.')
    require(math.isfinite(args.max_session_hours) and args.max_session_hours >= 0.0, '--max-session-hours must be finite and non-negative.')
    for distribution, expected_version in LOCKED_ADDITIONAL_PACKAGES.items():
        try:
            observed_version = importlib.metadata.version(distribution)
        except importlib.metadata.PackageNotFoundError as exc:
            raise RuntimeError(f'{distribution} is required for the fixed-training Cox imputation. Install the required package before continuing.') from exc
        require(observed_version == expected_version, f'{distribution} must be version {expected_version}; found {observed_version}.')

def nelson_aalen_at_observed_times(durations: Sequence[float], events: Sequence[int]) -> np.ndarray:
    time_values = np.asarray(durations, dtype=float).reshape(-1)
    event_values = np.asarray(events, dtype=int).reshape(-1)
    require(len(time_values) == len(event_values) and len(time_values) > 0, 'Nelson-Aalen inputs are empty or misaligned.')
    require(np.isfinite(time_values).all() and np.all(time_values > 0.0), 'Nelson-Aalen times must be positive and finite.')
    require(np.isin(event_values, [0, 1]).all(), 'Nelson-Aalen events must be binary.')
    event_times = np.unique(time_values[event_values == 1])
    require(event_times.size > 0, 'Nelson-Aalen estimation requires at least one event.')
    increments: list[float] = []
    for event_time in event_times:
        at_risk = int(np.sum(time_values >= event_time))
        event_count = int(np.sum((time_values == event_time) & (event_values == 1)))
        require(at_risk > 0 and 0 < event_count <= at_risk, 'Invalid Nelson-Aalen risk set.')
        increments.append(float(event_count / at_risk))
    cumulative = np.cumsum(np.asarray(increments, dtype=float))
    positions = np.searchsorted(event_times, time_values, side='right') - 1
    output = np.zeros(len(time_values), dtype=float)
    valid = positions >= 0
    output[valid] = cumulative[positions[valid]]
    require(np.isfinite(output).all() and np.all(output >= 0.0), 'Nelson-Aalen values are invalid.')
    return output

def validate_nelson_aalen_against_lifelines(durations: Sequence[float], events: Sequence[int], calculated: np.ndarray) -> None:
    from lifelines import NelsonAalenFitter
    fitter = NelsonAalenFitter(nelson_aalen_smoothing=False)
    fitter.fit(np.asarray(durations, dtype=float), event_observed=np.asarray(events, dtype=int))
    reference = np.asarray(fitter.predict(np.asarray(durations, dtype=float)), dtype=float)
    require(np.allclose(calculated, reference, rtol=0.0, atol=2e-12), 'The direct Nelson-Aalen estimator does not match lifelines with smoothing disabled.')

def outcome_aware_fixed_training_imputations(raw_training: pd.DataFrame, main_module: Any, m_imputations: int, max_iter: int, seed: int) -> tuple[list[pd.DataFrame], pd.DataFrame, np.ndarray]:
    import statsmodels.api as sm
    from statsmodels.imputation.mice import MICEData
    from statsmodels.tools.sm_exceptions import ConvergenceWarning as SMConvergenceWarning
    from statsmodels.tools.sm_exceptions import PerfectSeparationWarning
    predictors = list(main_module.PREDICTOR_COLUMNS)
    study_id = str(main_module.STUDY_ID_COL)
    time_col = str(main_module.TIME_COL)
    event_col = str(main_module.EVENT_COL)
    hazard = nelson_aalen_at_observed_times(raw_training[time_col].to_numpy(dtype=float), raw_training[event_col].to_numpy(dtype=int))
    validate_nelson_aalen_against_lifelines(raw_training[time_col].to_numpy(dtype=float), raw_training[event_col].to_numpy(dtype=int), hazard)
    matrix = raw_training[predictors].copy()
    matrix[EVENT_AUXILIARY_COLUMN] = raw_training[event_col].to_numpy(dtype=float)
    matrix[NELSON_AALEN_COLUMN] = hazard
    matrix_columns = predictors + [EVENT_AUXILIARY_COLUMN, NELSON_AALEN_COLUMN]
    missing_mask = matrix[predictors].isna()
    missing_counts = missing_mask.sum().astype(int).to_dict()
    incomplete_predictors = [predictor for predictor in predictors if int(missing_counts[predictor]) > 0]
    incomplete_continuous = [predictor for predictor in main_module.CONTINUOUS_COLUMNS if predictor in incomplete_predictors]
    incomplete_binary = [predictor for predictor in main_module.BINARY_COLUMNS if predictor in incomplete_predictors]
    require(incomplete_binary == ['fhxcvd'], f'The validated training set should have only fhxcvd as an incomplete binary predictor; found {incomplete_binary}.')
    observed_before = matrix[predictors].copy()
    completed_training_sets: list[pd.DataFrame] = []
    diagnostic_rows: list[dict[str, Any]] = []
    for imputation_index in range(1, m_imputations + 1):
        imputation_seed = int(seed + INFERENCE_IMPUTATION_SEED_OFFSET + imputation_index - 1)
        legacy_random_state = np.random.get_state()
        try:
            np.random.seed(imputation_seed)
            with warnings.catch_warnings(record=True) as captured:
                warnings.simplefilter('always')
                imputer = MICEData(matrix[matrix_columns].copy(), perturbation_method='gaussian', k_pmm=PMM_DONORS)
                for predictor in incomplete_continuous:
                    formula = ' + '.join((column for column in matrix_columns if column != predictor))
                    imputer.set_imputer(predictor, formula=formula, model_class=sm.OLS, k_pmm=PMM_DONORS, perturbation_method='gaussian')
                binary_predictor = incomplete_binary[0]
                binary_formula = ' + '.join((column for column in matrix_columns if column != binary_predictor))
                imputer.set_imputer(binary_predictor, formula=binary_formula, model_class=sm.GLM, init_kwds={'family': sm.families.Binomial()}, k_pmm=PMM_DONORS, perturbation_method='gaussian')
                for cycle_index in range(1, max_iter + 1):
                    imputer.update_all(n_iter=1)
                    fitted_results = getattr(imputer, 'results', {})
                    for predictor, fitted_result in fitted_results.items():
                        if predictor == binary_predictor:
                            require(hasattr(fitted_result, 'converged') and bool(fitted_result.converged), f'Imputation {imputation_index}, cycle {cycle_index}: the binary GLM did not provide a successful convergence flag.')
                        require(bool(getattr(fitted_result, 'converged', True)), f'Imputation {imputation_index}, cycle {cycle_index}: the {predictor} conditional model did not converge.')
                completed_matrix = imputer.data.copy()
        finally:
            np.random.set_state(legacy_random_state)
        completed_matrix = completed_matrix.loc[:, matrix_columns]
        completed_matrix.index = raw_training.index
        require(np.isfinite(completed_matrix.to_numpy(dtype=float)).all(), f'Imputation {imputation_index} produced non-finite values.')
        require(np.array_equal(completed_matrix[EVENT_AUXILIARY_COLUMN].to_numpy(dtype=float), matrix[EVENT_AUXILIARY_COLUMN].to_numpy(dtype=float)), f'Imputation {imputation_index} changed the event auxiliary variable.')
        require(np.allclose(completed_matrix[NELSON_AALEN_COLUMN].to_numpy(dtype=float), hazard, rtol=0.0, atol=0.0), f'Imputation {imputation_index} changed the Nelson-Aalen auxiliary variable.')
        for predictor in predictors:
            observed = ~missing_mask[predictor]
            require(np.array_equal(completed_matrix.loc[observed, predictor].to_numpy(dtype=float), observed_before.loc[observed, predictor].to_numpy(dtype=float)), f'Imputation {imputation_index} changed an observed value in {predictor}.')
        donor_checks: dict[str, bool] = {}
        imputed_hashes: dict[str, str] = {}
        imputed_summaries: dict[str, dict[str, float | int]] = {}
        for predictor in incomplete_predictors:
            imputed_values = completed_matrix.loc[missing_mask[predictor], predictor].to_numpy(dtype=float)
            observed_values = matrix.loc[~missing_mask[predictor], predictor].to_numpy(dtype=float)
            donor_checks[predictor] = bool(np.isin(imputed_values, observed_values).all())
            require(donor_checks[predictor], f'Imputation {imputation_index} produced a non-donor PMM value in {predictor}.')
            imputed_hashes[predictor] = hashlib.sha256(np.asarray(imputed_values, dtype='<f8').tobytes()).hexdigest()
            imputed_summaries[predictor] = {'n': int(len(imputed_values)), 'minimum': float(np.min(imputed_values)), 'maximum': float(np.max(imputed_values)), 'mean': float(np.mean(imputed_values)), 'unique': int(np.unique(imputed_values).size)}
        completed_predictors = completed_matrix[predictors].copy()
        for predictor in main_module.BINARY_COLUMNS:
            completed_predictors[predictor] = completed_predictors[predictor].astype(int)
        require(not completed_predictors.isna().any().any() and np.isfinite(completed_predictors.to_numpy(dtype=float)).all(), f'Imputation {imputation_index} retained an invalid predictor value.')
        for predictor in main_module.BINARY_COLUMNS:
            require(set(np.unique(completed_predictors[predictor].to_numpy(dtype=int))).issubset({0, 1}), f'Imputation {imputation_index} produced an invalid binary value in {predictor}.')
        completed_training = pd.concat([raw_training[[study_id, time_col, event_col]].reset_index(drop=True), completed_predictors.reset_index(drop=True)], axis=1)
        completed_training_sets.append(completed_training)
        warning_messages = [str(item.message) for item in captured]
        blocking_mice_warnings = [item for item in captured if issubclass(item.category, (SMConvergenceWarning, PerfectSeparationWarning, RuntimeWarning, np.ComplexWarning if hasattr(np, 'ComplexWarning') else RuntimeWarning))]
        require(not blocking_mice_warnings, f'Imputation {imputation_index} emitted a convergence/separation/numerical warning: ' + ' | '.join((str(item.message) for item in blocking_mice_warnings)))
        diagnostic_rows.append({'imputation': imputation_index, 'imputation_seed': imputation_seed, 'imputer': 'statsmodels_MICEData_type_specific_PMM', 'imputation_role': 'fixed_training_set_cox_coefficient_inference_only', 'predictor_count': len(predictors), 'auxiliary_variables': f'{EVENT_AUXILIARY_COLUMN}; {NELSON_AALEN_COLUMN}', 'continuous_imputation_model': 'OLS_with_gaussian_parameter_perturbation_and_PMM', 'binary_imputation_model': 'binomial_GLM_with_gaussian_parameter_perturbation_and_PMM', 'pmm_donors': PMM_DONORS, 'event_auxiliary_preserved': True, 'nelson_aalen_auxiliary_preserved': True, 'observed_predictor_cells_preserved': True, 'all_imputed_values_are_observed_donors': all(donor_checks.values()), 'missing_predictor_cells_before': int(missing_mask.to_numpy().sum()), 'missing_predictor_cells_after': int(completed_predictors.isna().to_numpy().sum()), 'mice_cycles_completed': int(max_iter), 'mice_cycles_planned': int(max_iter), 'warning_count': len(warning_messages), 'blocking_warning_count': len(blocking_mice_warnings), 'warnings_json': json.dumps(warning_messages, sort_keys=True), 'missing_counts_json': json.dumps(missing_counts, sort_keys=True), 'imputed_value_sha256_json': json.dumps(imputed_hashes, sort_keys=True), 'imputed_value_summary_json': json.dumps(imputed_summaries, sort_keys=True), 'prediction_pipeline_affected': False})
    require(len({int(row['imputation_seed']) for row in diagnostic_rows}) == m_imputations, 'Outcome-aware MICE seeds are not unique.')
    for predictor in incomplete_predictors:
        hashes = {json.loads(str(row['imputed_value_sha256_json']))[predictor] for row in diagnostic_rows}
        require(len(hashes) > 1, f'All {m_imputations} MICE datasets are identical for incomplete predictor {predictor}.')
    return (completed_training_sets, pd.DataFrame(diagnostic_rows), hazard)

def collect_cox_details(completed_training_sets: Sequence[pd.DataFrame], main_module: Any, enforce_precision_gate: bool=True) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    from lifelines.exceptions import ConvergenceWarning
    from scipy import stats
    detailed_rows: list[pd.DataFrame] = []
    for imputation_index, completed in enumerate(completed_training_sets, start=1):
        with warnings.catch_warnings(record=True) as captured:
            warnings.simplefilter('always')
            _, detail = main_module.fit_lifelines_cox_and_ph_test(completed)
        blocking = [item for item in captured if issubclass(item.category, (ConvergenceWarning, RuntimeWarning))]
        require(not blocking, f'Cox/PH fit for imputation {imputation_index} emitted a convergence or numerical warning: ' + ' | '.join((str(item.message) for item in blocking)))
        detail = detail.copy()
        detail.insert(0, 'imputation', imputation_index)
        detail['cox_fit_warning_count'] = len(captured)
        detail['cox_fit_warning_messages_json'] = json.dumps([str(item.message) for item in captured], sort_keys=True)
        detailed_rows.append(detail)
    by_imputation = pd.concat(detailed_rows, ignore_index=True)
    predictors = list(main_module.PREDICTOR_COLUMNS)
    expected_keys = {(imputation, predictor) for imputation in range(1, len(completed_training_sets) + 1) for predictor in predictors}
    observed_keys = set(zip(by_imputation['imputation'].astype(int), by_imputation['variable'].astype(str)))
    require(observed_keys == expected_keys, 'Cox imputation-level coefficient grid is incomplete.')
    ph_by_imputation = by_imputation[['imputation', 'Variable', 'variable', 'ph_test_statistic', 'ph_test_p', 'cox_fit_warning_count', 'cox_fit_warning_messages_json']].sort_values(['imputation', 'variable']).reset_index(drop=True)
    m_imputations = len(completed_training_sets)
    pooling_rows: list[dict[str, Any]] = []
    for predictor in predictors:
        subset = by_imputation.loc[by_imputation['variable'].eq(predictor)]
        coefficients = subset['beta'].to_numpy(dtype=float)
        within_variances = np.square(subset['SE_beta'].to_numpy(dtype=float))
        within = float(np.mean(within_variances))
        between = float(np.var(coefficients, ddof=1))
        between_adjusted = float((1.0 + 1.0 / m_imputations) * between)
        total = float(within + between_adjusted)
        relative_increase_variance = float(between_adjusted / within) if within > 0.0 else float('inf') if between_adjusted > 0.0 else 0.0
        lambda_missing_information = float(between_adjusted / total) if total > 0.0 else 0.0
        if relative_increase_variance > 0.0:
            rubin_df = float((m_imputations - 1) * (1.0 + 1.0 / relative_increase_variance) ** 2) if np.isfinite(relative_increase_variance) else float(m_imputations - 1)
            finite_df_fraction_missing_information = float(lambda_missing_information + 2.0 / (rubin_df + 3.0) * (1.0 - lambda_missing_information))
        else:
            rubin_df = float('inf')
            finite_df_fraction_missing_information = 0.0
        pooled_beta = float(np.mean(coefficients))
        pooled_se = float(np.sqrt(total))
        t_statistic = pooled_beta / pooled_se if pooled_se > 0.0 else np.nan
        critical = float(stats.t.ppf(0.975, df=rubin_df)) if np.isfinite(rubin_df) else 1.959963984540054
        p_value = float(2.0 * stats.t.sf(abs(t_statistic), df=rubin_df)) if np.isfinite(t_statistic) and np.isfinite(rubin_df) else float(2.0 * stats.norm.sf(abs(t_statistic))) if np.isfinite(t_statistic) else np.nan
        ci_low = pooled_beta - critical * pooled_se
        ci_high = pooled_beta + critical * pooled_se
        pooling_rows.append({'Variable': str(subset['Variable'].iloc[0]), 'variable': predictor, 'beta': pooled_beta, 'HR': float(np.exp(pooled_beta)), 'HR_95CI_lower': float(np.exp(ci_low)), 'HR_95CI_upper': float(np.exp(ci_high)), 'SE_beta': pooled_se, 't_statistic': t_statistic, 'rubin_df': rubin_df, 'p_value': p_value, 'ph_test_p_min': float(subset['ph_test_p'].min()), 'ph_test_p_median': float(subset['ph_test_p'].median()), 'ph_test_p_max': float(subset['ph_test_p'].max()), 'ph_test_n_below_0_05': int((subset['ph_test_p'] < 0.05).sum()), 'ph_test_n_imputations': int(len(subset)), 'within_imputation_variance': within, 'between_imputation_variance': between, 'total_variance': total, 'relative_increase_in_variance': relative_increase_variance, 'lambda_missing_information_large_sample': lambda_missing_information, 'fraction_missing_information_finite_df': finite_df_fraction_missing_information, 'monte_carlo_se_pooled_beta': float(np.sqrt(between / m_imputations)), 'monte_carlo_se_fraction_of_pooled_se': float(np.sqrt(between / m_imputations) / pooled_se if pooled_se > 0.0 else 0.0), 'n_imputations': int(m_imputations)})
    pooled = pd.DataFrame(pooling_rows)
    pooled['_order'] = pooled['variable'].map({predictor: index for index, predictor in enumerate(predictors)})
    pooled = pooled.sort_values('_order').drop(columns='_order').reset_index(drop=True)
    require(len(pooled) == len(predictors), 'Pooled Cox table has the wrong row count.')
    require(np.allclose(pooled['SE_beta'].to_numpy(dtype=float) ** 2, pooled['total_variance'].to_numpy(dtype=float), rtol=1e-12, atol=1e-14), 'Rubin pooled total variances do not reconcile with the pooled standard errors.')
    required_numeric = ['beta', 'HR', 'HR_95CI_lower', 'HR_95CI_upper', 'SE_beta', 'p_value', 'within_imputation_variance', 'between_imputation_variance', 'total_variance', 'relative_increase_in_variance', 'lambda_missing_information_large_sample', 'fraction_missing_information_finite_df', 'monte_carlo_se_pooled_beta', 'monte_carlo_se_fraction_of_pooled_se', 'ph_test_p_min', 'ph_test_p_median', 'ph_test_p_max']
    require(np.isfinite(pooled[required_numeric].apply(pd.to_numeric, errors='coerce').to_numpy(float)).all(), 'Pooled Cox table contains a non-finite required value.')
    require(np.all(pooled['HR_95CI_lower'].to_numpy(float) <= pooled['HR'].to_numpy(float)) and np.all(pooled['HR'].to_numpy(float) <= pooled['HR_95CI_upper'].to_numpy(float)), 'A pooled Cox confidence interval does not contain its hazard ratio.')
    if enforce_precision_gate:
        require(m_imputations == DEFAULT_M_IMPUTATIONS and float(pooled['monte_carlo_se_fraction_of_pooled_se'].max()) <= 0.1, 'Fifty imputations do not provide the prespecified Monte Carlo precision; increase the locked imputation count and rerun the final stage.')
    return (pooled, ph_by_imputation, by_imputation)

def cox_calibration_slope_at_horizon(data: pd.DataFrame, predicted_event_probability: Sequence[float], main_module: Any, horizon_months: float) -> float:
    from lifelines import CoxPHFitter
    probabilities = np.asarray(predicted_event_probability, dtype=float).reshape(-1)
    require(len(probabilities) == len(data), 'Calibration-slope inputs are misaligned.')
    require(np.isfinite(probabilities).all() and np.all((probabilities >= 0.0) & (probabilities <= 1.0)), 'Calibration-slope probabilities are invalid.')
    probabilities = np.clip(probabilities, 1e-06, 1.0 - 1e-06)
    complementary_log_log_risk = np.log(-np.log1p(-probabilities))
    observed_time = data[main_module.TIME_COL].to_numpy(dtype=float)
    observed_event = data[main_module.EVENT_COL].to_numpy(dtype=int)
    calibration_frame = pd.DataFrame({main_module.TIME_COL: np.minimum(observed_time, float(horizon_months)), main_module.EVENT_COL: ((observed_event == 1) & (observed_time <= float(horizon_months))).astype(int), 'complementary_log_log_risk': complementary_log_log_risk})
    fitter = CoxPHFitter()
    fitter.fit(calibration_frame, duration_col=main_module.TIME_COL, event_col=main_module.EVENT_COL)
    slope = float(fitter.params_['complementary_log_log_risk'])
    require(np.isfinite(slope), 'Cox calibration slope is non-finite.')
    return slope

def blocking_calibration_slope_warnings(captured: Sequence[warnings.WarningMessage]) -> list[warnings.WarningMessage]:
    from lifelines.exceptions import ConvergenceWarning
    return [item for item in captured if issubclass(item.category, (ConvergenceWarning, RuntimeWarning))]

def conditional_scalar_calibration_bootstrap(calibration_data: pd.DataFrame, main_module: Any, n_bootstrap: int, seed: int, horizon_months: float, expected_participants: int=EXPECTED_PARTICIPANTS, expected_minimum_appearances: int=21, expected_maximum_appearances: int=57) -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    study_id = str(main_module.STUDY_ID_COL)
    time_col = str(main_module.TIME_COL)
    event_col = str(main_module.EVENT_COL)
    required = {study_id, time_col, event_col, 'n_held_out_appearances', *EXPECTED_MODELS}
    require(set(calibration_data.columns) == required, 'Participant-level calibration input must contain only ID, outcomes, coverage, and seven model risks.')
    require(len(calibration_data) == expected_participants and calibration_data[study_id].nunique() == expected_participants and (not calibration_data[study_id].isna().any()), 'Participant-level calibration input must contain one row per participant.')
    calibration_data = calibration_data.assign(_deterministic_id_sort=calibration_data[study_id].astype(str)).sort_values('_deterministic_id_sort', kind='mergesort').drop(columns='_deterministic_id_sort').reset_index(drop=True)
    require(int(calibration_data['n_held_out_appearances'].min()) == expected_minimum_appearances and int(calibration_data['n_held_out_appearances'].max()) == expected_maximum_appearances, 'Participant-level calibration coverage is outside its locked range.')
    n = len(calibration_data)
    rng = np.random.default_rng(int(seed + CALIBRATION_BOOTSTRAP_SEED_OFFSET))
    bootstrap_indices = rng.integers(0, n, size=(n_bootstrap, n), endpoint=False)
    summary_rows: list[dict[str, Any]] = []
    long_rows: list[dict[str, Any]] = []
    completeness_rows: list[dict[str, Any]] = []
    minimum_successful = int(math.ceil(REQUIRED_BOOTSTRAP_SUCCESS_FRACTION * n_bootstrap))
    resample_metadata: list[dict[str, Any]] = []
    original_times = calibration_data[time_col].to_numpy(dtype=float)
    original_events = calibration_data[event_col].to_numpy(dtype=int)
    require(np.isfinite(original_times).all() and np.all(original_times > 0.0) and np.isin(original_events, [0, 1]).all(), 'Participant calibration outcomes are invalid.')
    for bootstrap_index, row_indices in enumerate(bootstrap_indices, start=1):
        sampled_times = original_times[row_indices]
        sampled_events = original_events[row_indices]
        events_by_horizon = int(np.sum((sampled_events == 1) & (sampled_times <= horizon_months)))
        require(events_by_horizon > 0, f'Bootstrap replicate {bootstrap_index} has no event by 120 months.')
        resample_metadata.append({'bootstrap_replicate': bootstrap_index, 'resample_sha256': hashlib.sha256(np.asarray(row_indices, dtype='<i8').tobytes()).hexdigest(), 'n_unique_participants_sampled': int(np.unique(row_indices).size), 'events_by_horizon': events_by_horizon})
    for model_name in EXPECTED_MODELS:
        probabilities = pd.to_numeric(calibration_data[model_name], errors='raise').to_numpy(float)
        require(np.isfinite(probabilities).all() and np.all((probabilities >= 0.0) & (probabilities <= 1.0)), f'Participant-level probabilities are invalid for {model_name}.')
        require(np.unique(probabilities).size >= 2, f'Participant-level probabilities have insufficient variation for {model_name}.')
        point_oe = main_module.calibration_oe_from_probs(calibration_data, probabilities, time_point=horizon_months)
        with warnings.catch_warnings(record=True) as point_slope_warnings:
            warnings.simplefilter('always')
            point_slope_value = cox_calibration_slope_at_horizon(calibration_data, probabilities, main_module, horizon_months)
        blocking_point_warnings = blocking_calibration_slope_warnings(point_slope_warnings)
        require(not blocking_point_warnings, f'Full-sample calibration slope emitted a convergence/numerical warning for {model_name}: ' + ' | '.join((str(item.message) for item in blocking_point_warnings)))
        point_oe_value = float(point_oe['calibration_oe_ratio'])
        require(np.isfinite([point_oe_value, point_slope_value]).all(), f'Full-sample scalar calibration is invalid for {model_name}.')
        successful_oe: list[float] = []
        successful_slope: list[float] = []
        slope_warning_replicates = 0
        for bootstrap_index, row_indices in enumerate(bootstrap_indices, start=1):
            sample = calibration_data.iloc[row_indices].reset_index(drop=True)
            sample_probabilities = probabilities[row_indices]
            metadata = resample_metadata[bootstrap_index - 1]
            oe_value = np.nan
            slope_value = np.nan
            oe_success = False
            slope_success = False
            oe_error_type = ''
            oe_error_message = ''
            slope_error_type = ''
            slope_error_message = ''
            slope_warning_messages: list[str] = []
            try:
                oe = main_module.calibration_oe_from_probs(sample, sample_probabilities, time_point=horizon_months)
                oe_value = float(oe['calibration_oe_ratio'])
                if not np.isfinite(oe_value):
                    raise ValueError('non-finite O:E estimate')
            except Exception as exc:
                oe_error_type = type(exc).__name__
                oe_error_message = str(exc)[:500]
            else:
                oe_success = True
                successful_oe.append(oe_value)
            try:
                if np.unique(sample_probabilities).size < 2:
                    raise ValueError('insufficient prediction variation')
                with warnings.catch_warnings(record=True) as captured_slope_warnings:
                    warnings.simplefilter('always')
                    slope_value = cox_calibration_slope_at_horizon(sample, sample_probabilities, main_module, horizon_months)
                slope_warning_messages = [str(item.message)[:500] for item in captured_slope_warnings]
                blocking_warnings = blocking_calibration_slope_warnings(captured_slope_warnings)
                if blocking_warnings:
                    raise RuntimeError('calibration-slope convergence/numerical warning: ' + ' | '.join((str(item.message) for item in blocking_warnings)))
                if not np.isfinite(slope_value):
                    raise ValueError('non-finite calibration-slope estimate')
            except Exception as exc:
                slope_error_type = type(exc).__name__
                slope_error_message = str(exc)[:500]
            else:
                slope_success = True
                successful_slope.append(slope_value)
                if slope_warning_messages:
                    slope_warning_replicates += 1
            long_rows.append({'model': model_name, 'bootstrap_replicate': bootstrap_index, 'resample_sha256': metadata['resample_sha256'], 'n_unique_participants_sampled': metadata['n_unique_participants_sampled'], 'events_by_horizon': metadata['events_by_horizon'], 'calibration_oe_ratio': oe_value, 'calibration_oe_success': oe_success, 'calibration_oe_error_type': oe_error_type, 'calibration_oe_error_message': oe_error_message, 'calibration_slope': slope_value, 'calibration_slope_success': slope_success, 'calibration_slope_error_type': slope_error_type, 'calibration_slope_error_message': slope_error_message, 'calibration_slope_warning_count': len(slope_warning_messages), 'calibration_slope_warning_messages_json': json.dumps(slope_warning_messages, sort_keys=True), 'resampling_unit': 'participant'})
        successful_oe_count = len(successful_oe)
        successful_slope_count = len(successful_slope)
        require(successful_oe_count >= minimum_successful and successful_slope_count >= minimum_successful, f'Too few scalar calibration bootstraps succeeded for {model_name}: O:E {successful_oe_count}/{n_bootstrap}; slope {successful_slope_count}/{n_bootstrap}.')
        require(successful_oe_count == n_bootstrap and successful_slope_count == n_bootstrap, f'The final release requires all {n_bootstrap} scalar calibration bootstraps to succeed for {model_name}; observed O:E {successful_oe_count} and slope {successful_slope_count}.')
        oe_low, oe_high = np.percentile(np.asarray(successful_oe), [2.5, 97.5], method='linear')
        slope_low, slope_high = np.percentile(np.asarray(successful_slope), [2.5, 97.5], method='linear')
        summary_rows.append({'model': model_name, 'n_participants': n, 'total_held_out_prediction_rows': int(calibration_data['n_held_out_appearances'].sum()), 'minimum_held_out_appearances': int(calibration_data['n_held_out_appearances'].min()), 'maximum_held_out_appearances': int(calibration_data['n_held_out_appearances'].max()), 'horizon_months': float(horizon_months), 'calibration_oe_ratio': point_oe_value, 'calibration_oe_ratio_bootstrap_95ci_low': float(oe_low), 'calibration_oe_ratio_bootstrap_95ci_high': float(oe_high), 'calibration_slope': point_slope_value, 'calibration_slope_definition': 'cox_coefficient_for_complementary_log_log_predicted_risk_with_followup_administratively_censored_at_horizon', 'calibration_slope_bootstrap_95ci_low': float(slope_low), 'calibration_slope_bootstrap_95ci_high': float(slope_high), 'observed_km_risk': float(point_oe['calibration_observed_km_risk']), 'mean_predicted_risk': float(point_oe['calibration_mean_predicted_risk']), 'n_bootstrap_requested': int(n_bootstrap), 'n_bootstrap_successful_oe': int(successful_oe_count), 'n_bootstrap_failed_oe': int(n_bootstrap - successful_oe_count), 'n_bootstrap_successful_slope': int(successful_slope_count), 'n_bootstrap_failed_slope': int(n_bootstrap - successful_slope_count), 'n_bootstrap_slope_warning_replicates': int(slope_warning_replicates), 'calibration_slope_full_sample_warning_count': int(len(point_slope_warnings)), 'bootstrap_seed': int(seed + CALIBRATION_BOOTSTRAP_SEED_OFFSET), 'bootstrap_interval_method': 'percentile_2.5th_and_97.5th_percentiles_numpy_linear', 'resampling_unit': 'participant', 'calibration_unit': 'one_averaged_held_out_prediction_per_participant', 'calibration_estimand': 'cross_fitted_repeated_split_ensemble_not_one_final_deployable_model', 'interval_interpretation': 'conditional_on_saved_participant_averaged_predictions;_excludes_model_development_and_split_stream_selection_uncertainty'})
        for metric, successful_count in [('calibration_oe_ratio', successful_oe_count), ('cox_calibration_slope_at_120_months', successful_slope_count)]:
            completeness_rows.append({'model': model_name, 'metric': metric, 'n_bootstrap_requested': int(n_bootstrap), 'minimum_successful_required': int(minimum_successful), 'n_bootstrap_successful': int(successful_count), 'n_bootstrap_failed': int(n_bootstrap - successful_count), 'n_nonblocking_warning_replicates': int(slope_warning_replicates if metric == 'cox_calibration_slope_at_120_months' else 0), 'complete': successful_count == n_bootstrap, 'bootstrap_seed': int(seed + CALIBRATION_BOOTSTRAP_SEED_OFFSET), 'resampling_unit': 'participant', 'interval_method': 'percentile_2.5th_and_97.5th_percentiles_numpy_linear'})
        print(f'  {model_name}: all {n_bootstrap} conditional O:E and slope bootstraps succeeded.', flush=True)
    summary = pd.DataFrame(summary_rows)
    long = pd.DataFrame(long_rows)
    completeness = pd.DataFrame(completeness_rows)
    require(set(summary['model']) == set(EXPECTED_MODELS), 'Scalar calibration summary is incomplete.')
    require(set(long['model']) == set(EXPECTED_MODELS) and len(long) == n_bootstrap * len(EXPECTED_MODELS) and (not long.duplicated(['model', 'bootstrap_replicate']).any()), 'Scalar calibration bootstrap audit is incomplete or duplicated.')
    for _, replicate_rows in long.groupby('bootstrap_replicate'):
        require(replicate_rows['resample_sha256'].nunique() == 1 and replicate_rows['n_unique_participants_sampled'].nunique() == 1 and (replicate_rows['events_by_horizon'].nunique() == 1), 'Models did not use the same participant resample within a bootstrap replicate.')
    return (summary, long, completeness)

def verify_scalar_bootstrap_summary(summary: pd.DataFrame, long: pd.DataFrame) -> None:
    for row in summary.to_dict('records'):
        model_name = str(row['model'])
        sub = long.loc[long['model'].eq(model_name)]
        require(len(sub) == int(row['n_bootstrap_requested']), f'Bootstrap row count changed for {model_name}.')
        for metric, success_name, summary_count_name, low_name, high_name in [('calibration_oe_ratio', 'calibration_oe_success', 'n_bootstrap_successful_oe', 'calibration_oe_ratio_bootstrap_95ci_low', 'calibration_oe_ratio_bootstrap_95ci_high'), ('calibration_slope', 'calibration_slope_success', 'n_bootstrap_successful_slope', 'calibration_slope_bootstrap_95ci_low', 'calibration_slope_bootstrap_95ci_high')]:
            success = sub[success_name].astype(str).str.lower().isin({'true', '1'})
            require(int(success.sum()) == int(row[summary_count_name]), f'Successful bootstrap count does not reconcile for {model_name} {metric}.')
            low, high = np.percentile(sub.loc[success, metric].to_numpy(dtype=float), [2.5, 97.5], method='linear')
            require(math.isclose(float(row[low_name]), float(low), rel_tol=1e-12, abs_tol=1e-12) and math.isclose(float(row[high_name]), float(high), rel_tol=1e-12, abs_tol=1e-12), f'Bootstrap percentiles do not reconcile for {model_name} {metric}.')

def format_p_value(value: float) -> str:
    if value < 0.001:
        return '<0.001'
    return f'{value:.3f}'

def make_formatted_table2(pooled: pd.DataFrame) -> pd.DataFrame:
    rows: list[dict[str, str]] = []
    for row in pooled.to_dict('records'):
        rows.append({'Variable': str(row['Variable']), 'HR (95% CI)': f"{float(row['HR']):.3f} ({float(row['HR_95CI_lower']):.3f} to {float(row['HR_95CI_upper']):.3f})", 'P value': format_p_value(float(row['p_value'])), 'PH test P, median [range]': f"{float(row['ph_test_p_median']):.3f} [{float(row['ph_test_p_min']):.3f} to {float(row['ph_test_p_max']):.3f}]"})
    return pd.DataFrame(rows)

def make_formatted_table3(original_formatted: pd.DataFrame, calibration_summary: pd.DataFrame) -> pd.DataFrame:
    output = original_formatted.copy()
    calibration = calibration_summary.set_index('model')
    for row_index, row in output.iterrows():
        model_name = str(row['Model'])
        if model_name not in calibration.index:
            output.at[row_index, 'Calibration_slope_120m'] = 'NA'
            output.at[row_index, 'O_E_120m'] = 'NA'
            continue
        cal = calibration.loc[model_name]
        output.at[row_index, 'Calibration_slope_120m'] = f"{float(cal['calibration_slope']):.3f} ({float(cal['calibration_slope_bootstrap_95ci_low']):.3f} to {float(cal['calibration_slope_bootstrap_95ci_high']):.3f})"
        output.at[row_index, 'O_E_120m'] = f"{float(cal['calibration_oe_ratio']):.3f} ({float(cal['calibration_oe_ratio_bootstrap_95ci_low']):.3f} to {float(cal['calibration_oe_ratio_bootstrap_95ci_high']):.3f})"
    return output.rename(columns={'Calibration_slope_120m': 'Cox_calibration_slope_120m'})

def set_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn('w:shd'))
    if shading is None:
        shading = OxmlElement('w:shd')
        properties.append(shading)
    shading.set(qn('w:fill'), fill)

def style_docx_table(table: Any, font_size: float=8.0) -> None:
    table.style = 'Table Grid'
    table.autofit = True
    for column_index, cell in enumerate(table.rows[0].cells):
        set_cell_shading(cell, 'D9EAF7')
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(font_size)
        cell.vertical_alignment = 1
    for row in table.rows[1:]:
        for column_index, cell in enumerate(row.cells):
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)

def add_frame_to_docx(document: Document, frame: pd.DataFrame, title: str, note: str | None, font_size: float=8.0) -> None:
    heading = document.add_paragraph()
    heading.style = document.styles['Heading 2']
    heading.add_run(title).bold = True
    table = document.add_table(rows=1, cols=len(frame.columns))
    for column_index, column in enumerate(frame.columns):
        table.cell(0, column_index).text = str(column)
    for row in frame.itertuples(index=False, name=None):
        cells = table.add_row().cells
        for column_index, value in enumerate(row):
            if pd.isna(value):
                text_value = 'NA'
            else:
                text_value = str(value)
            cells[column_index].text = text_value
    style_docx_table(table, font_size=font_size)
    if note:
        paragraph = document.add_paragraph()
        run = paragraph.add_run(note)
        run.italic = True
        run.font.size = Pt(8.5)

def prepare_document(title: str) -> Document:
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = (section.page_height, section.page_width)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    paragraph = document.add_paragraph()
    paragraph.style = document.styles['Title']
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(title)
    return document

def atomic_save_docx(document: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    handle, temporary_name = tempfile.mkstemp(prefix=f'.{path.name}.', suffix='.tmp.docx', dir=path.parent)
    os.close(handle)
    temporary = Path(temporary_name)
    try:
        document.save(temporary)
        reopened = Document(temporary)
        require(len(reopened.tables) > 0, f'DOCX contains no tables: {path.name}')
        os.replace(temporary, path)
    except Exception:
        temporary.unlink(missing_ok=True)
        raise

def write_results_documents(outdir: Path, main_output_dir: Path, formatted_table2: pd.DataFrame, formatted_table3: pd.DataFrame) -> list[Path]:
    docx_dir = outdir / 'docx_tables'
    table2_note = "Missing predictors in the fixed training set were multiply imputed in 50 independently seeded chained-equation datasets using predictive mean matching with 10 donors. The imputation matrix contained the baseline predictors, ASCVD event indicator and training-set Nelson-Aalen cumulative hazard. Outcome-related auxiliary variables were used only for this coefficient analysis and not for prediction. Log-hazard coefficients and their within-imputation variances were pooled with Rubin's rules and then exponentiated. Coefficient P values and PH-test P values are exploratory; no multiplicity-controlled confirmatory claims or variable rankings are made. PH-test P values are summarized descriptively across imputations and are not pooled or treated as a global test."
    table3_note = "Harrell's C, Uno's C, Brier score and IBS are means with empirical 2.5th-97.5th percentile ranges across 200 overlapping held-out splits; these ranges are not confidence intervals. O:E and the Cox calibration slope at 120 months use one averaged held-out risk per participant and are shown with percentile participant-bootstrap 95% intervals. These intervals are conditional on the saved predictions and exclude model-development and split-stream-selection uncertainty."
    table2_doc = prepare_document('Table 2')
    add_frame_to_docx(table2_doc, formatted_table2, 'Table 2. Multivariable Cox proportional hazards model for incident ASCVD in the fixed training set.', table2_note, font_size=8.5)
    table2_path = docx_dir / 'Table_2.docx'
    atomic_save_docx(table2_doc, table2_path)
    table3_doc = prepare_document('Table 3')
    add_frame_to_docx(table3_doc, formatted_table3, 'Table 3. Performance across 200 stratified random 80/20 splits.', table3_note, font_size=7.5)
    table3_path = docx_dir / 'Table_3.docx'
    atomic_save_docx(table3_doc, table3_path)
    combined = prepare_document('ASCVD Survival Models - Tables')
    subtitle = combined.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run('Primary 200-split analysis with fixed-split Cox inference and scalar calibration intervals').italic = True
    table_specs: list[tuple[pd.DataFrame, str, str | None, float]] = [(pd.read_csv(main_output_dir / 'tables' / 'table1_fixed_split.csv', keep_default_na=False), 'Table 1. Baseline characteristics of the overall cohort and the fixed training and test sets.', 'Values for variables with missing data are averaged across the 10 predictor-only fixed-split prediction datasets. Event counts and follow-up summaries are from the raw filtered dataset.', 8.0), (formatted_table2, 'Table 2. Multivariable Cox proportional hazards model for incident ASCVD in the fixed training set.', table2_note, 8.0), (formatted_table3, 'Table 3. Performance across 200 stratified random 80/20 splits.', table3_note, 7.0), (pd.read_csv(main_output_dir / 'tables' / 'supplementary_table_s1_missingness.csv'), 'Supplementary Table S1. Variables with missing baseline data before imputation.', None, 8.0), (pd.read_csv(main_output_dir / 'tables' / 'supplementary_table_s2_tuning_strategies.csv', keep_default_na=False), 'Supplementary Table S2. Prespecified tuning strategies used in internal validation.', "Six configurable absolute-risk models were selected by training-only validation IBS from 12 to 120 months. Survival SVM used Harrell's C. CoxPH and spline Cox had fixed specifications.", 6.5), (pd.read_csv(main_output_dir / 'tables' / 'repeated_200_paired_differences_vs_coxph_summary.csv', keep_default_na=False), 'Supplementary Table S3. Paired performance differences relative to CoxPH.', 'Differences are paired within the same held-out split. Empirical percentile ranges are descriptive and are not confidence intervals. No hypothesis tests or P values were calculated.', 6.5), (pd.read_csv(main_output_dir / 'tables' / 'supplementary_table_s5_fixed_split_metrics.csv', keep_default_na=False), 'Supplementary Table S5. Illustrative performance in the fixed train-test split.', 'This fixed partition is split 1 of the repeated stream and remains supplementary. Prediction metrics use the original predictor-only imputation pipeline and are unaffected by the outcome-aware coefficient imputation used for Table 2.', 7.5)]
    for frame, title, note, font_size in table_specs:
        add_frame_to_docx(combined, frame, title, note, font_size=font_size)
    combined_path = docx_dir / 'ASCVD_survival_models_tables.docx'
    atomic_save_docx(combined, combined_path)
    require(len(Document(combined_path).tables) == len(table_specs), 'Combined DOCX is incomplete.')
    return [table2_path, table3_path, combined_path]

def write_results_methods_notes(source_path: Path, destination_path: Path) -> None:
    text = source_path.read_text(encoding='utf-8')
    replacements = {'- No follow-up variable is supplied as a predictor or to an imputation model.': '- No outcome or follow-up variable is supplied as a predictor or to any imputation model used for prediction, tuning, held-out evaluation, calibration predictions or SHAP. For the supplementary fixed-training Cox coefficient table only, the event indicator and Nelson-Aalen cumulative hazard are included as auxiliary imputation variables, following White and Royston.', '- It retains 10 predictor-only stochastic imputations for coefficients, proportional-hazards checks, SHAP and fixed-split calibration.': '- The fixed-split prediction, SHAP and calibration components retain 10 predictor-only stochastic imputations. The separate Cox coefficient and proportional-hazards analysis uses 50 independently seeded, training-only chained-equation imputations with predictive mean matching (10 donors), containing the 15 predictors plus the event indicator and Nelson-Aalen cumulative hazard as auxiliary variables.', '- Each imputer is fit only on training predictors and then applied to held-out predictors.': '- Every prediction-pipeline imputer is fit only on training predictors and then applied to held-out predictors. The outcome-aware inferential imputer is fit and used only within the fixed training set; it is never applied to the held-out set.', '- O:E, calibration slope, five-group Kaplan-Meier calibration and flexible IPCW spline calibration are calculated once from the participant-level averaged predictions.': '- O:E, calibration slope, five-group Kaplan-Meier calibration and flexible IPCW spline calibration are calculated once from the participant-level averaged predictions. O:E and slope are reported with percentile 95% intervals from 1,000 participant bootstrap samples.', '- Flexible-calibration pointwise bands use participant bootstrap resampling and are conditional on the averaged predictions.': '- Scalar O:E and slope intervals and flexible-calibration pointwise bands use participant bootstrap resampling and are conditional on the averaged predictions.', '- Each split uses one seeded posterior-sampling completion as a fixed prediction-pipeline rule, not conventional multiple-imputation pooling. Split-level dispersion therefore includes allocation changes and the seeded imputation draw. Ten-imputation coefficient pooling is confined to the supplementary fixed split.': '- Each split uses one seeded posterior-sampling completion as a fixed prediction-pipeline rule, not conventional multiple-imputation pooling. Split-level dispersion therefore includes allocation changes and the seeded imputation draw. The separate supplementary coefficient/PH analysis uses 50 outcome-aware MICE/PMM imputations in the fixed training set only.', '- No P values or formal hypothesis tests are calculated, so no multiplicity adjustment is applied.': '- No P values or formal hypothesis tests compare predictive models or test calibration; therefore no multiplicity adjustment is claimed for those descriptive comparisons. Table 2 retains conventional pooled Cox coefficient P values, and PH-test P values are complete imputation-level diagnostics summarized descriptively rather than pooled.'}
    for old, new in replacements.items():
        require(text.count(old) == 1, f'Could not update the methods note: {old}')
        text = text.replace(old, new)
    header = 'FINAL ANALYSIS METHODS\n======================\nThis note describes the primary 200-split prediction analysis, the fixed-split Cox coefficient analysis, and the scalar calibration intervals.\n\n'
    addendum = 'Estimands and uncertainty\n--------------------------\n- The Cox calibration slope is the coefficient of complementary-log-log 120-month predicted risk in a Cox model administratively censored at 120 months. It is a hazard-scale calibration diagnostic and assumes a proportional score effect over that horizon.\n- No Cox-model P value for the null slope of zero is reported because it does not test the calibration target slope of one.\n- Scalar O:E and slope intervals are linear-interpolation percentile intervals from the prespecified 1,000 participant bootstrap replicates. They condition on the saved participant-averaged predictions.\n- The fixed-training Cox imputation uses statsmodels MICEData 0.14.4 with patsy 1.0.1, 20 prespecified cycles per independently seeded chain and 10 PMM donors. The cycle count is a fixed implementation rule, not a generic claim that a stochastic chain has converged.\n- The imputation analysis assumes missing at random conditional on the observed predictors, event indicator and training-set Nelson-Aalen H(T). It does not protect against missing-not-at-random mechanisms.\n- The fixed-split Cox coefficient and PH results are conditional on one outcome-stratified training split (992 participants, 92 events) with 15 terms. They are associational, not causal. PH tests have limited power; an absence of small P values does not establish proportional hazards.\n- The 15 pooled coefficient P values and the dependent PH diagnostics are exploratory/descriptive. No multiplicity-controlled confirmatory claim or variable ranking is made; effect estimates and confidence intervals are the primary coefficient presentation.\n- Fifty imputations must pass an ex-post precision gate requiring the Monte Carlo SE of every pooled log coefficient to be no more than 10% of its pooled standard error.\n\n'
    atomic_write_text(header + addendum + text, destination_path)

def output_paths(outdir: Path) -> dict[str, Path]:
    return {'pooled_cox': outdir / 'tables' / 'table2_cox_hazard_ratios_fixed_split_outcome_aware.csv', 'cox_by_imputation': outdir / 'tables' / 'cox_coefficients_by_imputation_outcome_aware.csv', 'ph': outdir / 'tables' / 'proportional_hazards_tests_by_imputation_outcome_aware.csv', 'imputation_diagnostics': outdir / 'tables' / 'fixed_split_outcome_aware_imputation_diagnostics.csv', 'calibration_summary': outdir / 'tables' / 'repeated_200_scalar_calibration_with_conditional_95ci.csv', 'calibration_bootstrap': outdir / 'tables' / 'repeated_200_scalar_calibration_bootstrap_estimates.csv', 'calibration_completeness': outdir / 'metadata' / 'scalar_calibration_bootstrap_completeness.csv', 'formatted_table2': outdir / 'tables' / 'table2_analysis_formatted.csv', 'formatted_table3': outdir / 'tables' / 'repeated_200_table3_formatted_with_calibration_95ci.csv', 'methods': outdir / 'metadata' / 'methods_notes.txt', 'method_metadata': outdir / 'metadata' / 'final_method_metadata.json', 'completion': outdir / 'analysis_complete.json'}

def validate_outputs(outdir: Path, expected_identity_sha256: str | None=None) -> None:
    paths = output_paths(outdir)
    for name, path in paths.items():
        require(path.is_file() and path.stat().st_size > 0, f'Final-stage output is missing: {name}')
    completion = json.loads(paths['completion'].read_text(encoding='utf-8'))
    require(completion.get('schema_version') == SCHEMA_VERSION and completion.get('status') == 'complete', 'Final-stage completion marker is invalid.')
    configuration = completion.get('configuration')
    require(isinstance(configuration, dict) and completion.get('identity_sha256') == hashlib.sha256(canonical_json(configuration).encode('utf-8')).hexdigest(), 'Final-stage completion configuration and identity do not reconcile.')
    if expected_identity_sha256 is not None:
        require(completion.get('identity_sha256') == expected_identity_sha256, 'Final-stage completion identity changed.')
    pooled = pd.read_csv(paths['pooled_cox'], float_precision='round_trip')
    require(len(pooled) == 15 and pooled['variable'].nunique() == 15, 'Cox table is incomplete.')
    require(np.isfinite(pooled[['HR', 'HR_95CI_lower', 'HR_95CI_upper', 'p_value']].apply(pd.to_numeric, errors='coerce').to_numpy(float)).all(), 'Cox table has invalid numeric values.')
    ph = pd.read_csv(paths['ph'], float_precision='round_trip')
    require(len(ph) == DEFAULT_M_IMPUTATIONS * 15 and (not ph.duplicated(['imputation', 'variable']).any()), 'PH table is incomplete.')
    diagnostics = pd.read_csv(paths['imputation_diagnostics'], keep_default_na=False)
    require(len(diagnostics) == DEFAULT_M_IMPUTATIONS and diagnostics['missing_predictor_cells_after'].astype(int).eq(0).all() and diagnostics['event_auxiliary_preserved'].astype(str).str.lower().isin({'true', '1'}).all() and diagnostics['nelson_aalen_auxiliary_preserved'].astype(str).str.lower().isin({'true', '1'}).all() and diagnostics['observed_predictor_cells_preserved'].astype(str).str.lower().isin({'true', '1'}).all() and diagnostics['all_imputed_values_are_observed_donors'].astype(str).str.lower().isin({'true', '1'}).all() and diagnostics['mice_cycles_completed'].astype(int).eq(DEFAULT_IMPUTE_MAX_ITER).all() and diagnostics['pmm_donors'].astype(int).eq(PMM_DONORS).all() and (diagnostics['imputation_seed'].astype(int).nunique() == DEFAULT_M_IMPUTATIONS), 'Outcome-aware imputation diagnostics failed.')
    calibration = pd.read_csv(paths['calibration_summary'], float_precision='round_trip')
    bootstrap = pd.read_csv(paths['calibration_bootstrap'], float_precision='round_trip')
    require(len(calibration) == len(EXPECTED_MODELS) and set(calibration['model']) == set(EXPECTED_MODELS), 'Scalar calibration summary is incomplete.')
    verify_scalar_bootstrap_summary(calibration, bootstrap)
    for prefix in ['calibration_oe_ratio', 'calibration_slope']:
        low = calibration[f'{prefix}_bootstrap_95ci_low'].to_numpy(dtype=float)
        high = calibration[f'{prefix}_bootstrap_95ci_high'].to_numpy(dtype=float)
        require(np.isfinite(low).all() and np.isfinite(high).all() and np.all(low <= high), f'Invalid bootstrap interval for {prefix}.')
    for name, expected_tables in [('Table_2.docx', 1), ('Table_3.docx', 1), ('ASCVD_survival_models_tables.docx', 7)]:
        path = outdir / 'docx_tables' / name
        require(path.is_file() and len(Document(path).tables) == expected_tables, f'DOCX failed validation: {name}')
    expected_artifacts = {path.relative_to(outdir).as_posix(): sha256_file(path) for path in [*[value for key, value in paths.items() if key != 'completion'], outdir / 'docx_tables' / 'Table_2.docx', outdir / 'docx_tables' / 'Table_3.docx', outdir / 'docx_tables' / 'ASCVD_survival_models_tables.docx']}
    require(completion.get('artifact_sha256') == dict(sorted(expected_artifacts.items())), 'Final-stage artifact inventory or a completed artifact hash changed.')
    methods = paths['methods'].read_text(encoding='utf-8')
    for phrase in ['event indicator and Nelson-Aalen cumulative hazard', '1,000 participant bootstrap samples', 'primary 200-split prediction analysis']:
        require(phrase in methods, f'Methods note lacks required wording: {phrase}')

def build_identity(args: argparse.Namespace, csv_path: Path, main_script_path: Path, main_output_dir: Path, secondary_output_dir: Path) -> dict[str, Any]:
    dependencies = {'raw_csv': csv_path, 'main_script': main_script_path, 'table1_fixed_split': main_output_dir / 'tables' / 'table1_fixed_split.csv', 'supplementary_table_s1_missingness': main_output_dir / 'tables' / 'supplementary_table_s1_missingness.csv', 'supplementary_table_s2_tuning_strategies': main_output_dir / 'tables' / 'supplementary_table_s2_tuning_strategies.csv', 'paired_differences_vs_coxph_summary': main_output_dir / 'tables' / 'repeated_200_paired_differences_vs_coxph_summary.csv', 'supplementary_table_s5_fixed_split_metrics': main_output_dir / 'tables' / 'supplementary_table_s5_fixed_split_metrics.csv', 'participant_mean_risk': main_output_dir / 'tables' / 'repeated_200_mean_held_out_risk_per_participant.csv', 'participant_calibration_summary': main_output_dir / 'tables' / 'repeated_200_participant_level_calibration_summary.csv', 'table3_formatted': main_output_dir / 'tables' / 'repeated_200_table3_formatted.csv', 'main_methods_notes': main_output_dir / 'metadata' / 'methods_notes.txt'}
    for label, path in dependencies.items():
        require(path.is_file(), f'Required final-stage dependency is missing ({label}): {path}')
    configuration = {'schema_version': SCHEMA_VERSION, 'purpose': 'fixed_cox_outcome_aware_imputation_and_scalar_calibration_intervals', 'source_sha256': sha256_file(Path(__file__).resolve()), 'dependency_sha256': {label: sha256_file(path) for label, path in sorted(dependencies.items())}, 'settings': {'seed': int(args.seed), 'm_imputations': int(args.m_imputations), 'impute_max_iter': int(args.impute_max_iter), 'imputer': 'statsmodels_MICEData_type_specific_predictive_mean_matching', 'pmm_donors': PMM_DONORS, 'inference_imputation_seed_offset': INFERENCE_IMPUTATION_SEED_OFFSET, 'calibration_bootstrap_reps': int(args.calibration_bootstrap_reps), 'calibration_bootstrap_seed_offset': CALIBRATION_BOOTSTRAP_SEED_OFFSET, 'horizon_months': float(args.horizon_months), 'required_bootstrap_success_fraction': REQUIRED_BOOTSTRAP_SUCCESS_FRACTION, 'prediction_pipeline_role': 'primary_200_split_predictor_only_imputation'}, 'additional_package_versions': dict(sorted(LOCKED_ADDITIONAL_PACKAGES.items()))}
    return {'configuration': configuration, 'identity_sha256': hashlib.sha256(canonical_json(configuration).encode('utf-8')).hexdigest()}

def validate_inputs(args: argparse.Namespace, main_module: Any, csv_path: Path, main_output_dir: Path) -> pd.DataFrame:
    data = main_module.load_dataset(str(csv_path))
    require(len(data) == EXPECTED_PARTICIPANTS and int(data[main_module.EVENT_COL].sum()) == EXPECTED_EVENTS, 'Filtered dataset counts changed.')
    train_indices, test_indices = main_module.primary_split(data, seed=args.seed, test_size=0.2)
    raw_training = data.iloc[train_indices].copy().reset_index(drop=True)
    raw_test = data.iloc[test_indices].copy().reset_index(drop=True)
    require(len(raw_training) == EXPECTED_FIXED_TRAINING_PARTICIPANTS and int(raw_training[main_module.EVENT_COL].sum()) == EXPECTED_FIXED_TRAINING_EVENTS and (len(raw_test) == EXPECTED_FIXED_TEST_PARTICIPANTS) and (int(raw_test[main_module.EVENT_COL].sum()) == EXPECTED_FIXED_TEST_EVENTS), 'Fixed split counts changed.')
    participant_path = main_output_dir / 'tables' / 'repeated_200_mean_held_out_risk_per_participant.csv'
    if participant_path.is_file():
        participant = pd.read_csv(participant_path, float_precision='round_trip')
        require(len(participant) == EXPECTED_PARTICIPANTS, 'Participant-averaged risk file has the wrong row count.')
    return raw_training

def run_functional_smoke_test(main_module: Any, seed: int) -> None:
    rng = np.random.default_rng(seed)
    n = 120
    time = rng.uniform(6.0, 180.0, size=n)
    event = rng.binomial(1, 0.28, size=n)
    event[:8] = 1
    synthetic = pd.DataFrame({main_module.STUDY_ID_COL: np.arange(n), main_module.TIME_COL: time, main_module.EVENT_COL: event})
    for predictor in main_module.PREDICTOR_COLUMNS:
        if predictor in main_module.BINARY_COLUMNS:
            synthetic[predictor] = rng.binomial(1, 0.45, size=n).astype(float)
        else:
            synthetic[predictor] = rng.normal(50.0, 8.0, size=n)
    for predictor in ['fhxcvd', 'hba1c', 'trig']:
        synthetic.loc[rng.choice(n, 8, replace=False), predictor] = np.nan
    completed, diagnostics, hazard = outcome_aware_fixed_training_imputations(synthetic, main_module, m_imputations=2, max_iter=3, seed=seed)
    require(len(completed) == 2 and len(diagnostics) == 2 and np.isfinite(hazard).all(), 'Outcome-aware imputation smoke test failed.')
    pooled, ph, by_imputation = collect_cox_details(completed, main_module, enforce_precision_gate=False)
    require(len(pooled) == 15 and len(ph) == 30 and (len(by_imputation) == 30), 'Cox pooling smoke test failed.')
    calibration = synthetic[[main_module.STUDY_ID_COL, main_module.TIME_COL, main_module.EVENT_COL]].copy()
    calibration['n_held_out_appearances'] = 30
    base = np.clip(0.03 + 0.002 * (synthetic['age'].fillna(synthetic['age'].median()) - 35), 0.01, 0.4)
    for model_index, model_name in enumerate(EXPECTED_MODELS):
        calibration[model_name] = np.clip(base * (0.94 + 0.02 * model_index), 0.005, 0.8)
    calibration_summary, bootstrap, completeness = conditional_scalar_calibration_bootstrap(calibration, main_module, n_bootstrap=20, seed=seed, horizon_months=120.0, expected_participants=n, expected_minimum_appearances=30, expected_maximum_appearances=30)
    verify_scalar_bootstrap_summary(calibration_summary, bootstrap)
    require(len(calibration_summary) == 7 and len(bootstrap) == 140 and (len(completeness) == 14), 'Conditional scalar-calibration bootstrap smoke test failed.')
    print('Final-stage functional smoke test passed.', flush=True)

def main() -> None:
    args = parse_args()
    validate_locked_args(args)
    csv_path = Path(args.csv).resolve()
    main_script_path = Path(args.main_script).resolve()
    main_output_dir = Path(args.main_output_dir).resolve()
    secondary_output_dir = Path(args.secondary_output_dir).resolve()
    outdir = Path(args.outdir).resolve()
    main_module = import_main_module(main_script_path)
    if args.functional_smoke_test:
        run_functional_smoke_test(main_module, args.seed)
        return
    raw_training = validate_inputs(args, main_module, csv_path, main_output_dir)
    if args.validate_only:
        print('Final-stage input and API validation passed.', flush=True)
        return
    identity = build_identity(args, csv_path, main_script_path, main_output_dir, secondary_output_dir)
    if args.resume and (outdir / 'analysis_complete.json').is_file():
        try:
            validate_outputs(outdir, expected_identity_sha256=identity['identity_sha256'])
        except Exception as exc:
            print(f'Existing final-stage outputs did not pass validation: {exc}', flush=True)
        else:
            print('Resumed the completed and validated final stage.', flush=True)
            return
    print('Creating outcome-aware fixed-training Cox imputations...', flush=True)
    completed_training_sets, imputation_diagnostics, hazard = outcome_aware_fixed_training_imputations(raw_training, main_module, m_imputations=args.m_imputations, max_iter=args.impute_max_iter, seed=args.seed)
    pooled_cox, ph_by_imputation, cox_by_imputation = collect_cox_details(completed_training_sets, main_module)
    participant_calibration_path = main_output_dir / 'tables' / 'repeated_200_mean_held_out_risk_per_participant.csv'
    participant_calibration = pd.read_csv(participant_calibration_path, float_precision='round_trip')
    raw_outcomes = main_module.load_dataset(str(csv_path))[[main_module.STUDY_ID_COL, main_module.TIME_COL, main_module.EVENT_COL]]
    outcome_check = participant_calibration[[main_module.STUDY_ID_COL, main_module.TIME_COL, main_module.EVENT_COL]].merge(raw_outcomes, on=main_module.STUDY_ID_COL, suffixes=('_saved', '_raw'), validate='one_to_one')
    require(len(outcome_check) == EXPECTED_PARTICIPANTS, 'Participant calibration outcomes do not match the raw cohort.')
    for column in [main_module.TIME_COL, main_module.EVENT_COL]:
        require(np.array_equal(outcome_check[f'{column}_saved'].to_numpy(), outcome_check[f'{column}_raw'].to_numpy()), f'Participant calibration {column} differs from the raw cohort.')
    print('Calculating conditional participant-bootstrap intervals for O:E and slope...', flush=True)
    calibration_summary, calibration_bootstrap, calibration_completeness = conditional_scalar_calibration_bootstrap(participant_calibration, main_module, n_bootstrap=args.calibration_bootstrap_reps, seed=args.seed, horizon_months=args.horizon_months)
    verify_scalar_bootstrap_summary(calibration_summary, calibration_bootstrap)
    original_calibration = pd.read_csv(main_output_dir / 'tables' / 'repeated_200_participant_level_calibration_summary.csv', float_precision='round_trip').set_index('model')
    for row in calibration_summary.to_dict('records'):
        model_name = str(row['model'])
        require(math.isclose(float(row['calibration_oe_ratio']), float(original_calibration.loc[model_name, 'calibration_oe_ratio']), rel_tol=1e-12, abs_tol=1e-12) and math.isclose(float(row['calibration_slope']), float(original_calibration.loc[model_name, 'calibration_slope']), rel_tol=1e-12, abs_tol=1e-12), f'Calibration point estimate does not match the primary output for {model_name}.')
    paths = output_paths(outdir)
    atomic_write_csv(pooled_cox, paths['pooled_cox'])
    atomic_write_csv(cox_by_imputation, paths['cox_by_imputation'])
    atomic_write_csv(ph_by_imputation, paths['ph'])
    atomic_write_csv(imputation_diagnostics, paths['imputation_diagnostics'])
    atomic_write_csv(calibration_summary, paths['calibration_summary'])
    atomic_write_csv(calibration_bootstrap, paths['calibration_bootstrap'])
    atomic_write_csv(calibration_completeness, paths['calibration_completeness'])
    formatted_table2 = make_formatted_table2(pooled_cox)
    original_table3 = pd.read_csv(main_output_dir / 'tables' / 'repeated_200_table3_formatted.csv', keep_default_na=False)
    formatted_table3 = make_formatted_table3(original_table3, calibration_summary)
    atomic_write_csv(formatted_table2, paths['formatted_table2'])
    atomic_write_csv(formatted_table3, paths['formatted_table3'])
    docx_paths = write_results_documents(outdir, main_output_dir, formatted_table2, formatted_table3)
    write_results_methods_notes(main_output_dir / 'metadata' / 'methods_notes.txt', paths['methods'])
    method_metadata = {'schema_version': SCHEMA_VERSION, 'software': {distribution: importlib.metadata.version(distribution) for distribution in sorted(LOCKED_ADDITIONAL_PACKAGES)}, 'fixed_split': {'training_participants': EXPECTED_FIXED_TRAINING_PARTICIPANTS, 'training_events': EXPECTED_FIXED_TRAINING_EVENTS, 'imputations': int(args.m_imputations), 'mice_cycles_per_chain': int(args.impute_max_iter), 'imputation_model': 'statsmodels MICEData with type-specific models and predictive mean matching', 'continuous_imputation_model': 'OLS with Gaussian parameter perturbation and predictive mean matching', 'binary_imputation_model': 'binomial GLM with Gaussian parameter perturbation and predictive mean matching', 'predictive_mean_matching_donors': PMM_DONORS, 'chains': 'independently_seeded', 'imputation_matrix': [*list(main_module.PREDICTOR_COLUMNS), EVENT_AUXILIARY_COLUMN, NELSON_AALEN_COLUMN], 'nelson_aalen_definition': 'sum_over_event_times_at_or_before_T_of_d(t)/Y(t)', 'nelson_aalen_cross_check': 'lifelines.NelsonAalenFitter(nelson_aalen_smoothing=False)', 'imputer_fit_population': 'fixed_training_set_only', 'held_out_outcomes_used': False, 'prediction_pipeline_role': 'primary_200_split_predictor_only_imputation', 'cox_pooling': 'Rubin_rules_with_t_reference_degrees_of_freedom', 'monte_carlo_precision_gate': 'maximum_sqrt_between_variance_over_m_divided_by_pooled_SE_0.10', 'missingness_assumption': 'missing_at_random_conditional_on_observed_predictors_event_and_training_set_Nelson_Aalen_H_T;_not_protection_against_MNAR', 'mice_cycle_interpretation': '20_prespecified_cycles_not_a_generic_stochastic_convergence_claim', 'fixed_split_inference_scope': 'conditional_on_one_outcome_stratified_992_participant_92_event_training_split_with_15_terms;_associational_not_causal', 'ph_test_handling': 'complete_imputation_level_results_and_descriptive_min_median_max;_not_pooled_as_independent', 'nelson_aalen_minimum': float(np.min(hazard)), 'nelson_aalen_maximum': float(np.max(hazard))}, 'scalar_calibration': {'models': EXPECTED_MODELS, 'participants': EXPECTED_PARTICIPANTS, 'bootstrap_reps': int(args.calibration_bootstrap_reps), 'resampling_unit': 'participant', 'interval': 'percentile_2.5th_and_97.5th_percentiles', 'predictions_refitted_inside_bootstrap': False, 'fit_warning_policy': 'convergence_and_runtime_warnings_fail_the_replicate;_all_other_warnings_are_recorded_in_the_replicate_audit', 'interpretation': 'conditional_on_saved_participant_averaged_predictions_and_excludes_model_development_and_split_stream_selection_uncertainty'}, 'references': [{'citation': 'White IR, Royston P. Imputing missing covariate values for the Cox model. Stat Med. 2009;28:1982-1998.', 'doi': '10.1002/sim.3618'}, {'citation': 'White IR, Royston P, Wood AM. Multiple imputation using chained equations: Issues and guidance for practice. Stat Med. 2011;30:377-399.', 'doi': '10.1002/sim.4067'}, {'citation': 'Morris TP, White IR, Royston P. Tuning multiple imputation by predictive mean matching and local residual draws. Medical Research Methodology. 2014;14:75.', 'doi': '10.1186/1471-2288-14-75'}, {'citation': 'von Hippel PT. How many imputations do you need? A two-stage calculation using a quadratic rule. Sociol Methods Res. 2020;49:699-718.', 'doi': '10.1177/0049124117747303'}, {'citation': 'Austin PC, Harrell FE Jr, van Klaveren D. Graphical calibration curves and the integrated calibration index for survival models. Stat Med. 2020;39:2714-2742.', 'doi': '10.1002/sim.8570'}, {'citation': 'Grambsch PM, Therneau TM. Proportional hazards tests and diagnostics based on weighted residuals. Biometrika. 1994;81:515-526.', 'doi': '10.1093/biomet/81.3.515'}, {'citation': 'Rubin DB. Multiple Imputation for Nonresponse in Surveys. Wiley; 1987.', 'role': 'multiple-imputation pooling'}]}
    atomic_write_json(method_metadata, paths['method_metadata'])
    completion_without_artifacts = {'schema_version': SCHEMA_VERSION, 'status': 'complete', 'identity_sha256': identity['identity_sha256'], 'configuration': identity['configuration'], 'primary_200_split_prediction_analysis': 'complete', 'extra_trees_sensitivity_analysis': 'complete', 'secondary_spline_analyses': 'complete', 'fixed_split_cox_inference': 'complete', 'conditional_scalar_calibration_intervals': 'complete', 'output_root': outdir.name}
    public_artifacts = [path for path in [*paths.values(), *docx_paths] if path != paths['completion'] and path.is_file()]
    completion_without_artifacts['artifact_sha256'] = {path.relative_to(outdir).as_posix(): sha256_file(path) for path in sorted(set(public_artifacts))}
    atomic_write_json(completion_without_artifacts, paths['completion'])
    validate_outputs(outdir, expected_identity_sha256=identity['identity_sha256'])
    print('Final stage completed and validated.', flush=True)
if __name__ == '__main__':
    main()
''',
}


REQUIRED_LAUNCHER_PACKAGES = {
    "statsmodels": "0.14.4",
    "patsy": "1.0.1",
}


def validate_launcher_dependencies() -> None:
    for distribution, expected_version in REQUIRED_LAUNCHER_PACKAGES.items():
        try:
            actual_version = importlib.metadata.version(distribution)
        except importlib.metadata.PackageNotFoundError as exc:
            raise SystemExit(f"Required package is missing: {distribution}=={expected_version}.") from exc
        if actual_version != expected_version:
            raise SystemExit(
                f"Package version mismatch for {distribution}: "
                f"expected {expected_version}, found {actual_version}."
            )


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Run the survival analysis and sensitivity analyses.")
    parser.add_argument("--csv", required=True, help="Analysis CSV.")
    parser.add_argument("--master", required=True, help="Source Excel workbook.")
    parser.add_argument("--sheet", required=True, help="Workbook sheet containing follow-up data.")
    parser.add_argument("--outdir", required=True, help="Folder for generated results.")
    parser.add_argument(
        "--stage",
        choices=["all", "main", "sensitivity", "secondary", "final"],
        default="all",
        help="Run all stages or one selected stage.",
    )
    parser.add_argument("--n-jobs", type=int, default=1, help="Parallel workers.")
    parser.add_argument(
        "--calibration-bootstrap-reps",
        type=int,
        default=1000,
        choices=[1000],
        help="Participant bootstrap samples for calibration; fixed at 1,000.",
    )
    parser.add_argument(
        "--max-session-hours",
        type=float,
        default=0.0,
        help="Pause safely after a completed checkpoint; zero is unlimited.",
    )
    parser.add_argument("--resume", action="store_true", help="Resume validated checkpoints.")
    parser.add_argument("--validate-only", action="store_true", help="Check packages and input data without fitting.")
    parser.add_argument("--smoke-test", action="store_true", help="Run short functional checks without the full analysis.")
    parser.add_argument("--skip-shap", action="store_true", help="Skip fixed-split SHAP outputs.")
    return parser.parse_args()


def materialize(root: Path) -> dict[str, Path]:
    paths = {}
    for name, source in _SOURCES.items():
        path = root / name
        path.write_text(source, encoding="utf-8", newline="\n")
        paths[name] = path
    return paths


def run_stage(name: str, command: list[str], environment: dict[str, str]) -> None:
    print(f"Running {name} stage.", flush=True)
    result = subprocess.run(command, env=environment, check=False)
    if result.returncode == 75:
        print("The analysis paused safely. Run the same command with --resume.", flush=True)
        raise SystemExit(75)
    if result.returncode != 0:
        raise SystemExit(result.returncode)


def common_options(args: argparse.Namespace) -> list[str]:
    options = [
        "--n-jobs",
        str(args.n_jobs),
        "--calibration-bootstrap-reps",
        str(args.calibration_bootstrap_reps),
        "--max-session-hours",
        str(args.max_session_hours),
    ]
    if args.resume:
        options.append("--resume")
    return options


def main() -> None:
    args = parse_args()
    if args.n_jobs == 0:
        raise SystemExit("--n-jobs cannot be zero.")
    if args.calibration_bootstrap_reps != 1000:
        raise SystemExit("--calibration-bootstrap-reps must equal 1000.")
    if args.max_session_hours < 0.0:
        raise SystemExit("--max-session-hours must be non-negative.")
    csv_path = Path(args.csv).expanduser().resolve()
    master_path = Path(args.master).expanduser().resolve()
    outdir = Path(args.outdir).expanduser().resolve()
    if not csv_path.is_file():
        raise SystemExit("The analysis CSV was not found.")
    if not master_path.is_file():
        raise SystemExit("The source workbook was not found.")
    if not args.sheet.strip():
        raise SystemExit("--sheet cannot be empty.")
    if args.validate_only and args.smoke_test:
        raise SystemExit("Use either --validate-only or --smoke-test, not both.")
    if args.validate_only and args.stage not in {"all", "main", "secondary"}:
        raise SystemExit("--validate-only can be used with all, main, or secondary.")
    validate_launcher_dependencies()
    outdir.mkdir(parents=True, exist_ok=True)
    main_out = outdir / "main"
    sensitivity_out = outdir / "sensitivity"
    secondary_out = outdir / "secondary"
    final_out = outdir / "final"
    with tempfile.TemporaryDirectory(prefix="analysis_runtime_") as temporary:
        runtime = Path(temporary)
        modules = materialize(runtime)
        environment = os.environ.copy()
        environment["ANALYSIS_MASTER_SHEET"] = args.sheet
        environment["PYTHONPATH"] = str(runtime) + os.pathsep + environment.get("PYTHONPATH", "")
        environment["MPLCONFIGDIR"] = str(runtime / ".matplotlib")
        environment["JOBLIB_TEMP_FOLDER"] = str(runtime / ".joblib")
        python = sys.executable
        common = common_options(args)
        commands = {
            "main": [
                python,
                str(modules["_analysis_main.py"]),
                "--csv",
                str(csv_path),
                "--outdir",
                str(main_out),
                *common,
                *(["--skip-shap"] if args.skip_shap else []),
            ],
            "sensitivity": [
                python,
                str(modules["_analysis_sensitivity.py"]),
                "--csv",
                str(csv_path),
                "--main-script",
                str(modules["_analysis_main.py"]),
                "--main-output-dir",
                str(main_out),
                "--outdir",
                str(sensitivity_out),
                *common,
            ],
            "secondary": [
                python,
                str(modules["_analysis_secondary.py"]),
                "--csv",
                str(csv_path),
                "--master",
                str(master_path),
                "--main-output-dir",
                str(main_out),
                "--outdir",
                str(secondary_out),
                *common,
            ],
            "final": [
                python,
                str(modules["_analysis_final.py"]),
                "--csv",
                str(csv_path),
                "--main-script",
                str(modules["_analysis_main.py"]),
                "--main-output-dir",
                str(main_out),
                "--secondary-output-dir",
                str(secondary_out),
                "--outdir",
                str(final_out),
                "--calibration-bootstrap-reps",
                str(args.calibration_bootstrap_reps),
                "--max-session-hours",
                str(args.max_session_hours),
                *(["--resume"] if args.resume else []),
            ],
        }
        if args.validate_only:
            selected = ["main", "secondary"] if args.stage == "all" else [args.stage]
            for name in selected:
                run_stage(name, [*commands[name], "--validate-only"], environment)
            print("Validation passed.", flush=True)
            return
        if args.smoke_test:
            smoke_flags = {
                "main": "--functional-smoke-test",
                "sensitivity": "--worker-import-smoke-test",
                "secondary": "--functional-smoke-test",
                "final": "--functional-smoke-test",
            }
            selected = ["main", "sensitivity", "secondary", "final"] if args.stage == "all" else [args.stage]
            for name in selected:
                run_stage(name, [*commands[name], smoke_flags[name]], environment)
            print("Smoke tests passed.", flush=True)
            return
        selected = ["main", "sensitivity", "secondary", "final"] if args.stage == "all" else [args.stage]
        for name in selected:
            run_stage(name, commands[name], environment)
    print("Analysis complete.", flush=True)


if __name__ == "__main__":
    main()
